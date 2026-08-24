from __future__ import annotations

import asyncio
import base64
import csv
import hashlib
import hmac
import io
import ipaddress
import json
import logging
import math
import os
import re
import secrets
import threading
import time
import zipfile
from collections import deque
from contextlib import asynccontextmanager
from datetime import datetime, timezone
from difflib import SequenceMatcher
from pathlib import Path
from typing import Any, Callable
from urllib.parse import quote

from fastapi import Body, FastAPI, File, Form, Header, HTTPException, Query, Request, UploadFile
from fastapi.responses import FileResponse, JSONResponse, StreamingResponse
from docx import Document
from docx.shared import Pt

from app.init_db import init_db
from app.repository import get_posting as repo_get_posting
from app.repository import list_postings as repo_list_postings
from app.repository import recommend_postings as repo_recommend_postings
from app.repository import record_audit_log
from app.repository import create_question_quality_run
from app.repository import create_review_session, get_review_session, prune_review_sessions
from app.repository import QuestionQualityReviewConflictError
from app.repository import get_question_quality_run
from app.repository import list_question_quality_eval_cases
from app.repository import list_question_quality_feedback
from app.repository import promote_question_quality_eval_case
from app.repository import question_quality_metrics
from app.repository import record_question_quality_review
from app.repository import rollback_question_quality_review
from app.repository import verify_question_quality_run_token
from app.services.ai_strategy import build_strategy_with_openai, rank_postings_with_openai
from app.services.external_api import fetch_ncs, fetch_ncs_highschool_course, fetch_public_inst, fetch_recruitment
from app.services.kordoc_parser import (
    KordocParseError,
    _is_full_ncs_code,
    _strip_full_ncs_code_prefix,
    parse_with_kordoc,
    structure_job_description,
    structure_job_notice,
)
from app.services.hwp_text_fallback import (
    extract_hwp_text,
    extract_hwpx_text,
    extract_linear_ncs_classification_terms,
)
from app.services.ncs_mcp_client import (
    NcsMcpError,
    classify_official_ability_unit_names,
    classify_official_detail_names,
    derive_detail_candidates_from_exact_ability_scopes,
    exact_official_units_by_name,
    ncs_mcp_status,
    search_units_by_detail,
    suggest_units_by_text,
    use_ncs_mcp_request_session,
)
from app.services.jd_strategy import (
    _planned_question_sequence_for_prompt,
    ai_extract_ncs_cl_codes,
    build_notice_context_from_jd,
    build_ncs_context_pack,
    build_strategy_with_openai as build_jd_strategy_with_openai,
    extract_detail_categories_from_jd,  # noqa: F401 - retained as a legacy monkeypatch boundary
    extract_small_categories_from_jd,
    extract_subcategory_text,
    extract_pdf_text,
    extract_focus_terms_from_pdf_vision,
    fetch_ncs_ksa_by_units,
    fetch_ncs_ksa_by_sclass_code,
    fetch_ncs_units_hrdk_by_cl_codes,
    fetch_ncs_units_hrdk_by_keywords,
    fetch_ncs_units_hrdk_by_sclass_names,
    fetch_ncs_units_hrdk_by_verified_sclass,
    generate_interview_questions_by_ncs_code,
    generate_personalized_interview_questions,
    generate_diverse_interview_questions,
    rank_ksa_factors_by_query,
    infer_keywords_from_subcategory_ai,
    is_similar_question_text,
    lookup_ncs_codes_by_sclass,
    normalize_question_dedup_key,
    review_ocr_terms_with_openai,
    rerank_ncs_matches,
)
from app.services.ncs import map_ncs
from app.services.provider_config import (
    GenerationCredentialError,
    OPENROUTER_PROVIDER,
    configured_generation_provider,
    generation_provider_config,
    normalize_generation_provider,
    openrouter_recovery_model,
    openrouter_reasoning_effort,
    provider_timeout_sec,
    resolve_generation_credential,
    resolve_generation_model,
    request_key_error_code,
    request_key_error_message,
    request_supported_generation_providers,
    sanitize_generation_model,
    use_generation_request,
)
from app.services.openai_quality_config import openai_role_model
from app.services.question_intent import (
    FOCUS_SCOPED_GENERAL_QUESTION_INTENTS,
    GENERAL_QUESTION_INTENTS,
    QUESTION_INTENT_PATTERNS,
    classify_question_intent,
)
from app.services.question_quality_ops import (
    QUALITY_POLICY_VERSION,
    derive_quality_control,
    feedback_prompt_context,
    sanitize_feedback_payload,
)
from app.services.question_quality_orchestrator import (
    RUNTIME_QUESTION_ORCHESTRATION_POLICY,
    evaluate_ksa_measurement,
    orchestrate_question_set,
)
from app.services.question_precision_grounding import (
    PRECISION_GROUNDING_POLICY,
    evaluate_question_precision_grounding,
)
from app.services.question_evaluation_alignment import (
    EVALUATION_ELICITATION_POLICY,
    evaluate_evaluation_elicitation_alignment,
)
from app.services.question_realism import (
    REALISM_POLICY_VERSION,
    evaluate_question_realism,
)
from app.services.question_surface import (
    build_question_task_frame,
    has_dangling_surface,
    normalize_ksa_type as _canonical_normalize_ksa_type,
    official_ksa_surface_aliases,
    public_task_object,
    replace_official_ksa_surface,
    stable_ksa_evidence_id,
)
from app.services.ai_question_quality_review import (
    review_interview_questions_with_ai,
)
from app.services.external_ai_privacy import sanitize_external_ai_source_text
from app.services.request_budget import use_request_budget
from app.services.auto_runner import start_auto_runner
from app.services.ax_readiness import assess_ax_readiness
from app.services.alio_ingestion import (
    AlioIngestionError,
    download_alio_attachment,
    inspect_alio_url,
)
from app.services.alio_sclass_profile import suggest_sclass_from_profile
from app.services.queue_manager import QueueManager
from app.services.sclass_pipeline import (
    extract_pdf_text_fallback,
    extract_sclass_from_pdf_bytes,
    extract_sclass_from_text,
    resolve_sclass_candidates_bundle,
)
from app.services.sync_workers import sync_ncs_units, sync_public_institutions
from app.settings import settings


class _RequestBodyTooLarge(Exception):
    pass


class RequestBodyLimitMiddleware:
    def __init__(self, app, limit_bytes: int | None = None):
        self.app = app
        self.limit_bytes = limit_bytes

    async def __call__(self, scope, receive, send):
        if scope.get("type") != "http":
            await self.app(scope, receive, send)
            return

        limit = int(self.limit_bytes or settings.max_request_body_bytes())
        headers = {
            bytes(key).lower(): bytes(value)
            for key, value in (scope.get("headers") or [])
        }
        raw_content_length = headers.get(b"content-length", b"").decode(
            "ascii", errors="ignore"
        )
        try:
            content_length = int(raw_content_length) if raw_content_length else 0
        except ValueError:
            content_length = 0
        if content_length > limit:
            response = JSONResponse(
                status_code=413,
                content={"detail": "request body exceeds MAX_REQUEST_BODY_MB"},
            )
            await response(scope, receive, send)
            return

        received = 0

        async def limited_receive():
            nonlocal received
            message = await receive()
            if message.get("type") == "http.request":
                received += len(message.get("body") or b"")
                if received > limit:
                    raise _RequestBodyTooLarge
            return message

        try:
            await self.app(scope, limited_receive, send)
        except _RequestBodyTooLarge:
            response = JSONResponse(
                status_code=413,
                content={"detail": "request body exceeds MAX_REQUEST_BODY_MB"},
            )
            await response(scope, receive, send)


class JsonCharsetMiddleware:
    """Make UTF-8 explicit for legacy Windows HTTP clients.

    JSON is UTF-8 by specification, but Windows PowerShell 5 can decode an
    ``application/json`` response using a legacy code page when no charset is
    present.  A reviewed Korean question then no longer matches its server-side
    SHA-256 hash when the client posts it back.
    """

    def __init__(self, app):
        self.app = app

    async def __call__(self, scope, receive, send):
        if scope.get("type") != "http":
            await self.app(scope, receive, send)
            return

        async def send_with_json_charset(message):
            if message.get("type") == "http.response.start":
                headers = []
                for raw_key, raw_value in message.get("headers") or []:
                    key = bytes(raw_key)
                    value = bytes(raw_value)
                    if key.lower() == b"content-type":
                        media_type = value.split(b";", 1)[0].strip().lower()
                        if (
                            media_type == b"application/json"
                            or media_type.endswith(b"+json")
                        ) and b"charset=" not in value.lower():
                            value = value + b"; charset=utf-8"
                    headers.append((key, value))
                message = {**message, "headers": headers}
            await send(message)

        await self.app(scope, receive, send_with_json_charset)


def _generation_request_budget_sec() -> float:
    try:
        configured = float(os.getenv("GENERATION_REQUEST_BUDGET_SEC", "285"))
    except (TypeError, ValueError):
        configured = 285.0
    # Keep 15 seconds for safe error mapping, serialization, and the response
    # hop below this deployment's 300-second function duration. Public
    # generation never uses that reserve to synthesize a server fallback.
    return max(30.0, min(285.0, configured))


class ExpensiveRequestLimitMiddleware:
    """Apply lightweight local backpressure to document/question APIs.

    This is intentionally process-local and only protects the application when
    it is reached directly. Production deployments should also enforce the
    same limits at the reverse proxy or shared rate-limit service.
    """

    _EXPENSIVE_PREFIXES = ("/api/jd/", "/api/questions/", "/api/alio/")
    _EXPENSIVE_PATHS = frozenset({"/api/ncs/units/options"})

    def __init__(self, app):
        self.app = app
        self._lock = threading.Lock()
        self._events: dict[tuple[str, str], deque[float]] = {}
        self._generation_slots = threading.BoundedSemaphore(settings.generation_max_concurrency())

    @classmethod
    def _is_expensive(cls, path: str) -> bool:
        return path.startswith(cls._EXPENSIVE_PREFIXES) or path in cls._EXPENSIVE_PATHS

    @staticmethod
    def _client_key(scope) -> str:
        client = scope.get("client")
        if isinstance(client, (tuple, list)) and client:
            return str(client[0] or "unknown")
        return "unknown"

    def _consume_rate_limit(self, key: tuple[str, str]) -> tuple[bool, int]:
        now = time.monotonic()
        window = settings.rate_limit_window_sec()
        limit = (
            settings.generation_rate_limit_requests_per_window()
            if key[1] == "generation"
            else settings.rate_limit_requests_per_window()
        )
        with self._lock:
            events = self._events.setdefault(key, deque())
            while events and now - events[0] >= window:
                events.popleft()
            if len(events) >= limit:
                retry_after = max(1, int(window - (now - events[0])))
                return False, retry_after
            events.append(now)
            if len(self._events) > 10_000:
                stale_keys = [
                    event_key
                    for event_key, values in self._events.items()
                    if not values or now - values[-1] >= window
                ]
                for event_key in stale_keys[:2_000]:
                    self._events.pop(event_key, None)
            return True, 0

    async def __call__(self, scope, receive, send):
        if scope.get("type") != "http" or not settings.rate_limit_enabled():
            await self.app(scope, receive, send)
            return

        path = str(scope.get("path") or "")
        if not self._is_expensive(path):
            await self.app(scope, receive, send)
            return

        is_generation = scope.get("method") == "POST" and (
            path.startswith("/api/jd/") or path.startswith("/api/questions/")
        )
        bucket = "generation" if is_generation else "expensive"
        allowed, retry_after = self._consume_rate_limit((self._client_key(scope), bucket))
        if not allowed:
            response = JSONResponse(
                status_code=429,
                content={"detail": "request rate limit exceeded"},
                headers={"Retry-After": str(retry_after)},
            )
            await response(scope, receive, send)
            return

        slot_acquired = False
        if is_generation:
            slot_acquired = self._generation_slots.acquire(blocking=False)
            if not slot_acquired:
                response = JSONResponse(
                    status_code=429,
                    content={"detail": "generation concurrency limit reached"},
                    headers={"Retry-After": "1"},
                )
                await response(scope, receive, send)
                return
        budgeted_generation = scope.get("method") == "POST" and (
            path == "/api/jd/strategy/upload"
            or path.startswith("/api/questions/generate")
        )
        try:
            if budgeted_generation:
                with use_request_budget(_generation_request_budget_sec()):
                    await self.app(scope, receive, send)
            else:
                await self.app(scope, receive, send)
        finally:
            if slot_acquired:
                self._generation_slots.release()


@asynccontextmanager
async def _lifespan(_app: FastAPI):
    init_db()
    start_auto_runner()
    yield


APP_VERSION = "1.4.8"
app = FastAPI(title="NCScope", version=APP_VERSION, lifespan=_lifespan)
app.add_middleware(RequestBodyLimitMiddleware)
app.add_middleware(JsonCharsetMiddleware)
app.add_middleware(ExpensiveRequestLimitMiddleware)
queue = QueueManager(max_retries=2)
logger = logging.getLogger("ncscope")


def _internal_http_error(
    code: str,
    exc: BaseException,
    *,
    status_code: int = 500,
) -> HTTPException:
    reference = secrets.token_hex(6)
    logger.error(
        "%s reference=%s",
        code,
        reference,
        exc_info=(type(exc), exc, exc.__traceback__),
    )
    return HTTPException(
        status_code=status_code,
        detail={"code": code, "reference": reference},
    )


# Middleware: Disable caching for all question generation endpoints
@app.middleware("http")
async def add_no_cache_headers(request, call_next):
    """Ensure no caching for dynamic question generation APIs"""
    # NCS matching and KSA collection make several MCP tool calls. Keep one
    # initialized, connection-pooled transport for the whole web request.
    with use_ncs_mcp_request_session():
        response = await call_next(request)

    response.headers["Referrer-Policy"] = "no-referrer"
    response.headers["X-Frame-Options"] = "DENY"
    response.headers["Permissions-Policy"] = "camera=(), microphone=(), geolocation=()"
    if request.url.path == "/":
        response.headers["Content-Security-Policy"] = (
            "default-src 'self'; base-uri 'none'; object-src 'none'; "
            "frame-ancestors 'none'; form-action 'self'; connect-src 'self'; "
            "img-src 'self' data:; style-src 'self' 'unsafe-inline'; "
            "script-src 'self' 'unsafe-inline'"
        )

    if (
        request.url.path.startswith("/api/questions/")
        or request.url.path.startswith("/api/jd/")
        or request.url.path == "/"
    ):
        response.headers["Cache-Control"] = "no-store, no-cache, must-revalidate, max-age=0"
        response.headers["Pragma"] = "no-cache"
        response.headers["Expires"] = "0"
        response.headers["X-Content-Type-Options"] = "nosniff"

    return response
BASE_DIR = Path(__file__).resolve().parent
UI_INDEX = BASE_DIR / "static" / "index.html"
_ALIO_CACHE: dict[str, dict] = {}
NCS_SCLASS_CSV = BASE_DIR.parent / "ncs_sclass_codes_with_code_no.csv"
_SCLASS_OPTIONS_CACHE: list[dict] | None = None
_NON_DETAIL_HIERARCHY_NAMES_CACHE: tuple[str, ...] | None = None
# 면접 질문 생성 최적 고정값 (요청당 최대 5개 기준)
FAST_NCS_TOP_K = 4          # NCS 매칭 상위 4개
FAST_KSA_UNITS = 3          # 한 배치의 공식 KSA 근거를 제공하는 능력단위 수
FAST_KSA_FACTORS_PER_UNIT = 2  # 단위당 KSA 2개 (총 6개)
# Keep one public generation request within a single bounded provider batch.
# More questions use the existing offset/history-based regeneration flow.
GENERATION_MAX_MAIN_QUESTIONS = 5
SUPPORTED_INTERVIEW_METHODS = (
    "경험면접",
    "상황면접",
    "발표면접",
    "토론면접",
    "인바스켓면접",
    "직무지식면접",
    "창의적 문제해결력면접",
)
OPTIONAL_INTERVIEW_METHODS: tuple[str, ...] = ()
QUALITY_INTERVIEW_METHODS = SUPPORTED_INTERVIEW_METHODS + OPTIONAL_INTERVIEW_METHODS
OPERATIONAL_REVIEW_NOTICE = (
    "NCScope output is a KSA-grounded structured-interview draft. "
    "A human reviewer must confirm final questions against blind-hiring rules "
    "and institution-specific evaluation standards."
)
MODEL_PRESERVED_QUESTION_SOURCES = {
    "openai_api",
    "openai_api_quality_repaired_fields",
    "openrouter_api",
    "openrouter_api_quality_repaired_fields",
    "codex_cli",
    "codex_cli_quality_repaired_fields",
    "claude_code",
    "claude_code_quality_repaired_fields",
    "model",
    "model_main_template_followups",
    "model_main_repaired_followups",
    "model_main_quality_repaired_fields",
}


def _subscription_cli_source_base(value: Any) -> str:
    """Return a trusted free-form model source with exact evidence metadata."""

    source = str(value or "").strip()
    for prefix in ("openai_api", "openrouter_api", "codex_cli", "claude_code"):
        if source == prefix or source.startswith(f"{prefix}_"):
            return prefix
    return ""


def _is_subscription_cli_source(value: Any) -> bool:
    return bool(_subscription_cli_source_base(value))
_BLIND_HIRING_CUE_RE = re.compile(
    r"(가족|부모|형제|배우자|자녀|나이|연령|출신\s*학교|학교명|학벌|출신\s*지역|출신지역|고향|"
    r"생년\s*월일|출생\s*(?:연도|년도|일|지)|몇\s*살|만\s*\d+\s*세|"
    r"혼인|결혼|기혼|미혼|결혼\s*여부|혼인\s*상태|임신|출산|자녀\s*계획|출산\s*계획|"
    r"외모|용모|(?:키|신장)\s*(?:가|는|를|와|및|/|,|:|：|\d)|체중|성별|종교|정치\s*성향|"
    r"병역|군필|미필|군\s*복무|복무\s*기간|전역|혈액형)"
)


def _contains_blind_hiring_cue(value: Any) -> bool:
    return bool(_BLIND_HIRING_CUE_RE.search(str(value or "")))


def _clamp_runtime_knobs(
    ncs_top_k: int | str | None,
    ksa_units: int | str | None,
    ksa_factors_per_unit: int | str | None,
) -> tuple[int, int, int]:
    def _to_int(v: int | str | None, default: int) -> int:
        try:
            return int(str(v).strip())
        except Exception:
            return default

    top_k = max(1, min(8, _to_int(ncs_top_k, FAST_NCS_TOP_K)))
    units = max(1, min(7, _to_int(ksa_units, FAST_KSA_UNITS)))
    factors = max(1, min(4, _to_int(ksa_factors_per_unit, FAST_KSA_FACTORS_PER_UNIT)))
    return top_k, units, factors


def _clamp_sclass_limit(value: int | str | None, default: int = 4) -> int:
    try:
        v = int(str(value).strip())
    except Exception:
        v = int(default)
    return max(1, min(6, v))


def _to_float_or(value: str | None, default: float) -> float:
    try:
        return float(str(value).strip())
    except Exception:
        return float(default)


def _clamp_int(value: int | str | None, default: int, lo: int, hi: int) -> int:
    try:
        v = int(str(value).strip())
    except Exception:
        v = int(default)
    return max(int(lo), min(int(hi), v))


def _coerce_bool_flag(value: Any, *, default: bool = False) -> bool:
    if isinstance(value, bool):
        return value
    if value is None:
        return default
    if isinstance(value, (int, float)):
        return int(value) != 0
    text = str(value).strip().lower()
    if not text:
        return default
    if text in {"1", "true", "t", "yes", "y", "on", "enabled"}:
        return True
    if text in {"0", "false", "f", "no", "n", "off", "disabled"}:
        return False
    return default


def _effective_partial_model_output(
    value: Any,
    *,
    default: bool | None = None,
) -> bool:
    return _coerce_bool_flag(
        value,
        default=bool(default if default is not None else False),
    )


def _coerce_positive_int(value: Any) -> int | None:
    try:
        parsed = int(str(value).strip())
    except Exception:
        return None
    if parsed <= 0:
        return None
    return parsed


def _norm_sclass_key(value: str) -> str:
    return re.sub(r"\s+", "", str(value or "")).strip().lower()


def _normalize_ncs_detail_term(value: Any) -> str:
    """Return a single human-confirmable NCS detail label.

    OCR/Kordoc sometimes returns a whole table row as the detail candidate,
    e.g. ``(전기안전관리) 01.... 05....``.  Passing that row downstream makes
    the model repeat eligibility text, page markers and unrelated legal notes
    in every interview question.  Only collapse long numbered rows; ordinary
    detail names (including names containing parentheses) are preserved.
    """

    text = re.sub(r"\s+", " ", str(value or "").strip())
    if not text:
        return ""
    numbered = re.search(r"(?:^|\s)0?1\s*[.)]", text) and re.search(
        r"(?:^|\s)0?2\s*[.)]", text
    )
    if len(text) >= 60 and numbered:
        parenthesized = re.search(r"\(([^()\n]{2,80})\)", text)
        if parenthesized:
            candidate = re.sub(r"\s+", " ", parenthesized.group(1)).strip(" -:：")
            if candidate:
                return candidate
        prefix = re.split(r"\s+(?=0?1\s*[.)])", text, maxsplit=1)[0]
        prefix = re.sub(r"^(?:공고\s*[·ㆍ./]?\s*직무\s*기술\s*서\s*상?|NCS\s*세분류\s*[:：]?)\s*", "", prefix, flags=re.IGNORECASE)
        prefix = prefix.strip(" -:：()")
        if 2 <= len(prefix) <= 80:
            return prefix
    return text


def _parse_sclass_terms(raw: str | None) -> list[str]:
    protected_slash = "\ufff0"
    split_value = re.sub(
        r"(?<=[A-Za-z])/(?=[A-Za-z])",
        protected_slash,
        str(raw or ""),
    )
    parts = [
        part.replace(protected_slash, "/")
        for part in re.split(r"[\n,;/|]+", split_value)
    ]
    out: list[str] = []
    seen: set[str] = set()
    for part in parts:
        term = _normalize_ncs_detail_term(part)
        if not term:
            continue
        key = _norm_sclass_key(term)
        if not key or key in seen:
            continue
        seen.add(key)
        out.append(term)
    return out


def _norm_detail_coverage_key(value: Any) -> str:
    return re.sub(r"[\s\-\_/|(),.·・]+", "", str(value or "")).strip().lower()


def _reviewed_ability_unit_names(values: Any) -> list[str]:
    """Normalize human-reviewed ability-unit cells for exact MCP matching.

    Official recruitment pages and institution JDs commonly render a unit as
    either ``01.문서 작성`` or as the final segment of a full NCS hierarchy.
    Only transport decoration is removed here; semantic/fuzzy rewriting is
    intentionally forbidden because these names become an exact unit lock.
    """

    raw_values = values if isinstance(values, list) else [values]
    output: list[str] = []
    seen: set[str] = set()
    for raw_value in raw_values:
        if isinstance(raw_value, dict):
            raw_value = raw_value.get("text") or raw_value.get("name") or ""
        for raw_line in re.split(r"[\n;|]+", str(raw_value or "")):
            text = re.sub(r"<[^>]+>", " ", raw_line).strip()
            text = re.sub(r"^(?:[-*•○●▪▫◦]+\s*)+", "", text).strip()
            text = re.sub(r"^(?:요구\s*)?능력단위(?:명)?\s*[:：\-]?\s*", "", text).strip()
            if ">" in text:
                text = text.rsplit(">", 1)[-1].strip()
            text = re.sub(
                r"^(?:NCS\s*)?\d{2}(?:[-\s]?\d{2}){4}(?:_[0-9A-Za-z]+)?\s*[:：.)\-]?\s*",
                "",
                text,
                flags=re.IGNORECASE,
            ).strip()
            text = re.sub(r"^\d{1,3}\s*[.)：:\-]\s*", "", text).strip()
            if not text or text in {"-", "해당없음", "없음", "미개발"}:
                continue
            key = _norm_detail_coverage_key(text)
            if len(key) < 2 or key in seen:
                continue
            seen.add(key)
            output.append(text)
    return output


def _reviewed_ability_unit_ordinals(fields: dict[str, Any]) -> dict[str, list[str]]:
    """Return parser-evidenced ability-unit ordinals keyed by exact name."""

    output: dict[str, list[str]] = {}
    positioned = fields.get("positioned_items") if isinstance(fields, dict) else []
    for item in positioned if isinstance(positioned, list) else []:
        if not isinstance(item, dict) or item.get("section") != "ability_units":
            continue
        name = str(item.get("text") or "").strip()
        ordinal_raw = str(item.get("ability_unit_ordinal") or "").strip()
        key = _norm_detail_coverage_key(name)
        if not key or not re.fullmatch(r"\d{1,2}", ordinal_raw):
            continue
        ordinal = ordinal_raw.zfill(2)
        bucket = output.setdefault(key, [])
        if ordinal not in bucket:
            bucket.append(ordinal)
    return output


def _lock_units_to_reviewed_ability_units(
    units: list[dict[str, Any]],
    reviewed_names: list[str],
) -> tuple[list[dict[str, Any]], list[str]]:
    """Return exact-name NCS units only when scope and base code are unique.

    An official ability-unit name can occur under multiple NCS details.  A
    reviewed label is therefore not enough to select a code globally: all
    exact-name rows must converge on one eight-digit detail-code prefix and one
    ten-digit ability-unit base code before the unit can be locked.  Ambiguous
    names remain unresolved and are never allowed to widen the downstream KSA
    lookup.  Multiple versions of the same base code remain compatible.
    """

    rows_by_name: dict[str, list[dict[str, Any]]] = {}
    for unit in units or []:
        if not isinstance(unit, dict):
            continue
        unit_name = str(unit.get("compeUnitName") or unit.get("unit_name") or "").strip()
        key = _norm_detail_coverage_key(unit_name)
        if key:
            rows_by_name.setdefault(key, []).append(unit)

    locked: list[dict[str, Any]] = []
    missing: list[str] = []
    seen_codes: set[str] = set()
    for reviewed_name in reviewed_names:
        matches = rows_by_name.get(_norm_detail_coverage_key(reviewed_name), [])
        if not matches:
            missing.append(reviewed_name)
            continue
        detail_scopes = {
            match.group(1)
            for unit in matches
            for match in [
                re.match(
                    r"^(\d{8})\d{2}(?:_|$)",
                    str(unit.get("ncsClCd") or unit.get("unit_code") or "").strip(),
                )
            ]
            if match
        }
        base_codes = {
            match.group(1)
            for unit in matches
            for match in [
                re.match(
                    r"^(\d{10})(?:_|$)",
                    str(unit.get("ncsClCd") or unit.get("unit_code") or "").strip(),
                )
            ]
            if match
        }
        # A missing code scope is not authoritative, while two or more scopes
        # mean the same reviewed name belongs to multiple selected details. A
        # second base code is also ambiguous even when both codes share a
        # detail, matching the official-catalog classifier contract.
        if len(detail_scopes) != 1 or len(base_codes) != 1:
            missing.append(reviewed_name)
            continue
        for match in matches:
            code = str(match.get("ncsClCd") or match.get("unit_code") or "").strip()
            if not code or code in seen_codes:
                continue
            seen_codes.add(code)
            locked.append(
                {
                    **match,
                    "requiredAbilityUnitName": reviewed_name,
                    "requiredAbilityUnitMatch": "exact",
                }
            )
    return locked, missing


def _scope_reviewed_ability_units_by_exact_detail_membership(
    detail_units_by_name: dict[str, list[dict[str, Any]]],
    reviewed_names: list[str],
) -> tuple[dict[str, list[str]], list[str]]:
    """Scope unassigned unit names only when one selected detail owns them.

    Some recruitment tables visually merge the detail cell across several
    ability-unit rows. Kordoc can still recover every unit label while the
    row-to-detail edge is absent. This helper reconstructs that edge from the
    authoritative unit lists of the details already extracted from the same
    document. It deliberately accepts normalized exact names only and leaves
    names that occur under zero or multiple details unresolved.
    """

    details_by_unit_name: dict[str, set[str]] = {}
    for detail_name, units in (detail_units_by_name or {}).items():
        detail = str(detail_name or "").strip()
        if not detail:
            continue
        for unit in units or []:
            if not isinstance(unit, dict):
                continue
            unit_name = str(
                unit.get("compeUnitName") or unit.get("unit_name") or ""
            ).strip()
            unit_key = _norm_detail_coverage_key(unit_name)
            if unit_key:
                details_by_unit_name.setdefault(unit_key, set()).add(detail)

    scoped: dict[str, list[str]] = {}
    unresolved: list[str] = []
    for reviewed_name in reviewed_names or []:
        matches = details_by_unit_name.get(
            _norm_detail_coverage_key(reviewed_name), set()
        )
        if len(matches) != 1:
            unresolved.append(reviewed_name)
            continue
        detail = next(iter(matches))
        scoped.setdefault(detail, []).append(reviewed_name)
    return scoped, unresolved


def _recover_ordinal_scoped_reviewed_ability_units(
    detail_units: list[dict[str, Any]],
    missing_names: list[str],
    ordinals_by_name: dict[str, list[str]],
    *,
    already_locked: list[dict[str, Any]] | None = None,
) -> tuple[list[dict[str, Any]], list[str]]:
    """Recover a renamed current unit from its table ordinal and code scope.

    Recruitment JDs often retain an earlier label while preserving the
    official two-digit unit ordinal. Recovery requires a unique base unit code
    in the already exact detail scope and a containment-only name change of at
    most six normalized characters. Ordinal-only or fuzzy matching is not
    accepted.
    """

    def code_parts(unit: dict[str, Any]) -> tuple[str, str]:
        code = str(unit.get("ncsClCd") or unit.get("unit_code") or "").strip()
        match = re.match(r"^(\d{8})(\d{2})(?:_|$)", code)
        return (match.group(1) + match.group(2), match.group(2)) if match else ("", "")

    detail_prefixes = {
        base_code[:8]
        for unit in detail_units or []
        if isinstance(unit, dict)
        for base_code, _ordinal in [code_parts(unit)]
        if base_code
    }
    if len(detail_prefixes) != 1:
        return [], list(missing_names or [])

    seen_codes = {
        str(unit.get("ncsClCd") or unit.get("unit_code") or "").strip()
        for unit in (already_locked or [])
        if isinstance(unit, dict)
    }
    recovered: list[dict[str, Any]] = []
    remaining: list[str] = []
    for required_name in missing_names or []:
        required_key = _norm_detail_coverage_key(required_name)
        ordinals = set(ordinals_by_name.get(required_key) or [])
        candidates: list[dict[str, Any]] = []
        for unit in detail_units or []:
            if not isinstance(unit, dict):
                continue
            base_code, ordinal = code_parts(unit)
            if not base_code or ordinal not in ordinals:
                continue
            current_name = str(unit.get("compeUnitName") or unit.get("unit_name") or "").strip()
            current_key = _norm_detail_coverage_key(current_name)
            if not current_key:
                continue
            shorter, longer = sorted((required_key, current_key), key=len)
            if shorter not in longer or len(longer) - len(shorter) > 6:
                continue
            candidates.append(unit)
        base_codes = {code_parts(unit)[0] for unit in candidates if code_parts(unit)[0]}
        if len(base_codes) != 1:
            remaining.append(required_name)
            continue
        accepted = False
        for unit in candidates:
            code = str(unit.get("ncsClCd") or unit.get("unit_code") or "").strip()
            if not code or code in seen_codes:
                continue
            seen_codes.add(code)
            accepted = True
            recovered.append(
                {
                    **unit,
                    "requiredAbilityUnitName": required_name,
                    "requiredAbilityUnitMatch": "ordinal_code_scope_current_name",
                    "source": "ncs-mcp-required-unit-ordinal-recovery",
                }
            )
        if not accepted:
            remaining.append(required_name)
    return recovered, remaining


def _recover_code_scoped_reviewed_ability_units(
    detail_units: list[dict[str, Any]],
    missing_names: list[str],
    candidate_units: list[dict[str, Any]],
    *,
    already_locked: list[dict[str, Any]] | None = None,
) -> tuple[list[dict[str, Any]], list[str]]:
    """Recover exact unit names when the MCP classification link is corrupt.

    The official NCS unit code starts with the eight-digit detail code. A small
    number of serving-DB rows can have a stale ``classification_id`` even
    though the unit code itself is correct. Recovery is allowed only when the
    required name remains exact and its code prefix belongs to an already
    exact-resolved detail scope. Semantic similarity and the returned path are
    never used as authority.
    """

    def detail_code(unit: dict[str, Any]) -> str:
        code = str(unit.get("ncsClCd") or unit.get("unit_code") or "").strip()
        match = re.match(r"^(\d{8})\d{2}(?:_|$)", code)
        return match.group(1) if match else ""

    allowed_detail_codes = {
        detail_code(unit)
        for unit in (detail_units or [])
        if isinstance(unit, dict) and detail_code(unit)
    }
    if not allowed_detail_codes:
        return [], list(missing_names or [])

    rows_by_name: dict[str, list[dict[str, Any]]] = {}
    for unit in candidate_units or []:
        if not isinstance(unit, dict) or detail_code(unit) not in allowed_detail_codes:
            continue
        name = str(unit.get("compeUnitName") or unit.get("unit_name") or "").strip()
        key = _norm_detail_coverage_key(name)
        if key:
            rows_by_name.setdefault(key, []).append(unit)

    seen_codes = {
        str(unit.get("ncsClCd") or unit.get("unit_code") or "").strip()
        for unit in (already_locked or [])
        if isinstance(unit, dict)
    }
    recovered: list[dict[str, Any]] = []
    remaining: list[str] = []
    for required_name in missing_names or []:
        matches = rows_by_name.get(_norm_detail_coverage_key(required_name), [])
        match_detail_codes = {
            detail_code(match) for match in matches if detail_code(match)
        }
        if len(match_detail_codes) != 1:
            remaining.append(required_name)
            continue
        accepted = False
        for match in matches:
            code = str(match.get("ncsClCd") or match.get("unit_code") or "").strip()
            if not code or code in seen_codes:
                continue
            seen_codes.add(code)
            accepted = True
            recovered.append(
                {
                    **match,
                    "requiredAbilityUnitName": required_name,
                    "requiredAbilityUnitMatch": "exact_name_code_scope_recovery",
                    "source": "ncs-mcp-required-unit-code-scope-recovery",
                }
            )
        if not accepted:
            remaining.append(required_name)
    return recovered, remaining


def _canonicalize_detail_lookup_terms(lookup_terms: list[str]) -> list[str]:
    """Use official NCS spellings for exact-equivalent reviewed labels.

    PDF tables often split an official label across lines (for example,
    ``프로젝트 \uad00\ub9ac``).  The NCS search service tokenizes that transport
    spelling differently from the official ``프로젝트\uad00\ub9ac`` label, even though
    they are the same classification.  Resolve only exact normalized matches
    from the bundled official catalog; unknown or merely similar labels remain
    untouched so they still require human selection.
    """

    parsed_terms = _parse_sclass_terms(
        "\n".join(str(term or "").strip() for term in (lookup_terms or []))
    )
    reviewed_terms: list[str] = []
    reviewed_keys: set[str] = set()
    for parsed_term in parsed_terms:
        # Manual review and structural fallback paths can bypass the Kordoc
        # candidate cleaner.  Apply the same transport-only cleanup at this
        # final public boundary; the code is trace metadata, not a detail name.
        term = re.sub(
            r"^(?:NCS\s*)?세분류(?:명|\(\s*직무(?:명)?\s*\))?\s*[:：\-]\s*",
            "",
            parsed_term,
            flags=re.IGNORECASE,
        ).strip()
        if _is_full_ncs_code(term):
            continue
        term = _strip_full_ncs_code_prefix(term)
        if _is_full_ncs_code(term):
            continue
        term = re.sub(r"^\d{1,3}\s*[.)：:\-]\s*", "", term).strip()
        if not term or _is_full_ncs_code(term):
            continue
        key = _norm_sclass_key(term)
        if not key or key in reviewed_keys:
            continue
        reviewed_keys.add(key)
        reviewed_terms.append(term)
    if not reviewed_terms:
        return []

    def catalog_variants(term: str) -> list[str]:
        variants = [term]
        without_header = re.sub(
            r"^(?:NCS\s*)?\uc138\ubd84\ub958(?:\uba85)?\s*[:：\-]\s*",
            "",
            term,
            flags=re.IGNORECASE,
        ).strip()
        without_number = re.sub(
            r"^\d{1,3}\s*[.)：:\-]\s*",
            "",
            without_header,
        ).strip()
        for value in (without_header, without_number):
            if value and value not in variants:
                variants.append(value)
        return variants

    variants_by_term = {
        term: catalog_variants(term)
        for term in reviewed_terms
    }
    catalog_queries = [
        variant
        for term in reviewed_terms
        for variant in variants_by_term[term]
    ]
    official_by_key: dict[str, str] = {}
    for row in lookup_ncs_codes_by_sclass(catalog_queries):
        if not isinstance(row, dict):
            continue
        official_name = str(row.get("sclass_name", "")).strip()
        key = _norm_detail_coverage_key(official_name)
        if official_name and key:
            official_by_key.setdefault(key, official_name)

    resolved_decorations_by_key: dict[str, str] = {}
    for row in classify_official_detail_names(reviewed_terms):
        if not isinstance(row, dict):
            continue
        official_names = [
            str(value or "").strip()
            for value in (row.get("officialDetailNames") or [])
            if str(value or "").strip()
        ]
        source_key = _norm_detail_coverage_key(str(row.get("sourceName") or ""))
        if (
            source_key
            and row.get("mappingState") == "official_current_exact"
            and row.get("resolvedCatalogExact") is True
            and len(official_names) == 1
        ):
            resolved_decorations_by_key[source_key] = official_names[0]

    canonical_terms: list[str] = []
    for term in reviewed_terms:
        official_name = next(
            (
                official_by_key[key]
                for variant in variants_by_term[term]
                if (key := _norm_detail_coverage_key(variant)) in official_by_key
            ),
            "",
        )
        if not official_name:
            official_name = resolved_decorations_by_key.get(
                _norm_detail_coverage_key(term),
                "",
            )
        canonical_terms.append(official_name or term)
    return canonical_terms


def _non_detail_hierarchy_names() -> tuple[str, ...]:
    """Return official 대분류/중분류 labels that cannot be treated as 세분류."""

    global _NON_DETAIL_HIERARCHY_NAMES_CACHE
    if _NON_DETAIL_HIERARCHY_NAMES_CACHE is not None:
        return _NON_DETAIL_HIERARCHY_NAMES_CACHE
    names: list[str] = []
    seen: set[str] = set()
    if NCS_SCLASS_CSV.exists():
        with NCS_SCLASS_CSV.open("r", encoding="utf-8-sig", newline="") as handle:
            for row in csv.DictReader(handle):
                for field in ("NCS_LCLAS_CDNM", "NCS_MCLAS_CDNM"):
                    name = str(row.get(field, "")).strip()
                    key = _norm_detail_coverage_key(name)
                    if name and key and key not in seen:
                        seen.add(key)
                        names.append(name)
    _NON_DETAIL_HIERARCHY_NAMES_CACHE = tuple(names)
    return _NON_DETAIL_HIERARCHY_NAMES_CACHE


def _used_local_hangul_text_fallback(parsed: dict[str, Any]) -> bool:
    metadata = parsed.get("metadata") if isinstance(parsed.get("metadata"), dict) else {}
    fallback = str(metadata.get("fallback") or "").casefold()
    if fallback in {"hwp-text", "hwpx-text"}:
        return True
    return any(
        str(member.get("fallback") or "").casefold() in {"hwp-text", "hwpx-text"}
        for member in (metadata.get("members") or [])
        if isinstance(member, dict)
    )


def _recover_hangul_fallback_detail_candidates(
    parsed: dict[str, Any],
    structured_fields: dict[str, Any],
) -> None:
    """Use official MCP exact matches to classify flattened document tables.

    HWP/HWPX and PDF fallback parsers can all preserve the table text while
    losing its cell boundaries.  The local extractor only proposes labels;
    this boundary releases a candidate solely when the NCS MCP resolves it as
    an official exact detail.
    """

    metadata = parsed.get("metadata") if isinstance(parsed.get("metadata"), dict) else {}
    supplemental_terms: list[str] = [
        str(value or "").strip()
        for value in (metadata.get("hangul_classification_terms") or [])
        if str(value or "").strip()
    ]
    generic_terms: list[str] = [
        str(value or "").strip()
        for value in (metadata.get("classification_terms") or [])
        if str(value or "").strip()
    ]
    for member in metadata.get("members") or []:
        if not isinstance(member, dict):
            continue
        supplemental_terms.extend(
            str(value or "").strip()
            for value in (member.get("hangul_classification_terms") or [])
            if str(value or "").strip()
        )
        generic_terms.extend(
            str(value or "").strip()
            for value in (member.get("classification_terms") or [])
            if str(value or "").strip()
        )
    existing_candidates = [
        str(value or "").strip()
        for value in (structured_fields.get("ncs_detail_candidates") or [])
        if str(value or "").strip()
    ]
    # A single PDF fallback already ran the structural table extractor. The
    # old enrichment pass queried every flattened classification term even
    # though it only appended results and never validated or removed an
    # existing candidate. Generation revalidates the human-selected label
    # against NCS MCP, so repeating broad MCP searches here adds latency but
    # no release authority. Keep recovery for HWP/HWPX and mixed ZIP inputs.
    if (
        str(metadata.get("fallback") or "").casefold() == "pdf-text"
        and not metadata.get("members")
        and existing_candidates
        and not supplemental_terms
    ):
        return
    used_fallback = _used_local_hangul_text_fallback(parsed)
    if not used_fallback and not supplemental_terms and not generic_terms:
        return
    possible_terms = _parse_sclass_terms(
        "\n".join(
            [
                *supplemental_terms,
                *generic_terms,
                *(
                    extract_linear_ncs_classification_terms(
                        str(parsed.get("markdown") or ""),
                        excluded_hierarchy_names=_non_detail_hierarchy_names(),
                        limit=40,
                    )
                    if used_fallback
                    else []
                ),
            ]
        )
    )
    if not possible_terms:
        return
    try:
        # The client stops once the global unit cap is reached. A broad ZIP
        # can contain many valid 세분류 labels, and popular early labels (총무,
        # 경영기획, 유지관리) may otherwise consume the fixed 200-unit budget
        # before later members are queried at all. Scale the cap with the term
        # count while keeping a hard bound on MCP response size.
        mcp_unit_limit = min(1000, max(200, len(possible_terms) * 40))
        units = search_units_by_detail(possible_terms, max_units=mcp_unit_limit)
    except NcsMcpError as exc:
        logger.warning("hwp_text_fallback_detail_mcp_failed: %s", exc)
        return

    recovered: list[str] = []
    seen: set[str] = set()
    for unit in units:
        if not isinstance(unit, dict):
            continue
        detail = str(unit.get("ncsSubdCdnm") or unit.get("resolvedDetailName") or "").strip()
        key = _norm_detail_coverage_key(detail)
        if detail and key and key not in seen:
            seen.add(key)
            recovered.append(detail)
    if not recovered:
        return

    existing = [
        str(value or "").strip()
        for value in (structured_fields.get("ncs_detail_candidates") or [])
        if str(value or "").strip()
    ]
    merged = _parse_sclass_terms("\n".join([*existing, *recovered]))
    structured_fields["ncs_detail_candidates"] = merged
    exact_source = (
        "hwp_text_mcp_exact"
        if used_fallback or supplemental_terms
        else "pdf_text_mcp_exact"
    )
    structured_fields["ncs_detail_source"] = exact_source
    structured_fields["ncs_detail_candidate_evidence"] = [
        {
            "detail": detail,
            "text": detail,
            "source": exact_source,
            "line": 0,
        }
        for detail in merged
    ]
    structured_fields["ncs_detail_absence_reason"] = ""
    structured_fields["ncs_detail_absence_state"] = ""
    structured_fields["ncs_detail_absence_evidence"] = ""
    structured_fields["ncs_detail_absence_filtered_candidate_reason"] = ""
    structured_fields["ncs_detail_absence_saw_ncs_table"] = False
    structured_fields["ncs_detail_absence_saw_detail_header"] = False
    structured_fields["ncs_detail_absence_blank_or_dash_detail_cell"] = False
    structured_fields["ncs_detail_absence_declared_no_mapping"] = False


def _has_non_marker_review_text(value: Any) -> bool:
    text = str(value or "").strip()
    if not text:
        return False
    text = re.sub(r"^(?:[-*•○●□■▪▫◦ㅇ¡]+\s*)+", "", text)
    text = re.sub(r"(?:\s*[-*•○●□■▪▫◦ㅇ¡])+\s*$", "", text).strip()
    return bool(text)


def _sanitize_parse_review_ability_artifacts(structured: dict[str, Any]) -> None:
    structured_fields = structured.get("fields")
    if not isinstance(structured_fields, dict):
        return
    ability_units = [
        str(value or "").strip()
        for value in (structured_fields.get("ability_units") or [])
        if _has_non_marker_review_text(value)
    ]
    structured_fields["ability_units"] = ability_units

    source_scopes = structured_fields.get("ability_units_by_detail")
    if isinstance(source_scopes, dict):
        structured_fields["ability_units_by_detail"] = {
            str(detail): [
                str(value or "").strip()
                for value in values
                if _has_non_marker_review_text(value)
            ]
            for detail, values in source_scopes.items()
            if isinstance(values, list)
            and any(_has_non_marker_review_text(value) for value in values)
        }

    positioned_items = structured_fields.get("positioned_items")
    if isinstance(positioned_items, list):
        structured_fields["positioned_items"] = [
            item
            for item in positioned_items
            if not (
                isinstance(item, dict)
                and item.get("section") == "ability_units"
                and not _has_non_marker_review_text(item.get("text"))
            )
        ]

    sections = structured.get("sections")
    if isinstance(sections, dict) and isinstance(sections.get("ability_units"), list):
        sections["ability_units"] = [
            item
            for item in sections["ability_units"]
            if not (
                isinstance(item, dict)
                and not _has_non_marker_review_text(item.get("text"))
            )
        ]


def _detail_lookup_coverage(
    lookup_terms: list[str],
    ncs_items: list[dict[str, Any]] | None,
) -> tuple[list[str], list[str]]:
    requested: dict[str, str] = {}
    normalized_lookup_terms = _parse_sclass_terms(
        "\n".join(str(term or "").strip() for term in (lookup_terms or []))
    )
    for term in normalized_lookup_terms:
        text = str(term or "").strip()
        key = _norm_detail_coverage_key(text)
        if text and key and key not in requested:
            requested[key] = text
    if not requested:
        return [], []

    covered_keys: set[str] = set()
    for row in ncs_items or []:
        if not isinstance(row, dict):
            continue
        for field in ("matchedDetailName", "reviewed_detail", "confirmed_detail", "ncs_detail", "ncsSubdCdnm"):
            key = _norm_detail_coverage_key(row.get(field))
            if key in requested:
                covered_keys.add(key)

    matched = [term for key, term in requested.items() if key in covered_keys]
    unmatched = [term for key, term in requested.items() if key not in covered_keys]
    return matched, unmatched


def _merge_sclass_terms(
    base_terms: list[str],
    add_terms: list[str] | None = None,
    remove_terms: list[str] | None = None,
) -> list[str]:
    out: list[str] = []
    seen: set[str] = set()
    for term in (base_terms or []) + (add_terms or []):
        key = _norm_sclass_key(term)
        if not key or key in seen:
            continue
        seen.add(key)
        out.append(str(term).strip())

    remove_keys = {_norm_sclass_key(x) for x in (remove_terms or []) if _norm_sclass_key(x)}
    if not remove_keys:
        return out
    return [x for x in out if _norm_sclass_key(x) not in remove_keys]


def _merge_review_text(*values: Any, max_chars: int = 3000) -> str:
    out: list[str] = []
    seen: set[str] = set()
    for value in values:
        parts = value if isinstance(value, list) else re.split(r"\n+", str(value or ""))
        for part in parts:
            text = str(part or "").strip()
            key = re.sub(r"\s+", "", text).lower()
            if not text or key in seen:
                continue
            seen.add(key)
            out.append(text)
    return "\n".join(out)[:max_chars].strip()


def _parse_question_plan_json(raw: str, reviewed_detail_terms: list[str]) -> dict[str, Any]:
    fallback_terms = _parse_sclass_terms("\n".join(str(term).strip() for term in (reviewed_detail_terms or []) if str(term).strip()))
    default_items = [
        {"detail": term, "enabled": True, "main_count": 1, "follow_up_count": 3}
        for term in fallback_terms
    ]
    if not str(raw or "").strip():
        items = default_items
    else:
        try:
            parsed = json.loads(raw)
        except json.JSONDecodeError as exc:
            raise HTTPException(status_code=400, detail=f"question_plan_json is invalid: {exc}") from exc
        candidates = parsed.get("items") if isinstance(parsed, dict) else parsed if isinstance(parsed, list) else []
        items = []
        for row in candidates or []:
            if not isinstance(row, dict):
                continue
            detail_value = str(
                row.get("detail") or row.get("name") or row.get("ncs_detail") or ""
            ).strip()
            details = _parse_sclass_terms(detail_value)
            if not details:
                continue
            enabled = row.get("enabled", True)
            enabled_bool = not (enabled is False or str(enabled).strip().lower() in {"0", "false", "no", "n"})
            try:
                main_count = int(row.get("main_count", row.get("question_count", 1)) or 0)
            except Exception:
                main_count = 1
            try:
                follow_up_count = int(row.get("follow_up_count", row.get("followups", 3)) or 0)
            except Exception:
                follow_up_count = 3
            for detail in details:
                items.append(
                    {
                        "detail": detail,
                        "enabled": enabled_bool,
                        "main_count": max(0, min(50, main_count)),
                        "follow_up_count": max(1, min(5, follow_up_count)),
                    }
                )

    seen: set[str] = set()
    normalized: list[dict[str, Any]] = []
    for item in items:
        key = _norm_sclass_key(str(item.get("detail", "")))
        if not key or key in seen:
            continue
        seen.add(key)
        main_count = max(0, min(50, int(item.get("main_count", 0) or 0)))
        normalized.append(
            {
                "detail": str(item.get("detail", "")).strip(),
                "enabled": bool(item.get("enabled", True)) and main_count > 0,
                "main_count": main_count,
                "follow_up_count": max(1, min(5, int(item.get("follow_up_count", 3) or 0))),
            }
        )
    selected = [item for item in normalized if item["enabled"]]
    if not selected and fallback_terms:
        selected = default_items
        normalized = default_items
    total_main = sum(int(item.get("main_count", 0) or 0) for item in selected)
    total_main = max(1, min(50, total_main)) if selected else 0
    selected_terms = [str(item.get("detail", "")).strip() for item in selected if str(item.get("detail", "")).strip()]
    follow_up_count = max([int(item.get("follow_up_count", 3) or 0) for item in selected] or [3])
    question_sequence: list[dict[str, Any]] = []
    for item in selected:
        for _ in range(max(0, int(item.get("main_count", 0) or 0))):
            question_sequence.append(
                {
                    "detail": str(item.get("detail", "")).strip(),
                    "follow_up_count": max(1, min(5, int(item.get("follow_up_count", 3) or 0))),
                }
            )
    return {
        "items": normalized,
        "selected_items": selected,
        "selected_terms": selected_terms,
        "question_sequence": question_sequence[:50],
        "total_main_count": total_main,
        "follow_up_count": max(1, min(5, follow_up_count)),
    }


def _generation_main_question_capacity() -> int:
    try:
        configured_maximum = int(
            str(
                os.getenv(
                    "GENERATION_MAX_MAIN_QUESTIONS",
                    str(GENERATION_MAX_MAIN_QUESTIONS),
                )
            ).strip()
        )
    except (TypeError, ValueError):
        configured_maximum = GENERATION_MAX_MAIN_QUESTIONS
    return max(1, min(GENERATION_MAX_MAIN_QUESTIONS, configured_maximum))


def _enforce_question_plan_capacity(question_plan: dict[str, Any]) -> None:
    """Reject work that cannot reliably finish inside the public request window."""

    selected_items = (
        question_plan.get("selected_items")
        if isinstance(question_plan, dict)
        else []
    )
    requested_main_questions = sum(
        max(0, int(item.get("main_count", 0) or 0))
        for item in (selected_items or [])
        if isinstance(item, dict) and bool(item.get("enabled", True))
    )
    selected_detail_count = sum(
        1
        for item in (selected_items or [])
        if isinstance(item, dict)
        and bool(item.get("enabled", True))
        and int(item.get("main_count", 0) or 0) > 0
    )
    if selected_detail_count > 1:
        raise HTTPException(
            status_code=422,
            detail={
                "code": "ncs_detail_capacity_exceeded",
                "message": "세분류는 한 번에 하나만 선택할 수 있습니다.",
                "requested_ncs_details": selected_detail_count,
                "max_ncs_details": 1,
                "retryable": False,
            },
        )
    maximum = _generation_main_question_capacity()
    if requested_main_questions <= maximum:
        return
    raise HTTPException(
        status_code=422,
        detail={
            "code": "question_plan_capacity_exceeded",
            "message": (
                "AI 안정성을 위해 주질문은 한 번에 최대 "
                f"{maximum}개까지 생성할 수 있습니다. "
                f"{maximum}개 이하로 줄인 뒤, 결과 화면의 다른 질문 생성 기능으로 이어서 생성해 주세요."
            ),
            "requested_main_questions": requested_main_questions,
            "max_main_questions": maximum,
            "retryable": False,
        },
    )


def _validated_auxiliary_generation_count(
    value: Any,
    *,
    field_name: str,
) -> int:
    """Validate auxiliary question counts before NCS or model work starts."""

    if isinstance(value, bool):
        parsed = 0
    else:
        try:
            parsed = int(str(value).strip())
        except (TypeError, ValueError):
            parsed = 0
    if parsed < 1:
        raise HTTPException(
            status_code=422,
            detail={
                "code": "question_count_invalid",
                "message": f"{field_name} must be a positive integer",
                "field": field_name,
                "retryable": False,
            },
        )
    _enforce_question_plan_capacity(
        {
            "selected_items": [
                {"enabled": True, "main_count": parsed},
            ]
        }
    )
    return parsed


_MAX_SELECTED_NCS_ITEMS = 5
_MAX_NCS_CODE_CHARS = 32
_MAX_NCS_NAME_CHARS = 120
_MAX_NCS_DEFINITION_CHARS = 1000
_MAX_STRENGTHS_CHARS = 2000
_MAX_NOTICE_TEXT_CHARS = 12000
_MAX_DUTY_TEXT_CHARS = 3000
_MAX_QUALIFICATION_TEXT_CHARS = 2400
_MAX_PREFERENCE_TEXT_CHARS = 2400
_MAX_EVALUATION_TEXT_CHARS = 2400
_MAX_ABILITY_UNIT_TEXT_CHARS = 6000
_MAX_PRESENTATION_MATERIAL_TEXT_CHARS = 6000
_MAX_PRESENTATION_MATERIAL_FILE_BYTES = 2 * 1024 * 1024
_MAX_GENERATION_AVOID_QUESTION_ITEMS = 50
_MAX_GENERATION_AVOID_QUESTION_CHARS = 300
_SAFE_NCS_CODE_RE = re.compile(r"[0-9]{4,20}(?:_[0-9]{2}v[0-9]{1,2})?", re.IGNORECASE)


def _raise_generation_input_capacity_error(
    *,
    field_name: str,
    limit: int,
    actual: int,
    reason: str = "capacity_exceeded",
    message: str = "",
) -> None:
    raise HTTPException(
        status_code=422,
        detail={
            "code": "generation_input_capacity_exceeded",
            "field": field_name,
            "limit": int(limit),
            "actual": int(actual),
            "reason": reason,
            "message": message or f"{field_name} exceeds the generation input limit of {limit}",
            "retryable": False,
        },
    )


def _validate_generation_text_input(
    value: Any,
    *,
    field_name: str,
    max_chars: int,
) -> str:
    raw = "" if value is None else str(value)
    if len(raw) > max_chars:
        _raise_generation_input_capacity_error(
            field_name=field_name,
            limit=max_chars,
            actual=len(raw),
        )
    return raw.strip()


def _validate_generation_text_collection_input(
    value: Any,
    *,
    field_name: str,
    max_chars: int,
) -> None:
    parts = value if isinstance(value, list) else [value]
    raw = "\n".join(str(part or "") for part in parts)
    _validate_generation_text_input(
        raw,
        field_name=field_name,
        max_chars=max_chars,
    )


def _validate_generation_ncs_code(value: Any, *, field_name: str = "ncs_code") -> str:
    code = _validate_generation_text_input(
        value,
        field_name=field_name,
        max_chars=_MAX_NCS_CODE_CHARS,
    )
    if code and not _SAFE_NCS_CODE_RE.fullmatch(code):
        _raise_generation_input_capacity_error(
            field_name=field_name,
            limit=_MAX_NCS_CODE_CHARS,
            actual=len(code),
            reason="unsafe_format",
            message=(
                f"{field_name} must contain a numeric NCS code and an optional "
                "version suffix such as _25v3"
            ),
        )
    return code


def _validate_selected_ncs_generation_input(value: Any) -> Any:
    if not isinstance(value, list):
        return value
    if len(value) > _MAX_SELECTED_NCS_ITEMS:
        _raise_generation_input_capacity_error(
            field_name="selected_ncs",
            limit=_MAX_SELECTED_NCS_ITEMS,
            actual=len(value),
        )
    for index, row in enumerate(value):
        if not isinstance(row, dict):
            continue
        prefix = f"selected_ncs[{index}]"
        _validate_generation_ncs_code(
            row.get("ncsClCd", ""),
            field_name=f"{prefix}.ncsClCd",
        )
        for key in ("compeUnitName", "ncsSubdCdnm"):
            _validate_generation_text_input(
                row.get(key, ""),
                field_name=f"{prefix}.{key}",
                max_chars=_MAX_NCS_NAME_CHARS,
            )
        _validate_generation_text_input(
            row.get("compeUnitDef", ""),
            field_name=f"{prefix}.compeUnitDef",
            max_chars=_MAX_NCS_DEFINITION_CHARS,
        )
    return value


def _generation_avoid_question_items(value: Any) -> list[Any]:
    if isinstance(value, str):
        raw = value.strip()
        if not raw:
            return []
        try:
            parsed = json.loads(raw)
        except Exception:
            parsed = [part.strip() for part in re.split(r"[\r\n]+", raw) if part.strip()]
    else:
        parsed = value
    return parsed if isinstance(parsed, list) else []


def _validate_and_extract_generation_avoid_questions(
    value: Any,
    *,
    field_name: str = "avoid_questions",
) -> list[str]:
    items = _generation_avoid_question_items(value)
    if len(items) > _MAX_GENERATION_AVOID_QUESTION_ITEMS:
        _raise_generation_input_capacity_error(
            field_name=field_name,
            limit=_MAX_GENERATION_AVOID_QUESTION_ITEMS,
            actual=len(items),
        )
    for index, item in enumerate(items):
        if isinstance(item, dict):
            text_value = item.get("question") or item.get("text") or ""
        else:
            text_value = item
        raw_text = "" if text_value is None else str(text_value)
        if len(raw_text) > _MAX_GENERATION_AVOID_QUESTION_CHARS:
            _raise_generation_input_capacity_error(
                field_name=f"{field_name}[{index}]",
                limit=_MAX_GENERATION_AVOID_QUESTION_CHARS,
                actual=len(raw_text),
            )
    return _extract_question_texts(items)


def _restrict_question_plan_to_terms(question_plan: dict[str, Any], allowed_terms: list[str]) -> dict[str, Any]:
    allowed: list[str] = []
    seen_allowed: set[str] = set()
    for term in allowed_terms or []:
        text = str(term or "").strip()
        key = _norm_sclass_key(text)
        if text and key and key not in seen_allowed:
            seen_allowed.add(key)
            allowed.append(text)
    if not allowed:
        return question_plan
    allowed_keys = {_norm_sclass_key(term) for term in allowed}
    kept = [
        dict(item)
        for item in (question_plan.get("items") or [])
        if isinstance(item, dict) and _norm_sclass_key(str(item.get("detail") or "")) in allowed_keys
    ]
    if kept:
        return _parse_question_plan_json(json.dumps({"items": kept}, ensure_ascii=False), allowed)
    follow_up_count = max(1, min(5, int(question_plan.get("follow_up_count", 3) or 3)))
    return _parse_question_plan_json(
        json.dumps(
            {
                "items": [
                    {"detail": term, "enabled": True, "main_count": 1, "follow_up_count": follow_up_count}
                    for term in allowed
                ]
            },
            ensure_ascii=False,
        ),
        allowed,
    )


def _parse_interview_methods(raw: str) -> list[str]:
    allowed = {
        "behavior": "경험면접",
        "behavioral": "경험면접",
        "행동형": "경험면접",
        "행동관찰면접": "경험면접",
        "행동관찰": "경험면접",
        "경험형": "경험면접",
        "경험면접": "경험면접",
        "experience": "경험면접",
        "situation": "상황면접",
        "situational": "상황면접",
        "상황형": "상황면접",
        "상황면접": "상황면접",
        "presentation": "발표면접",
        "pt": "발표면접",
        "pt면접": "발표면접",
        "발표": "발표면접",
        "발표형": "발표면접",
        "발표면접": "발표면접",
        "discussion": "토론면접",
        "debate": "토론면접",
        "토론": "토론면접",
        "토론형": "토론면접",
        "토론면접": "토론면접",
        "토의": "토론면접",
        "토의형": "토론면접",
        "토의면접": "토론면접",
        "inbasket": "인바스켓면접",
        "in-basket": "인바스켓면접",
        "인바스켓": "인바스켓면접",
        "인바스켓형": "인바스켓면접",
        "인바스켓면접": "인바스켓면접",
        "job_knowledge": "직무지식면접",
        "knowledge": "직무지식면접",
        "직무지식": "직무지식면접",
        "직무지식형": "직무지식면접",
        "직무지식면접": "직무지식면접",
        "지식": "직무지식면접",
        "지식형": "직무지식면접",
        "지식면접": "직무지식면접",
        "creative": "창의적 문제해결력면접",
        "creative_problem_solving": "창의적 문제해결력면접",
        "problem_solving": "창의적 문제해결력면접",
        "창의": "창의적 문제해결력면접",
        "창의형": "창의적 문제해결력면접",
        "창의적문제해결": "창의적 문제해결력면접",
        "창의적문제해결력": "창의적 문제해결력면접",
        "창의적문제해결력면접": "창의적 문제해결력면접",
        "창의적 문제해결력": "창의적 문제해결력면접",
        "창의적 문제해결력면접": "창의적 문제해결력면접",
    }
    text = str(raw or "").strip()
    values: list[str] = []
    if text:
        try:
            parsed = json.loads(text)
        except Exception:
            parsed = None
        if isinstance(parsed, list):
            values = [str(x).strip() for x in parsed]
        elif isinstance(parsed, dict):
            values = [str(x).strip() for x in (parsed.get("methods") or [])]
        else:
            values = [part.strip() for part in re.split(r"[\n,;/|]+", text) if part.strip()]
    values = [value for value in values if value]
    if len(values) > 1:
        raise HTTPException(
            status_code=422,
            detail={
                "code": "interview_method_capacity_exceeded",
                "message": "면접 형태는 한 번에 하나만 선택할 수 있습니다.",
                "requested_interview_methods": len(values),
                "max_interview_methods": 1,
                "retryable": False,
            },
        )
    out: list[str] = []
    seen: set[str] = set()
    for value in values:
        mapped = allowed.get(value) or allowed.get(value.lower()) or allowed.get(_norm_sclass_key(value))
        if mapped and mapped not in seen:
            seen.add(mapped)
            out.append(mapped)
    if text and not out:
        raise HTTPException(
            status_code=422,
            detail={
                "code": "interview_method_invalid",
                "message": "지원되는 면접 형태 하나를 선택해 주세요.",
                "retryable": False,
            },
        )
    return out or ["경험면접"]


def _group_interview_questions_for_response(questions: list[dict[str, Any]]) -> list[dict[str, Any]]:
    grouped: dict[tuple[str, str], list[dict[str, Any]]] = {}
    for q in questions or []:
        comp = str((q or {}).get("competency", "")).strip() or "핵심 직무"
        code = str((q or {}).get("ncsClCd", "")).strip()
        grouped.setdefault((comp, code), []).append(
            {
                "question": str((q or {}).get("question", "")).strip(),
                "follow_ups": list((q or {}).get("follow_ups", []) or []),
                "evaluation_points": list((q or {}).get("evaluation_points", []) or []),
                "method": str((q or {}).get("method") or (q or {}).get("type") or "").strip(),
                "question_focus": str((q or {}).get("question_focus") or "").strip(),
                "ncs_traceability": dict((q or {}).get("ncs_traceability") or {})
                if isinstance((q or {}).get("ncs_traceability"), dict)
                else {},
                "question_intent": str((q or {}).get("question_intent") or "").strip(),
                "question_repeat_signature": str((q or {}).get("question_repeat_signature") or "").strip(),
                "question_repeat_duplicate": bool((q or {}).get("question_repeat_duplicate") is True),
            }
        )
    return [
        {"competency": comp, "ncsClCd": code, "questions": qset}
        for (comp, code), qset in grouped.items()
    ]


def _clean_question_text(value: Any, max_chars: int = 90) -> str:
    text = re.sub(r"\s+", " ", str(value or "")).strip()
    return text[:max_chars].strip() if text else ""


def _ksa_terms_for_question(
    ncs_ksa: list[dict[str, Any]] | None,
    ncs_code: str,
    fallback_terms: list[str] | None = None,
    limit: int = 5,
) -> list[str]:
    if not str(ncs_code or "").strip():
        return []
    out: list[str] = []
    seen: set[str] = set()

    def add(value: Any) -> None:
        text = _clean_question_text(value, max_chars=60)
        key = _ksa_key(text)
        if _contains_blind_hiring_cue(text):
            return
        if text and key and key not in seen:
            seen.add(key)
            out.append(text)

    for row in ncs_ksa or []:
        if isinstance(row, dict) and ncs_code and str(row.get("ncsClCd", "")).strip() == ncs_code:
            add(row.get("factorName"))
            if len(out) >= limit:
                return out[:limit]
    if out:
        return out[:limit]
    for term in fallback_terms or []:
        add(term)
        if len(out) >= limit:
            return out[:limit]
    return out[:limit]


def _normalize_ksa_type(value: Any, factor_name: str = "") -> str:
    """Compatibility wrapper around the single canonical KSA classifier."""

    return _canonical_normalize_ksa_type(value, factor_name)


def _operational_focus_label(focus: Any, focus_type: str = "") -> str:
    """Convert an official KSA noun label into a natural task object.

    The official factor remains unchanged in ``question_focus`` and the
    evidence payload.  Only the wording shown inside a question is shortened,
    so labels such as ``소비자 패턴분석 능력`` do not produce the unnatural
    phrase ``능력을 직접 수행``.
    """
    label = _clean_question_text(focus, max_chars=60)
    kind = _normalize_ksa_type(focus_type, label)
    if kind == "기술":
        candidate = re.sub(r"\s*(?:관련\s*)?(?:능력|기술|스킬)\s*$", "", label).strip()
    elif kind == "지식":
        candidate = re.sub(r"\s*(?:관련\s*)?지식\s*$", "", label).strip()
    else:
        candidate = label
    if has_dangling_surface(candidate):
        return label
    return candidate if len(_ksa_key(candidate)) >= 2 else label


def _evidence_row_for_focus(
    ncs_ksa: list[dict[str, Any]] | None,
    ncs_code: str,
    focus: str,
) -> dict[str, Any]:
    code = str(ncs_code or "").strip()
    focus_key = _ksa_key(focus)
    same_code: list[dict[str, Any]] = []
    for row in ncs_ksa or []:
        if not isinstance(row, dict) or str(row.get("ncsClCd") or "").strip() != code:
            continue
        same_code.append(row)
        if focus_key and _ksa_key(row.get("factorName")) == focus_key:
            return dict(row)
    return dict(same_code[0]) if same_code else {}


def _evidence_row_for_id(
    ncs_ksa: list[dict[str, Any]] | None,
    ncs_code: str,
    evidence_id: Any,
) -> dict[str, Any]:
    """Resolve a model-selected evidence id without exposing KSA labels in prompts."""

    code = str(ncs_code or "").strip()
    expected = str(evidence_id or "").strip()
    if not expected:
        return {}
    for row in ncs_ksa or []:
        if not isinstance(row, dict):
            continue
        if code and str(row.get("ncsClCd") or "").strip() != code:
            continue
        if stable_ksa_evidence_id(row) == expected:
            return dict(row)
    return {}


def _question_task_frame(
    *,
    focus: str,
    focus_type: str,
    subject: str,
    detail: str,
    comp_def: str,
    evidence_row: dict[str, Any] | None = None,
) -> dict[str, str]:
    evidence = evidence_row if isinstance(evidence_row, dict) else {}
    context = _domain_context_pack(detail=detail, subject=subject, focus=focus, comp_def=comp_def)
    return build_question_task_frame(
        evidence_row=evidence or None,
        factor_name=focus,
        ksa_type=focus_type,
        element_name=evidence.get("elementName") or evidence.get("element_name") or "",
        competency_name=subject,
        competency_definition=comp_def,
        decision_dilemma=context.get("debate") or "",
    )


def _skill_focus_action(focus: Any, form: str = "connective") -> str:
    """Choose a natural observable-action verb for a skill factor object."""

    compact = re.sub(r"\s+", "", str(focus or "")).lower()
    if re.search(r"(?:적용|활용|수행|확인|검토|작성|검증|조정|선정)절차$", compact):
        verb = "수행"
    elif re.search(r"(?:기법|방법|절차|기준|규정|모형|모델)$", compact):
        verb = "적용"
    elif re.search(r"(?:도구|장비|시스템|프로그램|소프트웨어|양식)$", compact):
        verb = "활용"
    else:
        verb = "수행"
    forms = {
        "connective": {"적용": "적용해", "활용": "활용해", "수행": "수행해"},
        "future": {"적용": "적용할", "활용": "활용할", "수행": "수행할"},
        "must": {"적용": "적용해야", "활용": "활용해야", "수행": "수행해야"},
        "past": {"적용": "적용한", "활용": "활용한", "수행": "수행한"},
    }
    return forms.get(form, forms["connective"])[verb]


def _ksa_type_for_focus(
    ncs_ksa: list[dict[str, Any]] | None,
    ncs_code: str,
    focus: str,
) -> str:
    focus_key = _ksa_key(focus)
    for row in ncs_ksa or []:
        if not isinstance(row, dict) or str(row.get("ncsClCd", "")).strip() != str(ncs_code or "").strip():
            continue
        if _ksa_key(row.get("factorName", "")) != focus_key:
            continue
        return _normalize_ksa_type(
            row.get("ksaTypeName") or row.get("factorType") or row.get("ksa_type"),
            str(row.get("factorName") or ""),
        )
    return _normalize_ksa_type("", focus)


def _select_ksa_focus_for_method(
    ncs_ksa: list[dict[str, Any]] | None,
    ncs_code: str,
    method: str,
    fallback_terms: list[str] | None = None,
) -> str:
    code = str(ncs_code or "").strip()
    candidates: list[tuple[str, str, int]] = []
    seen: set[str] = set()
    for position, row in enumerate(ncs_ksa or []):
        if not isinstance(row, dict) or str(row.get("ncsClCd", "")).strip() != code:
            continue
        factor = _clean_question_text(row.get("factorName"), max_chars=60)
        key = _ksa_key(factor)
        if not factor or not key or key in seen or _contains_blind_hiring_cue(factor):
            continue
        seen.add(key)
        candidates.append(
            (
                factor,
                _normalize_ksa_type(
                    row.get("ksaTypeName") or row.get("factorType") or row.get("ksa_type"),
                    factor,
                ),
                position,
            )
        )
    if not candidates:
        for position, term in enumerate(fallback_terms or []):
            factor = _clean_question_text(term, max_chars=60)
            key = _ksa_key(factor)
            if not factor or not key or key in seen or _contains_blind_hiring_cue(factor):
                continue
            seen.add(key)
            candidates.append((factor, _normalize_ksa_type("", factor), position))
    if not candidates:
        return ""

    # Prefer the KSA type that the selected method can expose most directly.
    # The fallback order still permits a method when an official unit contains
    # only one KSA type; the runtime measurement gate then operationalizes it.
    preference = {
        "경험면접": ("기술", "태도", "지식", ""),
        "상황면접": ("태도", "기술", "지식", ""),
        "발표면접": ("지식", "기술", "태도", ""),
        "토론면접": ("태도", "기술", "지식", ""),
        # In-basket exercises primarily expose prioritization, judgment, and
        # rule application.  A psychomotor skill (for example food plating or
        # equipment manipulation) cannot be demonstrated by sorting documents.
        "인바스켓면접": ("태도", "지식", "기술", ""),
        "직무지식면접": ("지식", "기술", "태도", ""),
        # The current creative-problem format is a verbal/data-analysis task,
        # not a hands-on work sample, so knowledge and visible trade-off
        # behavior are more defensible than an arbitrary manual skill.
        "창의적 문제해결력면접": ("지식", "태도", "기술", ""),
    }.get(method, ("기술", "지식", "태도", ""))
    rank = {kind: index for index, kind in enumerate(preference)}
    candidates.sort(key=lambda item: (rank.get(item[1], len(rank)), item[2], len(item[0])))
    return candidates[0][0]


def _infer_model_focus_from_official_ksa(
    ncs_ksa: list[dict[str, Any]] | None,
    ncs_code: str,
    question: str,
    follow_ups: list[str],
) -> str:
    code = str(ncs_code or "").strip()
    if not code:
        return ""
    compact_text = re.sub(r"\s+", "", "\n".join([str(question or ""), *follow_ups])).lower()
    if not compact_text:
        return ""
    candidates: list[str] = []
    seen: set[str] = set()
    for row in ncs_ksa or []:
        if not isinstance(row, dict) or str(row.get("ncsClCd", "")).strip() != code:
            continue
        factor = _clean_question_text(row.get("factorName"), max_chars=60)
        key = _ksa_key(factor)
        if not factor or not key or key in seen or _contains_blind_hiring_cue(factor):
            continue
        seen.add(key)
        candidates.append(factor)
    for factor in sorted(candidates, key=len, reverse=True):
        if re.sub(r"\s+", "", factor).lower() in compact_text:
            return factor
    return ""


def _pick_unit_for_detail(
    target_detail: str,
    offset: int,
    ncs_matches: list[dict[str, Any]] | None,
) -> dict[str, Any]:
    rows = [x for x in (ncs_matches or []) if isinstance(x, dict)]
    if not rows:
        return {}
    detail_key = _norm_sclass_key(target_detail)
    exact: list[dict[str, Any]] = []
    fallback: list[dict[str, Any]] = []
    if detail_key:
        for row in rows:
            authoritative_detail_keys = {
                _norm_sclass_key(str(row.get("matchedDetailName", ""))),
                _norm_sclass_key(str(row.get("reviewed_detail", ""))),
                _norm_sclass_key(str(row.get("confirmed_detail", ""))),
                _norm_sclass_key(str(row.get("ncs_detail", ""))),
                _norm_sclass_key(str(row.get("ncsSubdCdnm", ""))),
            }
            authoritative_detail_keys.discard("")
            if detail_key in authoritative_detail_keys:
                exact.append(row)
                continue
            sclass = _norm_sclass_key(str(row.get("ncsSclasCdnm", "")))
            matched = [
                _norm_sclass_key(x)
                for x in (row.get("matched_keywords") or [])
                if str(x).strip()
            ] if isinstance(row.get("matched_keywords"), list) else []
            if sclass == detail_key or detail_key in matched:
                fallback.append(row)
    pool = exact or fallback or rows
    return dict(pool[offset % len(pool)])


def _unit_matches_planned_detail(row: dict[str, Any], target_detail: str) -> bool:
    detail_key = _norm_sclass_key(target_detail)
    if not detail_key or not isinstance(row, dict):
        return False
    authoritative = {
        _norm_sclass_key(str(row.get("matchedDetailName", ""))),
        _norm_sclass_key(str(row.get("reviewed_detail", ""))),
        _norm_sclass_key(str(row.get("confirmed_detail", ""))),
        _norm_sclass_key(str(row.get("ncs_detail", ""))),
        _norm_sclass_key(str(row.get("ncsSubdCdnm", ""))),
        _norm_sclass_key(str(row.get("ncsSclasCdnm", ""))),
    }
    authoritative.discard("")
    matched_keywords = {
        _norm_sclass_key(str(value))
        for value in (row.get("matched_keywords") or [])
        if str(value).strip()
    } if isinstance(row.get("matched_keywords"), list) else set()
    return detail_key in authoritative or detail_key in matched_keywords


def _ensure_question_plan_unit_coverage(
    question_plan: dict[str, Any],
    ranked_matches: list[dict[str, Any]] | None,
    candidate_units: list[dict[str, Any]] | None,
) -> list[dict[str, Any]]:
    """Keep at least one authoritative NCS unit for every planned detail."""

    covered = _dedupe_units_by_code(ranked_matches)
    selected_terms = [
        str(value).strip()
        for value in (question_plan.get("selected_terms") or [])
        if str(value).strip()
    ] if isinstance(question_plan, dict) else []
    for detail in selected_terms:
        if any(_unit_matches_planned_detail(row, detail) for row in covered):
            continue
        candidate = next(
            (
                dict(row)
                for row in (candidate_units or [])
                if isinstance(row, dict) and _unit_matches_planned_detail(row, detail)
            ),
            {},
        )
        if not candidate:
            continue
        candidate["matchedDetailName"] = (
            str(candidate.get("matchedDetailName") or "").strip() or detail
        )
        candidate["score"] = float(candidate.get("score", 1.0) or 1.0)
        candidate["matched_keywords"] = list(
            dict.fromkeys([detail, *list(candidate.get("matched_keywords") or [])])
        )
        covered.append(candidate)
        covered = _dedupe_units_by_code(covered)
    return covered


def _method_evaluation_points(
    method: str,
    ksa_terms: list[str],
    focus_type: str = "",
    surface_focus: str = "",
) -> list[str]:
    if method == "경험면접":
        kind = _normalize_ksa_type(
            focus_type,
            next((str(term or "").strip() for term in ksa_terms if str(term or "").strip()), ""),
        )
        if kind == "지식":
            ksa_evidence = "규정 적용 근거"
        elif kind == "태도":
            return [
                "당시 상황과 본인 역할",
                "선택 근거와 직접 행동",
                "선택으로 감수한 점",
                "결과를 입증하는 기록",
            ]
        else:
            ksa_evidence = "자료·도구를 사용한 수행 순서"
        return [
            "당시 상황과 본인 역할",
            ksa_evidence,
            "선택 근거와 직접 행동",
            "결과를 입증하는 기록",
        ]
    guide = {
        "상황면접": ["핵심 사실과 규정 확인", "대안별 위험을 반영한 판단", "행동·보고 순서", "후속조치와 예방"],
        "발표면접": ["자료 근거와 현황·원인 분석", "대안 비교와 우선순위", "실행계획의 구체성", "성과지표와 질의응답 대응"],
        "토론면접": ["사실·규정에 근거한 초기 입장", "대안별 일정·책임 영향 비교", "반대 근거 검토와 쟁점 조정", "공통안 또는 미합의 이송안의 실행 가능성"],
        "인바스켓면접": ["문서별 긴급도·영향도 판단", "보고·위임·직접처리 구분", "초기 행동과 시간 배분", "위험 통제와 후속 기록"],
        "직무지식면접": ["절차·기준의 근거", "실제 업무 적용", "예외상황 판단", "산출물 품질과 오류 예방"],
        "창의적 문제해결력면접": ["근거 기반 문제 정의", "복수 대안과 창의성", "검증 방법과 실현가능성", "의사결정·실행·위험 보완"],
    }
    points = list(guide.get(method, ["구체적 근거", "실제 행동", "산출물", "확인 결과"]))
    ksa_points: list[str] = []
    for term in ksa_terms:
        if _contains_blind_hiring_cue(term):
            continue
        term_type = _normalize_ksa_type(focus_type, term)
        if surface_focus:
            public_focus = str(surface_focus).strip()
        else:
            public_focus, _surface_source = public_task_object(
                factor_name=term,
                ksa_type=term_type,
            )
        if term_type == "태도":
            point = "압박 상황에서 드러난 선택 행동과 책임"
        elif term_type == "지식":
            point = "판단 근거와 적용 범위·예외 구분"
        else:
            point = "수행 순서·산출물·품질 확인"
        if point not in points and point not in ksa_points:
            ksa_points.append(point)
        if len(ksa_points) >= 1:
            break
    if ksa_points:
        # Keep all four method anchors while making the fourth dimension carry
        # the selected KSA's observable evidence. Replacing an anchor with a
        # fifth KSA-only point either broke method validity or exceeded the
        # shared exact-four scoring contract.
        points[-1] = f"{points[-1]} · {ksa_points[0]}"
    return points[:4]


def _behavior_anchored_evaluation(
    method: str,
    focus: str,
    evaluation_points: list[str] | None = None,
    focus_type: str = "",
    surface_focus: str = "",
) -> dict[str, Any]:
    """Build an observable five-level rubric paired with an interview task."""
    official_focus = str(focus or "핵심 수행기준").strip() or "핵심 수행기준"
    focus_label = str(surface_focus or official_focus).strip() or "핵심 수행기준"
    normalized_focus_type = _normalize_ksa_type(focus_type, official_focus)
    dimensions = [str(x).strip() for x in (evaluation_points or []) if str(x).strip()][:4]
    if method == "경험면접":
        if normalized_focus_type == "지식":
            ksa_evidence = "확인한 규정·문서·자료와 적용 범위를 판단 근거로 삼은 행동"
        elif normalized_focus_type == "태도":
            ksa_evidence = "압박·이해충돌 속에서 지킨 행동 기준과 선택 이유"
        else:
            ksa_evidence = "사용한 자료·도구, 실제 수행 순서와 직접 취한 조치"
        return {
            "scale": "5단계 행동기반 평정",
            "focus": focus_label,
            "focus_type": normalized_focus_type or "미분류",
            "dimensions": dimensions,
            "rating_levels": [
                {
                    "score": 5,
                    "label": "탁월",
                    "anchor": (
                        "탁월(5): 당시 사건의 조건과 수행 과제를 분명히 구분하고, "
                        f"{ksa_evidence}, 본인의 핵심 행동, 문서·수치·기록·피드백으로 확인한 결과를 "
                        "빠짐없이 구체적으로 설명한다."
                    ),
                },
                {
                    "score": 4,
                    "label": "우수",
                    "anchor": (
                        "우수(4): 사건의 조건, 판단 근거에 따른 직접 행동, 확인 가능한 결과를 "
                        "연결해 설명하지만 자료·순서·수치 중 한 부분의 세부 근거가 제한적이다."
                    ),
                },
                {
                    "score": 3,
                    "label": "보통",
                    "anchor": (
                        "보통(3): 실제 사례와 본인이 한 행동 및 결과는 제시하나, 당시 목표·판단 근거·"
                        "확인 자료 중 두 요소가 모호하거나 서로 충분히 연결되지 않는다."
                    ),
                },
                {
                    "score": 2,
                    "label": "미흡",
                    "anchor": (
                        "미흡(2): 사례의 맥락을 구체화하지 못하거나 직접 행동을 구분하지 못하며, "
                        "결과를 확인한 문서·수치·기록·피드백도 제시하지 못한다."
                    ),
                },
                {
                    "score": 1,
                    "label": "부족",
                    "anchor": (
                        "부족(1): 구체적인 과거 사건 없이 일반적인 생각이나 자기평가만 말해 직접 행동, "
                        "판단 근거와 확인 가능한 결과를 식별할 수 없다."
                    ),
                },
            ],
            "interviewer_instruction": (
                "질문과 동일한 사건·역할·판단 근거·직접 행동·결과 증거만 기록하고, 인상이나 "
                "KSA 자기평가가 아니라 답변에서 확인된 행동 근거로 평정하십시오."
            ),
        }
    behavior = {
        "경험면접": (
            "상황·과제·본인 역할·행동·결과·학습을 구분하고 선택 기준과 결과 지표를 근거로 설명한다",
            "사례와 본인 행동은 설명하지만 선택 기준, 결과 지표 또는 학습 중 일부가 구체적이지 않다",
            "일반론이나 팀 성과만 말하며 본인 행동과 결과를 확인할 근거가 없다",
        ),
        "상황면접": (
            "확인할 사실과 판단 기준을 밝히고 위험·이해관계자를 고려한 행동 순서와 후속조치를 제시한다",
            "기본 대응 방향은 타당하지만 사실 확인, 우선순위 또는 위험 통제의 연결이 일부 부족하다",
            "근거 없이 결론부터 제시하거나 위험·보고·후속조치를 설명하지 못한다",
        ),
        "발표면접": (
            "제공자료로 현황과 원인을 구분하고 대안 비교, 실행 우선순위, 성과지표와 질의응답 근거를 제시한다",
            "현황과 대안은 제시하지만 자료 근거, 실행계획 또는 성과지표 중 일부가 모호하다",
            "자료와 무관한 주장 위주이며 원인, 실행안, 성과 확인 방법이 연결되지 않는다",
        ),
        "토론면접": (
            "근거 있는 입장을 제시하고 상대 의견을 요약·검토하며 쟁점을 조정해 실행 가능한 공동 합의안을 만든다",
            "본인 입장과 일부 조정안은 제시하지만 경청 근거 또는 최종 합의의 실행 기준이 부족하다",
            "주장만 반복하거나 상대 의견을 다루지 못해 합의안과 후속 책임을 도출하지 못한다",
        ),
        "인바스켓면접": (
            "문서별 긴급도·영향도를 근거로 우선순위를 정하고 보고·위임·직접처리와 후속 기록을 구분한다",
            "대체로 우선순위를 정하지만 일부 판단 근거나 시간·위험 통제 계획이 구체적이지 않다",
            "도착 순서대로 처리하거나 보고·위임·보류 기준 없이 항목을 나열한다",
        ),
        "직무지식면접": (
            "절차와 기준의 근거를 산출물·예외상황·품질점검·오류 예방 방법까지 연결해 설명한다",
            "기본 절차와 용어는 설명하지만 예외 판단 또는 산출물 품질 확인 방법이 일부 부족하다",
            "용어를 단편적으로 나열하고 실제 절차, 예외 또는 오류 예방에 적용하지 못한다",
        ),
        "창의적 문제해결력면접": (
            "변화 신호로 문제를 재정의하고 복수 대안을 비교해 검증방법·실현가능성·리스크 보완책을 제시한다",
            "문제와 대안은 제시하지만 검증방법, 의사결정 기준 또는 실행계획 중 일부가 모호하다",
            "아이디어를 나열할 뿐 문제 근거, 대안 비교, 검증 및 실행 가능성을 설명하지 못한다",
        ),
    }.get(method, (
        "구체적인 근거와 관찰 가능한 행동, 결과를 연결해 설명한다",
        "기본 답변은 제시하지만 근거 또는 행동의 구체성이 일부 부족하다",
        "추상적인 주장만 제시해 실제 행동과 결과를 확인할 수 없다",
    ))
    ksa_behavior = {
        "지식": (
            "해당 지식을 판단 근거로 사용하고 적용 범위·예외와 잘못 적용했을 때의 영향을 정확히 연결한다",
            "관련 기준이나 원리는 설명하지만 판단 적용, 예외 또는 영향 중 일부의 연결이 불충분하다",
            "용어나 규정을 암기해 말할 뿐 판단과 적용 결과의 근거로 사용하지 못한다",
        ),
        "기술": (
            "해당 기술을 수행한 순서와 도구·조치를 밝히고 산출물의 품질 또는 결과를 검증한 증거를 제시한다",
            "수행 행동은 설명하지만 실행 순서, 산출물 또는 품질 검증 중 일부가 불명확하다",
            "기술 보유를 주장하거나 명칭만 반복하고 실제 수행 행동과 산출물을 제시하지 못한다",
        ),
        "태도": (
            "압박이나 이해 충돌 속에서 해당 태도를 드러낸 선택과 일관된 행동, 감수한 상충비용과 결과를 제시한다",
            "바람직한 의도와 행동은 말하지만 압박 상황의 선택, 상충비용 또는 지속성 중 일부가 모호하다",
            "가치나 의지를 선언할 뿐 어려운 조건에서 실제로 한 선택과 행동의 증거가 없다",
        ),
    }.get(normalized_focus_type, (
        "해당 KSA를 실제 판단과 행동, 확인 가능한 결과에 연결한다",
        "해당 KSA와 행동의 관련성은 보이지만 근거나 결과 중 일부가 구체적이지 않다",
        "해당 KSA의 명칭이나 자기평가만 말하고 관찰 가능한 증거를 제시하지 못한다",
    ))
    level4_behavior = {
        "경험면접": "본인 역할과 핵심 행동을 구체적으로 설명하고 선택 근거와 확인 가능한 결과를 연결한다",
        "상황면접": "주요 사실과 기준을 확인해 실행 가능한 행동·보고 순서와 결과 확인 방법을 제시한다",
        "발표면접": "자료에 근거해 현황과 원인을 구분하고 우선 대안, 실행계획과 성과지표를 연결한다",
        "토론면접": "상대 근거를 정확히 요약하고 대안별 영향을 비교해 공통안 또는 이송할 쟁점을 구체화한다",
        "인바스켓면접": "긴급도와 영향도를 근거로 문서를 분류하고 보고·위임·직접처리와 후속 기록을 제시한다",
        "직무지식면접": "절차와 기준의 근거를 실제 산출물, 주요 예외와 품질점검 방법에 연결한다",
        "창의적 문제해결력면접": "근거로 문제를 정의하고 복수 대안을 비교해 검증·실행 방법과 주요 위험을 제시한다",
    }.get(method, "판단 근거와 구체적 행동, 확인 가능한 결과를 연결한다")
    level2_behavior = {
        "경험면접": "팀이 한 일을 주로 설명해 본인 판단과 행동 또는 결과 근거를 확인하기 어렵다",
        "상황면접": "일반적인 대응 방향만 말하고 사실 확인, 행동 순서와 위험 통제 중 둘 이상이 빠져 있다",
        "발표면접": "자료 일부를 인용하지만 원인과 대안, 실행계획 또는 성과지표의 연결이 끊겨 있다",
        "토론면접": "자기 주장과 결론을 반복하고 상대 근거 검토, 쟁점 조정 또는 실행 책임을 제시하지 못한다",
        "인바스켓면접": "일부 우선순위만 정하고 보고·위임·보류의 근거나 후속 기록을 제시하지 못한다",
        "직무지식면접": "관련 용어를 나열하지만 절차의 근거, 예외 적용 또는 품질점검에 연결하지 못한다",
        "창의적 문제해결력면접": "아이디어를 나열하지만 문제 근거, 대안 비교와 검증·실행 방법을 연결하지 못한다",
    }.get(method, "관련 행동을 언급하지만 판단 근거와 결과 확인 방법을 제시하지 못한다")
    ksa_level4 = {
        "지식": "관련 기준을 실제 판단에 적용하고 범위와 주요 예외를 설명하지만 영향 검토의 일부가 제한적이다",
        "기술": "수행 순서와 산출물, 품질 확인 근거를 제시하지만 한 단계의 구체성이 부족하다",
        "태도": "압박 속 선택 행동과 책임을 설명하지만 감수한 상충비용 또는 지속성의 근거가 제한적이다",
    }.get(normalized_focus_type, "해당 KSA를 실제 판단과 행동에 적용하지만 결과 근거 한 부분이 제한적이다")
    ksa_level2 = {
        "지식": "관련 용어나 기준 일부를 언급하지만 적용 대상·예외·영향을 구분하지 못한다",
        "기술": "수행했다고 주장하지만 단계·산출물·품질 확인 중 둘 이상을 제시하지 못한다",
        "태도": "바람직한 태도를 선언하지만 압박 속 선택과 감수한 결과를 제시하지 못한다",
    }.get(normalized_focus_type, "KSA 명칭은 언급하지만 실제 적용 행동과 결과를 제시하지 못한다")
    excellence_extension = {
        "경험면접": "다른 행동 대안과 결과 차이를 비교하고 같은 직무에 전이할 구체적 개선점을 제시한다",
        "상황면접": "대안별 위험과 이해관계자 영향을 비교하고 실패 시 전환·보고 기준까지 제시한다",
        "발표면접": "자료의 한계와 반대 근거를 검토하고 실행 일정·담당·측정식을 일관되게 제시한다",
        "토론면접": "양측 근거의 타당성을 검증하고 공통안 또는 이송안의 책임·점검 기준까지 합의한다",
        "인바스켓면접": "권한과 의존관계를 반영해 처리 순서를 정하고 지연 시 재조정·보고 기준까지 제시한다",
        "직무지식면접": "기준 간 충돌을 해석하고 예외 승인·산출물 검증·오류 예방 절차까지 연결한다",
        "창의적 문제해결력면접": "대안의 가설과 검증 결과를 비교하고 중단·전환 기준과 후속 실험까지 제시한다",
    }.get(method, "대안과 위험을 비교하고 검증 지표와 보완책을 일관되게 제시한다")
    rating_levels = [
        {
            "score": 5,
            "label": "탁월",
            "anchor": (
                f"탁월(5): {behavior[0]}. {ksa_behavior[0]}. "
                f"{excellence_extension}."
            ),
        },
        {
            "score": 4,
            "label": "우수",
            "anchor": f"우수(4): {level4_behavior}. {ksa_level4}.",
        },
        {
            "score": 3,
            "label": "보통",
            "anchor": f"보통(3): {behavior[1]}. {ksa_behavior[1]}.",
        },
        {
            "score": 2,
            "label": "미흡",
            "anchor": f"미흡(2): {level2_behavior}. {ksa_level2}.",
        },
        {
            "score": 1,
            "label": "부족",
            "anchor": (
                f"부족(1): 질문의 핵심 과제를 오해하거나 근거 없는 결론·자기평가만 제시한다. "
                f"{behavior[2]}. {ksa_behavior[2]}."
            ),
        },
    ]
    return {
        "scale": "5단계 행동기반 평정",
        "focus": focus_label,
        "focus_type": normalized_focus_type or "미분류",
        "dimensions": dimensions,
        "rating_levels": rating_levels,
        "interviewer_instruction": (
            "모든 지원자에게 동일한 기본 과제와 허용된 후속질문 범위를 적용하고, "
            "답변의 인상이나 KSA 자기평가보다 제시된 판단 근거·행동·산출물·결과를 기록해 평정하십시오."
        ),
    }


def _behavior_anchors_ok(guide: Any) -> bool:
    if not isinstance(guide, dict):
        return False
    rating_levels = guide.get("rating_levels") if isinstance(guide.get("rating_levels"), list) else []
    expected_levels = [(5, "탁월"), (4, "우수"), (3, "보통"), (2, "미흡"), (1, "부족")]
    actual_levels: list[tuple[int, str, str]] = []
    for level in rating_levels:
        if not isinstance(level, dict):
            continue
        try:
            score = int(level.get("score") or 0)
        except (TypeError, ValueError):
            score = 0
        actual_levels.append(
            (
                score,
                str(level.get("label") or "").strip(),
                str(level.get("anchor") or "").strip(),
            )
        )
    adjacent_similarity_ok = bool(
        len(actual_levels) == 5
        and all(
            SequenceMatcher(None, actual_levels[index][2], actual_levels[index + 1][2]).ratio() < 0.82
            for index in range(len(actual_levels) - 1)
        )
    )
    return (
        str(guide.get("scale") or "").strip() == "5단계 행동기반 평정"
        and "anchors" not in guide
        and [(score, label) for score, label, _anchor in actual_levels] == expected_levels
        and all(len(anchor) >= 40 for _score, _label, anchor in actual_levels)
        and all(anchor.endswith(".") for _score, _label, anchor in actual_levels)
        and all("영향 검토 한 부분" not in anchor for _score, _label, anchor in actual_levels)
        and adjacent_similarity_ok
        and len(str(guide.get("interviewer_instruction") or "").strip()) >= 30
    )


def _context_item_list(value: Any) -> list[str]:
    return list(
        dict.fromkeys(
            item.strip()
            for item in re.split(r"[,，]", str(value or ""))
            if item.strip()
        )
    )


def _structured_case_materials(context: dict[str, Any], facts: list[str]) -> list[dict[str, str]]:
    explicit = context.get("case_materials")
    if isinstance(explicit, list):
        rows = [
            {
                "source": str(row.get("source") or "").strip(),
                "field": str(row.get("field") or "").strip(),
                "value": str(row.get("value") or "").strip(),
            }
            for row in explicit
            if isinstance(row, dict)
        ]
        if rows and all(all(row.values()) for row in rows):
            return rows
    sources = _context_item_list(context.get("evidence")) or ["사례 현황표"]
    return [
        {
            "source": sources[index % len(sources)],
            "field": f"판단 사실 {index + 1}",
            "value": fact,
        }
        for index, fact in enumerate(facts)
    ]


def _job_context_condition_pack(
    *,
    job_context_text: str,
    detail: str,
    subject: str,
    focus: str,
    comp_def: str,
) -> dict[str, Any]:
    """Build deterministic case facts from the uploaded job sources.

    Provider timeouts and quality repairs must retain the notice/JD evidence;
    falling back to a domain profile alone can otherwise produce plausible but
    unrelated scenarios.  Keep the excerpt bounded and use generic operating
    assumptions only for the missing operational quantities.
    """

    source_text = str(job_context_text or "").strip()
    layers = re.findall(r"\[([^\]]+)\]([^\[]*)", source_text)
    anchors = [detail, subject, focus, comp_def]
    chunks: list[str] = []
    source_names: list[str] = []
    material_rows: list[dict[str, str]] = []
    for raw_label, raw_value in layers:
        label = re.sub(r"\s+", " ", str(raw_label or "").strip())
        value = str(raw_value or "").strip()
        if not value:
            continue
        segments = _presentation_source_segments(value, anchors, limit=3)
        if not segments:
            segments = [re.sub(r"\s+", " ", value)[:360]]
        source_names.append(label or "직무자료")
        selected = " ".join(segments[:3]).strip()[:700]
        if selected:
            chunks.append(selected)
            material_rows.append(
                {"source": label or "직무자료", "field": "업무 근거 발췌", "value": selected}
            )
    chunks = list(dict.fromkeys(chunk for chunk in chunks if chunk))[:6]
    source_summary = " ".join(chunks[:2])[:560]
    subject_label = subject or detail or "해당 직무"
    facts = [
        f"입력자료 근거: {chunk[:360]}"
        for chunk in chunks[:3]
    ]
    facts.extend(
        [
            "운영 시나리오 가정: 핵심 확인자료 2종을 대조해야 함",
            "운영 시나리오 가정: 처리 마감까지 4시간이 남아 있음",
            "운영 시나리오 가정: 대안 2가지를 비교하고 1개를 우선 선택해야 함",
        ]
    )
    source_names = [*source_names, "운영 시나리오"]
    material_rows.extend(
        [
            {
                "source": "운영 시나리오",
                "field": "검토 기한",
                "value": "처리 마감까지 4시간이 남아 있음",
            },
            {
                "source": "운영 시나리오",
                "field": "자료 대조·대안 비교",
                "value": "핵심 확인자료 2종을 대조하고 대안 2가지를 비교해야 함",
            },
        ]
    )
    evidence = ", ".join(dict.fromkeys(source_names)) or "입력 직무자료"
    anchor_summary = source_summary or subject_label
    return {
        "evidence": evidence,
        "situation": f"{anchor_summary}에 관한 자료·기준이 서로 달라 사실 확인과 처리 우선순위를 정해야 하는 상황",
        "inbasket": f"{anchor_summary} 관련 요청서, 검토 기록, 보고 문서",
        "debate": (
            f"{subject_label}의 근거 자료를 모두 확인한 뒤 처리하자는 입장과 "
            "마감 내 저위험 범위를 조건부 처리하고 사후 검증하자는 입장"
        ),
        "stakeholders": "공고문·직무기술서에 명시된 담당자·협업부서·승인권자",
        "case_facts": " | ".join(facts),
        "case_materials": material_rows,
    }


def _task_conditions_for_method(
    method: str,
    subject: str = "",
    focus: str = "",
    detail: str = "",
    comp_def: str = "",
    focus_type: str = "",
    variation_index: int = 0,
    job_context_text: str = "",
) -> dict[str, Any]:
    """Return standardized candidate conditions paired with the task and rubric."""
    context = (
        _job_context_condition_pack(
            job_context_text=job_context_text,
            detail=detail,
            subject=subject,
            focus=focus,
            comp_def=comp_def,
        )
        if str(job_context_text or "").strip()
        else _domain_context_pack(detail=detail, subject=subject, focus=focus, comp_def=comp_def)
    )
    evidence_materials = _context_item_list(context.get("evidence"))
    inbasket_materials = _context_item_list(context.get("inbasket"))
    base: dict[str, dict[str, Any]] = {
        "경험면접": {
            "candidate_instruction": "한 가지 실제 사례를 선택해 상황·문제·본인 역할·행동·결과 순서로 답변하십시오.",
            "time_plan": [{"phase": "개별 답변", "minutes": 5}],
            "provided_materials": ["별도 자료 없음"],
            "required_outputs": ["상황·문제와 본인 역할", "판단 근거와 본인 행동", "확인 가능한 결과"],
        },
        "상황면접": {
            "candidate_instruction": "제시된 상황에서 먼저 확인할 사실, 판단 기준, 행동 순서, 보고 및 후속조치를 구분해 답변하십시오.",
            "time_plan": [{"phase": "개별 답변", "minutes": 5}],
            "provided_materials": ["상황 시나리오"],
            "required_outputs": ["사실 확인 항목", "판단 기준과 위험 통제", "행동·보고·후속조치 순서"],
        },
        "발표면접": {
            "candidate_instruction": "제공자료만을 근거로 현황과 원인을 구분하고, 비교 가능한 대안과 실행계획을 발표한 뒤 질의응답에 답하십시오.",
            "time_plan": [
                {"phase": "준비", "minutes": 20},
                {"phase": "발표", "minutes": 5},
                {"phase": "질의응답", "minutes": 5},
            ],
            "provided_materials": evidence_materials,
            "required_outputs": ["현황·원인 진단", "대안 2가지와 우선순위", "실행계획과 성과지표"],
        },
        "토론면접": {
            "candidate_instruction": "근거 있는 초기 입장을 밝히고 상대 의견을 검토하십시오. 공통 실행안을 찾되 합의가 어렵다면 남은 쟁점과 결정권자 이송 기준을 제시하십시오.",
            "time_plan": [
                {"phase": "개별 입장발표", "minutes": 1},
                {"phase": "전체 토론", "minutes": 20},
            ],
            "provided_materials": [*evidence_materials, "토론 쟁점과 상반된 입장"],
            "required_outputs": ["초기 입장과 확인 근거", "반대 근거 검토", "공통 실행안 또는 미합의 쟁점·이송 기준"],
        },
        "인바스켓면접": {
            "candidate_instruction": "동시에 접수된 문서를 분류해 처리 우선순위와 보고·위임·직접처리 결정을 기록하고 첫 행동을 제시하십시오.",
            "time_plan": [{"phase": "문서 검토 및 의사결정", "minutes": 30}],
            "provided_materials": [*inbasket_materials, "업무분장표", "전결규정"],
            "required_outputs": ["문서별 우선순위", "보고·위임·직접처리 판단", "초기 행동과 후속점검"],
        },
        "직무지식면접": {
            "candidate_instruction": "용어만 나열하지 말고 절차와 기준을 실제 산출물, 예외상황, 품질점검 및 오류 예방에 연결해 답변하십시오.",
            "time_plan": [{"phase": "개별 답변", "minutes": 5}],
            "provided_materials": ["별도 자료 없음"],
            "required_outputs": ["절차·판단 기준과 근거", "적용 범위와 핵심 예외", "산출물 품질점검과 오류 예방"],
        },
        "창의적 문제해결력면접": {
            "candidate_instruction": "제공자료에서 변화 신호를 찾아 문제를 재정의하고, 복수 대안을 비교해 검증·실행·보완 계획을 설명하십시오.",
            "time_plan": [
                {"phase": "준비", "minutes": 20},
                {"phase": "해결안 설명", "minutes": 7},
                {"phase": "질의응답", "minutes": 5},
            ],
            "provided_materials": evidence_materials,
            "required_outputs": ["문제 정의와 원인 가설", "대안 2가지와 선택 근거", "검증 방법·우선 실행계획·성과지표"],
        },
    }
    conditions = dict(base.get(method) or base["경험면접"])
    if method == "토론면접":
        focus_kind = _normalize_ksa_type(focus_type, focus)
        if focus_kind == "지식":
            conditions["required_outputs"] = [
                "초기 입장과 확인 근거",
                "반대 입장의 수용·불수용 기준",
                "공통안의 적용 범위·예외·검증·실행 책임 또는 미합의 이송 기준",
            ]
        elif focus_kind == "기술":
            conditions["required_outputs"] = [
                "초기 입장과 확인 근거",
                "반대 입장의 수용·불수용 기준",
                "공통안의 수행 절차·품질점검·실행 책임 또는 미합의 이송 기준",
            ]
        elif focus_kind == "태도":
            conditions["required_outputs"] = [
                "초기 입장과 확인 근거",
                "반대 입장의 수용·불수용 기준",
                "공통 행동·점검·실행 책임 또는 미합의 이송 기준",
            ]
    case_facts = [
        item.strip()
        for item in str(context.get("case_facts") or "").split("|")
        if item.strip()
    ]
    case_materials = _structured_case_materials(context, case_facts)
    if method in {"상황면접", "발표면접", "토론면접", "인바스켓면접", "창의적 문제해결력면접"}:
        variation = _question_variation_constraint(variation_index)
        if variation:
            case_facts = [*case_facts, f"추가 제약: {variation}"]
            case_materials = [
                *case_materials,
                {"source": "추가 제약 카드", "field": "운영 제약", "value": variation},
            ]
        topic_axis = _question_topic_axis(variation_index)
        case_facts = [*case_facts, f"주제 축: {topic_axis}"]
        case_materials = [
            *case_materials,
            {"source": "주제 축 카드", "field": "분석 초점", "value": topic_axis},
        ]
    authority_methods = {"상황면접", "토론면접", "인바스켓면접", "창의적 문제해결력면접"}
    if method in authority_methods:
        explicit_authority_facts = context.get("authority_facts")
        authority_facts = [
            str(value or "").strip()
            for value in explicit_authority_facts
            if str(value or "").strip()
        ] if isinstance(explicit_authority_facts, list) else [
            "응시자 역할: 실무 담당자이며 사실 확인·자료 수집·초안 작성·담당자 협조 요청은 직접처리할 수 있음",
            "대외 회신·일정 변경·예외 승인은 팀장 결재가 필요하고 자료 대조는 품질담당자에게 협조 요청할 수 있음",
        ]
        explicit_authority_materials = context.get("authority_materials")
        authority_materials = [
            {
                "source": str(row.get("source") or "").strip(),
                "field": str(row.get("field") or "").strip(),
                "value": str(row.get("value") or "").strip(),
            }
            for row in explicit_authority_materials
            if isinstance(row, dict)
            and all(str(row.get(field) or "").strip() for field in ("source", "field", "value"))
        ] if isinstance(explicit_authority_materials, list) else [
            {
                "source": "업무분장표",
                "field": "응시자 역할·직접처리 범위",
                "value": "실무 담당자, 사실 확인·자료 수집·초안 작성·담당자 협조 요청",
            },
            {
                "source": "전결규정",
                "field": "팀장 결재·협조 요청 범위",
                "value": "대외 회신·일정 변경·예외 승인 / 자료 대조 협조",
            },
        ]
        case_facts = [*case_facts, *authority_facts]
        case_materials = [*case_materials, *authority_materials]
    if method in {"상황면접", "발표면접", "토론면접", "인바스켓면접", "창의적 문제해결력면접"}:
        conditions["case_facts"] = case_facts
    if method in {"상황면접", "발표면접", "토론면접", "인바스켓면접", "창의적 문제해결력면접"}:
        material_sources = [
            str(row.get("source") or "").strip()
            for row in case_materials
            if isinstance(row, dict) and str(row.get("source") or "").strip()
        ]
        conditions["provided_materials"] = list(
            dict.fromkeys([*conditions.get("provided_materials", []), *material_sources])
        )
        conditions["case_materials"] = case_materials
    # Keep the topic axis on every method, including experience and knowledge
    # questions that do not receive a printable case packet.  The axis is an
    # audit/runtime field used by candidate selection and diversity checks; it
    # does not leak NCS labels into candidate-facing wording.
    conditions["topic_axis"] = _question_topic_axis(variation_index)
    conditions["standardization"] = (
        "모든 지원자에게 동일한 자료, 기본 과제, 시간 조건과 허용된 후속질문 범위를 적용합니다."
    )
    conditions["timing_basis"] = "기관의 공고·면접 운영계획에서 확정한 동일 시간을 모든 지원자에게 적용합니다."
    return conditions


def _task_conditions_ok(method: str, conditions: Any) -> bool:
    if not isinstance(conditions, dict):
        return False
    instruction = str(conditions.get("candidate_instruction") or "").strip()
    standardization = str(conditions.get("standardization") or "").strip()
    outputs = conditions.get("required_outputs") if isinstance(conditions.get("required_outputs"), list) else []
    materials = conditions.get("provided_materials") if isinstance(conditions.get("provided_materials"), list) else []
    time_plan = conditions.get("time_plan") if isinstance(conditions.get("time_plan"), list) else []
    timed_methods = {"발표면접", "토론면접", "인바스켓면접", "창의적 문제해결력면접"}
    time_ok = method not in timed_methods or bool(
        time_plan
        and all(
            isinstance(row, dict)
            and str(row.get("phase") or "").strip()
            and isinstance(row.get("minutes"), int)
            and int(row.get("minutes") or 0) > 0
            for row in time_plan
        )
    )
    return bool(
        len(instruction) >= 25
        and len(standardization) >= 25
        and len([item for item in outputs if str(item or "").strip()]) >= 2
        and any(str(item or "").strip() for item in materials)
        and time_ok
    )


def _case_materials_sufficient_ok(method: str, conditions: Any) -> bool:
    """Require decision-grade facts for methods that claim to provide a case pack."""
    material_methods = {"상황면접", "발표면접", "토론면접", "인바스켓면접", "창의적 문제해결력면접"}
    if method not in material_methods:
        return True
    if not isinstance(conditions, dict):
        return False
    materials = [
        str(item or "").strip()
        for item in (conditions.get("provided_materials") or [])
        if str(item or "").strip()
    ] if isinstance(conditions.get("provided_materials"), list) else []
    facts = [
        str(item or "").strip()
        for item in (conditions.get("case_facts") or [])
        if str(item or "").strip()
    ] if isinstance(conditions.get("case_facts"), list) else []
    case_materials = conditions.get("case_materials") if isinstance(conditions.get("case_materials"), list) else []
    structured_rows = [
        row
        for row in case_materials
        if isinstance(row, dict)
        and all(str(row.get(field) or "").strip() for field in ("source", "field", "value"))
    ]
    quantitative_pattern = re.compile(r"(?:\d|D\+|기한|마감|예정|건|명|시간|일정|금액|비율)", re.IGNORECASE)
    quantified_facts = sum(bool(quantitative_pattern.search(fact)) for fact in facts)
    quantified_rows = sum(
        bool(quantitative_pattern.search(str(row.get("value") or "")))
        for row in structured_rows
    )
    source_coverage = all(
        str(row.get("source") or "").strip() in materials
        for row in structured_rows
    )
    return bool(
        len(materials) >= 2
        and not all(material == "별도 자료 없음" for material in materials)
        and len(facts) >= 3
        and len(set(facts)) == len(facts)
        and quantified_facts >= 2
        and len(structured_rows) >= 3
        and len(structured_rows) == len(case_materials)
        and quantified_rows >= 2
        and source_coverage
    )


def _inbasket_authority_context_ok(method: str, conditions: Any) -> bool:
    if method != "인바스켓면접":
        return True
    if not isinstance(conditions, dict):
        return False
    visible = json.dumps(conditions, ensure_ascii=False)
    return bool(
        any(marker in visible for marker in ("응시자 역할", "실무 담당자", "직급"))
        and any(marker in visible for marker in ("직접처리", "처리 한도", "위임"))
        and any(marker in visible for marker in ("전결", "결재", "승인 권한", "팀장"))
    )


def _decision_authority_context_ok(method: str, conditions: Any) -> bool:
    """Require a common role and decision boundary for individual case work."""

    authority_methods = {"상황면접", "토론면접", "인바스켓면접", "창의적 문제해결력면접"}
    if method not in authority_methods:
        return True
    if not isinstance(conditions, dict):
        return False
    visible = json.dumps(conditions, ensure_ascii=False)
    return bool(
        any(marker in visible for marker in ("응시자 역할", "실무 담당자", "직급"))
        and any(marker in visible for marker in ("직접처리", "직접 수행", "직접 수행 범위", "위임"))
        and any(marker in visible for marker in ("전결", "결재", "승인 권한", "승인권자", "팀장"))
    )


def _debate_case_neutrality_ok(method: str, question: str, conditions: Any) -> bool:
    if method != "토론면접":
        return True
    if not isinstance(conditions, dict):
        return False
    fact_text = "\n".join(
        str(item or "").strip()
        for item in (conditions.get("case_facts") or [])
        if str(item or "").strip()
    )
    material_text = "\n".join(
        str(item.get("value") or "").strip()
        for item in (conditions.get("case_materials") or [])
        if isinstance(item, dict) and str(item.get("value") or "").strip()
    )
    evidence_text = f"{fact_text}\n{material_text}"
    combined = f"{question}\n{evidence_text}"
    # A generic operating condition such as "approval criteria changed" does not
    # itself create an approval-scope dilemma. Apply this gate only when the case
    # asks candidates to decide what may happen before approval.
    approval_scope_dispute = any(
        marker in combined
        for marker in (
            "미승인",
            "승인 전",
            "승인전",
            "사전 승인",
            "승인 대기",
            "승인이 대기",
            "변경 승인이 대기",
            "승인 없이",
        )
    )
    if not approval_scope_dispute:
        return True
    answer_labels_absent = not re.search(
        r"(?:승인\s*전|사전\s*승인\s*없이|미승인(?:\s*상태)?(?:에서|로)?)"
        r"[^.\n]{0,40}(?:허용(?:됨|함)?|가능(?:함|하다)?|(?:착수|진행|수행|처리)해도\s*됨)",
        evidence_text,
    )
    ambiguity_visible = any(
        marker in combined
        for marker in ("명시하지", "명시되지", "해석", "불명확", "절차 공백", "확정될 때까지")
    )
    return bool(answer_labels_absent and ambiguity_visible)


def _focus_context_overlay(focus: str) -> dict[str, str]:
    key = re.sub(r"\s+", "", str(focus or "")).lower()
    if not key:
        return {}
    document_focus = any(token in key for token in ("문서", "기록", "자료"))
    if document_focus and any(token in key for token in ("보안", "법규", "법령", "규정", "정보보호")):
        return {
            "_replace_context": "true",
            "evidence": "문서 보안 규정, 접근권한 목록, 보존·폐기 기록, 예외 승인 내역",
            "situation": "접근권한 예외 요청, 구버전 보안 규정 혼재, 민감 문서 오배포 위험",
            "inbasket": "접근권한 변경 요청, 보존기간 만료 문서 목록, 오배포 신고, 예외 승인 검토서",
            "debate": "문서 접근·보존 통제를 엄격히 적용하는 입장과 업무 연속성을 위한 제한적 예외를 허용하는 입장",
        }
    if document_focus and any(token in key for token in ("오류", "검증", "점검", "정확", "품질")):
        return {
            "_replace_context": "true",
            "evidence": "문서 초안, 원자료, 오류 대조표, 검토 이력",
            "situation": "원자료와 문서 기준 불일치, 필수 항목 누락, 결재 마감 임박",
            "inbasket": "오류 정정 요청, 원자료 확인 메일, 검토 이력, 결재 보류 문서",
            "debate": "오류 검증과 재확인을 우선하는 입장과 결재 마감 준수를 우선하는 입장",
        }
    if any(
        token in key
        for token in ("세법", "조세", "부가가치세", "법인세", "소득세", "취득세", "재산세")
    ):
        return {
            "_replace_context": "true",
            "evidence": "과세대상 자산대장, 세액 산정표, 신고기한 자료, 예외·감면 검토서",
            "situation": "과세대상 분류 이견, 신고기한 임박, 감면 적용 근거 누락",
            "inbasket": "세액 산정 검토 요청, 자산 분류 이견, 신고기한 알림, 감면 증빙 보완 문서",
            "debate": "세법 기준을 보수적으로 적용하는 입장과 근거가 있는 감면·예외 적용을 검토하는 입장",
        }
    overlays: list[tuple[tuple[str, ...], dict[str, str]]] = [
        (
            ("보고서", "문서", "자료", "기록"),
            {
                "evidence": "문서 초안, 누락 자료 목록, 검토 의견, 결재 일정",
                "situation": "자료 누락과 문서 기준 불일치",
                "inbasket": "문서 수정 요청, 누락 자료 확인 메일, 결재 일정 변경 문서",
                "debate": "문서 정확성을 우선하는 입장과 처리 속도를 우선하는 입장",
            },
        ),
        (
            ("예산", "일정", "계획", "자원", "원가", "비용", "절감"),
            {
                "evidence": "예산 배정표, 일정표, 자원 사용 내역, 변경 요청서",
                "situation": "일정 변경과 예산·자원 제약",
                "inbasket": "일정 변경 요청, 예산 조정 검토서, 우선순위 재배정 문서",
                "debate": "계획 준수를 우선하는 입장과 현장 변경 대응을 우선하는 입장",
            },
        ),
        (
            ("고객", "민원", "예절", "서비스", "방문객"),
            {
                "evidence": "고객 응대 기록, 민원 접수 내역, 서비스 기준표",
                "situation": "고객 불만과 서비스 기준 충돌",
                "inbasket": "민원 접수 문서, 고객 응대 기록, 현장 조치 요청",
                "debate": "고객 편의를 우선하는 입장과 서비스 기준 준수를 우선하는 입장",
            },
        ),
        (
            ("안전", "위험", "점검", "품질", "오류"),
            {
                "evidence": "점검 결과표, 오류 사례, 위험요인 목록, 개선 조치 기록",
                "situation": "안전·품질 위험과 처리 일정 압박",
                "inbasket": "점검 보완 요청, 오류 정정 문서, 긴급 위험 보고",
                "debate": "위험 예방을 우선하는 입장과 일정 준수를 우선하는 입장",
            },
        ),
        (
            ("기물", "파지", "운반", "적재", "보관", "취급"),
            {
                "_replace_context": "true",
                "evidence": "기물별 취급 기준, 파손·오염 기록, 작업 동선, 인력 배치표",
                "situation": "기물 파손·오염 위험과 피크시간 운반 지연",
                "inbasket": "기물 파손 보고, 긴급 운반 요청, 작업 동선 변경, 위생 점검 문서",
                "debate": "기물별 안전한 파지·운반 절차를 우선하는 입장과 피크시간 처리 속도를 우선하는 입장",
            },
        ),
        (
            ("법규", "규정", "기준", "절차"),
            {
                "evidence": "관련 규정, 절차서, 예외 승인 기록, 적용 사례",
                "situation": "기준 해석 차이와 예외 처리 요청",
                "inbasket": "기준 확인 요청, 예외 승인 검토서, 절차 보완 문서",
                "debate": "규정 준수를 엄격히 적용하는 입장과 예외 허용을 검토하는 입장",
            },
        ),
        (
            ("데이터", "분석", "지표", "통계"),
            {
                "evidence": "지표 추이, 원자료, 분석 결과표, 이상값 목록",
                "situation": "지표 이상값과 원인 불명확성",
                "inbasket": "분석 재검토 요청, 원자료 확인 메일, 지표 보고 일정",
                "debate": "추가 검증을 우선하는 입장과 신속 보고를 우선하는 입장",
            },
        ),
        (
            ("보안", "출입", "신분증", "여권", "감별", "통제"),
            {
                "_replace_context": "true",
                "evidence": "출입 로그, 신분 확인 기록, 보안 통제 기준, 예외 요청서",
                "situation": "본인 확인 예외 요청과 보안 기준 충돌",
                "inbasket": "출입 예외 요청, 신분 확인 기록, 보안 책임자 보고 문서",
                "debate": "보안 통제를 우선하는 입장과 방문객 편의를 우선하는 입장",
            },
        ),
    ]
    for keywords, overlay in overlays:
        if any(keyword in key for keyword in keywords):
            return dict(overlay)
    return {}


def _merge_context_overlay(base: dict[str, Any], overlay: dict[str, str]) -> dict[str, Any]:
    if not overlay:
        return base
    merged = dict(base)
    replace_context = str(overlay.get("_replace_context") or "").strip().lower() == "true"
    for field in ("evidence", "situation", "inbasket", "debate"):
        extra = str(overlay.get(field) or "").strip()
        current = str(merged.get(field) or "").strip()
        if not extra:
            continue
        # A focus overlay supplies a more precise policy choice. Combining two
        # complete A-vs-B pairs creates four positions and an incoherent task.
        if replace_context or field == "debate":
            merged[field] = extra
            continue
        if current and normalize_question_dedup_key(extra) not in normalize_question_dedup_key(current):
            phrases: list[str] = []
            phrase_keys: list[str] = []
            phrase_sources = (
                [*current.split(","), *extra.split(",")]
                if field in {"evidence", "inbasket"}
                else [*extra.split(","), *current.split(",")]
            )
            for phrase in phrase_sources:
                cleaned = str(phrase or "").strip()
                key = normalize_question_dedup_key(cleaned)
                if not cleaned or not key:
                    continue
                if any(key in prior or prior in key for prior in phrase_keys):
                    continue
                phrases.append(cleaned)
                phrase_keys.append(key)
            merged[field] = ", ".join(phrases)
        elif not current:
            merged[field] = extra
    return merged


def _domain_context_pack(detail: str, subject: str, focus: str, comp_def: str) -> dict[str, Any]:
    # Select the work domain from the competency context first.  A KSA factor
    # can contain an overloaded word (for example, document *security*) that
    # must not turn a document-management task into a physical guard/patrol
    # simulation.  The factor is only a fallback when the competency itself
    # provides no recognizable domain, and remains available for the narrower
    # focus overlay below.
    identity_source = " ".join(str(x or "") for x in (detail, subject))
    identity_key = re.sub(r"\s+", "", identity_source).lower()
    definition_key = re.sub(r"\s+", "", str(comp_def or "")).lower()
    focus_key = re.sub(r"\s+", "", str(focus or "")).lower()
    default = {
        "evidence": "실적자료, 민원·오류 사례, 업무 기준",
        "situation": "자료 오류, 일정 지연, 이해관계자 요청",
        "inbasket": "긴급 민원, 상급자 보고 요청, 자료 오류 정정, 일정 충돌 문서",
        "debate": (
            "자료와 기준의 오류 가능성이 남으면 처리를 보류하고 검증을 완료하자는 입장과 "
            "서비스 지연을 막기 위해 저위험 건부터 조건부 처리한 뒤 사후 검증하자는 입장"
        ),
        "stakeholders": "상급자, 협업 부서, 민원인",
        "case_facts": (
            "긴급 요청 1건과 일반 요청 2건이 동시에 접수됨 | 핵심 확인자료 1건이 누락됨 | "
            "처리 마감까지 4시간이 남았고 예외 건은 팀장 보고가 필요함"
        ),
    }
    packs: list[tuple[tuple[str, ...], dict[str, Any]]] = [
        (
            ("공적개발원조", "개발원조", "국별협력", "oda사업"),
            {
                "evidence": "국가협력전략, 사업기획서, 정책 검토자료, 성과지표, 이해관계자 협의 기록",
                "situation": (
                    "협력국이 긴급 수요 반영을 요청했지만 정책 적합성 검토와 변경 승인이 대기 중이고, "
                    "미승인 활동 집행·예산 확약에는 사전 승인이 필요합니다. 영향분석과 잠정 일정안 작성이 "
                    "미승인 활동 착수에 포함되는지는 절차에 명시되지 않아 일정과 책임 근거가 충돌하는 상황"
                ),
                "inbasket": "협력국 요청서, 정책 검토 의견, 사업 일정 변경안, 예산·성과지표 검토 문서",
                "debate": (
                    "절차 해석이 확정될 때까지 기존 승인 범위의 활동과 인력을 유지하자는 입장과 "
                    "미승인 활동 집행·예산 확약은 하지 않되 제한된 인력을 영향분석과 잠정 일정안 작성에 배정하자는 입장"
                ),
                "stakeholders": "협력국 담당기관, 사업 수행기관, 정책 담당자, 승인권자, 평가 담당자",
                "case_facts": (
                    "정책 검토 완료 예정: D+3 | 기존 사업 일정 확정 시점: D+5 | "
                    "현행 절차: 미승인 활동 집행·예산 확약은 사전 승인 필요 | "
                    "절차 공백: 영향분석·잠정 일정안 작성의 미승인 활동 착수 해당 여부는 명시되지 않음"
                ),
            },
        ),
        (
            ("프로젝트관리", "프로젝트 인적자원", "프로젝트인적자원", "변경관리"),
            {
                "evidence": "변경요청서, 승인 기록, 일정·원가 영향자료, 역할·책임표",
                "situation": (
                    "범위 확대 승인 결정은 D+3, 기존 범위 검수 준비 마감은 D+5이며 핵심 인력 1명은 기존 업무 중입니다. "
                    "절차상 변경 실행·배치 확정·비용 집행은 승인 대상이나 사전 분석의 착수 해당 여부는 불명확한 상황"
                ),
                "situation_prompt": (
                    "범위 확대 승인 결정은 D+3, 기존 범위 검수 준비 마감은 D+5이며 핵심 인력 1명은 기존 업무 중입니다. "
                    "변경관리 절차서에는 변경 범위 착수·최종 인력배치·비용 집행이 사전 승인 대상으로 제시된 상황"
                ),
                "inbasket": "변경요청서, 승인 대기 문서, 핵심 인력 재배치 요청, 일정 영향 보고서",
                "debate": (
                    "해석이 확정될 때까지 핵심 인력을 기존 업무에 전담하자는 입장과 "
                    "변경 실행·확정·비용 집행은 보류하되 하루 2시간을 사전 분석에 배정하자는 입장"
                ),
                "stakeholders": "프로젝트 관리자, 승인권자, 수행팀, 요청 부서, 품질 담당자",
                "case_facts": (
                    "변경심의 예정: D+3 | 기존 범위 검수 준비 마감: D+5 | 핵심 인력 1명은 기존 승인 범위 작업 중 | "
                    "기존 검수 준비 잔여 작업량: 36인시, 핵심 인력 가용시간: 8시간/일 | "
                    "사전 분석에 하루 2시간 배정 시 D+3까지 기존 작업 투입량 6인시 감소 | "
                    "현행 절차: 변경 범위 착수·최종 인력배치·비용 집행은 사전 승인 필요 | "
                    "절차 공백: 영향분석·가용성 확인·잠정안 작성이 변경 범위 착수에 포함되는지는 명시되지 않음"
                ),
                "case_materials": [
                    {"source": "변경요청서", "field": "심의 상태", "value": "승인 대기, 결정 예정 D+3"},
                    {"source": "기존 범위 일정표", "field": "검수 준비 마감", "value": "D+5"},
                    {"source": "역할·책임표", "field": "핵심 인력 가용성", "value": "1명, 기존 승인 범위 작업 중"},
                    {"source": "작업량 현황표", "field": "검수 준비 잔여량·가용시간", "value": "36인시 / 8시간·일"},
                    {"source": "일정 영향표", "field": "사전 분석 배정 영향", "value": "하루 2시간, D+3까지 기존 작업 6인시 감소"},
                    {"source": "변경관리 절차서", "field": "사전 승인 대상", "value": "변경 범위 착수·최종 인력배치·비용 집행"},
                    {"source": "변경관리 절차서", "field": "해석 공백", "value": "영향분석·가용성 확인·잠정안 작성의 착수 해당 여부 미기재"},
                ],
                "authority_facts": [
                    "응시자 역할: 프로젝트 인적자원관리 실무 담당자로 사실 확인·자료 수집·승인요청안 작성·위험 보고를 직접 수행할 수 있음",
                    "기존 범위 일정 변경·대외 확약은 팀장 결재가 필요하고 변경 범위 착수·최종 인력배치·비용 집행은 변경 승인권자 결정이 필요함",
                ],
                "authority_materials": [
                    {
                        "source": "업무분장표",
                        "field": "응시자 역할·직접 수행 범위",
                        "value": "프로젝트 인적자원관리 실무 담당자 / 사실 확인·자료 수집·승인요청안 작성·위험 보고",
                    },
                    {
                        "source": "전결규정",
                        "field": "결재·승인 권한",
                        "value": "기존 범위 일정 변경·대외 확약은 팀장 결재 / 변경 범위 착수·최종 인력배치·비용 집행은 변경 승인권자 결정",
                    },
                ],
            },
        ),
        (
            ("조리", "식음료", "메뉴", "식재료", "스톡", "영업장"),
            {
                "evidence": "메뉴별 판매량, 식재료 재고표, 위생점검 결과, 고객 불만 사례",
                "situation": "식재료 수급 차질, 위생 기준 이슈, 조리 일정 지연",
                "inbasket": "식재료 재고표, 위생점검 요청, 고객 불만 접수, 조리 일정 변경 문서",
                "debate": "위생·품질 기준을 강화하는 입장과 조리 속도·영업 효율을 우선하는 입장",
                "stakeholders": "조리장, 홀 담당자, 위생 담당자, 고객",
            },
        ),
        (
            ("청소", "환경미화", "미화", "품질검증"),
            {
                "evidence": "구역별 청소 점검표, 민원 접수 내역, 오염도 확인 사진, 작업 배치표",
                "situation": "청소 범위 변경, 반복 민원, 안전사고 위험 구역 발견",
                "inbasket": "구역별 작업지시서, 민원 접수 문서, 안전주의 요청, 인력 배치 변경 문서",
                "debate": "청소 품질 기준을 강화하는 입장과 제한된 인력·시간 내 처리 효율을 우선하는 입장",
                "stakeholders": "현장 관리자, 이용자, 협업 근무자, 안전 담당자",
            },
        ),
        (
            ("화물", "운송", "운임", "화주", "배차", "차량"),
            {
                "evidence": "운송 의뢰서, 배차표, 운임 산정 자료, 화주 요청 변경 내역",
                "situation": "화주 요청 변경, 배차 지연, 운임 산정 오류",
                "inbasket": "운송 의뢰서, 배차 변경 요청, 화주 불만 접수, 운임 검토 문서",
                "debate": "운송 기준과 안전·정확성을 강화하는 입장과 납기·비용 효율을 우선하는 입장",
                "stakeholders": "화주, 운전 담당자, 배차 담당자, 협력 운송사",
            },
        ),
        (
            ("사회복지", "사례관리", "대상자", "욕구", "상담"),
            {
                "evidence": "초기상담 기록, 욕구 사정표, 서비스 연계 현황, 대상자 동의 기록",
                "situation": "대상자 욕구 변경, 보호자 요청, 서비스 연계 일정 지연",
                "inbasket": "상담 기록지, 서비스 연계 요청서, 대상자 긴급 연락, 기관 회신 문서",
                "debate": "대상자 자기결정권을 우선하는 입장과 기관 자원 배분 기준을 우선하는 입장",
                "stakeholders": "대상자, 보호자, 사례관리자, 연계기관 담당자",
            },
        ),
        (
            ("병원", "간호", "요양", "보건", "환자", "진료", "의료", "산업보건"),
            {
                "evidence": "환자 안내 기록, 보건교육 계획서, 건강상담 기록, 안전보건 점검 결과",
                "situation": "환자·근로자 문의 증가, 교육 일정 변경, 건강정보 기록 오류",
                "inbasket": "교육 일정표, 상담 기록지, 점검 요청 문서, 긴급 안내 요청",
                "debate": "건강·안전 기준을 엄격히 적용하는 입장과 현장 수용성·업무 연속성을 우선하는 입장",
                "stakeholders": "대상자, 의료진, 안전보건 담당자, 부서 관리자",
            },
        ),
        (
            ("정보기술", "it프로젝트", "it비즈니스", "시스템", "장애티켓", "sla"),
            {
                "evidence": "요구사항 정의서, 장애 티켓, SLA 현황, 비용편익분석표, 일정·리스크 등록부",
                "situation": "요구사항 변경, 장애 재발, SLA 위반 위험",
                "inbasket": "변경 요청서, 장애 재발 보고, 사용자 문의, 일정 조정 문서",
                "debate": "IT 거버넌스와 표준 준수를 우선하는 입장과 현업 요청 처리 속도를 우선하는 입장",
                "stakeholders": "현업 부서, IT PM, 개발 담당자, 운영 담당자, 공급업체, 보안 담당자",
            },
        ),
        (
            ("화력발전", "원자력발전", "발전설비", "전기제어", "환경설비", "보호계전기", "원자로"),
            {
                "evidence": "운전일지, 설비 알람 로그, 정비이력, 절차서, 작업허가서, 장애 티켓",
                "situation": "설비 이상 징후, 절차서와 현장 상태 불일치, 환경·안전 기준 접근",
                "inbasket": "교대조 인수인계 문서, 정비 우선순위 요청, 장애 재발 보고, 변경관리 승인 문서",
                "debate": "안전보수성을 우선하는 입장과 가동률·사업 일정 준수를 우선하는 입장",
                "stakeholders": "운전원, 정비팀, 안전품질팀, 환경팀, 규제기관, 현업부서",
            },
        ),
        (
            ("하수", "상수", "관로", "수질", "누수", "블록시스템"),
            {
                "evidence": "수질측정값, 설비 운전 로그, 관망도, 유량·수압 기록, 안전점검표",
                "situation": "수질 경보, 계측기 이상값, 현장 작업 일정 충돌",
                "inbasket": "약품투입 기록 정정 요청, 계측기 교정 요청, 누수 신고, 법정보고자료 제출 문서",
                "debate": "수질·안전 기준을 보수적으로 적용하는 입장과 처리 효율·주민 불편 최소화를 우선하는 입장",
                "stakeholders": "주민, 관제실, 수질 담당자, 현장 작업자, 지자체, 감독기관",
            },
        ),
        (
            ("객실", "유원", "스포츠", "레저", "어트랙션", "체크인", "체크 인"),
            {
                "evidence": "예약·운영 현황, 시설 점검표, 안전관리 기록, 이용객 민원 내역",
                "situation": "예약 변경, 시설 점검 지연, 이용객 안전 안내 필요 상황",
                "inbasket": "예약 현황표, 시설 점검 요청, 이용객 민원, 안전 안내 문서",
                "debate": "고객 안전·서비스 기준을 강화하는 입장과 회전율·운영 효율을 우선하는 입장",
                "stakeholders": "이용객, 프론트 담당자, 현장 운영자, 시설 담당자, 안전 담당자",
            },
        ),
        (
            ("설비", "정비", "관로", "하수", "상수", "시설", "건설", "건축", "해체"),
            {
                "evidence": "설비 점검 기록, 도면, 작업허가서, 이상 징후 로그, 안전점검 결과",
                "situation": "설비 이상 징후, 작업 일정 충돌, 안전 기준 미충족 가능성",
                "inbasket": "점검 기록표, 작업허가 요청, 장애 신고, 안전조치 확인 문서",
                "debate": "안전·품질 기준을 강화하는 입장과 공정 일정·운영 연속성을 우선하는 입장",
                "stakeholders": "현장 작업자, 안전 담당자, 설비 운영자, 협력업체",
            },
        ),
        (
            ("보안", "경비", "순찰"),
            {
                "evidence": "순찰 기록, 출입 통제 로그, 경비 배치표, 이상 상황 보고서",
                "situation": "출입 통제 예외 요청, 순찰 공백 가능성, 이상 상황 신고",
                "inbasket": "순찰 기록표, 출입 승인 요청, 이상 상황 보고서, 근무 배치 변경 문서",
                "debate": "보안 통제 기준을 강화하는 입장과 방문객 편의·운영 효율을 우선하는 입장",
                "stakeholders": "방문객, 현장 경비원, 시설 담당자, 보안 책임자",
            },
        ),
        (
            ("객실", "유원", "스포츠", "시설운영", "레저"),
            {
                "evidence": "이용객 민원 내역, 시설 점검표, 예약·운영 현황, 안전관리 기록",
                "situation": "이용객 불만, 시설 점검 지연, 안전 안내 필요 상황",
                "inbasket": "예약 현황표, 시설 점검 요청, 이용객 민원, 안전 안내 문서",
                "debate": "고객 안전·서비스 기준을 강화하는 입장과 회전율·운영 효율을 우선하는 입장",
                "stakeholders": "이용객, 현장 운영자, 시설 담당자, 안전 담당자",
            },
        ),
        (
            ("사무", "총무", "문서", "부동산", "행정", "프로젝트", "정보기술", "it"),
            {
                "evidence": "요구사항 목록, 결재 문서, 회의록, 일정표, 이해관계자 요청 내역",
                "situation": "요구사항 변경, 결재 지연, 문서 기준 불일치",
                "inbasket": "결재 대기 문서, 회의 요청, 요구사항 변경 메일, 마감 임박 업무 목록",
                "debate": "문서·절차 기준을 강화하는 입장과 처리 속도·협업 효율을 우선하는 입장",
                "stakeholders": "요청 부서, 결재권자, 협업 담당자, 외부 관계자",
            },
        ),
    ]
    specific_packs = packs[:-1]
    generic_packs = packs[-1:]

    # A specific signal in the competency definition must beat a broad word
    # such as "행정" or "프로젝트" in its title.  The focus remains a final
    # fallback so overloaded factor words cannot override the job identity.
    for candidate_key in (identity_key, definition_key):
        if not candidate_key:
            continue
        for keywords, pack in specific_packs:
            if any(keyword.lower() in candidate_key for keyword in keywords):
                return _merge_context_overlay({**default, **pack}, _focus_context_overlay(focus))
    for candidate_key in (identity_key, definition_key):
        if not candidate_key:
            continue
        for keywords, pack in generic_packs:
            if any(keyword.lower() in candidate_key for keyword in keywords):
                return _merge_context_overlay({**default, **pack}, _focus_context_overlay(focus))
    if focus_key:
        for keywords, pack in packs:
            if any(keyword.lower() in focus_key for keyword in keywords):
                return _merge_context_overlay({**default, **pack}, _focus_context_overlay(focus))
    return _merge_context_overlay(default, _focus_context_overlay(focus))


def _presentation_source_segments(text: str, anchors: list[str], limit: int = 6) -> list[str]:
    compact = re.sub(r"\s+", " ", str(text or "").strip())
    # Review merging adds internal provenance tags such as
    # ``[담당업무-우선]``. They are useful for server tracing but are not
    # candidate-facing facts and can make a generated prompt look like a
    # template or expose implementation metadata.
    compact = re.sub(
        r"\[(?:공고문|직무기술서|담당업무|지원자격|우대사항|면접평가항목|발표자료)[^\]]*\]\s*",
        "",
        compact,
    )
    if not compact:
        return []
    segments = [
        re.sub(r"\s+", " ", chunk).strip(" -·•○◦▪")
        for chunk in re.split(r"(?<=[.!?。])\s+|(?=[○◦•▪●]\s)|(?=\d+[.)]\s)", compact)
        if len(chunk.strip()) >= 8
    ]
    normalized_anchors = [re.sub(r"\s+", "", str(value or "").lower()) for value in anchors if str(value or "").strip()]
    matched = []
    for segment in segments:
        normalized = re.sub(r"\s+", "", segment.lower())
        if normalized_anchors and any(anchor in normalized for anchor in normalized_anchors if len(anchor) >= 2):
            matched.append(segment)
    selected = matched or segments
    return list(dict.fromkeys(selected))[:limit]


def _presentation_work_segments(text: str, anchors: list[str], limit: int = 6) -> list[str]:
    """Extract operational duty/JD statements for a presentation packet.

    Notice parsers often return a single long paragraph containing the
    institution introduction, eligibility rules, and the actual duty.  A
    presentation task must lead with the duty section, so split known source
    headings and rank statements containing the selected NCS/job anchors.
    """

    compact = re.sub(r"\s+", " ", str(text or "").strip())
    compact = re.sub(
        r"\[(?:공고문|직무기술서|담당업무|지원자격|우대사항|면접평가항목|발표자료)[^\]]*\]\s*",
        "",
        compact,
    )
    if not compact:
        return []

    # Prefer the actual duty section when the source is a full recruitment
    # notice.  The lookahead stops before eligibility/benefits/evaluation text.
    duty_match = re.search(
        r"(?:\d\s*)?(?:채용\s*분야\s*및\s*담당\s*업무|담당\s*업무|수행\s*업무)"
        r"\s*[:：]?\s*(.*?)(?=(?:\d\s*)?(?:응시\s*자격|자격\s*요건|우대\s*사항|전형\s*방법|필요\s*지식|필요\s*기술|직무\s*수행\s*태도|직업\s*기초\s*능력)|$)",
        compact,
        flags=re.IGNORECASE,
    )
    working_text = duty_match.group(1).strip() if duty_match else compact
    working_text = re.sub(r"(?:※|참고\s*사항).*?$", "", working_text).strip()
    working_text = re.sub(r"(?=\d+[.)]\s*)", "\n", working_text)
    working_text = re.sub(r"(?=[○◦•▪●])", "\n", working_text)
    working_text = re.sub(
        r"(?=(?:능력단위\s*[·.]?\s*정의|필요\s*지식|필요\s*기술|직무\s*수행\s*태도|필요\s*자격|직업\s*기초\s*능력))",
        "\n",
        working_text,
    )
    segments = [
        re.sub(r"\s+", " ", part).strip(" -·•○◦▪●")
        for part in re.split(r"\n+|(?<=[.!?。])\s+", working_text)
        if len(part.strip()) >= 10
    ]
    if not segments:
        return []
    normalized_anchors = [
        re.sub(r"\s+", "", str(value or "").casefold())
        for value in anchors
        if str(value or "").strip()
    ]
    operational_markers = (
        "담당", "수행", "관리", "운영", "점검", "유지", "보수", "작성", "분석",
        "설계", "조작", "운전", "기록", "보고", "검토", "계획", "절차", "안전",
        "대응", "복구", "협업", "민원", "성과", "지표",
    )
    scored: list[tuple[int, int, str]] = []
    for index, segment in enumerate(segments):
        normalized = re.sub(r"\s+", "", segment.casefold())
        anchor_score = sum(2 for anchor in normalized_anchors if len(anchor) >= 2 and anchor in normalized)
        marker_score = sum(1 for marker in operational_markers if marker in segment)
        # Introductory institution/eligibility paragraphs are not operational
        # duty evidence unless they also contain a selected anchor.
        penalty = 3 if any(token in segment for token in ("설립", "비전", "지원 바랍니다", "응시 자격")) and not anchor_score else 0
        scored.append((anchor_score + marker_score - penalty, -index, segment))
    ranked = [row[2] for row in sorted(scored, key=lambda row: (row[0], row[1]), reverse=True)]
    return list(dict.fromkeys(ranked))[: max(1, int(limit or 6))]


def _generation_job_context_text(
    *,
    duty_text: str = "",
    jd_text: str = "",
    notice_text: str = "",
    evaluation_text: str = "",
) -> str:
    """Build a bounded source bundle for deterministic question recovery.

    The provider prompt already receives these layers separately.  The same
    source bundle is passed to template/fallback repairs so a timeout never
    downgrades a job-specific question into a domain-profile example.
    """

    layers = (
        ("담당업무", duty_text),
        ("직무기술서", jd_text),
        ("공고문", notice_text),
        ("면접평가항목", evaluation_text),
    )
    parts: list[str] = []
    for label, value in layers:
        raw_value = str(value or "").strip()
        if not raw_value:
            continue
        compact_value = re.sub(r"\s+", " ", raw_value)[:1800]
        parts.append(f"[{label}]{compact_value}")
    return "\n".join(parts)[:6000]


def _presentation_task_excerpt(value: str, limit: int = 360) -> str:
    """Keep the candidate-facing task concrete without mid-sentence clipping."""

    compact = re.sub(r"\s+", " ", str(value or "").strip())
    if len(compact) <= limit:
        return compact
    first_sentence = re.split(r"(?<=[.!?。])\s+", compact, maxsplit=1)[0].strip()
    if 80 <= len(first_sentence) <= limit:
        return first_sentence
    return compact[:limit].rsplit(" ", 1)[0].rstrip(" ,·") + "…"


def _build_presentation_material_packet(
    *,
    interview_methods: list[str],
    jd_text: str,
    notice_text: str,
    duty_text: str,
    question_plan: dict[str, Any],
    ncs_matches: list[dict[str, Any]],
    ncs_ksa: list[dict[str, Any]],
    supplemental_text: str = "",
) -> dict[str, Any] | None:
    """Build a deterministic, reviewable presentation task packet."""
    if "발표면접" not in {str(value or "").strip() for value in (interview_methods or [])}:
        return None
    selected_items = [
        item for item in (question_plan.get("selected_items") or [])
        if isinstance(item, dict)
        and bool(item.get("enabled", True))
        and int(item.get("main_count", 0) or 0) > 0
    ] if isinstance(question_plan, dict) else []
    selected_detail = str((selected_items[0] if selected_items else {}).get("detail") or "").strip()
    unit = next(
        (
            row for row in (ncs_matches or [])
            if isinstance(row, dict)
            and (not selected_detail or str(row.get("ncsSubdCdnm") or row.get("compeUnitName") or "").strip() == selected_detail)
        ),
        next((row for row in (ncs_matches or []) if isinstance(row, dict)), {}),
    )
    code = str(unit.get("ncsClCd") or "").strip()
    unit_name = str(unit.get("compeUnitName") or selected_detail or "선정 NCS 능력단위").strip()
    detail = str(unit.get("ncsSubdCdnm") or selected_detail or "확정 세분류").strip()
    comp_def = str(unit.get("compeUnitDef") or "").strip()
    factors = list(dict.fromkeys(
        str(row.get("factorName") or row.get("factor_name") or row.get("factor") or row.get("ksa") or "").strip()
        for row in (ncs_ksa or [])
        if isinstance(row, dict)
        and (not code or str(row.get("ncsClCd") or "").strip() in {"", code})
    ))[:6]
    factors = [value for value in factors if value]
    focus = factors[0] if factors else "핵심 업무 수행기준"
    # Presentation packets must be built from the selected job materials.  Do
    # not route this path through the old domain-profile table: those profiles
    # contain example scenarios for other products and can silently inject
    # unrelated, domain-specific facts into a candidate's assignment.
    context: dict[str, Any] = {}
    # Keep all supplied source layers in the packet.  A duty-text override is
    # useful for editing, but it must not hide the notice/JD evidence that the
    # interviewer is expected to review.  Each layer is searched separately so
    # a long notice cannot crowd out the actual duty and NCS wording.
    source_anchors = [detail, unit_name, *factors[:3]]
    source_layers = [
        ("공고문", notice_text),
        ("직무기술서", jd_text),
        ("담당업무 보완", duty_text),
    ]
    source_chunks_by_layer: list[tuple[str, list[str]]] = []
    for source_label, source_value in source_layers:
        layer_chunks = _presentation_work_segments(
            str(source_value or ""),
            source_anchors,
            limit=3,
        )
        if layer_chunks:
            source_chunks_by_layer.append((source_label, layer_chunks))
    source_chunks: list[str] = []
    seen_source_chunks: set[str] = set()
    for _, chunks in source_chunks_by_layer:
        for chunk in chunks:
            source_key = re.sub(r"\s+", "", str(chunk or "")).casefold()
            if not source_key or source_key in seen_source_chunks:
                continue
            seen_source_chunks.add(source_key)
            source_chunks.append(chunk)
    source_summary = " ".join(source_chunks)[:900]
    if not source_summary:
        source_summary = " ".join(
            re.sub(r"\s+", " ", str(value or "").strip())
            for _, value in source_layers
            if str(value or "").strip()
        )[:900]
    duty_summary = source_summary or f"{detail}의 {unit_name} 수행내용"
    # The main assignment should lead with the most operational source. Keep
    # the notice in the evidence table, but prefer the reviewed duty override
    # and JD excerpts over a notice's institutional introduction when forming
    # the short candidate-facing task sentence.
    task_source_chunks: list[str] = []
    seen_task_source_chunks: set[str] = set()
    for source_label in ("담당업무 보완", "직무기술서", "공고문"):
        for layer, chunks in source_chunks_by_layer:
            if layer != source_label:
                continue
            for chunk in chunks:
                source_key = re.sub(r"\s+", "", str(chunk or "")).casefold()
                if not source_key or source_key in seen_task_source_chunks:
                    continue
                seen_task_source_chunks.add(source_key)
                task_source_chunks.append(chunk)
    task_duty_summary = _presentation_task_excerpt(
        " ".join(task_source_chunks) or duty_summary
    )
    facts = list(dict.fromkeys([
        *[
            f"{source_label}에서 확인된 업무 근거: {chunk[:420]}"
            for source_label, chunks in source_chunks_by_layer
            for chunk in chunks[:2]
        ],
        f"능력단위 정의: {comp_def[:500] or unit_name}",
        f"KSA 평가 근거: {focus}",
        (
            f"발표 시나리오(가정): {task_duty_summary}를 수행하는 중 "
            f"'{focus}' 적용에 필요한 확인자료가 일부 누락되거나 값이 충돌한 경우"
        ),
        (
            f"판단 과제: {unit_name}의 수행 범위·예외·보고 기준을 확인하고 "
            "안전·품질·일정 기준으로 실행 순서를 정함"
        ),
        "입력 자료에 없는 기관 고유 수치·사실은 가정과 사실을 구분해 발표함",
    ]))[:10]
    presentation_task = str(context.get("presentation_task") or "").strip()
    if not presentation_task:
        presentation_task = (
            f"공고문·직무기술서에 제시된 '{task_duty_summary}' 업무를 수행하는 담당자라고 가정합니다. "
            f"능력단위 '{unit_name}'의 정의({comp_def[:260] or '직무 수행 절차와 기준'})와 "
            f"KSA 근거 '{focus}'를 적용해, 필요한 확인자료·판단 기준·적용 범위·예외·보고 순서를 발표하십시오. "
            "자료 일부가 누락되거나 값이 충돌하는 경우에는 사실과 가정을 구분하고, "
            "안전·품질·일정·권한 기준으로 대안 2가지를 비교한 뒤 실행안과 검증 가능한 결과자료를 제시하십시오."
        )
    constraints = [
        str(value).strip()
        for value in (context.get("presentation_constraints") or [])
        if str(value).strip()
    ][:6]
    if not constraints:
        constraints = [
            "공고문·직무기술서·NCS KSA에 없는 기관 고유 수치나 사실은 추정하지 않습니다.",
            f"'{focus}'를 판단 근거로 명시하고, 적용 범위와 예외를 구분합니다.",
            "입력 자료의 불일치·누락·제약은 확인 필요 사항으로 표시하고, 확인 전후의 판단을 구분합니다.",
        ]
    deliverables = [
        str(value).strip()
        for value in (context.get("presentation_deliverables") or [])
        if str(value).strip()
    ][:6]
    if not deliverables:
        deliverables = [
            f"{task_duty_summary[:260]} 업무의 핵심 사실·위험·우선순위 표",
            f"'{focus}' 적용을 위한 확인자료·판단 기준·적용 범위·예외 목록",
            f"{unit_name} 수행 대안 2가지·선택 근거·역할·보고 순서",
            "검증 가능한 산출물·기록·피드백과 재발 방지 보완 조치",
        ]
    material_rows = [
        {
            "source": "공고문·직무기술서",
            "field": "발표 대상 업무",
            "value": source_summary or "확정된 직무·공고 맥락",
        },
        {
            "source": "NCS 능력단위",
            "field": "능력단위·정의",
            "value": " · ".join(value for value in (unit_name, comp_def) if value)[:700],
        },
    ]
    for source_label, chunks in source_chunks_by_layer:
        material_rows.append(
            {
                "source": source_label,
                "field": "핵심 근거 발췌",
                "value": " ".join(chunks)[:700],
            }
        )
    if factors:
        material_rows.append({"source": "NCS KSA", "field": "평가 근거", "value": " · ".join(factors)[:700]})
    material_rows.extend([
        {"source": "입력자료 종합", "field": "수행내용 근거", "value": duty_summary[:900]},
        {"source": "입력자료 종합", "field": "검토 제약", "value": "자료에 직접 확인되는 사실·누락·불일치와 확인이 필요한 항목을 구분"},
        {"source": "입력자료 종합", "field": "판단권한", "value": "공고문·직무기술서에 명시된 담당 범위와 보고·승인 기준을 우선 확인"},
    ])
    if supplemental_text.strip():
        material_rows.append({"source": "사용자 보완자료", "field": "추가 확인자료", "value": supplemental_text.strip()[:700]})
    material_rows = [
        row for row in material_rows
        if isinstance(row, dict) and all(str(row.get(field) or "").strip() for field in ("source", "field", "value"))
    ][:10]
    return {
        "title": f"{detail} · {unit_name} 발표 과제 자료",
        "generated": True,
        "source": "server_job_notice_ncs_ksa",
        "scenario_label": "NCS 기반 예시 시나리오 · 기관 사실·수치는 최종 확인 필요",
        "ncs_code": code,
        "ncs_detail": detail,
        "competency": unit_name,
        "focus": focus,
        "provided_materials": list(dict.fromkeys(row["source"] for row in material_rows)),
        "case_materials": material_rows,
        "case_facts": facts,
        "task_prompt": presentation_task,
        "constraints": constraints,
        "required_deliverables": deliverables,
        "slide_outline": [
            {"slide": 1, "title": "담당업무·핵심 문제", "instruction": f"'{task_duty_summary[:260]}'에서 실제 수행 대상, 위험요인, 누락·충돌 자료와 즉시 영향을 구분해 우선순위를 정리합니다."},
            {"slide": 2, "title": "NCS KSA 판단 근거", "instruction": f"능력단위 정의와 KSA '{focus}'를 실제 업무 단계·확인자료·적용 범위·예외 판단에 연결합니다."},
            {"slide": 3, "title": "대안 비교·선택", "instruction": f"{unit_name} 수행 대안 2가지를 안전·품질·일정·권한 기준으로 비교하고 선택 근거와 조건을 제시합니다."},
            {"slide": 4, "title": "실행 순서·검증", "instruction": "담당자별 행동, 보고·승인 순서, 산출물과 기록·피드백으로 결과를 확인하고 재발 방지 조치를 제시합니다."},
        ],
        "use_rules": [
            "제공자료에 없는 정밀 수치·법령 조항은 사실처럼 만들지 않습니다.",
            "자료 간 값이 충돌하면 확인 필요 사항으로 표시하고 판단 근거를 설명합니다.",
            "기관 사실과 자동 구성된 운영 시나리오는 최종 면접 전 사람이 검토합니다.",
        ],
        "review_required": True,
    }


def _presentation_material_prompt_text(packet: dict[str, Any] | None) -> str:
    if not isinstance(packet, dict):
        return ""
    lines = [
        "[서버 자동 생성 발표자료]",
        f"제목: {str(packet.get('title') or '').strip()}",
        f"시나리오 성격: {str(packet.get('scenario_label') or '').strip()}",
        f"세분류: {str(packet.get('ncs_detail') or '').strip()} / 능력단위: {str(packet.get('competency') or '').strip()}",
        f"평가 초점: {str(packet.get('focus') or '').strip()}",
        f"발표 메인 과제: {str(packet.get('task_prompt') or '').strip()}",
        "근거 자료:",
    ]
    for row in (packet.get("case_materials") or [])[:10]:
        if isinstance(row, dict):
            lines.append(f"- {row.get('source')}: {row.get('field')} = {row.get('value')}")
    lines.append("슬라이드 구성:")
    for slide in (packet.get("slide_outline") or [])[:4]:
        if isinstance(slide, dict):
            lines.append(f"- {slide.get('slide')}. {slide.get('title')}: {slide.get('instruction')}")
    lines.append("발표 제약조건:")
    lines.extend(f"- {str(value).strip()}" for value in (packet.get("constraints") or [])[:6] if str(value).strip())
    lines.append("필수 산출물:")
    lines.extend(f"- {str(value).strip()}" for value in (packet.get("required_deliverables") or [])[:6] if str(value).strip())
    lines.append("운영 규칙: 자료에 없는 정밀 수치·법령 조항은 추정하지 말고, 충돌 값은 확인 대상으로 남깁니다.")
    return "\n".join(lines)[:7000]


def _attach_presentation_material_packet(strategy: dict[str, Any], packet: dict[str, Any] | None) -> dict[str, Any]:
    if not isinstance(strategy, dict) or not isinstance(packet, dict):
        return strategy
    for question in strategy.get("interview_questions") or []:
        if not isinstance(question, dict) or str(question.get("type") or question.get("method") or "").strip() != "발표면접":
            continue
        conditions = dict(question.get("task_conditions") or {}) if isinstance(question.get("task_conditions"), dict) else {}
        packet_materials = [
            str(value).strip()
            for value in (packet.get("provided_materials") or [])
            if str(value).strip()
        ]
        # The packet's source labels are authoritative for a presentation
        # task. Do not retain the generic fallback labels (for example,
        # "접수 현황표") from an unrelated example scenario.
        conditions["provided_materials"] = list(dict.fromkeys(packet_materials)) or list(
            dict.fromkeys(
                str(value).strip()
                for value in (conditions.get("provided_materials") or [])
                if str(value).strip()
            )
        )
        conditions["case_materials"] = list(packet.get("case_materials") or [])[:10]
        conditions["presentation_task"] = str(packet.get("task_prompt") or "").strip()
        conditions["presentation_constraints"] = list(packet.get("constraints") or [])[:6]
        conditions["required_outputs"] = list(packet.get("required_deliverables") or [])[:6] or conditions.get("required_outputs", [])
        packet_facts = [
            str(value).strip()
            for value in (packet.get("case_facts") or [])
            if str(value).strip()
        ]
        # The fallback renderer has a generic example case pack so every
        # method remains answerable when no presentation packet exists. Once a
        # packet is built from this job's notice/JD/NCS inputs, those examples
        # must not leak into the candidate-facing assignment (they can look
        # like hardcoded facts from an unrelated occupation).
        conditions["case_facts"] = list(dict.fromkeys(packet_facts)) or list(
            dict.fromkeys(
                str(value).strip()
                for value in (conditions.get("case_facts") or [])
                if str(value).strip()
            )
        )
        question["task_conditions"] = conditions
        question["presentation_material"] = packet
    strategy["presentation_material"] = packet
    strategy["presentation_material_generated"] = True
    strategy["presentation_material_review_required"] = True
    return strategy


def _has_korean_final_consonant(text: str) -> bool:
    cleaned = re.sub(r"[\s\]\)\}\"'.,!?…:;]+$", "", str(text or ""))
    for ch in reversed(cleaned):
        code = ord(ch)
        if 0xAC00 <= code <= 0xD7A3:
            return ((code - 0xAC00) % 28) != 0
    return False


def _with_josa(text: str, final_consonant_josa: str, no_final_consonant_josa: str) -> str:
    value = str(text or "").strip()
    return value + (final_consonant_josa if _has_korean_final_consonant(value) else no_final_consonant_josa)


def _quoted_with_josa(text: str, final_consonant_josa: str, no_final_consonant_josa: str) -> str:
    """Legacy alias kept for callers; candidate-facing labels are no longer quoted."""
    value = str(text or "").strip()
    return value + (final_consonant_josa if _has_korean_final_consonant(value) else no_final_consonant_josa)


def _scenario_condition_text(value: Any) -> str:
    """Render either a full scenario clause or a comma-list as one condition."""
    scenario = re.sub(r"\s+", " ", str(value or "")).strip(" ,.;")
    if not scenario:
        return "업무 기준과 일정이 충돌한 상황"
    if re.search(r"(?:상황|조건|상태)$", scenario):
        return scenario
    return f"{_with_josa(scenario, '이', '가')} 동시에 발생한 상황"


def _experience_scenario_condition(value: Any) -> str:
    """Keep an experience prompt concrete without copying an entire case brief."""
    scenario = re.sub(r"\s+", " ", str(value or "")).strip(" ,.;")
    parts = [part.strip() for part in scenario.split(",") if part.strip()]
    selected = parts[:2] if parts else ["업무 기준과 일정 충돌"]
    connective_endings = (
        (r"\s*조건이고$", " 조건"),
        (r"\s*조건이며$", " 조건"),
        (r"중이고$", "중인 상황"),
        (r"있고$", "있는 상황"),
        (r"이며$", "인 상황"),
        (r"이고$", "인 상황"),
    )
    normalized: list[str] = []
    for part in selected:
        rendered = part
        for pattern, replacement in connective_endings:
            if re.search(pattern, rendered):
                rendered = re.sub(pattern, replacement, rendered)
                break
        if not re.search(r"(?:상황|조건|상태)$", rendered):
            rendered = f"{_with_josa(rendered, '이', '가')} 발생한 상황"
        normalized.append(rendered)
    return ", ".join(normalized)


_QUESTION_VARIATION_CONSTRAINTS = (
    "",
    "마감시간이 임박하고 확인 자료 일부가 누락된 조건",
    "관련 부서의 요구와 현장 처리 기준이 충돌하는 조건",
    "예외 요청과 품질 오류 위험이 동시에 발견된 조건",
    "인력과 시간이 제한되어 우선순위를 다시 정해야 하는 조건",
    "처리 속도 요구와 오류 예방 기준이 충돌하는 조건",
    "기존 절차로 해결되지 않아 추가 확인과 보고가 필요한 조건",
    "인수인계 자료와 최신 규정이 일치하지 않는 조건",
    "고객 민원 확대 가능성과 내부 승인 지연이 함께 있는 조건",
    "안전·보안 기준을 지키면서 긴급 요청을 처리해야 하는 조건",
    "외부 협력업체 산출물에서 반복 오류가 발견된 조건",
    "성과지표는 악화됐지만 원인 자료가 불충분한 조건",
    "업무량이 급증해 기존 처리 순서로 마감을 지키기 어려운 조건",
    "시스템 장애로 수기 기록과 사후 대조가 필요한 조건",
    "승인 기준이 바뀌었지만 변경 공지가 일부 담당자에게 전달되지 않은 조건",
    "이해관계자마다 성공 기준과 허용 위험이 다른 조건",
    "초기 조치가 효과가 없어 대안을 재설계해야 하는 조건",
    "핵심 자료의 신뢰성이 낮아 교차검증이 필요한 조건",
    "단기 성과와 장기 재발 방지 중 선택해야 하는 조건",
)

_QUESTION_VARIATION_STAKEHOLDER_PRESSURES = (
    "요청 부서는 즉시 처리를 요구하는 조건",
    "검토 담당자는 증빙 완결을 우선하는 조건",
    "서비스 담당자는 지연 안내와 대안 제시를 요구하는 조건",
    "품질 담당자는 추가 검증 전까지 처리를 보류하라고 요구하는 조건",
    "의사결정자는 비용 최소화를 요구하는 조건",
    "외부 협력자가 오류 책임에 이견을 제기한 조건",
    "신규 담당자가 즉시 지원을 요청한 조건",
    "최종 승인권자가 자리를 비운 조건",
)

# Operating pressure and topic are separate diversity axes.  This keeps a
# seven-question set from becoming the same new-project/budget scenario with
# only the deadline wording changed.
_QUESTION_TOPIC_AXES = (
    "협업·이해관계자 조정",
    "정책·규정 준수",
    "리스크·품질관리",
    "성과지표·검증",
    "이용자·형평성",
    "디지털·프로세스 개선",
    "자원·일정 조정",
    "조직학습·인수인계",
    "보안·개인정보 보호",
    "지속가능성·환경 영향",
)


def _question_topic_axis(variation_index: int) -> str:
    try:
        index = max(0, int(variation_index or 0))
    except (TypeError, ValueError):
        index = 0
    return _QUESTION_TOPIC_AXES[index % len(_QUESTION_TOPIC_AXES)]


def _question_variation_constraint(variation_index: int) -> str:
    try:
        index = max(0, int(variation_index or 0))
    except (TypeError, ValueError):
        index = 0
    if index < len(_QUESTION_VARIATION_CONSTRAINTS):
        return _QUESTION_VARIATION_CONSTRAINTS[index]

    # A single modulo-based pool eventually repeats during long review/regenerate
    # sessions. Combine an operational constraint with one or two distinct
    # stakeholder pressures. This yields 18 * 8 * 8 = 1,152 deterministic
    # combinations beyond the initial direct constraints.
    base_constraints = _QUESTION_VARIATION_CONSTRAINTS[1:]
    offset = index - len(_QUESTION_VARIATION_CONSTRAINTS)
    base = base_constraints[offset % len(base_constraints)]
    pressure_round = offset // len(base_constraints)
    pressure_count = len(_QUESTION_VARIATION_STAKEHOLDER_PRESSURES)
    primary_index = pressure_round % pressure_count
    primary = _QUESTION_VARIATION_STAKEHOLDER_PRESSURES[primary_index]
    secondary_slot = (pressure_round // pressure_count) % pressure_count
    if secondary_slot == 0:
        return f"{base}이고, {primary}"
    secondary_candidates = [
        pressure
        for position, pressure in enumerate(_QUESTION_VARIATION_STAKEHOLDER_PRESSURES)
        if position != primary_index
    ]
    secondary = secondary_candidates[secondary_slot - 1]
    return f"{base}이고, {primary}이며, {secondary}"


def _focus_measurement_instruction(focus: str, focus_type: str) -> str:
    focus = str(focus or "핵심 수행기준").strip() or "핵심 수행기준"
    kind = _normalize_ksa_type(focus_type, focus)
    if kind == "지식":
        return "사용한 판단 근거와 적용 범위·예외를 밝혀 주세요."
    if kind == "태도":
        return (
            "상충하는 요구와 압박 속에서 선택한 행동과 그 선택으로 감수할 "
            "상충비용 또는 불이익을 밝혀 주세요."
        )
    return (
        "실제로 수행할 단계·조치, 산출물과 품질 확인 방법을 제시해 주세요."
    )


def _compact_material_reference(value: str, limit: int = 5, focus: str = "") -> str:
    """Keep task prompts readable while full materials remain in task_conditions."""

    limit = max(1, int(limit or 1))
    items = [
        item.strip()
        for item in re.split(r"[,，]", str(value or ""))
        if item.strip()
    ]
    items = list(dict.fromkeys(items))
    if len(items) <= limit:
        return ", ".join(items)
    if limit == 1:
        return items[0]
    stopwords = {"확인", "판단", "기준", "절차", "행동", "검토", "업무", "수행", "관련"}
    focus_tokens = [
        token
        for token in re.findall(r"[0-9A-Za-z가-힣]{2,}", str(focus or ""))
        if token not in stopwords
    ]
    priority_indices = [
        index
        for index, item in enumerate(items)
        if any(token in item for token in focus_tokens)
    ][:limit]
    remaining_indices = [index for index in range(len(items)) if index not in priority_indices]
    remaining_slots = limit - len(priority_indices)
    if remaining_slots >= len(remaining_indices):
        sampled_indices = remaining_indices
    elif remaining_slots == 1:
        sampled_indices = [remaining_indices[len(remaining_indices) // 2]]
    else:
        sampled_indices = [
            remaining_indices[(position * (len(remaining_indices) - 1)) // (remaining_slots - 1)]
            for position in range(remaining_slots)
        ]
    selected = [items[index] for index in sorted({*priority_indices, *sampled_indices})]
    return ", ".join(dict.fromkeys(selected))


def _ensure_question_material_reference(
    question: Any,
    method: str,
    conditions: dict[str, Any] | None,
) -> str:
    """Make concrete case rows visible in presentation/debate questions."""

    text = str(question or "").strip()
    if method not in {"발표면접", "토론면접"} or not text:
        return text
    if "[제공자료]" in text or "[공통자료]" in text:
        return text
    if not isinstance(conditions, dict):
        return text
    rows = conditions.get("case_materials") if isinstance(conditions.get("case_materials"), list) else []
    concrete = [
        {
            "source": str(row.get("source") or "").strip(),
            "field": str(row.get("field") or "").strip(),
            "value": str(row.get("value") or "").strip(),
        }
        for row in rows
        if isinstance(row, dict)
        and all(str(row.get(field) or "").strip() for field in ("source", "field", "value"))
    ]
    if not concrete:
        return text
    label = "[제공자료]" if method == "발표면접" else "[공통자료]"
    material_text = "; ".join(
        f"{row['source']}/{row['field']}={row['value'][:120]}"
        for row in concrete[:5]
    )
    return f"{text} {label} {material_text}"


def _strip_question_material_reference(value: Any) -> str:
    """Remove the printable material appendix for repeat/dedup comparison."""

    text = str(value or "").strip()
    return re.sub(r"\s+\[(?:제공자료|공통자료)\].*$", "", text).strip()


def _freeform_ksa_anchor(
    focus: Any,
    focus_type: Any = "",
    task_frame: dict[str, str] | None = None,
) -> str:
    """Turn an internal KSA surface into a short, readable question anchor.

    ``question_focus_surface`` is deliberately verbose because it is also an
    audit key.  It must not become the candidate-facing sentence.  Fallback
    questions use only the semantic noun, leaving the provider prompt free to
    choose its own incident and wording.
    """

    candidates = [
        str((task_frame or {}).get("task_object") or "").strip(),
        str(focus or "").strip(),
    ]
    kind = _normalize_ksa_type(focus_type, str(focus or ""))
    for raw in candidates:
        anchor = _clean_question_text(raw, max_chars=80)
        if not anchor:
            continue
        anchor = re.sub(
            r"\s*(?:관련\s*)?(?:실무\s*)?(?:적용·검증\s*절차|수행·검증\s*절차|"
            r"확인·판단\s*기준|행동\s*기준|확인\s*절차|검증\s*절차)\s*$",
            "",
            anchor,
        ).strip(" ·ㆍ/")
        if kind == "기술":
            anchor = re.sub(r"\s*(?:능력|기술|스킬)\s*$", "", anchor).strip()
        elif kind == "지식":
            anchor = re.sub(r"\s*지식\s*$", "", anchor).strip()
        if anchor and len(_ksa_key(anchor)) >= 2:
            return anchor
    return "핵심 업무"


def _freeform_ksa_question(
    *,
    method: str,
    subject: str,
    focus: str,
    detail: str,
    comp_def: str,
    focus_type: str = "",
    variation_index: int = 0,
    task_frame: dict[str, str] | None = None,
    job_context_text: str = "",
) -> str:
    """Build a concise KSA-anchored repair without the stock STAR scaffold.

    This path is used only when an OpenRouter candidate fails the editorial
    gate.  It keeps the evidence anchor and method shape, but does not append
    a fixed situation/role/document/result checklist to the candidate's text.
    """

    frame = dict(task_frame or _question_task_frame(
        focus=focus,
        focus_type=focus_type,
        subject=subject,
        detail=detail,
        comp_def=comp_def,
    ))
    label = str(subject or detail or "해당 직무").strip() or "해당 직무"
    anchor = _freeform_ksa_anchor(focus, focus_type, frame)
    variant = max(0, int(variation_index or 0)) % 3

    if method == "경험면접":
        variants = (
            f"{label}에서 {anchor}가 필요했던 실제 경험을 하나 들려 주세요. "
            "그때 본인이 내린 결정과 확인된 결과는 무엇이었나요?",
            f"{label} 업무 중 {anchor}와 관련해 예상과 다른 결과를 만난 사례가 있나요? "
            "무엇을 근거로 대응했는지 설명해 주세요.",
            f"{label}에서 {anchor}를 활용해 문제를 해결했던 일을 말씀해 주세요. "
            "본인의 역할과 가장 중요한 결과는 무엇이었나요?",
        )
        return variants[variant]
    if method == "상황면접":
        return (
            f"{label}에서 {anchor}와 관련해 예상과 다른 정보가 나온 상황입니다. "
            "무엇을 먼저 확인하고 어떤 조치를 선택하시겠습니까?"
        )
    if method == "직무지식면접":
        return (
            f"{label}의 {anchor}를 실제 업무에 적용할 때, 먼저 확인할 근거와 "
            "예외가 생기면 조정할 부분을 설명해 주세요."
        )
    if method == "인바스켓면접":
        return (
            f"{label}에서 {anchor}와 관련한 요청이 동시에 들어왔습니다. "
            "무엇을 먼저 판단하고 어떤 조치를 남기시겠습니까?"
        )
    if method == "발표면접":
        return (
            f"[발표과제] 제공된 직무 자료에서 {anchor}와 관련한 문제 한 가지를 골라, "
            "핵심 근거와 선택한 대응을 발표해 주세요."
        )
    if method == "토론면접":
        return (
            f"[토론과제] {label}에서 {anchor}와 관련해 서로 다른 처리안이 필요한 상황입니다. "
            "각 안의 근거와 위험을 비교한 뒤 공통 실행안을 제시해 주세요."
        )
    if method == "창의적 문제해결력면접":
        return (
            f"[창의적 문제해결력과제] {label}에서 {anchor}와 관련한 문제가 반복되고 있습니다. "
            "문제를 새롭게 정의하고 대안과 검증 방법을 제시해 주세요."
        )
    return _question_for_method(
        method=method,
        subject=subject,
        focus=focus,
        detail=detail,
        comp_def=comp_def,
        focus_type=focus_type,
        variation_index=variation_index,
        task_frame=task_frame,
        job_context_text=job_context_text,
    )


def _question_for_method(
    method: str,
    subject: str,
    focus: str,
    detail: str,
    comp_def: str,
    focus_type: str = "",
    variation_index: int = 0,
    task_frame: dict[str, str] | None = None,
    job_context_text: str = "",
    freeform_ksa: bool = False,
) -> str:
    if freeform_ksa:
        return _freeform_ksa_question(
            method=method,
            subject=subject,
            focus=focus,
            detail=detail,
            comp_def=comp_def,
            focus_type=focus_type,
            variation_index=variation_index,
            task_frame=task_frame,
            job_context_text=job_context_text,
        )
    label = subject or detail or "해당 직무"
    # Carry a short source-derived anchor into deterministic repairs so a
    # provider timeout cannot erase the actual duty from the candidate-facing
    # prompt.  The full source bundle remains in the audit payload/materials.
    source_anchor = ""
    if str(job_context_text or "").strip():
        for _source_label, _source_value in re.findall(
            r"\[([^\]]+)\]([^\[]*)", str(job_context_text)
        ):
            source_parts = _presentation_source_segments(
                _source_value,
                [detail, subject, focus],
                limit=2,
            )
            if source_parts:
                source_anchor = re.sub(r"\s+", " ", source_parts[0]).strip()
                if len(source_anchor) > 84:
                    source_anchor = source_anchor[:84].rsplit(" ", 1)[0].rstrip(" ,·") + "…"
                break
        if source_anchor:
            label = f"{label} · 공고·직무기술서 근거: {source_anchor}"
    focus = focus or "핵심 수행기준"
    focus_type = _normalize_ksa_type(focus_type, focus)
    frame = dict(task_frame or _question_task_frame(
        focus=focus,
        focus_type=focus_type,
        subject=subject,
        detail=detail,
        comp_def=comp_def,
    ))
    prompt_focus = str(frame.get("task_object") or _operational_focus_label(focus, focus_type)).strip()
    context = _domain_context_pack(detail=detail, subject=subject, focus=focus, comp_def=comp_def)
    variation = _question_variation_constraint(variation_index)
    variation_sentence = f"또한 {variation}입니다. " if variation else ""
    variation_clause = f"{variation}일 때, " if variation else ""
    measurement_instruction = _focus_measurement_instruction(prompt_focus, focus_type)
    prompt_evidence = _compact_material_reference(context["evidence"], focus=prompt_focus)
    prompt_inbasket = _compact_material_reference(context["inbasket"], focus=prompt_focus)
    if method == "발표면접":
        if focus_type == "지식":
            presentation_material_context = (
                f"{_quoted_with_josa(prompt_focus, '을', '를')} 적용해 현황을 진단하는 데 필요한"
            )
        elif focus_type == "태도":
            presentation_material_context = (
                f"{_quoted_with_josa(prompt_focus, '이', '가')} 드러나는 선택을 판단하기 위한"
            )
        else:
            presentation_material_context = (
                f"{_quoted_with_josa(prompt_focus, '을', '를')} {_skill_focus_action(prompt_focus)} 얻은 결과와 품질을 검증하기 위한"
            )
        return (
            f"[발표과제] {label} 업무에서 {presentation_material_context} "
            f"{_with_josa(prompt_evidence, '이', '가')} 주어졌습니다. {variation_sentence}"
            "자료를 바탕으로 현황 문제를 진단하고 개선안을 발표한 뒤 질의응답에 답해 주세요. "
            "발표에는 현황 진단, 원인 분석, 대안 2가지, 실행 우선순위, 성과지표와 질의응답의 판단 근거를 포함하세요. "
            f"{measurement_instruction}"
        )
    if method == "토론면접":
        dilemma = str(frame.get("decision_dilemma") or context["debate"]).strip()
        situation = _scenario_condition_text(context["situation"])
        debate_variation = variation_sentence if len(situation) < 90 else ""
        if focus_type == "태도":
            debate_focus_context = (
                f"압박과 이해 충돌 속에서도 "
                f"{_quoted_with_josa(prompt_focus, '을', '를')} 지켜야 하는 가운데"
            )
            decision_requirement = "제시안에는 공동 행동·점검·실행 책임을 명시하세요."
        elif focus_type == "지식":
            debate_focus_context = (
                f"{_quoted_with_josa(prompt_focus, '을', '를')} 적용해야 하는 가운데"
            )
            decision_requirement = "제시안에는 적용 범위·예외·검증·실행 책임을 명시하세요."
        else:
            debate_focus_context = (
                f"{prompt_focus}에 따라 실제 수행 단계·산출물·품질 검증 방법을 정해야 하는 가운데"
            )
            decision_requirement = "제시안에는 수행 단계·산출물·품질 검증·실행 책임을 명시하세요."
        return (
            f"[토론과제] {label} 업무에서 {debate_focus_context} "
            f"{situation}입니다. {_with_josa(dilemma, '이', '가')} 충돌합니다. "
            f"{debate_variation}"
            "각 입장의 근거·위험·타당성을 검토하세요. 합의할 수 있다면 공통 실행안을, "
            "합의가 어렵다면 미합의 쟁점과 결정권자 이송 기준을 제시해 주세요. "
            f"{decision_requirement}"
        )
    if method == "인바스켓면접":
        if focus_type == "태도":
            return (
                f"[인바스켓과제] {label} 관련 문서·요청이 동시에 들어왔습니다. 제공 항목은 {prompt_inbasket}입니다. "
                f"{variation_sentence}"
                f"상충하는 요구와 압박 속에서도 {_quoted_with_josa(prompt_focus, '이', '가')} 드러나는 처리 우선순위와 "
                "상급자 보고, 위임, 직접처리 판단 및 초기 행동을 제시하고, 그 선택으로 감수할 "
                "상충비용 또는 불이익을 밝혀 주세요."
            )
        elif focus_type == "지식":
            return (
                f"[인바스켓과제] {label} 관련 문서·요청이 동시에 들어왔습니다. 제공 항목은 {prompt_inbasket}입니다. "
                f"{variation_sentence}"
                f"{_quoted_with_josa(prompt_focus, '을', '를')} 적용해 처리 우선순위와 상급자 보고, "
                "위임, 직접처리 판단 및 초기 행동을 제시하고, 그 판단 근거와 적용 범위·예외를 밝혀 주세요."
            )
        else:
            return (
                f"[인바스켓과제] {label} 관련 문서·요청이 동시에 들어왔습니다. 제공 항목은 {prompt_inbasket}입니다. "
                f"{variation_sentence}"
                f"{_quoted_with_josa(prompt_focus, '을', '를')} 실제로 {_skill_focus_action(prompt_focus)} 처리 우선순위와 상급자 보고, "
                "위임, 직접처리 판단 및 초기 행동을 제시하고, 수행 단계·조치·산출물·품질 확인 방법을 밝혀 주세요."
            )
    if method == "상황면접":
        situation_text = _scenario_condition_text(context.get("situation_prompt") or context["situation"])
        if focus_type == "태도":
            situation_focus_context = (
                f"{situation_text}에서 {_quoted_with_josa(prompt_focus, '을', '를')} 지키기 어려운 조건"
            )
        elif focus_type == "지식":
            situation_focus_context = (
                f"{_quoted_with_josa(prompt_focus, '을', '를')} 적용해야 하는 가운데 "
                f"{situation_text}"
            )
        else:
            situation_focus_context = (
                f"{_quoted_with_josa(prompt_focus, '을', '를')} 실제로 {_skill_focus_action(prompt_focus, 'must')} 하는 가운데 "
                f"{situation_text}"
            )
        return (
            f"{label} 업무 중 {situation_focus_context}입니다. "
            f"{variation_sentence}"
            "어떤 기준으로 판단하고 위험요인을 어떻게 통제하며, 사실 확인부터 보고와 실행까지 어떤 순서로 행동하시겠습니까? "
            f"{measurement_instruction}"
        )
    if method == "직무지식면접":
        if focus_type == "지식":
            knowledge_task = (
                f"{_quoted_with_josa(prompt_focus, '을', '를')} 토대로 절차·판단 기준과 근거, "
                "적용 범위·핵심 예외상황, 산출물 품질점검과 오류 예방 방법"
            )
        elif focus_type == "태도":
            knowledge_task = (
                f"{_quoted_with_josa(prompt_focus, '이', '가')} 요구되는 상황의 절차·판단 기준과 핵심 산출물, "
                "상충하는 요구나 예외상황 속 선택 행동, 감수할 결과와 오류 예방 방법"
            )
        else:
            knowledge_task = (
                f"{_quoted_with_josa(prompt_focus, '을', '를')} {_skill_focus_action(prompt_focus, 'future')} 때의 절차·판단 기준·수행 순서, "
                "핵심 산출물과 예외상황, 품질 확인과 오류 예방 방법"
            )
        return (
            f"{label}에서 {variation_clause}{_with_josa(knowledge_task, '을', '를')} 말씀해 주세요."
        )
    if method == "창의적 문제해결력면접":
        if focus_type == "지식":
            creative_focus_context = (
                f"{_quoted_with_josa(prompt_focus, '을', '를')} 적용해야 하는 복합 문제"
            )
        elif focus_type == "태도":
            creative_focus_context = (
                f"압박 속에서 {_quoted_with_josa(prompt_focus, '이', '가')} 드러나는 선택을 요구하는 복합 문제"
            )
        else:
            creative_focus_context = (
                f"{_quoted_with_josa(prompt_focus, '을', '를')} 실제로 {_skill_focus_action(prompt_focus)} 해결해야 하는 복합 문제"
            )
        return (
            f"[창의적 문제해결력과제] {label} 업무에서 {creative_focus_context}가 발생했습니다. "
            f"{variation_sentence}"
            f"제공된 {_with_josa(prompt_evidence, '을', '를')} 바탕으로 미래예측 관점에서 핵심 문제를 정의하고 해결안을 설명한 뒤 "
            "질의응답에 답해 주세요. 해결안에는 원인 가설이 포함된 문제 정의, 비교 가능한 대안 2가지, "
            f"선택안의 검증 방법, 우선 실행계획과 성과지표를 포함하세요. {measurement_instruction}"
        )
    experience_context = (
        f"{_with_josa(label, '과', '와')} 관련해 본인이 판단하고 행동한 실제 경험 한 가지를 선택해 주세요. "
        "직무 경험이 없다면 본인 역할이 분명한 프로젝트나 교육실습 사례도 가능합니다."
    )
    if variation:
        experience_variation = _experience_scenario_condition(variation)
        experience_context += f" 특히 {experience_variation}과 유사한 사례를 우선해 주세요."
    if focus_type == "태도":
        evidence_task = (
            f"그 경험에서 압박이나 상충 요구가 있었던 장면을 골라 {_quoted_with_josa(prompt_focus, '과', '와')} 관련해 "
            "본인이 선택한 행동과 그 결과"
        )
    elif focus_type == "지식":
        evidence_task = (
            f"그 경험에서 {_quoted_with_josa(prompt_focus, '을', '를')} 실제 판단에 사용한 장면을 골라 무엇을 확인했고 어떤 기준으로 판단해 "
            "본인이 어떻게 행동했는지"
        )
    else:
        evidence_task = (
            f"그 경험에서 {_quoted_with_josa(prompt_focus, '이', '가')} 요구된 장면을 골라 어떤 순서와 조치로 "
            "산출물을 만들고 결과를 확인했는지"
        )
    return (
        f"{experience_context} {_with_josa(evidence_task, '을', '를')} 설명해 주세요. "
        "당시 상황과 문제, 본인 역할, 실제 행동과 결과를 포함해 주세요."
    )


def _followups_for_method(
    method: str,
    subject: str,
    focus: str,
    count: int,
    variant_index: int = 0,
    focus_type: str = "",
    task_frame: dict[str, str] | None = None,
) -> list[str]:
    if count <= 0:
        return []
    label = subject or "해당 업무"
    focus = focus or "핵심 수행기준"
    focus_type = _normalize_ksa_type(focus_type, focus)
    frame = dict(task_frame or _question_task_frame(
        focus=focus,
        focus_type=focus_type,
        subject=subject,
        detail="",
        comp_def="",
    ))
    prompt_focus = str(frame.get("task_object") or _operational_focus_label(focus, focus_type)).strip()
    context = _domain_context_pack(detail="", subject=subject, focus=focus, comp_def="")
    if focus_type == "태도":
        experience_focus_probe = (
            f"그 선택에서 {_quoted_with_josa(prompt_focus, '이', '가')} 드러난 실제 행동과 감수한 상충비용은 무엇이었습니까?"
        )
        situation_focus_probe = (
            f"그 선택에서 {_quoted_with_josa(prompt_focus, '이', '가')} 드러난 구체적 행동과 예상 위험은 무엇입니까?"
        )
        inbasket_focus_probe = (
            f"압박 속에서도 {_quoted_with_josa(prompt_focus, '이', '가')} 드러나게 문서와 요청을 어떻게 분류하겠습니까?"
        )
        creative_focus_probe = (
            f"대안 선택 과정에서 {_quoted_with_josa(prompt_focus, '이', '가')} 드러날 행동과 상충비용은 무엇입니까?"
        )
    elif focus_type == "지식":
        experience_focus_probe = (
            f"{_quoted_with_josa(prompt_focus, '을', '를')} 적용할 때 어떤 자료를 확인했고 적용 범위와 예외를 어떻게 판별했습니까?"
        )
        situation_focus_probe = (
            f"{_quoted_with_josa(prompt_focus, '을', '를')} 적용할 때 근거 자료를 확인한 뒤 범위와 예외를 어떻게 구분하겠습니까?"
        )
        inbasket_focus_probe = (
            f"{_quoted_with_josa(prompt_focus, '을', '를')} 적용해 문서와 요청을 분류할 기준·범위·예외는 무엇입니까?"
        )
        creative_focus_probe = (
            f"{_quoted_with_josa(prompt_focus, '을', '를')} 적용해 원인 가설을 세울 때 어떤 자료로 예외 가능성을 검증하겠습니까?"
        )
    else:
        experience_focus_probe = (
            f"{_quoted_with_josa(prompt_focus, '을', '를')} 실제로 {_skill_focus_action(prompt_focus, 'past')} 순서·도구·조치와 산출물은 무엇이었습니까?"
        )
        situation_focus_probe = (
            f"{_quoted_with_josa(prompt_focus, '을', '를')} 실제로 {_skill_focus_action(prompt_focus, 'future')} 순서·조치·산출물과 예상 위험은 무엇입니까?"
        )
        inbasket_focus_probe = (
            f"{_quoted_with_josa(prompt_focus, '을', '를')} 실제로 {_skill_focus_action(prompt_focus)} 문서와 요청을 분류할 절차와 산출물은 무엇입니까?"
        )
        creative_focus_probe = (
            f"{_quoted_with_josa(prompt_focus, '을', '를')} 실제로 {_skill_focus_action(prompt_focus, 'future')} 분석 절차·산출물과 검증 방법은 무엇입니까?"
        )
    banks = {
        "경험면접": [
            "당시 상황과 본인이 맡은 역할을 구체적으로 설명해 주세요.",
            experience_focus_probe,
            "다른 선택지와 비교해 그 행동을 선택한 기준은 무엇이었습니까?",
            "성과를 어떤 기준이나 지표로 확인했습니까?",
            "같은 상황이 다시 주어진다면 어떤 점을 개선하시겠습니까?",
        ],
        "상황면접": [
            "판단 전에 먼저 확인해야 할 사실과 기준은 무엇입니까?",
            situation_focus_probe,
            f"{context['stakeholders']} 등 이해관계자에게 어떤 순서와 방식으로 설명하시겠습니까?",
            "결과가 기대와 다르게 나오면 어떤 후속 조치를 하시겠습니까?",
            "같은 문제가 반복되지 않도록 어떤 예방 장치를 두시겠습니까?",
        ],
        "발표면접": [
            f"{_with_josa(prompt_focus, '을', '를')} 발표에서 진단할 때 핵심 근거 자료는 무엇입니까?",
            "대안 중 우선순위를 가장 높게 둔 방안과 그 이유는 무엇입니까?",
            "면접위원이 반대 의견을 제시한다면 어떤 근거로 답변하시겠습니까?",
            "실행 일정, 필요 자원, 성과지표를 어떻게 설정하겠습니까?",
            f"{label} 현장에 적용할 때 가장 큰 리스크와 보완책은 무엇입니까?",
        ],
        "토론면접": [
            f"{label}의 {prompt_focus}에 관한 초기 입장을 정하기 전에 어떤 문서와 사실을 확인하겠습니까?",
            "반대 입장에서 수용할 부분과 수용하지 않을 부분을 어떤 기준으로 구분하겠습니까?",
            "두 입장의 근거가 충돌할 때 쟁점을 어떻게 조정하시겠습니까?",
            "공통 실행안의 적용 범위·예외·검증 기준은 무엇이며, 합의가 어렵다면 어떤 쟁점을 이송하겠습니까?",
            "결정 이후 실행 책임과 후속 점검은 어떻게 정리하시겠습니까?",
        ],
        "인바스켓면접": [
            inbasket_focus_probe,
            "가장 먼저 처리할 항목과 보류할 항목을 각각 무엇으로 보겠습니까?",
            f"{context['stakeholders']} 중 누구에게 보고, 위임, 직접 처리할지 어떻게 선택하겠습니까?",
            "마감 지연이나 민원 확대 가능성은 어떻게 통제하겠습니까?",
            "30분 이후 후속 확인과 기록은 어떻게 남기겠습니까?",
        ],
        "직무지식면접": [
            f"{_with_josa(prompt_focus, '을', '를')} 업무에 활용할 때 반드시 확인할 기준이나 규정은 무엇입니까?",
            "그 기준을 실제 업무에 적용할 때 자주 발생하는 예외상황은 무엇입니까?",
            "해당 산출물의 품질을 어떻게 점검하겠습니까?",
            "잘못 적용했을 때 발생할 수 있는 리스크와 보완책은 무엇입니까?",
            "신규 담당자에게 이 절차를 설명한다면 어떤 순서로 교육하겠습니까?",
        ],
        "창의적 문제해결력면접": [
            "핵심 문제정의를 위해 먼저 확인할 사실과 기준은 무엇입니까?",
            creative_focus_probe,
            "대안 중 실행 우선순위를 높게 둘 방안과 그 이유는 무엇입니까?",
            "선택한 대안의 리스크와 보완책은 어떻게 정리하겠습니까?",
            "성과지표와 후속 점검 기준은 무엇으로 설정하겠습니까?",
        ],
    }
    bank = banks.get(method, banks["경험면접"])
    limit = max(0, min(5, count))
    if limit <= 0:
        return []
    if method == "토론면접":
        # Debate probes are a deliberate sequence: evidence for the initial
        # position, treatment of the opposing argument, issue adjustment,
        # then outcome and accountability. Rotating this order makes a panel
        # ask about execution before it has tested listening or adjustment.
        return list(bank[:limit])
    if limit >= len(bank):
        return list(bank[:limit])

    focus_slot = _FOLLOW_UP_FOCUS_SLOT_INDEX.get(method, 1)
    locked_indices = [0]
    if 0 <= focus_slot < len(bank) and focus_slot not in locked_indices:
        locked_indices.append(focus_slot)
    locked_indices = [idx for idx in locked_indices if idx < len(bank)]

    selected: list[str] = []
    selected_indices: set[int] = set()
    for bank_index in locked_indices:
        if len(selected) >= limit:
            break
        selected.append(bank[bank_index])
        selected_indices.add(bank_index)

    candidates = [idx for idx in range(len(bank)) if idx not in selected_indices]
    if candidates:
        offset = max(0, int(variant_index or 0)) % len(candidates)
        rotated = candidates[offset:] + candidates[:offset]
        for bank_index in rotated:
            if len(selected) >= limit:
                break
            selected.append(bank[bank_index])
    return selected[:limit]


def _clean_question_items(values: Any, limit: int) -> list[str]:
    if not isinstance(values, list):
        return []
    out: list[str] = []
    seen: set[str] = set()
    for value in values:
        text = _clean_question_text(value, max_chars=140)
        key = normalize_question_dedup_key(text)
        if not text or not key or key in seen:
            continue
        if _contains_blind_hiring_cue(text):
            continue
        seen.add(key)
        out.append(text)
        if len(out) >= limit:
            break
    return out


def _merge_question_items(primary: list[str], fallback: list[str], limit: int) -> list[str]:
    out: list[str] = []
    seen: set[str] = set()
    for value in primary + fallback:
        text = _clean_question_text(value, max_chars=140)
        key = normalize_question_dedup_key(text)
        if not text or not key or key in seen:
            continue
        if _contains_blind_hiring_cue(text):
            continue
        seen.add(key)
        out.append(text)
        if len(out) >= limit:
            break
    return out


_GENERAL_QUESTION_INTENTS = GENERAL_QUESTION_INTENTS
_FOCUS_SCOPED_GENERAL_QUESTION_INTENTS = FOCUS_SCOPED_GENERAL_QUESTION_INTENTS
_QUESTION_INTENT_PATTERNS = QUESTION_INTENT_PATTERNS


def _compact_question_text(value: Any) -> str:
    return re.sub(r"\s+", "", normalize_question_dedup_key(str(value or "")))


def _compact_question_intent_text(value: Any) -> str:
    raw = str(value or "").strip().lower()
    raw = re.sub(r"[^0-9a-z가-힣]+", " ", raw)
    return re.sub(r"\s+", "", raw)


def _question_intent_key(question: str) -> str:
    return classify_question_intent(question, unknown="other")


def _question_intent_for_repeat_signature(question: str) -> str:
    intent = _question_intent_key(question)
    return "" if intent == "other" else intent


def _question_focus_signature(value: Any) -> str:
    return _compact_question_text(value)[:80]


def _question_text_similarity(left: Any, right: Any) -> float:
    left_key = _compact_question_text(left)
    right_key = _compact_question_text(right)
    if not left_key or not right_key:
        return 0.0
    if left_key == right_key:
        return 1.0
    return SequenceMatcher(None, left_key, right_key).ratio()


def _question_repeat_signature(item: dict[str, Any]) -> str:
    question = str((item or {}).get("question") or "").strip()
    intent = _question_intent_for_repeat_signature(question)
    if not intent:
        return ""
    method = str((item or {}).get("method") or (item or {}).get("type") or "").strip()
    method_key = _compact_question_text(method)
    focus = str((item or {}).get("question_focus") or "").strip()
    refs = item.get("ksa_refs")
    if not focus and isinstance(refs, list) and refs:
        focus = str(refs[0] or "").strip()
    variation_axis = str((item or {}).get("question_variation_axis") or "").strip()
    if variation_axis:
        # The provider-free fallback exposes a deliberate, structured angle
        # (priority, evidence comparison, reporting, etc.) while preserving the
        # same official KSA factor. Include that axis in dedup identity so a
        # five-slot fallback set is not incorrectly reported as a repeat merely
        # because the competency and KSA are shared.
        focus = f"{focus} · {variation_axis}".strip(" ·")
    has_focus_scope = bool(focus)
    if intent in _GENERAL_QUESTION_INTENTS and not (
        intent in _FOCUS_SCOPED_GENERAL_QUESTION_INTENTS and has_focus_scope
    ):
        return f"{intent}|general"
    if focus:
        return f"{intent}|{method_key}|focus:{_question_focus_signature(focus)}"
    subject = str((item or {}).get("ncsClCd") or (item or {}).get("competency") or "").strip()
    return f"{intent}|{method_key}|{_question_focus_signature(subject)}"


def _question_near_repeat(item: dict[str, Any], previous: dict[str, Any]) -> bool:
    signature = _question_repeat_signature(item)
    if not signature or signature != _question_repeat_signature(previous):
        return False
    if signature.endswith("|general"):
        return True
    question = str((item or {}).get("question") or "").strip()
    previous_question = str((previous or {}).get("question") or "").strip()
    if not question or not previous_question:
        return False
    if normalize_question_dedup_key(question) == normalize_question_dedup_key(previous_question):
        return True
    if min(len(_compact_question_text(question)), len(_compact_question_text(previous_question))) < 24:
        return False
    return _question_text_similarity(question, previous_question) >= 0.88


def _raw_model_scenarios_are_distinct(
    item: dict[str, Any],
    previous_items: list[dict[str, Any]],
) -> bool:
    """Keep distinct model scenarios from collapsing after template repair."""

    raw_question = str(item.get("model_question_raw") or "").strip()
    if not raw_question:
        return False
    current_probe = dict(item)
    current_probe["question"] = raw_question
    previous_probes: list[dict[str, Any]] = []
    for previous in previous_items:
        previous_raw = str(previous.get("model_question_raw") or "").strip()
        if not previous_raw:
            continue
        probe = dict(previous)
        probe["question"] = previous_raw
        previous_probes.append(probe)
    return bool(
        previous_probes
        and not any(_question_near_repeat(current_probe, previous) for previous in previous_probes)
    )


def _refresh_question_repeat_metadata(items: list[dict[str, Any]]) -> None:
    seen_by_signature: dict[str, list[dict[str, Any]]] = {}
    for item in items:
        if not isinstance(item, dict):
            continue
        signature = _question_repeat_signature(item)
        previous_items = seen_by_signature.get(signature, []) if signature else []
        repeat_duplicate = any(_question_near_repeat(item, previous) for previous in previous_items)
        if repeat_duplicate and _raw_model_scenarios_are_distinct(item, previous_items):
            repeat_duplicate = False

        item["question_intent"] = _question_intent_key(str(item.get("question") or ""))
        item["question_repeat_signature"] = signature
        item["question_repeat_duplicate"] = bool(repeat_duplicate)
        if signature:
            seen_by_signature.setdefault(signature, []).append(dict(item))


def _pick_alternate_question_focus(
    current_focus: str,
    ksa_terms: list[str],
    used_focuses: set[str],
) -> str:
    current_key = _question_focus_signature(current_focus)
    base_focus = _clean_question_text(current_focus, max_chars=44)
    fallback_terms = [
        f"{base_focus} 사실 확인",
        f"{base_focus} 판단 기준",
        f"{base_focus} 실행 우선순위",
        f"{base_focus} 오류 예방",
        f"{base_focus} 이해관계자 조정",
        f"{base_focus} 결과 점검",
    ] if base_focus else []
    for term in [*(ksa_terms or []), *fallback_terms]:
        text = _clean_question_text(term, max_chars=60)
        key = _question_focus_signature(text)
        if not text or not key or key == current_key or key in used_focuses:
            continue
        if _contains_blind_hiring_cue(text):
            continue
        return text
    return ""


def _compact_contains_term(text: str, term: str) -> bool:
    compact_text = re.sub(r"\s+", "", str(text or "")).lower()
    compact_term = re.sub(r"\s+", "", str(term or "")).lower()
    return bool(compact_text and compact_term and compact_term in compact_text)


def _follow_up_job_context(q: dict[str, Any], focus: str = "") -> str:
    focus = _clean_question_text(focus, max_chars=60)
    for key in (
        "competency",
        "compeUnitName",
        "required_job_context",
        "ncs_detail",
        "ncsSubdCdnm",
        "matchedDetailName",
        "ncsSclasCdnm",
    ):
        context = _clean_question_text(q.get(key), max_chars=60)
        if not context:
            continue
        if focus and _compact_contains_term(context, focus):
            continue
        if focus and _compact_contains_term(focus, context):
            continue
        return context
    return ""


def _inject_focus_into_follow_up(method: str, focus: str, follow_up: str, job_context: str = "") -> str:
    text = _clean_question_text(follow_up, max_chars=140)
    focus = _clean_question_text(focus, max_chars=60)
    if not text or not focus or _ksa_factor_relevant_to_text(focus, text):
        return text
    context = _clean_question_text(job_context, max_chars=60)
    context_part = ""
    if context and not _compact_contains_term(text, context) and not _compact_contains_term(focus, context):
        context_part = f" {context} 상황에서" if method == "상황면접" else f" {context} 업무에서"
    prefix_by_method = {
        "경험면접": f"{_quoted_with_josa(focus, '을', '를')} 적용하는 과정에서{context_part} 본인 행동과 선택 이유를 중심으로",
        "상황면접": f"{_quoted_with_josa(focus, '과', '와')} 관련해{context_part}",
        "발표면접": f"{_quoted_with_josa(focus, '을', '를')} 발표 쟁점으로 볼 때{context_part}",
        "토론면접": f"{_quoted_with_josa(focus, '을', '를')} 토론 쟁점으로 볼 때{context_part}",
        "인바스켓면접": f"{_quoted_with_josa(focus, '을', '를')} 처리 기준으로 삼고 보고·위임·직접처리 판단까지 포함할 때{context_part}",
        "직무지식면접": f"{_quoted_with_josa(focus, '과', '와')} 관련한 기준과 예외상황·산출물 품질까지 포함해{context_part}",
        "창의적 문제해결력면접": f"{_quoted_with_josa(focus, '과', '와')} 관련한 원인과 대안 관점에서{context_part}",
    }
    prefix = prefix_by_method.get(method, f"{_quoted_with_josa(focus, '과', '와')} 관련해{context_part}")
    body_limit = max(35, 138 - len(prefix))
    body = _clean_question_text(text, max_chars=body_limit)
    return _clean_question_text(f"{prefix} {body}", max_chars=140)


def _follow_ups_non_focus_shape_ok(method: str, follow_ups: list[str]) -> bool:
    clean = [str(item or "").strip() for item in follow_ups if str(item or "").strip()]
    if len(clean) < 3:
        return False
    if _contains_blind_hiring_cue("\n".join(clean)):
        return False
    keys = [normalize_question_dedup_key(item) for item in clean]
    if any(not key for key in keys) or len(set(keys)) != len(keys):
        return False

    compact_items = [re.sub(r"\s+", "", item) for item in clean]
    anchors = _FOLLOW_UP_METHOD_ANCHORS.get(method, ())
    anchor_hits = {
        anchor
        for anchor in anchors
        if any(anchor in compact for compact in compact_items)
    }
    if len(anchor_hits) < 2:
        return False

    open_prompt_hits = sum(
        1
        for item in clean[:3]
        if re.search(r"(무엇|어떤|어떻게|얼마|어땠|어떠|왜|기준|이유|설명|말씀|제시|확인|선택|평가|점검|정리)", item)
    )
    return open_prompt_hits >= 3


def _repair_model_followups_with_focus(
    method: str,
    q: dict[str, Any],
    follow_ups: list[str],
    limit: int,
) -> list[str]:
    focus = _clean_question_text(
        q.get("question_focus_surface") or q.get("question_focus"),
        max_chars=60,
    )
    count = max(0, min(5, int(limit or 0)))
    clean = _merge_question_items(follow_ups, [], count)
    if count < 3 or len(clean) < 3 or not focus:
        return []
    if _follow_ups_quality_ok(method, q, clean):
        return []

    job_context = _follow_up_job_context(q, focus)
    preferred = min(_FOLLOW_UP_FOCUS_SLOT_INDEX.get(method, 1), len(clean) - 1)
    candidate_indices = list(dict.fromkeys([preferred, 1 if len(clean) > 1 else 0, 0, 2 if len(clean) > 2 else 0]))
    for target_index in candidate_indices:
        repaired = list(clean)
        repaired[target_index] = _inject_focus_into_follow_up(method, focus, repaired[target_index], job_context)
        if repaired[target_index] == clean[target_index]:
            continue
        if _contains_blind_hiring_cue("\n".join(repaired)):
            continue
        if _follow_ups_quality_ok(method, q, repaired):
            return repaired
    return []


def _repair_candidate_surface_text(value: Any, official_factor: str, surface_focus: str) -> tuple[str, bool]:
    """Replace one internal NCS label in candidate-visible text only."""

    text = str(value or "").strip()
    factor = str(official_factor or "").strip()
    surface = str(surface_focus or "").strip()
    if not text:
        return "", False
    if not factor or not surface or _ksa_key(factor) == _ksa_key(surface):
        return _collapse_repeated_question_phrases(text)
    repaired_text, surface_repaired = replace_official_ksa_surface(
        text,
        factor,
        surface,
    )
    repaired_text, phrase_repaired = _collapse_repeated_question_phrases(
        repaired_text
    )
    return repaired_text, bool(surface_repaired or phrase_repaired)


_REPEATED_MIDDLE_DOT_PHRASE_RE = re.compile(
    r"(?P<phrase>[0-9A-Za-z가-힣]+(?:\s+[0-9A-Za-z가-힣]+){0,3})"
    r"\s*·\s*(?P=phrase)"
    r"(?=(?:\s*·|[이가은는을를에에서로와과의,.;:!?\s]|$))"
)
_REPEATED_TWO_TOKEN_PHRASE_RE = re.compile(
    r"(?P<first>[0-9A-Za-z가-힣]+(?:[·ㆍ/][0-9A-Za-z가-힣]+)*)\s+"
    r"(?P<second>[0-9A-Za-z가-힣]+(?:[·ㆍ/][0-9A-Za-z가-힣]+)*)\s+"
    r"(?P=first)\s+(?P=second)"
    r"(?P<particle>은|는|이|가|을|를|에|에서|로|으로|와|과|의)?"
    r"(?=[\s,.;:!?]|$)"
)


def _collapse_repeated_question_phrases(value: Any) -> tuple[str, bool]:
    """Remove accidental consecutive phrase copies without rewriting meaning."""

    text = str(value or "").strip()
    if not text:
        return "", False
    original = text
    for _ in range(4):
        collapsed = _REPEATED_MIDDLE_DOT_PHRASE_RE.sub(r"\g<phrase>", text)
        collapsed = _REPEATED_TWO_TOKEN_PHRASE_RE.sub(
            r"\g<first> \g<second>\g<particle>",
            collapsed,
        )
        if collapsed == text:
            break
        text = collapsed
    return text, text != original


def _experience_surface_core(value: Any) -> str:
    """Keep the approved public KSA surface intact for candidate-facing text."""

    return _clean_question_text(value, max_chars=100)


def _experience_star_followups(
    *,
    focus_type: str,
    surface_focus: str,
    count: int,
) -> list[str]:
    """Create concise, answer-linked STAR probes for a provider-authored incident."""

    surface = _experience_surface_core(surface_focus) or "해당 업무"
    kind = str(focus_type or "").strip()
    if kind == "지식":
        action_probe = (
            f"앞서 언급한 판단에서 {surface}에 따라 어떤 규정·문서·자료를 확인했고, "
            "적용 범위를 그렇게 정한 이유는 무엇인가요?"
        )
    elif kind == "태도":
        action_probe = (
            f"앞서 언급한 압박이나 이해충돌 속에서 {_with_josa(surface, '을', '를')} 위해 "
            "어떤 행동을 직접 선택했고, 그 선택 근거와 감수한 점은 무엇인가요?"
        )
    else:
        action_probe = (
            f"앞서 언급한 {surface} 수행에서 사용한 자료·도구와 실제 순서, "
            "그 순서를 선택한 이유와 본인이 직접 취한 조치는 무엇인가요?"
        )
    probes = [
        "방금 말씀하신 사례에서 당시 본인이 맡은 역할과 구체적인 목표는 무엇이었나요?",
        action_probe,
        "답변에 결과 근거가 없다면, 결과를 확인한 문서·수치·기록·피드백과 이후 보완한 점을 설명해 주세요.",
        "방금 설명한 행동 외에 당시 고려한 다른 방법이 있었다면, 선택하지 않은 이유는 무엇인가요?",
        "앞서 언급한 경험을 같은 업무에 다시 적용한다면 무엇을 유지하거나 바꾸시겠습니까?",
    ]
    return probes[: max(0, min(5, int(count or 0)))]


def _experience_has_task_role(compact_text: str) -> bool:
    """Require an explicit candidate role/goal, not a generic '담당 업무' noun."""

    return bool(
        re.search(
            r"(?:본인|당시|맡은|담당한).{0,14}(?:역할|목표|책임)|"
            r"(?:역할|목표|책임).{0,10}(?:맡|담당)",
            str(compact_text or ""),
        )
    )


def _complete_experience_question_star(
    question: Any,
    *,
    focus_type: str,
    focus_surface: str = "",
) -> tuple[str, bool]:
    """Preserve a model-authored incident while filling missing STAR evidence slots."""

    text = _clean_question_text(question, max_chars=260)
    if not text:
        return "", False
    text, phrase_repaired = _collapse_repeated_question_phrases(text)
    compact = re.sub(r"\s+", "", text)
    has_task = _experience_has_task_role(compact)
    has_action = any(token in compact for token in ("본인이직접", "직접수행", "직접선택", "핵심행동"))
    has_result_evidence = (
        any(token in compact for token in ("결과", "성과", "영향"))
        and any(token in compact for token in ("문서", "수치", "기록", "피드백", "증빙"))
    )
    if has_task and has_action and has_result_evidence:
        return text, phrase_repaired

    sentences = [
        part.strip()
        for part in re.split(r"(?<=[.!?？])\s+", text)
        if part.strip()
    ]
    stem_sentence = next(
        (
            sentence
            for sentence in sentences
            if re.search(r"(경험|사례|말씀|설명|어떤|무엇|어떻게)", sentence)
        ),
        sentences[0] if sentences else text,
    )
    stem = stem_sentence.rstrip(" .?!。")
    stem_compact = re.sub(r"\s+", "", stem)
    has_task = _experience_has_task_role(stem_compact)
    has_action = any(
        token in stem_compact
        for token in ("본인이직접", "직접수행", "직접선택", "핵심행동")
    )
    has_result_evidence = (
        any(token in stem_compact for token in ("결과", "성과", "영향"))
        and any(
            token in stem_compact
            for token in ("문서", "수치", "기록", "피드백", "증빙")
        )
    )

    kind = str(focus_type or "").strip()
    surface = _experience_surface_core(focus_surface)
    needs_surface = bool(
        surface
        and _ksa_key(surface)
        and _ksa_key(surface) not in _ksa_key(text)
    )
    if kind == "지식":
        evidence_clause = (
            (f"{surface}에 따라 " if needs_surface else "")
            + "확인한 규정·문서·자료와 적용 범위의 판단 근거, 그에 따라 직접 취한 행동"
        )
    elif kind == "태도":
        evidence_clause = (
            (f"{_with_josa(surface, '이', '가')} 드러난 " if needs_surface else "")
            + "압박이나 이해충돌 상황에서 직접 선택한 행동과 선택 근거, 감수한 점"
        )
    else:
        evidence_clause = (
            (f"{surface}에 " if needs_surface else "")
            + "사용한 자료·도구, 수행 순서와 그 순서를 선택한 이유, 직접 취한 행동"
        )

    missing_slots: list[str] = []
    if not has_task:
        missing_slots.append("당시 맡은 역할과 목표")
    if not has_action:
        missing_slots.append(evidence_clause)
    if not has_result_evidence:
        missing_slots.append(
            "문서·수치·기록·피드백으로 확인한 결과"
        )
    if not missing_slots:
        return text, phrase_repaired
    requested_slots = _with_josa(", ".join(missing_slots), "을", "를")
    completion = f"{stem}. {requested_slots} 구체적으로 설명해 주세요."
    completion, _ = _collapse_repeated_question_phrases(completion)
    return _clean_question_text(completion, max_chars=280), True


def _rewrite_preserved_experience_duplicate(
    item: dict[str, Any],
    *,
    focus_type: str,
    surface_focus: str,
    variation_index: int,
) -> str:
    """Turn a duplicate provider draft into a distinct, evidence-locked STAR prompt."""

    surface = _experience_surface_core(surface_focus) or "해당 업무"
    variants = (
        (
            "기존 자료와 최신 기록이 서로 달라 그대로 처리하기 어려웠던",
            "차이를 대조하고 적용 범위를 결정해 직접 수정한 행동",
            "수정 전후 기록이나 오류 감소",
        ),
        (
            "마감 직전에 예외 항목이 발견되어 처리 기준을 다시 정해야 했던",
            "예외 근거와 영향 범위를 확인해 직접 내린 결정",
            "승인 문서나 최종 반영 내역",
        ),
        (
            "관련 부서의 요구가 충돌해 한 가지 처리 방향을 선택해야 했던",
            "상충하는 요구를 비교해 직접 제시하고 실행한 조정안",
            "협의 기록이나 후속 업무 결과",
        ),
        (
            "필수 근거나 증빙 일부가 누락되어 진행 여부를 판단해야 했던",
            "추가 확인 범위와 보류 기준을 정해 직접 취한 조치",
            "보완 문서나 재작업 방지 결과",
        ),
        (
            "기존 절차만으로 해결되지 않는 문제가 생겨 처리 방식을 바꿔야 했던",
            "대체 절차의 근거를 확인해 직접 실행한 핵심 행동",
            "검토 기록이나 처리 품질 변화",
        ),
        (
            "담당자 변경으로 인수인계 정보가 불완전해 사실관계를 복원해야 했던",
            "누락 정보를 확인하고 우선순위를 정해 직접 마무리한 행동",
            "인수인계 기록이나 일정 준수 결과",
        ),
    )
    situation, action, result = variants[max(0, int(variation_index)) % len(variants)]
    kind = str(focus_type or "").strip()
    if kind == "지식":
        action = f"확인한 규정·문서와 적용 범위의 판단 근거, {action}"
    elif kind == "기술":
        action = f"사용한 자료·도구와 수행 순서, {action}"
    elif kind == "태도":
        action = f"압박 속에서도 지킨 행동 기준과 선택 이유, {action}"
    return _clean_question_text(
        f"{_with_josa(surface, '이', '가')} 필요한 업무에서 {situation} 경험을 말씀해 주세요. "
        f"당시 본인의 역할과 목표, {action}, "
        f"{_with_josa(result, '으로', '로')} 성과를 확인한 방법을 구체적으로 설명해 주세요.",
        max_chars=280,
    )


def _adjust_generated_questions(
    strategy: dict[str, Any],
    question_plan: dict[str, Any],
    interview_methods: list[str],
    ncs_matches: list[dict[str, Any]] | None = None,
    ncs_ksa: list[dict[str, Any]] | None = None,
    job_context_text: str = "",
) -> dict[str, Any]:
    if not isinstance(strategy, dict):
        return strategy
    questions = strategy.get("interview_questions")
    if not isinstance(questions, list):
        questions = []
    target_total = int(question_plan.get("total_main_count", 0) or 0)
    default_follow_count = max(0, min(5, int(question_plan.get("follow_up_count", 3) or 0)))
    methods = interview_methods or list(SUPPORTED_INTERVIEW_METHODS)
    sequence = [item for item in (question_plan.get("question_sequence") or []) if isinstance(item, dict)]
    if sequence:
        target_total = len(sequence)
    elif target_total <= 0:
        target_total = len([q for q in questions if isinstance(q, dict)])

    source_questions = [dict(q) for q in questions if isinstance(q, dict)]
    while len(source_questions) < target_total:
        source_questions.append({})
    source_questions = source_questions[:target_total] if target_total > 0 else source_questions

    adjusted: list[dict[str, Any]] = []
    fallback_rows: list[dict[str, Any]] = []
    detail_offsets: dict[str, int] = {}
    seen_repeat_items_by_signature: dict[str, list[dict[str, Any]]] = {}
    used_focus_by_context: dict[tuple[str, str], set[str]] = {}
    for idx, row in enumerate(source_questions):
        item = dict(row)
        planned = sequence[idx] if idx < len(sequence) else {}
        planned_evidence_id = str(planned.get("evidence_id") or "").strip()
        planned_ncs_code = str(planned.get("ncsClCd") or "").strip()
        planned_evidence = (
            _evidence_row_for_id(ncs_ksa, planned_ncs_code, planned_evidence_id)
            if planned_evidence_id
            else {}
        )
        source_hint = str(item.get("question_source") or "").strip()
        cli_source_base = _subscription_cli_source_base(source_hint)
        is_subscription_cli_candidate = bool(cli_source_base)
        provider_evidence = (
            planned_evidence
            or _evidence_row_for_id(ncs_ksa, "", item.get("question_evidence_id"))
            if is_subscription_cli_candidate
            else {}
        )
        if planned_evidence_id:
            item["question_evidence_assignment_valid"] = bool(
                planned_evidence
                and str(row.get("question_evidence_id") or "").strip()
                == planned_evidence_id
            )
        target_detail = str(planned.get("detail", "")).strip()
        detail_key = _norm_sclass_key(target_detail)
        offset = detail_offsets.get(detail_key, 0)
        unit = _pick_unit_for_detail(target_detail, offset, ncs_matches)
        detail_offsets[detail_key] = offset + 1

        provider_code = str(provider_evidence.get("ncsClCd") or "").strip()
        provider_unit = next(
            (
                dict(candidate)
                for candidate in (ncs_matches or [])
                if isinstance(candidate, dict)
                and provider_code
                and str(candidate.get("ncsClCd") or "").strip() == provider_code
            ),
            {},
        )
        selected_unit = (
            provider_unit
            if is_subscription_cli_candidate and provider_evidence
            else unit
        )
        if selected_unit:
            item["ncsClCd"] = str(selected_unit.get("ncsClCd", "")).strip() or provider_code or str(item.get("ncsClCd", "")).strip()
            item["competency"] = (
                str(provider_evidence.get("compeUnitName", "")).strip()
                or str(selected_unit.get("compeUnitName", "")).strip()
                or str(item.get("competency", "")).strip()
            )
            item["compeUnitDef"] = str(selected_unit.get("compeUnitDef", "")).strip() or str(item.get("compeUnitDef", "")).strip()
            item["ncsSubdCdnm"] = str(selected_unit.get("ncsSubdCdnm", "")).strip() or str(item.get("ncsSubdCdnm", "")).strip()
            item["ncsSclasCdnm"] = str(selected_unit.get("ncsSclasCdnm", "")).strip() or str(item.get("ncsSclasCdnm", "")).strip()
            item["matchedDetailName"] = str(selected_unit.get("matchedDetailName", "")).strip() or target_detail
            item["ncs_detail"] = (
                target_detail
                or str(selected_unit.get("matchedDetailName", "")).strip()
                or str(selected_unit.get("ncsSubdCdnm", "")).strip()
                or str(selected_unit.get("ncsSclasCdnm", "")).strip()
                or str(item.get("ncs_detail", "")).strip()
            )
        elif target_detail:
            item["ncs_detail"] = target_detail
            if not str(item.get("competency", "")).strip():
                item["competency"] = target_detail

        row_follow_count = max(0, min(5, int(planned.get("follow_up_count", default_follow_count) or 0)))
        planned_method = str(planned.get("type") or "").strip()
        method = (
            planned_method
            if planned_method in methods and planned_method in SUPPORTED_INTERVIEW_METHODS
            else methods[idx % len(methods)]
        )
        item["method"] = method
        item["type"] = method

        ncs_code = str(item.get("ncsClCd", "")).strip()
        subject = str(item.get("competency", "")).strip() or target_detail or "해당 직무"
        existing_refs = list(item.get("ksa_refs", []) or []) if isinstance(item.get("ksa_refs"), list) else []
        ksa_terms = _ksa_terms_for_question(
            ncs_ksa=ncs_ksa,
            ncs_code=ncs_code,
            fallback_terms=existing_refs,
        )
        raw_question = _sanitize_candidate_document_leaks(
            item.get("question", ""),
            subject=subject,
        )
        item["model_question_raw"] = raw_question
        raw_followups = [
            _sanitize_candidate_document_leaks(value, subject=subject)
            for value in _clean_question_items(item.get("follow_ups"), limit=5)
        ]
        raw_followups = [value for value in raw_followups if value]
        if not raw_followups and str(item.get("follow_up", "")).strip():
            raw_followups = _clean_question_items([item.get("follow_up")], limit=1)
        original_model_followups = list(raw_followups)
        provided_evidence = _evidence_row_for_id(
            ncs_ksa,
            ncs_code,
            planned_evidence_id or item.get("question_evidence_id"),
        )
        if (
            not provided_evidence
            and is_subscription_cli_candidate
            and str(item.get("provider_evidence_alias") or "").strip() == "element_id"
        ):
            hinted_element = str(item.get("element_id") or "").strip()
            hinted_type = _ksa_key(item.get("ksa_type"))
            for candidate in ncs_ksa or []:
                if not isinstance(candidate, dict):
                    continue
                if str(candidate.get("ncsClCd") or "").strip() != ncs_code:
                    continue
                candidate_element = str(
                    candidate.get("elementId") or candidate.get("element_id") or ""
                ).strip()
                candidate_type = _ksa_key(
                    candidate.get("ksaTypeName")
                    or candidate.get("factorType")
                    or candidate.get("ksa_type")
                )
                if hinted_element and candidate_element != hinted_element:
                    continue
                if hinted_type and candidate_type and hinted_type != candidate_type:
                    continue
                provided_evidence = dict(candidate)
                item["provider_evidence_alias_resolved"] = True
                break
        requested_focus = _clean_question_text(item.get("question_focus"), max_chars=60)
        if not planned_evidence_id and not provided_evidence and requested_focus:
            requested_focus_row = _evidence_row_for_focus(ncs_ksa, ncs_code, requested_focus)
            if _ksa_key(requested_focus_row.get("factorName")) == _ksa_key(requested_focus):
                provided_evidence = requested_focus_row
        provided_focus = _clean_question_text(provided_evidence.get("factorName"), max_chars=60)
        inferred_focus = (
            ""
            if planned_evidence_id
            else _infer_model_focus_from_official_ksa(
                ncs_ksa,
                ncs_code,
                raw_question,
                raw_followups,
            )
        )
        focus = (
            provided_focus
            or inferred_focus
            or _select_ksa_focus_for_method(
                ncs_ksa=ncs_ksa,
                ncs_code=ncs_code,
                method=method,
                fallback_terms=ksa_terms or existing_refs,
            )
            or _clean_question_text(item.get("competency") or target_detail or "핵심 수행기준")
        )
        item["question_focus"] = focus
        focus_type = _ksa_type_for_focus(ncs_ksa, ncs_code, focus)
        item["question_focus_type"] = focus_type
        focus_evidence = provided_evidence or _evidence_row_for_focus(ncs_ksa, ncs_code, focus)
        task_frame = _question_task_frame(
            focus=focus,
            focus_type=focus_type,
            subject=subject,
            detail=target_detail,
            comp_def=str(item.get("compeUnitDef", "")).strip(),
            evidence_row=focus_evidence,
        )
        item["question_focus_surface"] = task_frame["task_object"]
        item["question_task_frame"] = task_frame
        item["question_evidence_id"] = task_frame.get("evidence_id", "")
        item["question_evidence_required"] = bool(focus_evidence)
        normalized_model_question = _normalize_model_task_marker(method, raw_question)
        normalized_model_question = _sanitize_candidate_document_leaks(
            normalized_model_question,
            subject=subject,
        )
        # Subscription CLI providers receive the JD, notice and exact KSA
        # evidence and are expected to
        # translate them into a real work incident.  Injecting the NCS
        # competency/detail label here turns that behavioral question back into
        # the mechanical "<module> 업무에서 ..." wording we are avoiding.
        if not is_subscription_cli_candidate:
            normalized_model_question = _normalize_model_job_context(
                method, item, normalized_model_question
            )
        raw_evaluation_points = [
            _sanitize_candidate_document_leaks(value, subject=subject)
            for value in _clean_question_items(item.get("evaluation_points"), limit=6)
        ]
        raw_evaluation_points = [value for value in raw_evaluation_points if value]
        original_model_evaluation_points = list(raw_evaluation_points)
        candidate_surface_repaired_fields: set[str] = set()
        if is_subscription_cli_candidate:
            repaired = False
        else:
            normalized_model_question, repaired = _repair_candidate_surface_text(
                normalized_model_question,
                focus,
                task_frame["task_object"],
            )
        if repaired:
            candidate_surface_repaired_fields.add("question")
        # OpenRouter already receives the KSA semantics and is instructed to
        # author a natural, single-focus question.  Do not append the legacy
        # STAR checklist here; doing so was the source of the repeated
        # "당시 역할·문서·수치·피드백" boilerplate seen by candidates.
        if is_subscription_cli_candidate and cli_source_base != "openrouter_api" and method == "경험면접":
            normalized_model_question, star_completed = (
                _complete_experience_question_star(
                    normalized_model_question,
                    focus_type=focus_type,
                    focus_surface=task_frame["task_object"],
                )
            )
            if star_completed:
                candidate_surface_repaired_fields.add("question")
        repaired_followups_for_surface: list[str] = []
        for value in raw_followups:
            if is_subscription_cli_candidate:
                repaired_value, repaired = _collapse_repeated_question_phrases(
                    value
                )
            else:
                repaired_value, repaired = _repair_candidate_surface_text(
                    value, focus, task_frame["task_object"]
                )
            repaired_followups_for_surface.append(repaired_value)
            if repaired:
                candidate_surface_repaired_fields.add("follow_ups")
        raw_followups = repaired_followups_for_surface
        if (
            is_subscription_cli_candidate
            and cli_source_base != "openrouter_api"
            and method == "경험면접"
            and row_follow_count >= 3
        ):
            star_followups = _experience_star_followups(
                focus_type=focus_type,
                surface_focus=task_frame["task_object"],
                count=row_follow_count,
            )
            if star_followups and star_followups != raw_followups[: len(star_followups)]:
                raw_followups = [
                    *star_followups,
                    *raw_followups[len(star_followups) : row_follow_count],
                ][:row_follow_count]
                candidate_surface_repaired_fields.add("follow_ups")
        repaired_evaluation_points: list[str] = []
        for value in raw_evaluation_points:
            if is_subscription_cli_candidate:
                repaired_value, repaired = _collapse_repeated_question_phrases(
                    value
                )
            else:
                repaired_value, repaired = _repair_candidate_surface_text(
                    value, focus, task_frame["task_object"]
                )
            repaired_evaluation_points.append(repaired_value)
            if repaired:
                candidate_surface_repaired_fields.add("evaluation_points")
        raw_evaluation_points = repaired_evaluation_points
        if candidate_surface_repaired_fields:
            item["candidate_surface_repairs"] = sorted(candidate_surface_repaired_fields)
        item["model_followups_raw"] = original_model_followups
        item["model_evaluation_points_raw"] = original_model_evaluation_points
        raw_merged = "\n".join(
            [raw_question, *original_model_followups, *original_model_evaluation_points]
        )
        main_replacement_reasons: list[str] = []
        followup_replacement_reasons: list[str] = []
        raw_followups_final = _merge_question_items(raw_followups, [], row_follow_count)
        repaired_followups: list[str] = []
        if not raw_question:
            main_replacement_reasons.append("no_model_question")
        else:
            if _contains_blind_hiring_cue(raw_merged):
                main_replacement_reasons.append("blind_hiring_cue")
            if not _operating_conditions_separated(method, normalized_model_question):
                main_replacement_reasons.append("operating_conditions_in_main")
            if not _method_shape_ok(method, normalized_model_question):
                main_replacement_reasons.append("main_question_method_shape")
            if not _main_question_task_marker_ok(method, normalized_model_question):
                main_replacement_reasons.append("main_question_official_sample_shape")
            if not _decision_dilemma_quality_ok(method, normalized_model_question):
                main_replacement_reasons.append("decision_dilemma_quality")
            if not _debate_option_defensibility_ok(method, normalized_model_question):
                main_replacement_reasons.append("debate_option_defensibility")
            if not _debate_outcome_flexibility_ok(method, normalized_model_question):
                main_replacement_reasons.append("debate_outcome_flexibility")
            context_item = dict(item)
            context_item["type"] = method
            context_item["method"] = method
            context_item["question"] = normalized_model_question
            context_item["ksa_refs"] = ksa_terms or existing_refs
            if not evaluate_ksa_measurement(context_item).get("passed"):
                main_replacement_reasons.append("ksa_measurement_task")
            if not _follow_ups_quality_ok(method, context_item, raw_followups_final):
                if not main_replacement_reasons:
                    repaired_followups = _repair_model_followups_with_focus(
                        method=method,
                        q=context_item,
                        follow_ups=raw_followups,
                        limit=row_follow_count,
                    )
                if repaired_followups:
                    followup_replacement_reasons.append("follow_up_focus_injected")
                else:
                    followup_replacement_reasons.append("follow_up_quality")
        cli_hard_replacement_reasons = {
            "no_model_question",
            "blind_hiring_cue",
        }
        # A non-empty provider draft remains model-authored.  Editorial or KSA
        # quality failures must stay visible to the quality gate so the outer
        # provider pass can regenerate them; replacing them here with a server
        # template makes every KSA sound like the same canned STAR question.
        cli_requires_replacement = bool(
            is_subscription_cli_candidate
            and any(
                reason in cli_hard_replacement_reasons
                for reason in main_replacement_reasons
            )
        )
        openai_blocking_replacement_reasons = set(main_replacement_reasons) & {
            "no_model_question",
            "blind_hiring_cue",
        }
        use_model_question = bool(
            raw_question
            and (
                (not main_replacement_reasons)
                or (
                    is_subscription_cli_candidate
                    and not cli_requires_replacement
                )
                or (
                    not is_subscription_cli_candidate
                    and not openai_blocking_replacement_reasons
                )
            )
        )
        use_raw_model_followups = bool(
            use_model_question
            and (not followup_replacement_reasons or is_subscription_cli_candidate)
        )
        use_repaired_model_followups = bool(
            use_model_question
            and repaired_followups
            and followup_replacement_reasons == ["follow_up_focus_injected"]
        )
        model_replacement_reasons = [*main_replacement_reasons, *followup_replacement_reasons]

        template_question = _question_for_method(
            method=method,
            subject=subject,
            focus=focus,
            detail=target_detail,
            comp_def=str(item.get("compeUnitDef", "")).strip(),
            focus_type=focus_type,
            variation_index=idx,
            task_frame=task_frame,
            job_context_text=job_context_text,
            freeform_ksa=(cli_source_base == "openrouter_api"),
        )
        template_followups = _followups_for_method(
            method=method,
            subject=subject,
            focus=focus,
            count=row_follow_count,
            variant_index=idx,
            focus_type=focus_type,
            task_frame=task_frame,
        )
        focus_first_ksa_terms = [focus, *[term for term in ksa_terms if _ksa_key(term) != _ksa_key(focus)]]
        method_eval_points = _method_evaluation_points(
            method,
            focus_first_ksa_terms,
            focus_type,
            surface_focus=task_frame["task_object"],
        )

        item["question"] = normalized_model_question if use_model_question else template_question
        if use_model_question and is_subscription_cli_candidate:
            item["question_source"] = (
                f"{cli_source_base}_quality_repaired_fields"
                if candidate_surface_repaired_fields
                else cli_source_base
            )
        elif use_model_question and use_raw_model_followups:
            item["question_source"] = (
                "model_main_quality_repaired_fields"
                if candidate_surface_repaired_fields
                else "model"
            )
        elif use_model_question and use_repaired_model_followups:
            item["question_source"] = "model_main_repaired_followups"
        elif use_model_question:
            item["question_source"] = "model_main_template_followups"
        else:
            item["question_source"] = "template_fallback"
        item["model_question_preserved"] = bool(use_model_question)
        if is_subscription_cli_candidate and use_model_question:
            item["model_quality_warnings"] = model_replacement_reasons
            item["model_replacement_reasons"] = []
        else:
            item["model_replacement_reasons"] = (
                [] if use_model_question and use_raw_model_followups else model_replacement_reasons
            )
        if use_raw_model_followups:
            item["follow_ups"] = raw_followups_final
        elif use_repaired_model_followups:
            item["follow_ups"] = _merge_question_items(repaired_followups, template_followups, row_follow_count)
        else:
            item["follow_ups"] = template_followups
        item["follow_up"] = item["follow_ups"][0] if item["follow_ups"] else ""
        if use_model_question and is_subscription_cli_candidate:
            # The provider authors the candidate question, follow-ups and
            # observable evaluation points from the assigned KSA.  Preserve an
            # invalid count/content so the final gate can request a model retry
            # instead of laundering it through deterministic server wording.
            item["evaluation_points"] = list(raw_evaluation_points)
        elif use_model_question:
            item["evaluation_points"] = _merge_question_items(
                raw_evaluation_points, method_eval_points, 4
            )
        else:
            item["evaluation_points"] = method_eval_points
        item["task_conditions"] = _task_conditions_for_method(
            method=method,
            subject=subject,
            focus=focus,
            detail=target_detail,
            comp_def=str(item.get("compeUnitDef", "")).strip(),
            focus_type=focus_type,
            variation_index=idx,
            job_context_text=job_context_text,
        )
        item["question_topic_axis"] = str(
            (item.get("task_conditions") or {}).get("topic_axis")
            or _question_topic_axis(idx)
        ).strip()
        item["question_variation_index"] = idx
        item["assessment_guide"] = _behavior_anchored_evaluation(
            method,
            focus,
            item["evaluation_points"],
            focus_type,
            surface_focus=task_frame["task_object"],
        )
        if not str(item.get("question") or "").strip():
            item["question"] = (
                str(template_question or "").strip()
                or "주어진 근거를 바탕으로 면접 질문을 제시해 주세요."
            )
            if not item.get("follow_ups"):
                item["follow_ups"] = list(template_followups or [])
            if not item.get("evaluation_points"):
                item["evaluation_points"] = list(method_eval_points or [])
            if not item.get("task_conditions"):
                item["task_conditions"] = _task_conditions_for_method(
                    method=method,
                    subject=subject,
                    focus=focus,
                    detail=target_detail,
                    comp_def=str(item.get("compeUnitDef", "")).strip(),
                    focus_type=focus_type,
                    variation_index=idx,
                    job_context_text=job_context_text,
                )
            if not item.get("assessment_guide"):
                item["assessment_guide"] = _behavior_anchored_evaluation(
                    method,
                    focus,
                    item.get("evaluation_points") or method_eval_points,
                    focus_type,
                    surface_focus=task_frame["task_object"],
                )
            item["follow_up"] = item["follow_ups"][0] if item["follow_ups"] else ""
            item["question_source"] = "template_fallback"
            item["model_question_preserved"] = False
            item["model_replacement_reasons"] = list(
                dict.fromkeys(
                    [
                        *(
                            str(reason).strip()
                            for reason in (item.get("model_replacement_reasons") or [])
                            if str(reason).strip()
                        ),
                        "question_content_missing",
                    ]
                )
            )
        focus_context_key = (ncs_code or subject, method)
        used_focuses = used_focus_by_context.setdefault(focus_context_key, set())
        repeat_signature = _question_repeat_signature(item)
        previous_repeat_items = seen_repeat_items_by_signature.get(repeat_signature, []) if repeat_signature else []
        # Within one generated set, a repeated method+KSA intent should consume a
        # different KSA (or a different observable angle of the same KSA). A
        # cosmetic operating-condition variation is not enough diversity.
        repeat_near_duplicate = bool(previous_repeat_items)
        if repeat_near_duplicate and raw_question:
            previous_model_questions = [
                str(previous.get("model_question_raw") or "").strip()
                for previous in previous_repeat_items
                if str(previous.get("model_question_raw") or "").strip()
            ]
            # A valid model question that explicitly selected this official KSA
            # outranks an earlier empty/template slot. Distinct model scenarios
            # may also assess the same KSA; repeated model scenarios may not.
            if (
                (use_model_question and not previous_model_questions)
                or _raw_model_scenarios_are_distinct(item, previous_repeat_items)
            ):
                repeat_near_duplicate = False
        duplicate_replaced = False
        if (
            repeat_signature
            and repeat_near_duplicate
            and not is_subscription_cli_candidate
        ):
            alternate_focus = _pick_alternate_question_focus(
                current_focus=focus,
                ksa_terms=ksa_terms,
                used_focuses=used_focuses,
            )
            if alternate_focus:
                focus = alternate_focus
                item["question_focus"] = focus
                focus_type = _ksa_type_for_focus(ncs_ksa, ncs_code, focus)
                item["question_focus_type"] = focus_type
                focus_evidence = _evidence_row_for_focus(ncs_ksa, ncs_code, focus)
                task_frame = _question_task_frame(
                    focus=focus,
                    focus_type=focus_type,
                    subject=subject,
                    detail=target_detail,
                    comp_def=str(item.get("compeUnitDef", "")).strip(),
                    evidence_row=focus_evidence,
                )
                item["question_focus_surface"] = task_frame["task_object"]
                item["question_task_frame"] = task_frame
                item["question_evidence_id"] = task_frame.get("evidence_id", "")
                item["question_evidence_required"] = bool(focus_evidence)
                template_question = _question_for_method(
                    method=method,
                    subject=subject,
                    focus=focus,
                    detail=target_detail,
                    comp_def=str(item.get("compeUnitDef", "")).strip(),
                    focus_type=focus_type,
                    variation_index=idx,
                    task_frame=task_frame,
                    job_context_text=job_context_text,
                )
                template_followups = _followups_for_method(
                    method=method,
                    subject=subject,
                    focus=focus,
                    count=row_follow_count,
                    variant_index=idx,
                    focus_type=focus_type,
                    task_frame=task_frame,
                )
                method_eval_points = _method_evaluation_points(
                    method,
                    [focus, *[x for x in ksa_terms if x != focus]],
                    focus_type,
                    surface_focus=task_frame["task_object"],
                )
                item["question"] = template_question
                item["follow_ups"] = template_followups
                item["follow_up"] = item["follow_ups"][0] if item["follow_ups"] else ""
                item["evaluation_points"] = method_eval_points
                item["task_conditions"] = _task_conditions_for_method(
                    method=method,
                    subject=subject,
                    focus=focus,
                    detail=target_detail,
                    comp_def=str(item.get("compeUnitDef", "")).strip(),
                    focus_type=focus_type,
                    variation_index=idx,
                    job_context_text=job_context_text,
                )
                item["assessment_guide"] = _behavior_anchored_evaluation(
                    method,
                    focus,
                    method_eval_points,
                    focus_type,
                    surface_focus=task_frame["task_object"],
                )
                item["question_source"] = "template_fallback"
                item["model_question_preserved"] = False
                reasons = [
                    str(reason).strip()
                    for reason in (item.get("model_replacement_reasons") or [])
                    if str(reason).strip()
                ] if isinstance(item.get("model_replacement_reasons"), list) else []
                item["model_replacement_reasons"] = list(dict.fromkeys([*reasons, "duplicate_question_intent"]))
                repeat_signature = _question_repeat_signature(item)
                duplicate_replaced = True
                previous_repeat_items = seen_repeat_items_by_signature.get(repeat_signature, []) if repeat_signature else []
                repeat_near_duplicate = any(_question_near_repeat(item, previous) for previous in previous_repeat_items)
        item["question_intent"] = _question_intent_key(str(item.get("question") or ""))
        item["question_repeat_signature"] = repeat_signature
        item["question_repeat_duplicate"] = bool(repeat_near_duplicate and not duplicate_replaced)
        if repeat_signature:
            seen_repeat_items_by_signature.setdefault(repeat_signature, []).append(dict(item))
        focus_key = _question_focus_signature(focus)
        if focus_key:
            used_focuses.add(focus_key)
        adjusted.append(item)
        fallback_rows.append(
            {
                "question": template_question,
                "follow_ups": template_followups,
                "evaluation_points": method_eval_points,
                "task_conditions": _task_conditions_for_method(
                    method=method,
                    subject=subject,
                    focus=focus,
                    detail=target_detail,
                    comp_def=str(item.get("compeUnitDef", "")).strip(),
                    focus_type=focus_type,
                    variation_index=idx,
                    job_context_text=job_context_text,
                ),
                "assessment_guide": _behavior_anchored_evaluation(
                    method,
                    focus,
                    method_eval_points,
                    focus_type,
                    surface_focus=task_frame["task_object"],
                ),
            }
        )

    probe_strategy = dict(strategy)
    probe_strategy["interview_questions"] = [dict(q) for q in adjusted]
    probe_strategy["question_plan_used"] = question_plan
    probe_strategy = _attach_ksa_evidence_to_strategy(probe_strategy, ncs_ksa)
    probe_items = {
        int(item.get("index") or 0): item
        for item in (probe_strategy.get("question_quality_report") or {}).get("items", [])
        if isinstance(item, dict)
    }
    field_repair_issue_fields: dict[str, set[str]] = {
        "follow_up_depth": {"follow_ups"},
        "follow_up_quality": {"follow_ups"},
        "evaluation_points": {"evaluation_points"},
        "evaluation_points_quality": {"evaluation_points"},
        "blind_hiring_safe": {"question"},
        "candidate_surface_safe": {"question", "follow_ups"},
        "ncs_grounded": {"question", "task_conditions"},
        "detail_grounded": {"question", "task_conditions"},
        "ksa_grounded": {"question", "task_conditions"},
        "main_question_method_shape": {"question"},
        "main_question_job_context": {"question"},
        "method_shape": {"question", "follow_ups", "evaluation_points"},
        "specific_context": {"question"},
        "job_specific_context": {"question", "task_conditions"},
        "natural_wording": {"question"},
        "focus_scenario_coherence": {"question", "task_conditions"},
        "decision_dilemma_quality": {"question", "follow_ups", "evaluation_points"},
        "debate_option_defensibility": {"question", "follow_ups", "evaluation_points"},
        "debate_outcome_flexibility": {"question", "follow_ups", "evaluation_points"},
        "debate_case_neutrality": {"question", "follow_ups", "evaluation_points", "task_conditions"},
        "operating_conditions_separated": {"question", "task_conditions"},
        "field_realism": {"question", "follow_ups", "evaluation_points", "task_conditions", "assessment_guide"},
        "precision_grounding": {"question", "follow_ups", "evaluation_points", "task_conditions", "assessment_guide"},
        "standardized_task_conditions": {"task_conditions"},
        "case_materials_sufficient": {"task_conditions"},
        "decision_authority_context": {"task_conditions"},
        "inbasket_authority_context": {"task_conditions"},
        "behavior_anchored_evaluation": {"assessment_guide"},
        "official_sample_format": {"follow_ups", "evaluation_points"},
    }
    for pos, item in enumerate(adjusted):
        source = str(item.get("question_source") or "").strip()
        if source not in MODEL_PRESERVED_QUESTION_SOURCES:
            continue
        probe_item = probe_items.get(pos + 1) or {}
        if probe_item.get("ready") is True:
            continue
        raw_issues = {
            str(issue).strip()
            for issue in (probe_item.get("issues") or [])
            if str(issue).strip()
        } if isinstance(probe_item.get("issues"), list) else set()
        if _is_subscription_cli_source(source):
            existing_warnings = [
                str(reason).strip()
                for reason in (item.get("model_quality_warnings") or [])
                if str(reason).strip()
            ] if isinstance(item.get("model_quality_warnings"), list) else []
            quality_warnings = [
                f"quality_gate_{issue}"
                for issue in (probe_item.get("issues") or [])
                if str(issue).strip()
            ] if isinstance(probe_item.get("issues"), list) else []
            item["model_quality_warnings"] = list(
                dict.fromkeys([*existing_warnings, *quality_warnings])
            )
            item["model_question_preserved"] = True
            continue
        fallback = fallback_rows[pos] if pos < len(fallback_rows) else {}
        repairable_fields: set[str] = set()
        if raw_issues and raw_issues.issubset(field_repair_issue_fields):
            for issue in raw_issues:
                repairable_fields.update(field_repair_issue_fields[issue])
        if repairable_fields and "question" not in repairable_fields:
            candidate = dict(item)
            if "follow_ups" in repairable_fields:
                candidate["follow_ups"] = list(fallback.get("follow_ups") or [])
                candidate["follow_up"] = candidate["follow_ups"][0] if candidate["follow_ups"] else ""
            if "evaluation_points" in repairable_fields:
                candidate["evaluation_points"] = list(fallback.get("evaluation_points") or [])
            if "question" in repairable_fields:
                candidate["question"] = str(
                    fallback.get("question") or candidate.get("question") or ""
                ).strip()
            if "task_conditions" in repairable_fields:
                candidate["task_conditions"] = dict(fallback.get("task_conditions") or {})
            if "assessment_guide" in repairable_fields or "evaluation_points" in repairable_fields:
                candidate["assessment_guide"] = dict(fallback.get("assessment_guide") or {})
                repairable_fields.add("assessment_guide")

            candidate_probe = {
                "interview_questions": [candidate],
                "question_plan_used": {"total_main_count": 1},
            }
            candidate_probe = _attach_ksa_evidence_to_strategy(candidate_probe, ncs_ksa)
            candidate_report = (
                candidate_probe.get("question_quality_report")
                if isinstance(candidate_probe.get("question_quality_report"), dict)
                else {}
            )
            if candidate_report.get("passed") is True:
                if source == "model":
                    candidate["question_source"] = "model_main_quality_repaired_fields"
                candidate["model_question_preserved"] = True
                candidate["quality_repaired_fields"] = sorted(repairable_fields)
                candidate["quality_repair_reasons"] = list(
                    dict.fromkeys(
                        [
                            *[f"quality_field_repair_{issue}" for issue in sorted(raw_issues)],
                        ]
                    )
                )
                adjusted[pos] = candidate
                continue
        existing_reasons = [
            str(reason).strip()
            for reason in (item.get("model_replacement_reasons") or [])
            if str(reason).strip()
        ] if isinstance(item.get("model_replacement_reasons"), list) else []
        quality_reasons = [
            str(issue).strip()
            for issue in (probe_item.get("issues") or [])
            if str(issue).strip()
        ] if isinstance(probe_item.get("issues"), list) else []
        item["question"] = str(fallback.get("question") or item.get("question") or "").strip()
        item["question_source"] = "template_fallback"
        item["model_question_preserved"] = False
        item["model_replacement_reasons"] = list(dict.fromkeys([*existing_reasons, *quality_reasons]))
        item["follow_ups"] = list(fallback.get("follow_ups") or [])
        item["follow_up"] = item["follow_ups"][0] if item["follow_ups"] else ""
        item["evaluation_points"] = list(fallback.get("evaluation_points") or [])
        item["task_conditions"] = dict(fallback.get("task_conditions") or {})
        item["assessment_guide"] = dict(fallback.get("assessment_guide") or {})
    _refresh_question_repeat_metadata(adjusted)
    strategy["interview_questions"] = adjusted
    strategy["interview_by_competency"] = _group_interview_questions_for_response(adjusted)
    strategy["question_plan_used"] = question_plan
    strategy["interview_methods_used"] = methods
    strategy["question_customization_policy"] = "model_preserve_with_guidebook_template_fallback_followup_gate"
    return strategy


def _repeat_count_from_weight(weight: float, default: int = 1, max_repeat: int = 6) -> int:
    try:
        v = int(round(float(weight)))
    except Exception:
        v = int(default)
    return max(1, min(int(max_repeat), v))


def _build_priority_notice_text(
    notice_text: str,
    duty_text: str = "",
    qualification_text: str = "",
    preference_text: str = "",
    evaluation_text: str = "",
) -> str:
    notice = str(notice_text or "").strip()
    duty = str(duty_text or "").strip()
    qualification = str(qualification_text or "").strip()
    preference = str(preference_text or "").strip()
    evaluation = str(evaluation_text or "").strip()

    parts: list[str] = []
    if duty:
        parts.append(f"[담당업무-우선]\n{duty[:2500]}")
    if qualification:
        parts.append(f"[지원자격-우선]\n{qualification[:1800]}")
    if preference:
        parts.append(f"[우대사항-우선]\n{preference[:1800]}")
    if evaluation:
        parts.append(f"[면접평가항목-우선]\n{evaluation[:1800]}")
    if notice:
        parts.append(f"[공고문-보조]\n{notice[:2500]}")
    return "\n\n".join(parts).strip()


def _build_priority_query_text(
    base_text: str,
    duty_text: str = "",
    qualification_text: str = "",
    preference_text: str = "",
    evaluation_text: str = "",
) -> str:
    base = str(base_text or "").strip()[:5000]
    duty = str(duty_text or "").strip()[:2500]
    qualification = str(qualification_text or "").strip()[:1600]
    preference = str(preference_text or "").strip()[:1600]
    evaluation = str(evaluation_text or "").strip()[:1500]

    base_w = _to_float_or(os.getenv("JD_BASE_TEXT_WEIGHT", "1.0"), 1.0)
    duty_w = _to_float_or(os.getenv("DUTY_TEXT_WEIGHT", "3.0"), 3.0)
    qualification_w = _to_float_or(os.getenv("QUALIFICATION_TEXT_WEIGHT", "1.4"), 1.4)
    preference_w = _to_float_or(os.getenv("PREFERENCE_TEXT_WEIGHT", "1.2"), 1.2)
    eval_w = _to_float_or(os.getenv("EVALUATION_TEXT_WEIGHT", "2.5"), 2.5)

    base_rep = _repeat_count_from_weight(base_w, default=1, max_repeat=4)
    duty_rep = _repeat_count_from_weight(duty_w, default=3, max_repeat=6)
    qualification_rep = _repeat_count_from_weight(qualification_w, default=1, max_repeat=3)
    preference_rep = _repeat_count_from_weight(preference_w, default=1, max_repeat=3)
    eval_rep = _repeat_count_from_weight(eval_w, default=2, max_repeat=6)

    chunks: list[str] = []
    if duty:
        chunks.extend([f"[담당업무]{duty}"] * duty_rep)
    if qualification:
        chunks.extend([f"[지원자격]{qualification}"] * qualification_rep)
    if preference:
        chunks.extend([f"[우대사항]{preference}"] * preference_rep)
    if evaluation:
        chunks.extend([f"[면접평가항목]{evaluation}"] * eval_rep)
    if base:
        chunks.extend([f"[기본텍스트]{base}"] * base_rep)
    return "\n".join(chunks).strip()


def _collect_ksa_candidate_units(
    primary_units: list[dict[str, Any]] | None,
    secondary_units: list[dict[str, Any]] | None = None,
    max_units: int = 12,
) -> list[dict[str, Any]]:
    out: list[dict[str, Any]] = []
    seen_codes: set[str] = set()

    for bucket in ((primary_units or []), (secondary_units or [])):
        for row in bucket:
            if not isinstance(row, dict):
                continue
            code = str(row.get("ncsClCd", "")).strip()
            if not code or code in seen_codes:
                continue
            seen_codes.add(code)
            try:
                score = float(row.get("score", 1.0) or 1.0)
            except Exception:
                score = 1.0
            out.append(
                {
                    "ncsClCd": code,
                    "compeUnitName": str(row.get("compeUnitName", "")).strip(),
                    "compeUnitLevel": str(row.get("compeUnitLevel", "")).strip(),
                    "ncsSubdCdnm": str(row.get("ncsSubdCdnm", "")).strip(),
                    "compeUnitDef": str(row.get("compeUnitDef", "")).strip(),
                    "score": score,
                    "matched_keywords": list(row.get("matched_keywords", []) or []),
                }
            )
            if len(out) >= max(1, int(max_units or 12)):
                return out
    return out


def _fetch_ncs_ksa_or_502(
    ncs_matches: list[dict[str, Any]],
    max_units: int,
    max_factors_per_unit: int,
) -> list[dict[str, Any]]:
    try:
        return fetch_ncs_ksa_by_units(
            ncs_matches=ncs_matches,
            max_units=max_units,
            max_factors_per_unit=max_factors_per_unit,
        )
    except NcsMcpError as exc:
        # An otherwise healthy MCP can legitimately return no KSA rows when a
        # manually submitted/expired NCS unit is not present in the official
        # database.  Treat that as an actionable input selection error rather
        # than an upstream outage so the request stops before provider work.
        if "no official ksa rows" in str(exc).casefold():
            logger.warning("ncs_mcp_ksa_unavailable selected_units=%s", len(ncs_matches or []))
            raise HTTPException(
                status_code=422,
                detail={
                    "code": "ncs_ksa_unavailable",
                    "message": "선택한 NCS 세분류에서 공식 KSA 근거를 찾지 못했습니다. 목록에서 다른 세분류를 선택해 주세요.",
                    "retryable": False,
                },
            ) from exc
        raise _internal_http_error(
            "ncs_mcp_ksa_failed",
            exc,
            status_code=502,
        ) from exc


def _dedupe_units_by_code(units: list[dict[str, Any]] | None) -> list[dict[str, Any]]:
    out: list[dict[str, Any]] = []
    seen: set[str] = set()
    for row in units or []:
        if not isinstance(row, dict):
            continue
        code = str(row.get("ncsClCd", "")).strip()
        if not code or code in seen:
            continue
        seen.add(code)
        out.append(dict(row))
    return out


def _select_units_for_question_plan(
    question_plan: dict[str, Any],
    ncs_matches: list[dict[str, Any]] | None,
) -> list[dict[str, Any]]:
    if not isinstance(question_plan, dict):
        return []
    sequence = [item for item in (question_plan.get("question_sequence") or []) if isinstance(item, dict)]
    if not sequence:
        return _dedupe_units_by_code(ncs_matches)

    selected: list[dict[str, Any]] = []
    detail_offsets: dict[str, int] = {}
    for planned in sequence:
        target_detail = str(planned.get("detail", "")).strip()
        detail_key = _norm_sclass_key(target_detail)
        offset = detail_offsets.get(detail_key, 0)
        unit = _pick_unit_for_detail(target_detail, offset, ncs_matches)
        detail_offsets[detail_key] = offset + 1
        if unit:
            selected.append(unit)
    return _dedupe_units_by_code(selected)


def _merge_ksa_rows(existing: list[dict[str, Any]] | None, added: list[dict[str, Any]] | None) -> list[dict[str, Any]]:
    out: list[dict[str, Any]] = []
    seen: set[tuple[str, str]] = set()
    for row in (existing or []) + (added or []):
        if not isinstance(row, dict):
            continue
        code = str(row.get("ncsClCd", "")).strip()
        factor = str(row.get("factorName", "")).strip()
        key = (code, _ksa_key(factor))
        if not code or not factor or key in seen:
            continue
        seen.add(key)
        out.append(dict(row))
    return out


def _supplement_ksa_for_question_plan(
    question_plan: dict[str, Any],
    ncs_matches: list[dict[str, Any]] | None,
    ncs_ksa: list[dict[str, Any]] | None,
    max_factors_per_unit: int,
) -> list[dict[str, Any]]:
    selected_units = _select_units_for_question_plan(question_plan, ncs_matches)
    if not selected_units:
        return list(ncs_ksa or [])
    covered_codes = {
        str(row.get("ncsClCd", "")).strip()
        for row in (ncs_ksa or [])
        if isinstance(row, dict) and str(row.get("ncsClCd", "")).strip() and str(row.get("factorName", "")).strip()
    }
    missing_units = [
        unit
        for unit in selected_units
        if str(unit.get("ncsClCd", "")).strip() and str(unit.get("ncsClCd", "")).strip() not in covered_codes
    ]
    if not missing_units:
        return list(ncs_ksa or [])
    fetched = _fetch_ncs_ksa_or_502(
        ncs_matches=missing_units,
        max_units=len(missing_units),
        max_factors_per_unit=max_factors_per_unit,
    )
    return _merge_ksa_rows(ncs_ksa, fetched)


def _require_ncs_mcp_url() -> str:
    endpoint = settings.ncs_mcp_endpoint()
    if endpoint:
        return endpoint
    raise HTTPException(
        status_code=503,
        detail=(
            "NCS_MCP_URL is required for NCScope. Start NCS_MCP, the read-only "
            "local NCS DB search server, with the compact serving DB and set NCS_MCP_URL."
        ),
    )


def _public_question_rows(result: Any) -> list[dict[str, Any]]:
    if not isinstance(result, dict):
        return []
    questions: list[dict[str, Any]] = []
    for key in ("main_questions", "questions", "interview_questions"):
        rows = result.get(key)
        if isinstance(rows, list):
            questions.extend(row for row in rows if isinstance(row, dict))
    return questions


def _public_questions_precision_grounded(result: Any) -> bool:
    """Recheck every public question without trusting provider-owned evidence.

    A future server-issued material registry can be passed explicitly at this
    boundary.  Until that registry exists, fields embedded in model output are
    intentionally ignored so a provider cannot self-attest that a hidden
    amount, formula, or clause excerpt was supplied to the candidate.
    """

    questions = _public_question_rows(result)
    return bool(questions) and all(
        evaluate_question_precision_grounding(question).get("passed") is True
        for question in questions
    )


def _require_official_ksa_result(result: dict[str, Any]) -> None:
    """Fail closed unless every public question is traceable and structurally safe.

    Official KSA labels are internal evidence, not candidate-facing copy.  A
    stable evidence identifier proves only which official row was used.  It
    never proves that the wording measures the KSA; the independent AI review
    owns that semantic decision.
    """

    if result.get("ncs_ksa_available") is not True:
        raise HTTPException(
            status_code=502,
            detail="Official NCS KSA is unavailable from NCS_MCP; question generation was stopped.",
        )

    questions = _public_question_rows(result)
    def _supporting_evidence_ids(question: dict[str, Any]) -> set[str]:
        """Return IDs attested by the service-owned NCS evidence registry.

        Do not inspect ``question_task_frame`` or question-local evidence rows
        here.  Those fields may originate in model output and therefore cannot
        prove their own authenticity.
        """

        evidence_ids: set[str] = set()
        rows = result.get("official_ksa_evidence")
        if not isinstance(rows, list):
            return evidence_ids
        question_code = str(
            question.get("ncsClCd") or question.get("ncs_code") or ""
        ).strip()
        refs = {
            _ksa_key(value)
            for value in (question.get("ksa_refs") or [])
            if str(value or "").strip()
        }
        for row in rows:
            if not isinstance(row, dict):
                continue
            row_code = str(row.get("ncsClCd") or row.get("unit_code") or "").strip()
            row_factor = str(row.get("factorName") or row.get("factor_name") or "").strip()
            if not row_code or not row_factor:
                continue
            if not question_code or row_code != question_code:
                continue
            if not refs or _ksa_key(row_factor) not in refs:
                continue
            row_id = str(row.get("evidence_id") or "").strip()
            computed_id = stable_ksa_evidence_id(row)
            if row_id == computed_id:
                evidence_ids.add(computed_id)
        return evidence_ids

    invalid_indices: list[int] = []
    for index, question in enumerate(questions, start=1):
        refs = [
            str(value).strip()
            for value in (question.get("ksa_refs") or [])
            if str(value).strip()
        ] if isinstance(question.get("ksa_refs"), list) else []
        follow_ups = [
            value
            for value in (question.get("follow_ups") or [])
            if str(value).strip() or isinstance(value, dict)
        ] if isinstance(question.get("follow_ups"), list) else []
        evaluation_points = question.get("evaluation_points")
        if not isinstance(evaluation_points, list):
            evaluation_points = question.get("eval_points")
        evaluation_points_ok = bool(
            isinstance(evaluation_points, list)
            and 1 <= len(evaluation_points) <= 5
            and all(str(value).strip() for value in evaluation_points)
        )
        evidence_id = str(question.get("question_evidence_id") or "").strip()
        supporting_ids = _supporting_evidence_ids(question)
        stable_evidence_id = bool(re.fullmatch(r"ksa_[0-9a-f]{24}", evidence_id))
        evidence_consistent = bool(evidence_id and evidence_id in supporting_ids)
        task_frame = question.get("question_task_frame")
        task_frame_evidence_id = (
            str(task_frame.get("evidence_id") or "").strip()
            if isinstance(task_frame, dict)
            else ""
        )
        task_frame_consistent = bool(
            task_frame_evidence_id and task_frame_evidence_id == evidence_id
        )
        core_grounding_ok = (
            str(question.get("question_focus_source") or "").strip() == "official_ksa"
            and refs
            and str(question.get("ncsClCd") or question.get("ncs_code") or "").strip()
            and _subscription_cli_source_base(question.get("question_source"))
            == "openai_api"
            and stable_evidence_id
            and evidence_consistent
            and task_frame_consistent
            and isinstance(evaluation_points, list)
            and evaluation_points_ok
            and 1 <= len(follow_ups) <= 5
        )
        if not core_grounding_ok:
            invalid_indices.append(index)

    if not questions or invalid_indices:
        raise HTTPException(
            status_code=502,
            detail="Question generation produced unverified or invalid NCS KSA grounding.",
        )


def _require_legacy_ncs_api_enabled() -> None:
    if not settings.enable_legacy_ncs_api():
        raise HTTPException(
            status_code=410,
            detail="legacy NCS API endpoints are disabled; use NCS_MCP_URL-backed local NCS DB endpoints",
        )


def _check_upload_size(data: bytes, label: str) -> None:
    max_bytes = settings.max_upload_bytes()
    if len(data or b"") > max_bytes:
        raise HTTPException(
            status_code=413,
            detail=f"{label} exceeds MAX_UPLOAD_MB ({max_bytes // (1024 * 1024)} MB)",
        )


async def _read_upload_limited(upload: UploadFile, label: str) -> bytes:
    max_bytes = settings.max_upload_bytes()
    data = await upload.read(max_bytes + 1)
    _check_upload_size(data, label)
    return data


_ARCHIVE_MEMBER_LIMIT = 12
_SUPPORTED_ARCHIVE_DOC_SUFFIXES = {
    ".pdf",
    ".hwp",
    ".hwpx",
    ".docx",
    ".txt",
    ".png",
    ".jpg",
    ".jpeg",
    ".webp",
}
_BLOCKED_DOC_SUFFIXES = {".hwp", ".hwpx"}
_REVIEW_SESSION_TTL_SEC = 30 * 60
_REVIEW_SESSION_MAX = 100
_REVIEW_SESSION_SIGNATURE_VERSION = "v1"
_REVIEW_SESSION_MAX_CLOCK_SKEW_SEC = 60
_REVIEW_SESSION_SHA256_RE = re.compile(r"^[0-9a-f]{64}$")
_REVIEW_SESSION_ID_RE = re.compile(r"^[A-Za-z0-9_-]{16,128}$")
_REVIEW_SESSION_TOKEN_RE = re.compile(r"^v1\.[A-Za-z0-9_-]{43}$")
_REVIEW_SESSION_LOCK = threading.Lock()
_REVIEW_SESSION_BY_ID: dict[str, dict[str, Any]] = {}


def _suffix_of(name: str) -> str:
    return Path(str(name or "").replace("\\", "/")).suffix.lower()


def _safe_member_label(name: str) -> str:
    value = str(name or "").replace("\\", "/").split("/")[-1].strip()
    value = re.sub(r"[\r\n\t]+", " ", value)
    if len(value) > 160:
        suffix = Path(value).suffix[:16]
        value = f"{value[: max(1, 160 - len(suffix))]}{suffix}"
    return value or "archive_member"


def _reject_hwp_upload(label: str, filename: str | None = None) -> None:
    if _suffix_of(filename or "") in _BLOCKED_DOC_SUFFIXES:
        raise HTTPException(
            status_code=400,
            detail=f"{label}은(는) HWP/HWPX 형식입니다. PDF로 변환 후 탑재해 주세요.",
        )


def _parse_single_document_upload(data: bytes, filename: str, label: str) -> dict[str, Any]:
    name = str(filename or "")
    suffix = _suffix_of(name)
    if suffix == ".txt":
        return {
            "markdown": data.decode("utf-8", errors="ignore"),
            "metadata": {"filename": name},
            "parser": "plain_text",
        }
    try:
        parsed = parse_with_kordoc(
            data,
            filename=name,
            ocr=os.getenv("KORDOC_OCR", "true").strip().lower() in {"1", "true", "yes", "y"},
        )
        if suffix in {".hwp", ".hwpx"}:
            try:
                local_text = extract_hwpx_text(data) if suffix == ".hwpx" else extract_hwp_text(data)
            except Exception:
                # This is a supplemental classifier pass. A working Kordoc
                # result must remain usable if the local Hangul reader cannot
                # handle a particular binary/ZIP variant.
                logger.warning("hangul_supplemental_classification_unavailable", exc_info=True)
                local_text = ""
            local_terms = extract_linear_ncs_classification_terms(
                local_text,
                excluded_hierarchy_names=_non_detail_hierarchy_names(),
                limit=40,
            ) if local_text.strip() else []
            if local_terms:
                parsed = dict(parsed)
                metadata = (
                    dict(parsed.get("metadata"))
                    if isinstance(parsed.get("metadata"), dict)
                    else {}
                )
                metadata["filename"] = str(metadata.get("filename") or name)
                metadata["hangul_classification_terms"] = local_terms
                parsed["metadata"] = metadata
        return parsed
    except KordocParseError as exc:
        if suffix in {".hwp", ".hwpx"}:
            try:
                text = extract_hwpx_text(data) if suffix == ".hwpx" else extract_hwp_text(data)
            except Exception:
                logger.warning("hangul_fallback_text_extraction_failed", exc_info=True)
                text = ""
            if text.strip():
                local_terms = extract_linear_ncs_classification_terms(
                    text,
                    excluded_hierarchy_names=_non_detail_hierarchy_names(),
                    limit=40,
                )
                return {
                    "markdown": text,
                    "parser": f"{suffix[1:]}_text_fallback",
                    "metadata": {
                        "filename": name,
                        "fallback": f"{suffix[1:]}-text",
                        "hangul_classification_terms": local_terms,
                    },
                    "warnings": ["Kordoc parse failed; used local Hangul text fallback."],
                }
        if suffix == ".pdf":
            text = extract_pdf_text(data)
            if not text.strip():
                try:
                    text = extract_pdf_text_fallback(data, max_pages=6)
                except Exception:
                    text = ""
            if text.strip():
                local_terms = extract_linear_ncs_classification_terms(
                    text,
                    excluded_hierarchy_names=_non_detail_hierarchy_names(),
                    limit=40,
                )
                # A single PDF or a ZIP member can combine an undeveloped job
                # and other fully classified jobs. Preserve structural detail
                # cells even when the fallback text has lost table boundaries;
                # the later MCP exact lookup rejects proprietary labels.
                try:
                    structural_result = extract_sclass_from_pdf_bytes(
                        data,
                        filename=name,
                    )
                    local_terms.extend(
                        str(value or "").strip()
                        for value in (structural_result.get("detail_candidates") or [])
                        if str(value or "").strip()
                    )
                except Exception:
                    logger.warning(
                        "pdf_fallback_structural_classification_unavailable",
                        exc_info=True,
                    )
                local_terms = list(dict.fromkeys(local_terms))[:40]
                return {
                    "markdown": text,
                    "parser": "pdf_text_fallback",
                    "metadata": {
                        "filename": name,
                        "fallback": "pdf-text",
                        "classification_terms": local_terms,
                    },
                    "warnings": ["Kordoc parse failed; used PDF text fallback."],
                }
        if suffix in {".png", ".jpg", ".jpeg", ".webp"}:
            raise HTTPException(
                status_code=422,
                detail=(
                    f"{label} 이미지에서 텍스트를 읽지 못했습니다. "
                    "세분류 검증을 위해 PDF로 변환해 다시 탑재해 주세요."
                ),
            ) from exc
        raise HTTPException(
            status_code=422,
            detail=f"{label} could not be parsed by Kordoc",
        ) from exc


def _parse_upload_document(data: bytes, filename: str, label: str) -> dict[str, Any]:
    """Parse one upload or a ZIP of supported documents into one markdown payload."""

    name = str(filename or "")
    if _suffix_of(name) != ".zip":
        return _parse_single_document_upload(data, name, label)

    max_bytes = settings.max_upload_bytes()
    members: list[dict[str, Any]] = []
    chunks: list[str] = []
    warnings: list[str] = []
    total_uncompressed = 0
    try:
        with zipfile.ZipFile(io.BytesIO(data)) as archive:
            for info in archive.infolist():
                if info.is_dir():
                    continue
                member_name = info.filename
                suffix = _suffix_of(member_name)
                if suffix not in _SUPPORTED_ARCHIVE_DOC_SUFFIXES:
                    continue
                total_uncompressed += int(info.file_size or 0)
                if total_uncompressed > max_bytes:
                    raise HTTPException(
                        status_code=413,
                        detail=f"{label} archive contents exceed MAX_UPLOAD_MB ({max_bytes // (1024 * 1024)} MB)",
                    )
                if len(members) >= _ARCHIVE_MEMBER_LIMIT:
                    warnings.append(f"archive member limit reached: {_ARCHIVE_MEMBER_LIMIT}")
                    break
                member_label = _safe_member_label(member_name)
                if info.flag_bits & 0x1:
                    warnings.append(f"{member_label}: encrypted ZIP member is not supported")
                    continue
                try:
                    member_bytes = archive.read(info)
                except (RuntimeError, OSError, zipfile.BadZipFile):
                    warnings.append(f"{member_label}: ZIP member could not be read")
                    continue
                try:
                    parsed = _parse_single_document_upload(member_bytes, member_label, label)
                except HTTPException as exc:
                    warnings.append(f"{member_label}: {exc.detail}")
                    continue
                markdown = str(parsed.get("markdown") or "").strip()
                if not markdown:
                    warnings.append(f"{member_label}: empty parse result")
                    continue
                member_metadata = parsed.get("metadata") if isinstance(parsed.get("metadata"), dict) else {}
                member_parser = str(parsed.get("parser") or "unknown").strip()
                member_version = str(parsed.get("parser_version") or "").strip()
                member_record = {
                    "filename": member_label,
                    "suffix": suffix,
                    "parser": member_parser,
                }
                if member_parser == "kordoc" and member_version:
                    member_record["parser_version"] = member_version
                member_fallback = str(member_metadata.get("fallback") or "").strip()
                if member_fallback:
                    member_record["fallback"] = member_fallback
                member_terms = [
                    str(value or "").strip()
                    for value in (member_metadata.get("hangul_classification_terms") or [])
                    if str(value or "").strip()
                ]
                if member_terms:
                    member_record["hangul_classification_terms"] = member_terms
                classification_terms = [
                    str(value or "").strip()
                    for value in (member_metadata.get("classification_terms") or [])
                    if str(value or "").strip()
                ]
                classification_terms = list(dict.fromkeys(classification_terms))[:40]
                if classification_terms:
                    member_record["classification_terms"] = classification_terms
                members.append(member_record)
                chunks.append(f"# ZIP member: {member_label}\n\n{markdown}")
                warnings.extend(str(x) for x in (parsed.get("warnings") or []) if str(x).strip())
    except zipfile.BadZipFile as exc:
        raise HTTPException(status_code=422, detail=f"{label} is not a readable ZIP archive") from exc

    if not chunks:
        raise HTTPException(
            status_code=422,
            detail=f"{label} ZIP contains no parseable PDF/HWP/HWPX/DOCX/TXT/image job-description files",
        )
    member_parsers = list(
        dict.fromkeys(
            str(member.get("parser") or "unknown").strip()
            for member in members
        )
    )
    archive_parser = member_parsers[0] if len(member_parsers) == 1 else "mixed_document_parsers"
    archive_versions = list(
        dict.fromkeys(
            str(member.get("parser_version") or "").strip()
            for member in members
            if str(member.get("parser") or "").strip() == "kordoc"
            and str(member.get("parser_version") or "").strip()
        )
    )
    result = {
        "markdown": "\n\n---\n\n".join(chunks),
        "metadata": {"filename": name, "archive": True, "members": members},
        "warnings": warnings,
        "parser": archive_parser,
    }
    if archive_parser == "kordoc" and len(archive_versions) == 1:
        result["parser_version"] = archive_versions[0]
    return result


def _sha256_bytes(data: bytes) -> str:
    return hashlib.sha256(data or b"").hexdigest()


def _sha256_text(text: str) -> str:
    return hashlib.sha256(str(text or "").encode("utf-8")).hexdigest()


def _quality_feedback_context(ncs_matches: list[dict[str, Any]] | None) -> str:
    codes = list(
        dict.fromkeys(
            str(row.get("ncsClCd") or "").strip()
            for row in (ncs_matches or [])
            if isinstance(row, dict) and str(row.get("ncsClCd") or "").strip()
        )
    )
    if not codes:
        return ""
    try:
        events = list_question_quality_feedback(codes, limit=40)
    except Exception:
        logger.warning("question_quality_feedback_unavailable", exc_info=True)
        return ""
    contexts = [feedback_prompt_context(events, code, max_items=4) for code in codes[:6]]
    return "\n\n".join(context for context in contexts if context)


def _join_generation_context(*parts: str) -> str:
    return "\n\n".join(str(part or "").strip() for part in parts if str(part or "").strip())[:12000]


def _register_question_quality_evidence(
    strategy: dict[str, Any],
    *,
    source_endpoint: str,
    ncs_matches: list[dict[str, Any]] | None,
) -> dict[str, Any]:
    report = strategy.get("question_quality_report") if isinstance(strategy.get("question_quality_report"), dict) else {}
    control = derive_quality_control(strategy)
    policy_version = str(report.get("policy") or QUALITY_POLICY_VERSION).strip() or QUALITY_POLICY_VERSION
    questions = [item for item in (strategy.get("interview_questions") or []) if isinstance(item, dict)]
    report_items = {
        int(item.get("index") or 0): item
        for item in (report.get("items") or [])
        if isinstance(item, dict)
    }
    question_items: list[dict[str, Any]] = []
    for index, question in enumerate(questions, start=1):
        text = str(question.get("question") or "").strip()
        question_hash = _sha256_text(text)
        question["question_hash"] = question_hash
        item_report = report_items.get(index) or {}
        question_items.append(
            {
                "index": index,
                "question_hash": question_hash,
                "method": str(question.get("type") or question.get("method") or "").strip(),
                "ncs_code": str(question.get("ncsClCd") or "").strip(),
                "question_focus": str(question.get("question_focus") or "").strip(),
                "question_focus_surface": str(question.get("question_focus_surface") or "").strip(),
                "question_evidence_id": str(question.get("question_evidence_id") or "").strip(),
                "question_source": str(question.get("question_source") or "").strip(),
                "model_replacement_reasons": [
                    str(reason).strip()
                    for reason in (question.get("model_replacement_reasons") or [])
                    if str(reason).strip()
                ][:30]
                if isinstance(question.get("model_replacement_reasons"), list)
                else [],
                "quality_repair_reasons": [
                    str(reason).strip()
                    for reason in (question.get("quality_repair_reasons") or [])
                    if str(reason).strip()
                ][:30]
                if isinstance(question.get("quality_repair_reasons"), list)
                else [],
                "ready": bool(item_report.get("ready")),
                "issues": list(item_report.get("issues") or [])[:30],
                "checks": dict(item_report.get("checks") or {}),
                "check_statuses": dict(item_report.get("check_statuses") or {}),
            }
        )
    codes = list(
        dict.fromkeys(
            [
                str(row.get("ncsClCd") or "").strip()
                for row in (ncs_matches or [])
                if isinstance(row, dict) and str(row.get("ncsClCd") or "").strip()
            ]
            + [item["ncs_code"] for item in question_items if item["ncs_code"]]
        )
    )
    names = list(
        dict.fromkeys(
            str(row.get("compeUnitName") or "").strip()
            for row in (ncs_matches or [])
            if isinstance(row, dict) and str(row.get("compeUnitName") or "").strip()
        )
    )
    run_id = f"qqr_{secrets.token_hex(16)}"
    # Hex avoids accidental collisions with API-key-shaped substrings checked
    # in user-authored feedback fields while retaining 192 bits of entropy.
    review_token = f"qqt_{secrets.token_hex(24)}"
    payload = {
        "id": run_id,
        "review_token": review_token,
        "source_endpoint": source_endpoint,
        "ncs_codes": codes,
        "competency_names": names,
        "quality_policy_version": policy_version,
        "generator_version": str(strategy.get("question_customization_policy") or ""),
        "question_count": int((report.get("summary") or {}).get("question_count") or len(questions)),
        "ready_count": int((report.get("summary") or {}).get("ready_count") or 0),
        "review_required": bool(control.get("review_required", True)),
        "escalation_required": bool(control.get("escalation_required", False)),
        "exception_allowed": bool(control.get("exception_allowed", True)),
        "trigger_codes": list(control.get("trigger_codes") or []),
        "evidence": {
            "quality_policy": policy_version,
            "question_items": question_items,
            "source_reference": "NCS 채용모델 면접문항 자료실",
            "recorded_at": datetime.now(timezone.utc).isoformat(),
        },
    }
    try:
        persisted = create_question_quality_run(payload)
        quality_control = {
            **control,
            "policy_version": policy_version,
            "run_id": persisted["id"],
            "review_token": persisted["review_token"],
            "evidence_recorded": True,
        }
    except Exception:
        logger.error("question_quality_evidence_persistence_failed", exc_info=True)
        trigger_codes = list(dict.fromkeys([*(control.get("trigger_codes") or []), "quality_evidence_store_unavailable"]))
        quality_control = {
            **control,
            "policy_version": policy_version,
            "run_id": "",
            "review_token": "",
            "review_required": True,
            "escalation_required": True,
            "trigger_codes": trigger_codes,
            "evidence_recorded": False,
        }
    strategy["quality_control"] = quality_control
    return quality_control


def _request_ip_hash(request: Request | None) -> str:
    host = ""
    try:
        host = str((request.client.host if request and request.client else "") or "").strip()
    except Exception:
        host = ""
    if not host:
        return ""
    return _sha256_text(host)


_CODEX_PROXY_HEADER_NAMES = frozenset(
    {
        "client-ip",
        "forwarded",
        "x-forwarded-for",
        "x-forwarded-host",
        "x-forwarded-proto",
        "x-real-ip",
    }
)


_GENERATION_PROVIDER_ALIASES = {
    "codex": "codex_cli",
    "codex_cli": "codex_cli",
    "claude": "claude_code",
    "claude_code": "claude_code",
    "openai": "openai_api",
    "openai_api": "openai_api",
    "openrouter": "openrouter_api",
    "openrouter_api": "openrouter_api",
}
_SUBSCRIPTION_CLI_PROVIDERS = frozenset({"codex_cli", "claude_code"})


def _configured_generation_provider() -> str:
    return configured_generation_provider()


def _request_generation_provider(value: Any = "") -> str:
    provider = normalize_generation_provider(value, default=_configured_generation_provider())
    if provider not in request_supported_generation_providers():
        raise HTTPException(
            status_code=400,
            detail={
                "code": "generation_provider_unsupported",
                "provider": str(provider or ""),
                "message": (
                    "generation_provider는 'openai_api'만 지원합니다. "
                    "공개 질문 생성에서는 OpenRouter 및 개인 구독 CLI 로그인을 사용할 수 없습니다."
                ),
                "retryable": False,
            },
        )
    return provider


def _inferred_generation_provider_from_key(request_key: str, provider: Any = "") -> str:
    try:
        resolved = resolve_generation_credential(
            generation_api_key=request_key,
            requested_provider=provider,
        )
    except GenerationCredentialError:
        return normalize_generation_provider(
            provider,
            default=_configured_generation_provider(),
        )
    return resolved.provider


def _require_local_subscription_cli_request(
    request: Request | None,
    provider: str,
) -> None:
    """Keep cached subscription logins inside the local workstation boundary."""

    if provider not in _SUBSCRIPTION_CLI_PROVIDERS:
        return
    try:
        has_proxy_headers = bool(request) and any(
            name in request.headers for name in _CODEX_PROXY_HEADER_NAMES
        )
    except Exception:
        # If the request metadata cannot be inspected, do not risk exposing the
        # workstation's cached subscription login through an intermediary.
        has_proxy_headers = True
    if has_proxy_headers:
        raise HTTPException(
            status_code=403,
            detail=f"{provider} generation is available only from this local workstation",
        )
    host = ""
    try:
        host = str((request.client.host if request and request.client else "") or "").strip()
    except Exception:
        host = ""
    normalized_host = host.strip("[]").casefold()
    try:
        is_loopback = bool(ipaddress.ip_address(normalized_host).is_loopback)
    except ValueError:
        # Starlette's in-process TestClient uses the synthetic host
        # ``testclient``. It cannot originate a network request and is needed
        # for exercising the real upload path in local regression tests.
        is_loopback = normalized_host in {"localhost", "testclient"}
    if not is_loopback:
        raise HTTPException(
            status_code=403,
            detail=f"{provider} generation is available only from this local workstation",
        )


def _require_local_codex_request(request: Request | None) -> None:
    """Backward-compatible guard used by existing Codex boundary tests."""

    _require_local_subscription_cli_request(
        request,
        "codex_cli",
    )


def _record_audit_event(
    request: Request | None,
    *,
    action: str,
    resource_type: str,
    resource_id: str,
) -> None:
    try:
        record_audit_log(
            actor_id="anonymous",
            action=action,
            resource_type=resource_type,
            resource_id=resource_id,
            ip_hash=_request_ip_hash(request),
        )
    except Exception:
        return


def _prune_review_sessions(now: float | None = None, *, persist: bool = False) -> None:
    current = float(now if now is not None else time.time())
    expired = [
        session_id
        for session_id, session in _REVIEW_SESSION_BY_ID.items()
        if current - float(session.get("created_at", 0.0) or 0.0) > _REVIEW_SESSION_TTL_SEC
    ]
    for session_id in expired:
        _REVIEW_SESSION_BY_ID.pop(session_id, None)
    if len(_REVIEW_SESSION_BY_ID) > _REVIEW_SESSION_MAX:
        oldest = sorted(
            _REVIEW_SESSION_BY_ID.items(),
            key=lambda item: float(item[1].get("created_at", 0.0) or 0.0),
        )
        for session_id, _ in oldest[: max(0, len(_REVIEW_SESSION_BY_ID) - _REVIEW_SESSION_MAX)]:
            _REVIEW_SESSION_BY_ID.pop(session_id, None)
    if persist:
        try:
            prune_review_sessions(
                now=current,
                ttl_sec=_REVIEW_SESSION_TTL_SEC,
                max_count=_REVIEW_SESSION_MAX,
            )
        except Exception:
            logger.warning("review_session_metadata_prune_failed", exc_info=True)


def _review_session_signing_key() -> bytes:
    raw = os.getenv("REVIEW_SESSION_SIGNING_KEY", "")
    if not raw or not raw.strip():
        return b""
    return raw.encode("utf-8")


def _review_session_timestamp(value: Any, field: str) -> float:
    if isinstance(value, bool) or not isinstance(value, (int, float)):
        raise ValueError(f"{field} must be a finite timestamp")
    timestamp = float(value)
    if not math.isfinite(timestamp):
        raise ValueError(f"{field} must be a finite timestamp")
    return timestamp


def _review_session_signature_message(session: dict[str, Any]) -> bytes:
    created_at = _review_session_timestamp(session.get("created_at"), "created_at")
    expires_at = _review_session_timestamp(session.get("expires_at"), "expires_at")
    claims = {
        "created_at": created_at.hex(),
        "document_sha256": str(session.get("document_sha256") or ""),
        "expires_at": expires_at.hex(),
        "filename": str(session.get("filename") or ""),
        "id": str(session.get("id") or ""),
        "markdown_sha256": str(session.get("markdown_sha256") or ""),
        "version": _REVIEW_SESSION_SIGNATURE_VERSION,
    }
    return json.dumps(
        claims,
        ensure_ascii=False,
        sort_keys=True,
        separators=(",", ":"),
    ).encode("utf-8")


def _sign_review_session(session: dict[str, Any], signing_key: bytes) -> str:
    digest = hmac.new(
        signing_key,
        _review_session_signature_message(session),
        hashlib.sha256,
    ).digest()
    encoded = base64.urlsafe_b64encode(digest).rstrip(b"=").decode("ascii")
    return f"{_REVIEW_SESSION_SIGNATURE_VERSION}.{encoded}"


def _verified_signed_review_session(
    review_session_payload: dict[str, Any],
    session_id: str,
    signing_key: bytes,
    *,
    now: float,
    field_name: str = "jd_review_json",
) -> dict[str, Any]:
    """Validate public review metadata and reconstruct trusted session claims."""

    try:
        signed_id = review_session_payload.get("id")
        document_sha256 = review_session_payload.get("document_sha256")
        markdown_sha256 = review_session_payload.get("markdown_sha256")
        filename = review_session_payload.get("filename")
        token = review_session_payload.get("token")
        if not isinstance(signed_id, str) or not _REVIEW_SESSION_ID_RE.fullmatch(signed_id):
            raise ValueError("invalid id")
        if signed_id != session_id:
            raise ValueError("mismatched id")
        if not isinstance(document_sha256, str) or not _REVIEW_SESSION_SHA256_RE.fullmatch(
            document_sha256
        ):
            raise ValueError("invalid document hash")
        if not isinstance(markdown_sha256, str) or not _REVIEW_SESSION_SHA256_RE.fullmatch(
            markdown_sha256
        ):
            raise ValueError("invalid markdown hash")
        if (
            not isinstance(filename, str)
            or not filename
            or len(filename) > 160
            or filename != _safe_member_label(filename)
        ):
            raise ValueError("unsafe filename")
        if not isinstance(token, str) or not _REVIEW_SESSION_TOKEN_RE.fullmatch(token):
            raise ValueError("invalid token")

        created_at = _review_session_timestamp(
            review_session_payload.get("created_at"),
            "created_at",
        )
        expires_at = _review_session_timestamp(
            review_session_payload.get("expires_at"),
            "expires_at",
        )
        if expires_at != created_at + _REVIEW_SESSION_TTL_SEC:
            raise ValueError("invalid expiry")

        claims = {
            "id": signed_id,
            "document_sha256": document_sha256,
            "markdown_sha256": markdown_sha256,
            "filename": filename,
            "created_at": created_at,
            "expires_at": expires_at,
        }
        expected_token = _sign_review_session(claims, signing_key)
        if not hmac.compare_digest(token, expected_token):
            raise ValueError("invalid signature")
        if created_at > now + _REVIEW_SESSION_MAX_CLOCK_SKEW_SEC:
            raise ValueError("future timestamp")
        if now > expires_at:
            raise ValueError("expired timestamp")
    except (TypeError, ValueError, OverflowError) as exc:
        raise HTTPException(
            status_code=409,
            detail=f"{field_name}.review_session is expired or invalid",
        ) from exc

    return {
        "id": signed_id,
        "document_sha256": document_sha256,
        "markdown_sha256": markdown_sha256,
        "filename": filename,
        "created_at": created_at,
    }


def _public_review_session(session: dict[str, Any]) -> dict[str, Any]:
    public = {
        "id": session["id"],
        "document_sha256": session["document_sha256"],
        "markdown_sha256": session["markdown_sha256"],
        "filename": session.get("filename", ""),
        "created_at": session["created_at"],
        "expires_at": session["created_at"] + _REVIEW_SESSION_TTL_SEC,
    }
    signing_key = _review_session_signing_key()
    if signing_key:
        public["token"] = _sign_review_session(public, signing_key)
    return public


def _create_review_session(upload_bytes: bytes, structured: dict[str, Any], filename: str) -> dict[str, Any]:
    document = structured.get("document") if isinstance(structured.get("document"), dict) else {}
    markdown = str(document.get("markdown") or "")
    session = {
        "id": secrets.token_urlsafe(24),
        "filename": _safe_member_label(filename),
        "created_at": time.time(),
        "document_sha256": _sha256_bytes(upload_bytes),
        "markdown_sha256": _sha256_text(markdown),
        "markdown_size": len(markdown.encode("utf-8")),
    }
    with _REVIEW_SESSION_LOCK:
        _prune_review_sessions(session["created_at"], persist=True)
        _REVIEW_SESSION_BY_ID[session["id"]] = session
        _prune_review_sessions(session["created_at"])
    try:
        create_review_session(session)
    except Exception:
        # Keep the existing local fallback if the optional persistence DB is
        # temporarily unavailable; normal single-process behavior is unchanged.
        logger.warning("review_session_metadata_persist_failed", exc_info=True)
    return _public_review_session(session)


def _validate_review_session(
    review_payload: dict[str, Any],
    upload_bytes: bytes,
    filename: str | None = None,
    *,
    field_name: str = "jd_review_json",
    upload_label: str = "jd_file",
    parse_endpoint: str = "/api/jd/parse-review",
) -> dict[str, Any]:
    review_session_payload = review_payload.get("review_session")
    if not isinstance(review_session_payload, dict):
        review_session_payload = {}
    session_id = str(
        review_payload.get("review_session_id")
        or review_session_payload.get("id")
        or ""
    ).strip()
    if not session_id:
        raise HTTPException(
            status_code=400,
            detail=(
                f"{field_name}.review_session_id is required; "
                f"call {parse_endpoint} before generation"
            ),
        )
    now = time.time()
    signing_key = _review_session_signing_key()
    signed_session: dict[str, Any] = {}
    if signing_key:
        signed_session = _verified_signed_review_session(
            review_session_payload,
            session_id,
            signing_key,
            now=now,
            field_name=field_name,
        )
    with _REVIEW_SESSION_LOCK:
        _prune_review_sessions(now)
        session = dict(_REVIEW_SESSION_BY_ID.get(session_id, {}) or {})
    if not session:
        try:
            session = dict(
                get_review_session(
                    session_id,
                    now=now,
                    ttl_sec=_REVIEW_SESSION_TTL_SEC,
                )
                or {}
            )
        except Exception:
            logger.warning("review_session_metadata_load_failed", exc_info=True)
            session = {}
    if not session and signed_session:
        # Vercel instances do not share memory or /tmp. The signed public
        # metadata is sufficient to recover the hash-only session safely.
        session = dict(signed_session)
    if not session:
        raise HTTPException(status_code=409, detail=f"{field_name}.review_session_id is expired or unknown")
    if signed_session and any(
        session.get(field) != signed_session.get(field)
        for field in (
            "id",
            "filename",
            "document_sha256",
            "markdown_sha256",
            "created_at",
        )
    ):
        raise HTTPException(
            status_code=409,
            detail=f"{field_name}.review_session does not match server parse session",
        )
    if session.get("document_sha256") != _sha256_bytes(upload_bytes):
        raise HTTPException(
            status_code=409,
            detail=f"{field_name} review session does not match uploaded {upload_label}",
        )
    if signing_key and filename is not None and session.get("filename") != _safe_member_label(filename):
        raise HTTPException(
            status_code=409,
            detail=f"{field_name} review session does not match uploaded {upload_label} filename",
        )
    payload_document = review_payload.get("document") if isinstance(review_payload.get("document"), dict) else {}
    payload_markdown = str(payload_document.get("markdown") or "")
    if signing_key and "markdown" not in payload_document:
        raise HTTPException(
            status_code=400,
            detail=f"{field_name}.document.markdown is required for signed review session validation",
        )
    if (payload_markdown or signing_key) and _sha256_text(payload_markdown) != session.get("markdown_sha256"):
        raise HTTPException(
            status_code=400,
            detail=f"{field_name}.document.markdown does not match server parse session",
        )
    if payload_markdown:
        session["markdown"] = payload_markdown
    return session


def _sanitize_request_openai_key(value: str | None, *, provider: str = "openai_api") -> str:
    """Validate a BYOK credential without retaining or reflecting it."""

    key = str(value or "").strip()
    if not key:
        return ""
    if len(key) > 512 or any(ord(char) < 33 or ord(char) > 126 for char in key):
        raise HTTPException(
            status_code=400,
            detail={
                "code": request_key_error_code(provider, "key_invalid"),
                "provider": _inferred_generation_provider_from_key(key, provider),
                "message": request_key_error_message(provider, invalid=True),
                "retryable": False,
            },
        )
    return key


def _openai_key_source(request_key: str, request: Request | None = None) -> str:
    del request
    return settings.openai_key_source(request_key)


def _generation_key_source(
    provider: str,
    *,
    generation_api_key: Any = "",
    openai_api_key: Any = "",
    openrouter_api_key: Any = "",
) -> str:
    if any(
        str(value or "").strip()
        for value in (generation_api_key, openai_api_key, openrouter_api_key)
    ):
        return "request"
    if normalize_generation_provider(provider) == OPENROUTER_PROVIDER:
        return settings.openrouter_key_source("")
    return settings.openai_key_source("")


def _sanitize_request_generation_model(
    value: Any,
    *,
    provider: str = "",
) -> str:
    try:
        requested_model = sanitize_generation_model(value)
    except ValueError as exc:
        raise HTTPException(
            status_code=400,
            detail={
                "code": "generation_model_invalid",
                "provider": _request_generation_provider(provider or _configured_generation_provider()),
                "message": "generation_model 형식이 올바르지 않습니다.",
                "retryable": False,
            },
        ) from exc
    if requested_model:
        raise HTTPException(
            status_code=400,
            detail={
                "code": "generation_model_override_disabled",
                "provider": _request_generation_provider(
                    provider or _configured_generation_provider()
                ),
                "message": (
                    "공개 질문 생성 모델은 역할별 승인 구성으로 고정되어 있습니다. "
                    "OpenAI API 키만 입력해 주세요."
                ),
                "retryable": False,
            },
        )
    return ""


def _resolve_request_generation(
    *,
    generation_api_key: Any = "",
    openai_api_key: Any = "",
    openrouter_api_key: Any = "",
    provider: Any = "",
    generation_model: Any = "",
) -> tuple[str, str, str]:
    sanitized_generation_key = _sanitize_request_openai_key(
        generation_api_key,
        provider=provider or "",
    )
    sanitized_openai_key = _sanitize_request_openai_key(
        openai_api_key,
        provider="openai_api",
    )
    sanitized_openrouter_key = _sanitize_request_openai_key(
        openrouter_api_key,
        provider="openrouter",
    )
    try:
        credential = resolve_generation_credential(
            generation_api_key=sanitized_generation_key,
            openai_api_key=sanitized_openai_key,
            openrouter_api_key=sanitized_openrouter_key,
            requested_provider=provider,
        )
    except GenerationCredentialError as exc:
        raise HTTPException(
            status_code=400,
            detail={
                "code": exc.code,
                "provider": exc.provider or "",
                "message": exc.message,
                "retryable": False,
            },
        ) from exc
    requested_model = _sanitize_request_generation_model(
        generation_model,
        provider=credential.provider,
    )
    resolved_model = resolve_generation_model(
        provider=credential.provider,
        explicit_model=requested_model,
    )
    return credential.provider, resolved_model, credential.api_key


def _require_allowed_openai_key(
    request_key: str,
    request: Request | None = None,
    *,
    provider: str = "openai_api",
) -> None:
    del request
    if str(request_key or "").strip():
        return
    resolved_provider = _inferred_generation_provider_from_key(request_key, provider)
    raise HTTPException(
        status_code=400,
        detail={
            "code": request_key_error_code(resolved_provider, "key_required"),
            "provider": resolved_provider,
            "message": request_key_error_message(resolved_provider),
            "retryable": False,
        },
    )


_SENSITIVE_QUERY_PARAMS = (
    "generation_api_key",
    "openai_api_key",
    "openrouter_api_key",
    "job_posting",
    "user_profile",
    "notice_text",
    "duty_text",
    "qualification_text",
    "preference_text",
    "evaluation_text",
    "strengths",
    "jd_review_json",
    "notice_review_json",
    "question_plan_json",
    "interview_methods_json",
    "avoid_questions",
    "avoid_questions_json",
)


def _reject_sensitive_query_params(request: Request, *, destination: str) -> None:
    blocked = [name for name in _SENSITIVE_QUERY_PARAMS if name in request.query_params]
    if not blocked:
        return
    raise HTTPException(
        status_code=400,
        detail=f"{', '.join(blocked)} must be sent in the {destination}, not the query string",
    )


def _ksa_key(value: str) -> str:
    return re.sub(r"\s+", "", str(value or "")).lower()


def _criteria_values(row: dict[str, Any]) -> list[str]:
    """Return normalized NCS performance-criteria text without exposing it in prompts."""

    raw = (
        row.get("performanceCriteria")
        or row.get("performance_criteria")
        or row.get("performanceCriteriaTexts")
        or row.get("criteria")
        or []
    )
    if isinstance(raw, dict):
        raw = raw.get("items") or raw.get("criteria") or raw.get("performance_criteria") or [raw]
    if isinstance(raw, str):
        raw = [raw]
    if not isinstance(raw, (list, tuple)):
        return []
    values: list[str] = []
    seen: set[str] = set()
    for value in raw:
        if isinstance(value, dict):
            value = (
                value.get("text")
                or value.get("description")
                or value.get("criterion")
                or value.get("criteria")
                or value.get("performance_criteria")
                or value.get("performanceCriteria")
            )
        text = _clean_question_text(value, max_chars=240)
        if text and text not in seen:
            seen.add(text)
            values.append(text)
    return values[:12]


def _clean_ksa_evidence_row(row: dict[str, Any]) -> dict[str, Any]:
    return {
        "evidence_id": stable_ksa_evidence_id(row),
        "ncsClCd": str(row.get("ncsClCd", "")).strip(),
        "compeUnitName": str(row.get("compeUnitName", "")).strip(),
        "factorName": str(row.get("factorName", "")).strip(),
        "ksaTypeName": str(
            row.get("ksaTypeName") or row.get("factorType") or row.get("ksa_type") or ""
        ).strip(),
        "elementName": str(row.get("elementName", "")).strip(),
        "factorLevel": str(row.get("factorLevel", "")).strip(),
        "factorSource": str(row.get("factorSource", "")).strip(),
        "ksaStatus": str(row.get("ksaStatus", "")).strip(),
        "elementId": str(row.get("elementId") or row.get("element_id") or "").strip(),
        "factorNo": str(row.get("factorNo") or row.get("ksaNo") or row.get("number") or "").strip(),
        "performanceCriteria": _criteria_values(row),
    }


def _ncs_traceability_payload(
    question: dict[str, Any],
    evidence: dict[str, Any] | None,
) -> dict[str, Any]:
    """Attach an internal, auditable link from a question to NCS evidence.

    The payload is deliberately separate from candidate-facing wording.  It
    records the ability unit, element, performance criteria (when returned by
    NCS_MCP), KSA factor and stable evidence id so reviewers can trace the
    question back to the official source without copying NCS labels into the
    question itself.
    """

    row = evidence if isinstance(evidence, dict) else {}
    criteria = _criteria_values(row)
    return {
        "ncs_code": str(question.get("ncsClCd") or row.get("ncsClCd") or "").strip(),
        "ability_unit_name": str(
            question.get("competency")
            or row.get("compeUnitName")
            or ""
        ).strip(),
        "element_id": str(row.get("elementId") or "").strip(),
        "element_name": str(row.get("elementName") or "").strip(),
        "performance_criteria": criteria,
        "performance_criteria_linked": bool(criteria),
        "ksa_type": str(row.get("ksaTypeName") or "").strip(),
        "ksa_factor": str(row.get("factorName") or "").strip(),
        "ksa_factor_no": str(row.get("factorNo") or "").strip(),
        "evidence_id": str(row.get("evidence_id") or "").strip(),
        "source": str(row.get("factorSource") or row.get("source") or "").strip(),
    }


_METHOD_MAIN_QUESTION_REQUIRED_TERMS: dict[str, tuple[str, ...]] = {
    "경험면접": ("경험", "상황", "본인", "행동", "결과"),
    "상황면접": ("상황", "판단", "기준", "순서", "위험"),
    "발표면접": ("발표", "진단", "대안", "실행", "성과지표", "질의응답"),
    "토론면접": ("토론", "충돌", "입장", "근거", "합의"),
    "인바스켓면접": ("인바스켓", "문서", "우선순위", "보고", "위임", "직접처리"),
    "직무지식면접": ("절차", "기준", "산출물", "예외상황"),
    "창의적 문제해결력면접": ("창의적", "미래예측", "문제", "정의", "대안", "검증", "실행"),
}

_METHOD_EVALUATION_ANCHORS: dict[str, tuple[str, ...]] = {
    "경험면접": ("구체적상황", "본인역할", "행동", "결과", "성과", "학습", "판단근거"),
    "상황면접": (
        "사실확인", "핵심사실", "규정확인", "판단기준", "대안별위험", "행동순서",
        "보고순서", "위험요인", "이해관계자", "첫조치", "후속조치", "예방",
    ),
    "발표면접": ("자료분석", "논리적구조화", "대안", "실행계획", "성과지표", "질의응답"),
    "토론면접": (
        "초기입장", "근거", "경청", "상호작용", "갈등조정", "쟁점조정",
        "합의안", "공통안", "미합의", "이송안", "공동합의", "반대의견", "타당성검토",
    ),
    "인바스켓면접": ("우선순위", "문서", "요청분류", "보고", "위임", "직접처리", "시간관리", "리스크", "후속점검"),
    "직무지식면접": ("절차", "기준", "직무지식", "예외상황", "산출물", "품질", "오류예방"),
    "창의적 문제해결력면접": ("미래예측", "문제정의", "창의적사고", "원인가설", "대안", "검증", "실현가능성", "의사결정", "실행계획", "성과지표", "리스크"),
}

_FOLLOW_UP_METHOD_ANCHORS: dict[str, tuple[str, ...]] = {
    "경험면접": ("상황", "역할", "행동", "선택", "기준", "이유", "성과", "개선", "학습", "교훈"),
    "상황면접": ("확인", "기준", "이유", "위험", "이해관계자", "순서", "후속", "예방"),
    "발표면접": ("근거자료", "대안", "우선순위", "반대의견", "답변", "질의응답", "일정", "자원", "성과지표", "리스크"),
    "토론면접": ("초기입장", "입장발표", "근거", "반대의견", "수용", "조정", "합의안", "실행책임", "후속점검"),
    "인바스켓면접": ("문서", "요청", "분류", "우선순위", "먼저처리", "첫조치", "기준", "행동", "조치", "보류", "보고", "위임", "직접처리", "통제", "기록"),
    "직무지식면접": ("기준", "규정", "예외상황", "산출물", "품질", "리스크", "보완책", "교육", "순서"),
    "창의적 문제해결력면접": ("미래예측", "문제정의", "문제", "정의", "가설", "검증", "대안", "실현가능성", "의사결정", "실행계획", "우선순위", "리스크", "보완책", "성과지표", "후속점검"),
}

_FOLLOW_UP_REQUIRED_ANCHOR_GROUPS: dict[str, tuple[tuple[str, ...], ...]] = {
    "인바스켓면접": (
        ("문서", "요청", "분류", "우선순위", "먼저처리", "첫조치", "보류"),
        ("보고", "위임", "직접처리", "기록", "통제", "후속점검"),
    ),
    "직무지식면접": (
        ("절차", "기준", "규정", "순서"),
        ("예외상황", "산출물", "품질", "오류", "리스크", "보완책"),
    ),
}

_FOLLOW_UP_FOCUS_SLOT_INDEX: dict[str, int] = {
    "경험면접": 1,
    "상황면접": 1,
    "발표면접": 0,
    "토론면접": 0,
    "인바스켓면접": 0,
    "직무지식면접": 0,
    "창의적 문제해결력면접": 1,
}

_ASSESSABLE_EVALUATION_TERMS = (
    "상황",
    "역할",
    "행동",
    "결과",
    "성과",
    "문제",
    "정의",
    "가설",
    "판단",
    "기준",
    "근거",
    "위험",
    "대응",
    "분석",
    "구조",
    "대안",
    "실행",
    "지표",
    "소통",
    "합의",
    "조정",
    "분류",
    "우선순위",
    "보고",
    "위임",
    "처리",
    "시간",
    "절차",
    "산출물",
    "품질",
    "예외",
    "점검",
    "자료",
    "문서",
    "요청",
    "계획",
    "오류",
    "검증",
    "리스크",
    "보완",
    "기록",
    "검토",
    "확인",
    "적용",
    "파악",
)

_VAGUE_EVALUATION_POINT_KEYS = {
    "성실성",
    "태도",
    "열정",
    "자신감",
    "인성",
    "적극성",
    "책임감",
    "표현력",
}

_JOB_CONTEXT_STOPWORDS = {
    "해당",
    "직무",
    "업무",
    "관련",
    "기준",
    "핵심",
    "수행",
    "상황",
    "질문",
    "설명",
    "제시",
    "과정",
    "본인",
    "결과",
    "경험",
    "절차",
    "적용",
    "근거",
    "능력",
    "단위",
    "관리",
}

_KSA_RELEVANCE_STOPWORDS = _JOB_CONTEXT_STOPWORDS | {
    "지식",
    "기술",
    "능력",
    "태도",
    "자세",
    "의지",
    "관련",
    "해당",
    "정확히",
    "성실하고",
    "꼼꼼한",
    "바른",
    "적극적",
    "객관적",
}

_UNRESOLVED_KSA_PLACEHOLDER_RE = re.compile(r"(?<![A-Za-z])KSA(?![A-Za-z])", re.IGNORECASE)


_OFFICIAL_SAMPLE_FORMAT_RULES: dict[str, dict[str, tuple[str, ...]]] = {
    "경험면접": {
        "task_any": ("경험", "사례"),
        "task_all": ("상황", "본인", "행동", "결과"),
            "eval_any": (
                "본인역할과행동", "성과와학습", "구체적상황설명",
                "구체적상황과본인역할", "판단근거와실제행동", "결과확인근거", "학습과전이",
                "당시상황과본인역할", "선택근거와직접행동", "결과를입증하는기록",
            ),
    },
    "상황면접": {
        "task_any": ("상황",),
        "task_all": ("판단", "기준", "순서", "위험"),
        "eval_any": (
            "판단기준", "위험요인인식", "이해관계자대응",
            "핵심사실과규정확인", "대안별위험을반영한판단", "행동·보고순서", "후속조치와예방",
        ),
    },
    "발표면접": {
        "task_any": ("[발표과제]", "발표과제"),
        "task_all": ("발표", "진단", "대안", "실행", "성과지표", "질의응답"),
        "eval_any": (
            "자료분석력", "논리적구조화", "대안의실행가능성", "질의응답대응",
            "자료근거와현황·원인분석", "대안비교와우선순위", "성과지표와질의응답대응",
        ),
    },
    "토론면접": {
        "task_any": ("[토론과제]", "토론과제"),
        "task_all": ("충돌", "입장", "근거", "합의"),
        "eval_any": (
            "입장발표근거", "경청과상호작용", "갈등조정", "최종합의안도출",
            "확인근거의적절성", "상대근거의타당성검토", "쟁점조정", "실행가능한공동합의",
            "사실·규정에근거한초기입장", "대안별일정·책임영향비교", "반대근거검토와쟁점조정",
            "공통안또는미합의이송안의실행가능성",
        ),
    },
    "인바스켓면접": {
        "task_any": ("[인바스켓과제]", "인바스켓과제"),
        "task_all": ("문서", "우선순위", "보고", "위임", "직접처리"),
        "eval_any": (
            "우선순위판단", "문서·요청분류", "시간관리",
            "문서별긴급도·영향도판단", "보고·위임·직접처리구분", "초기행동과시간배분",
        ),
    },
    "직무지식면접": {
        "task_any": ("절차", "기준"),
        "task_all": ("절차", "기준", "산출물", "예외상황"),
        "eval_any": (
            "절차·기준이해", "직무지식적용", "산출물품질",
            "절차·기준의근거", "실제업무적용", "산출물품질과오류예방",
        ),
    },
    "창의적 문제해결력면접": {
        "task_any": ("[창의적문제해결력과제]", "창의적문제해결력과제", "창의적문제해결력"),
        "task_all": ("미래예측", "문제", "정의", "대안", "검증", "실행"),
        "eval_any": (
            "미래예측과문제정의", "창의적사고와대안도출", "검증방법", "실현가능성", "의사결정과실행계획",
            "근거기반문제정의", "복수대안과창의성", "검증방법과실현가능성", "의사결정·실행·위험보완",
        ),
    },
}


def _method_shape_ok(method: str, text: str) -> bool:
    compact = re.sub(r"\s+", "", str(text or ""))
    required = _METHOD_MAIN_QUESTION_REQUIRED_TERMS.get(method)
    if not required:
        return False
    if not all(term in compact for term in required):
        return False
    if method == "인바스켓면접" and not any(
        marker in compact
        for marker in ("동시", "여러", "요청", "메일", "보고요청", "일정충돌", "첫조치", "보류", "마감")
    ):
        return False
    if method == "직무지식면접":
        if any(marker in compact for marker in ("경험을말씀", "당시상황", "본인역할")):
            return False
        if not any(marker in compact for marker in ("설명", "확인", "적용", "점검", "예방", "제시")):
            return False
    return True


def _main_question_task_marker_ok(method: str, text: str) -> bool:
    marker_by_method = {
        "발표면접": ("[발표과제]", "발표과제"),
        "토론면접": ("[토론과제]", "토론과제"),
        "인바스켓면접": ("[인바스켓과제]", "인바스켓과제"),
        "창의적 문제해결력면접": ("[창의적 문제해결력과제]", "창의적문제해결력과제"),
    }
    markers = marker_by_method.get(method)
    if not markers:
        return True
    compact = re.sub(r"\s+", "", str(text or ""))
    return any(re.sub(r"\s+", "", marker) in compact for marker in markers)


def _normalize_model_task_marker(method: str, text: str) -> str:
    question = str(text or "").strip()
    if not question or _main_question_task_marker_ok(method, question):
        return question

    if method == "발표면접":
        prefix = "[발표과제]"
    elif method == "토론면접":
        prefix = "[토론과제]"
    elif method == "인바스켓면접":
        prefix = "[인바스켓과제]"
    elif method == "창의적 문제해결력면접":
        prefix = "[창의적 문제해결력과제]"
    else:
        prefix = ""
    if not prefix:
        return question
    if not _method_shape_ok(method, question) and not _method_shape_ok(method, f"{prefix} {question}"):
        return question
    return f"{prefix} {question}"


_TASK_METHODS_WITH_SEPARATE_OPERATING_CONDITIONS = {
    "발표면접",
    "토론면접",
    "인바스켓면접",
    "창의적 문제해결력면접",
}
_OPERATING_CONDITION_IN_QUESTION_PATTERNS = (
    re.compile(r"(?:준비|발표|질의\s*응답|토론|설명|문서\s*검토)\s*시간"),
    re.compile(r"제한\s*시간"),
    re.compile(r"첫\s*\d+\s*분\s*(?:행동|조치|계획)"),
    re.compile(r"(?:준비|발표|질의\s*응답|토론|설명|검토)\s*\d+\s*분"),
    re.compile(
        r"\d+\s*분\s*(?:동안|후|안에|이내)?\s*"
        r"(?:준비|발표|질의\s*응답|토론|설명|검토|답변)"
    ),
)


def _operating_conditions_separated(method: str, question: str) -> bool:
    """Keep exam timing/instructions in task_conditions, not in the substantive task."""
    if method not in _TASK_METHODS_WITH_SEPARATE_OPERATING_CONDITIONS:
        return True
    text = str(question or "").strip()
    return bool(text) and not any(
        pattern.search(text) for pattern in _OPERATING_CONDITION_IN_QUESTION_PATTERNS
    )


def _normalize_model_job_context(method: str, q: dict[str, Any], text: str) -> str:
    question = str(text or "").strip()
    if not question:
        return ""
    probe = dict(q)
    probe["model_question_preserved"] = True
    if _main_question_job_context_ok(probe, question):
        return question
    terms = _primary_job_context_terms(probe)
    context = _clean_question_text(terms[0] if terms else "", max_chars=60)
    if not context or _compact_contains_term(question, context):
        return question

    marker = ""
    body = question
    for candidate_marker in (
        "[발표과제]",
        "[토론과제]",
        "[인바스켓과제]",
        "[창의적 문제해결력과제]",
    ):
        if body.startswith(candidate_marker):
            marker = candidate_marker
            body = body[len(candidate_marker):].strip()
            break

    if marker:
        candidate = f"{marker} {context} 업무에서 {body}".strip()
    elif method == "상황면접":
        candidate = f"{context} 업무 중 {question}"
    else:
        candidate = f"{context}에서 {question}"

    if _contains_blind_hiring_cue(candidate):
        return question
    if not _method_shape_ok(method, candidate):
        return question
    if not _main_question_task_marker_ok(method, candidate):
        return question
    if not _main_question_job_context_ok(probe, candidate):
        return question
    return candidate


def _sanitize_candidate_document_leaks(value: Any, *, subject: str = "") -> str:
    """Remove copied notice/JD boilerplate from candidate-facing text.

    Model output must never expose legal disclaimers, page markers or a whole
    NCS table row as if it were the job context.  Keep the actual work label
    and sentence structure intact so this is a bounded repair, not a rewrite.
    """

    text = re.sub(r"\s+", " ", str(value or "").strip())
    if not text:
        return ""
    label = _normalize_ncs_detail_term(subject) or "해당 직무"
    # Models may copy the notice/JD preamble and attach a Korean particle to
    # the copied sentence ("...채용한다.에서", "...채용한다.의", or
    # "...채용한다.을 수행하던").  Replace the whole sentence while keeping
    # the particle's grammatical role in a short, candidate-facing label.
    leaked_sentence = re.compile(
        r"공고\s*[·ㆍ./]?\s*직무\s*기술\s*서\s*상?\s*"
        r"[^.。!?]{1,520}?[.。!?]\s*"
        r"(에서|으로|부터|까지|에게|의|에|은|는|이|가|을|를)?",
        flags=re.IGNORECASE,
    )

    def replace_leaked_sentence(match: re.Match[str]) -> str:
        particle = str(match.group(1) or "").strip()
        particle_map = {"을": "를", "를": "를"}
        normalized_particle = particle_map.get(particle, particle)
        return f"{label} 업무{normalized_particle}"

    text = leaked_sentence.sub(replace_leaked_sentence, text)
    # If the model stopped before a sentence boundary, remove the same copied
    # row up to the first candidate-facing experience/situation phrase.
    text = re.sub(
        r"공고\s*[·ㆍ./]?\s*직무\s*기술\s*서\s*상?\s*[^.。!?]{0,520}?(?=(?:실제\s*상황|경험\s*사례|상황에서|사례를))",
        f"{label} 업무에서 ",
        text,
        flags=re.IGNORECASE,
    )
    text = re.sub(r"\[NCS[^\]\n]{0,220}(?:\]|…)", " ", text, flags=re.IGNORECASE)
    text = re.sub(r"※[^\n]{0,320}(?:국가유공자|보훈|가점|법률|제\s*\d+\s*조)[^\n]*", " ", text)
    text = re.sub(r"\s*-\s*\d+\s*-\s*", " ", text)
    text = re.sub(r"\s+", " ", text).strip(" -")
    return text


def _evaluation_points_quality_ok(method: str, evaluation_points: list[str]) -> bool:
    if len(evaluation_points) != 4:
        return False
    compact_points = [re.sub(r"\s+", "", str(point or "")) for point in evaluation_points]
    if any(point in _VAGUE_EVALUATION_POINT_KEYS for point in compact_points):
        return False
    anchors = _METHOD_EVALUATION_ANCHORS.get(method, ())
    anchor_hits = {
        anchor
        for anchor in anchors
        if any(anchor in compact for compact in compact_points)
    }
    foreign_anchor_counts = {
        other_method: sum(
            1
            for anchor in other_anchors
            if any(anchor in compact for compact in compact_points)
        )
        for other_method, other_anchors in _METHOD_EVALUATION_ANCHORS.items()
        if other_method != method
    }
    if any(
        count >= 2
        and (
            count > len(anchor_hits)
            or (count == len(anchor_hits) and len(anchor_hits) < 3)
        )
        for count in foreign_anchor_counts.values()
    ):
        return False
    assessable_count = sum(
        1
        for compact in compact_points
        if any(term in compact for term in _ASSESSABLE_EVALUATION_TERMS)
    )
    return len(anchor_hits) >= 2 and assessable_count >= 3


def _job_context_terms(q: dict[str, Any]) -> list[str]:
    focus = str(q.get("question_focus") or "")
    surface_focus = str(q.get("question_focus_surface") or "")
    focus_type = str(q.get("question_focus_type") or "")
    raw_values: list[str] = [
        str(q.get("competency") or ""),
        str(q.get("ncs_detail") or ""),
        str(q.get("ncsSubdCdnm") or ""),
        str(q.get("ncsSclasCdnm") or ""),
        surface_focus,
        focus,
        _operational_focus_label(focus, focus_type),
    ]
    if isinstance(q.get("ksa_refs"), list):
        raw_values.extend(str(x or "") for x in q.get("ksa_refs") or [])
    terms: list[str] = []
    seen: set[str] = set()
    for value in raw_values:
        for token in re.findall(r"[가-힣A-Za-z0-9]{2,}", value):
            token = token.strip()
            key = token.lower()
            if key in _JOB_CONTEXT_STOPWORDS or len(token) < 2:
                continue
            if key not in seen:
                seen.add(key)
                terms.append(token)
    return terms[:8]


def _job_specific_context_ok(q: dict[str, Any], question: str, follow_ups: list[str]) -> bool:
    primary_terms = _primary_job_context_terms(q)
    if not primary_terms:
        return False
    compact_text = re.sub(r"\s+", "", "\n".join([question, *follow_ups])).lower()
    primary_hits = [
        term
        for term in primary_terms
        if re.sub(r"\s+", "", term).lower() in compact_text
    ]
    if not primary_hits and _is_subscription_cli_source(q.get("question_source")):
        primary_hits = [
            term
            for term in primary_terms
            if any(
                len(variant) >= 2 and variant in compact_text
                for variant in {
                    re.sub(r"\s+", "", term).lower(),
                    re.sub(
                        r"(?:관리|운영|지원|기획|평가|수립)$",
                        "",
                        re.sub(r"\s+", "", term).lower(),
                    ),
                }
            )
        ]
    if not primary_hits:
        return False
    if _is_subscription_cli_source(q.get("question_source")):
        return True
    terms = _job_context_terms(q)
    hits = [term for term in terms if re.sub(r"\s+", "", term).lower() in compact_text]
    required = 1 if len(terms) == 1 else 2
    return len(hits) >= required


def _primary_job_context_terms(q: dict[str, Any]) -> list[str]:
    raw_values = [
        str(q.get("competency") or ""),
        str(q.get("ncs_detail") or ""),
        str(q.get("ncsSubdCdnm") or ""),
        str(q.get("ncsSclasCdnm") or ""),
    ]
    terms: list[str] = []
    seen: set[str] = set()
    for value in raw_values:
        clean = _clean_question_text(value, max_chars=80)
        candidates = [clean]
        candidates.extend(re.findall(r"[가-힣A-Za-z0-9]{2,}", clean))
        for token in candidates:
            key = re.sub(r"\s+", "", str(token or "")).lower()
            if not key or key in _JOB_CONTEXT_STOPWORDS or len(key) < 2:
                continue
            if key not in seen:
                seen.add(key)
                terms.append(token)
    return terms[:8]


def _main_question_job_context_ok(q: dict[str, Any], question: str) -> bool:
    if not bool(q.get("model_question_preserved")):
        return True
    terms = _primary_job_context_terms(q)
    if not terms:
        return True
    compact_question = re.sub(r"\s+", "", str(question or "")).lower()
    for term in terms:
        compact_term = re.sub(r"\s+", "", term).lower()
        variants = {
            compact_term,
            re.sub(r"(?:관리|운영|지원|기획|평가|수립)$", "", compact_term),
        }
        if any(len(variant) >= 2 and variant in compact_question for variant in variants):
            return True
    return False


def _follow_ups_quality_ok(method: str, q: dict[str, Any], follow_ups: list[str]) -> bool:
    clean = [str(item or "").strip() for item in follow_ups if str(item or "").strip()]
    if len(clean) < 3:
        return False
    keys = [normalize_question_dedup_key(item) for item in clean]
    if any(not key for key in keys) or len(set(keys)) != len(keys):
        return False

    compact_items = [re.sub(r"\s+", "", item) for item in clean]
    merged = "\n".join(compact_items).lower()
    anchors = _FOLLOW_UP_METHOD_ANCHORS.get(method, ())
    anchor_hits = {
        anchor
        for anchor in anchors
        if any(anchor in compact for compact in compact_items)
    }
    if len(anchor_hits) < 2:
        return False
    required_groups = _FOLLOW_UP_REQUIRED_ANCHOR_GROUPS.get(method, ())
    for group in required_groups:
        if not any(anchor in compact for anchor in group for compact in compact_items):
            return False

    open_prompt_hits = sum(
        1
        for item in clean[:3]
        if re.search(r"(무엇|어떤|어떻게|얼마|어땠|어떠|왜|기준|이유|설명|말씀|제시|확인|선택|평가|점검|정리)", item)
    )
    if open_prompt_hits < 3:
        return False

    context_terms = _job_context_terms(q)
    if context_terms:
        context_hits = [
            term
            for term in context_terms
            if re.sub(r"\s+", "", term).lower() in merged
        ]
        if not context_hits:
            return False
    focus = str(q.get("question_focus_surface") or q.get("question_focus") or "").strip()
    if focus and not _ksa_factor_relevant_to_text(focus, "\n".join(clean)):
        return False
    return True


def _ksa_factor_relevant_to_text(factor_name: str, text: str) -> bool:
    factor = str(factor_name or "").strip()
    if not factor:
        return False
    compact_factor = re.sub(r"\s+", "", factor).lower()
    compact_text = re.sub(r"\s+", "", str(text or "")).lower()
    if compact_factor and compact_factor in compact_text:
        return True
    operational_factor = re.sub(
        r"\s*(?:관련\s*)?(?:능력|기술|스킬|지식)\s*$",
        "",
        factor,
    ).strip()
    compact_operational_factor = re.sub(r"\s+", "", operational_factor).lower()
    if len(compact_operational_factor) >= 2 and compact_operational_factor in compact_text:
        return True

    tokens: list[str] = []
    seen: set[str] = set()
    for token in re.findall(r"[가-힣A-Za-z0-9]{2,}", factor):
        key = token.lower()
        if key in _KSA_RELEVANCE_STOPWORDS:
            continue
        if key not in seen:
            seen.add(key)
            tokens.append(token)
    if not tokens:
        return False
    hits = [
        token
        for token in tokens
        if re.sub(r"\s+", "", token).lower() in compact_text
    ]
    required = 1 if len(tokens) == 1 else 2
    return len(hits) >= required


def _ksa_evidence_relevance_ok(
    question: str,
    follow_ups: list[str],
    evaluation_points: list[str],
    q: dict[str, Any],
    matching_ksa_evidence: list[dict[str, Any]],
) -> bool:
    if not matching_ksa_evidence:
        return False
    placeholder_text = "\n".join(
        [
            str(question or ""),
            *[str(x or "") for x in follow_ups],
            *[str(x or "") for x in evaluation_points],
        ]
    )
    if _UNRESOLVED_KSA_PLACEHOLDER_RE.search(placeholder_text):
        return False
    visible_prompt_text = "\n".join(
        [
            str(question or ""),
            *[str(x or "") for x in follow_ups],
        ]
    )
    evidence_id = str(q.get("question_evidence_id") or "").strip()
    surface_focus = str(q.get("question_focus_surface") or "").strip()
    if evidence_id:
        linked_rows = [
            row
            for row in matching_ksa_evidence
            if isinstance(row, dict)
            and str(row.get("evidence_id") or stable_ksa_evidence_id(row)).strip()
            == evidence_id
        ]
        if not linked_rows:
            return False
        semantic_targets = [
            surface_focus,
            *[str(row.get("factorName") or "").strip() for row in linked_rows],
        ]
        return any(
            target and _ksa_factor_relevant_to_text(target, visible_prompt_text)
            for target in semantic_targets
        )
    return any(
        _ksa_factor_relevant_to_text(str(row.get("factorName") or ""), visible_prompt_text)
        for row in matching_ksa_evidence
        if isinstance(row, dict)
    )


def _candidate_surface_ksa_safe(
    q: dict[str, Any],
    question: str,
    follow_ups: list[str],
    evaluation_points: list[str],
) -> bool:
    """Keep official evidence labels out of candidate-facing question copy."""

    visible = "\n".join([str(question or ""), *follow_ups, *evaluation_points])
    public_surface = str(q.get("question_focus_surface") or "").strip()
    if public_surface:
        # The validated public task object may deliberately retain meaningful
        # words from the source factor. Remove that complete surface before
        # checking whether an un-repaired official label leaked elsewhere.
        visible = visible.replace(public_surface, "")
    compact_visible = re.sub(r"\s+", "", visible).casefold()
    evidence_rows = q.get("ksa_evidence") if isinstance(q.get("ksa_evidence"), list) else []
    factors = [
        str(row.get("factorName") or "").strip()
        for row in evidence_rows
        if isinstance(row, dict) and str(row.get("factorName") or "").strip()
    ]
    if not factors and isinstance(q.get("ksa_refs"), list):
        factors = [str(value or "").strip() for value in q.get("ksa_refs") or [] if str(value or "").strip()]
    for factor in factors:
        for alias in official_ksa_surface_aliases(factor):
            compact_alias = re.sub(r"\s+", "", alias).casefold()
            if compact_alias and compact_alias in compact_visible:
                return False
    return True


def _natural_question_wording_ok(q: dict[str, Any], question: str, follow_ups: list[str]) -> bool:
    main = str(question or "").strip()
    method = str((q or {}).get("type") or (q or {}).get("method") or "").strip()
    max_main_chars = {
        "경험면접": 360,
        "상황면접": 420,
        "발표면접": 440,
        "토론면접": 480,
        "인바스켓면접": 380,
        "직무지식면접": 360,
        "창의적 문제해결력면접": 460,
    }.get(method, 400)
    if not main or len(main) > max_main_chars:
        return False
    if re.search(r" {2,}", main):
        return False
    if _collapse_repeated_question_phrases(main)[1]:
        return False
    if re.search(r"(?:설명하고|제시하고|말씀하고)\s+또한\b", main):
        return False
    # Do not let a provider fall back to the stock, label-free experience
    # opening.  It is technically grammatical but hides the uploaded duty and
    # encourages the rest of the sentence to become a checklist of fixed STAR
    # clauses (the exact failure mode seen in production).
    if re.search(
        r"^해당\s*직무(?:에서|중에?)\s*업무를\s*수행하던\s*실제\s*상황(?:\s*하나)?를\s*골라",
        main,
    ):
        return False
    scaffold_markers = (
        "실제 상황 하나를 골라",
        "요구사항과 기준",
        "당시 맡은 역할",
        "직접 맡은 범위",
        "문서·수치·기록·피드백",
        "이후 무엇을 개선",
    )
    scaffold_hits = sum(marker in main for marker in scaffold_markers)
    if method in {"", "경험면접"} and scaffold_hits >= 3 and len(main) >= 220:
        return False
    # A purpose-clause fragment must never be exposed as a KSA surface. It is
    # the characteristic failure mode of rows such as
    # ``...도와주고자 하는 태도`` when the source label is stripped.
    if re.search(
        r"(?:하고자|고자)\s*(?:관련\s*)?(?:행동\s*기준|확인·판단\s*기준|수행·검증\s*절차)",
        "\n".join([main, *[str(item or "") for item in follow_ups]]),
    ):
        return False
    if re.search(r"(?:고|며|지만|는데|면서)(?:이|가)\s*발생", main):
        return False
    if re.search(r"상황(?:이|가)\s*(?:동시에\s*)?발생한\s*상황", main):
        return False
    if re.search(
        r"발생하고[^.]{0,180}(?:조건이고|조건이며)[^.]{0,180}조건인\s*상황에서",
        main,
    ):
        return False
    visible = "\n".join([main, *[str(item or "") for item in follow_ups]])
    if re.search(r"(?:능력|기술|스킬)\s*(?:을|를)\s*(?:직접\s*)?수행", visible):
        return False
    if re.search(
        r"(?:기법|방법|도구|장비|시스템|프로그램|소프트웨어)\s*"
        r"['\"”’]?\s*(?:을|를)\s*(?:(?:직접|실제로)\s*)?수행",
        visible,
    ):
        return False
    if re.search(r"지식\s*(?:을|를)\s*판단\s*근거", visible):
        return False
    focus = str(q.get("question_focus") or "").strip()
    focus_surface = str(q.get("question_focus_surface") or "").strip()
    # A public task surface is an internal bridge, not prose to paste into a
    # sentence as ``... 기준에 따라``. Reject that mechanical attachment so
    # the provider candidate is regenerated or replaced by the concise server
    # template instead of being shown as a finished question.
    if focus_surface:
        escaped_surface = re.escape(focus_surface)
        source = str((q or {}).get("question_source") or "").strip()
        mechanical_surface_check = source not in {
            "template_fallback",
            "rule_fallback",
            "quality_orchestrator_repair",
            "simulation_candidate",
            "deterministic",
            "deterministic_template",
        }
        if method == "경험면접" and mechanical_surface_check and re.search(
            rf"{escaped_surface}\s*(?:에\s*따라|관련\s*(?:행동|확인·판단)\s*기준)",
            main,
        ):
            return False
    # A legitimate NCS factor can itself contain the word "관련" (for example,
    # "도로교통 관련 법규").  Repeating that exact factor is evidence
    # traceability, not conjunction stuffing, so only count scaffold wording.
    scaffold_visible = visible
    for protected_focus in sorted(
        {value for value in (focus, focus_surface) if value},
        key=len,
        reverse=True,
    ):
        scaffold_visible = scaffold_visible.replace(protected_focus, "")
    if scaffold_visible.count("관련") >= 5:
        return False
    if re.search(r"\([^)]{0,100}(?:하는|위한|통해|하여)\)\.?$", main):
        return False

    focus_type = str(q.get("question_focus_type") or "").strip() or _normalize_ksa_type("", focus)
    if focus and focus_type == "태도":
        compact_visible = re.sub(r"\s+", "", visible)
        compact_focus = re.sub(r"\s+", "", focus)
        if re.search(re.escape(compact_focus) + r".{0,12}적용", compact_visible):
            return False
    return True


def _focus_scenario_coherence_ok(method: str, q: dict[str, Any], question: str) -> bool:
    """Require discussion topics to drive the conflict, not appear as a trailing label only."""
    if method != "토론면접":
        return True
    focus = str(q.get("question_focus_surface") or q.get("question_focus") or "").strip()
    if not focus:
        return False
    # Anything after this phrase is an output constraint, not the scenario itself.
    scenario_text = re.split(r"(?:합의 기준에는|최종 합의안에는)", str(question or ""), maxsplit=1)[0]
    compact_scenario = re.sub(r"\s+", "", scenario_text).lower()
    compact_focus = re.sub(r"\s+", "", focus).lower()
    if compact_focus and compact_focus in compact_scenario:
        return True
    stopwords = {
        "관련", "능력", "기술", "지식", "태도", "활용", "적용", "업무", "수행",
        "기준", "대한", "위한", "이해", "및", "과", "와", "의", "을", "를",
    }
    focus_tokens = [
        token
        for token in re.findall(r"[0-9A-Za-z가-힣]{2,}", focus)
        if token not in stopwords
    ]
    if not focus_tokens:
        return True
    return any(token.replace(" ", "").lower() in compact_scenario for token in focus_tokens)


_DILEMMA_CONSEQUENCE_DIMENSIONS: tuple[tuple[str, ...], ...] = (
    (
        "승인",
        "규정",
        "절차",
        "근거",
        "책임",
        "보안",
        "안전",
        "품질",
        "정확",
        "오류",
        "위험",
        "검증",
    ),
    ("일정", "납기", "마감", "지연", "속도", "신속", "SLA", "연속성", "가동률", "회전율"),
    ("비용", "예산", "정산", "인력", "자원", "효율", "수익"),
    (
        "민원",
        "불편",
        "편의",
        "수요",
        "요청",
        "협업",
        "자기결정",
        "수용성",
        "이용자",
        "고객",
        "만족",
        "참여",
    ),
)
_DILEMMA_DECISION_ACTIONS = (
    "금지", "허용", "보류", "착수", "집행", "처리", "강화", "완화", "우선",
    "준수", "공유", "배분", "배정", "유지", "확보", "준비", "적용", "예외", "조건부", "중단", "검증",
    "확정", "합의", "조정", "선택", "결정",
)


def _decision_dilemma_quality_ok(method: str, question: str) -> bool:
    """Require a debate to contain a real choice with at least two consequence axes."""
    if method != "토론면접":
        return True
    text = str(question or "").strip()
    compact = re.sub(r"\s+", "", text)
    opposing_options = compact.count("입장") >= 2 or all(marker in compact for marker in ("A안", "B안"))
    decision_action = any(marker in compact for marker in _DILEMMA_DECISION_ACTIONS)
    consequence_dimensions = sum(
        1 for markers in _DILEMMA_CONSEQUENCE_DIMENSIONS if any(marker in compact for marker in markers)
    )
    has_event = any(
        marker in compact
        for marker in (
            "발생",
            "요청",
            "요구",
            "대기",
            "불일치",
            "누락",
            "재발",
            "임박",
            "변경",
            "부족",
            "지연",
            "낮",
            "하락",
            "달성",
        )
    )
    return bool(opposing_options and decision_action and consequence_dimensions >= 2 and has_event)


def _debate_outcome_flexibility_ok(method: str, question: str) -> bool:
    """Do not reward consensus at any cost when escalation is the safer outcome."""
    if method != "토론면접":
        return True
    compact = re.sub(r"\s+", "", str(question or ""))
    permits_non_consensus = any(
        marker in compact
        for marker in ("합의가어렵", "미합의", "결정권자", "이송", "상신", "남은쟁점")
    )
    forced_consensus_only = any(
        marker in compact
        for marker in ("반드시합의", "공동합의안을도출", "최종합의안을도출")
    ) and not permits_non_consensus
    return bool(permits_non_consensus and not forced_consensus_only)


def _debate_option_defensibility_ok(method: str, question: str) -> bool:
    """Reject an apparent shortcut that asks candidates to defend unauthorized work."""
    if method != "토론면접":
        return True
    text = str(question or "")
    compact = re.sub(r"\s+", "", text)
    if "승인" not in compact:
        return True
    option_match = re.search(
        r"([^.!?]{0,500}입장과[^.!?]{0,500}입장)(?:이|가)?\s*충돌",
        text,
    )
    option_text = re.sub(r"\s+", "", option_match.group(1) if option_match else text)
    shortcut_markers = ("사후승인", "선착수", "먼저집행", "조건부집행", "승인전에착수")
    if not any(marker in option_text for marker in shortcut_markers):
        return True
    # An explicit rule allowing emergency work can make both positions
    # defensible. Merely calling the work low-risk is not such a rule.
    return any(
        marker in option_text
        for marker in ("사전착수허용조항", "긴급착수허용", "승인전허용", "예외승인근거")
    )


def _official_sample_format_ok(
    method: str,
    question: str,
    follow_ups: list[str],
    evaluation_points: list[str],
) -> bool:
    task_text = re.sub(r"\s+", "", "\n".join([str(question or ""), *follow_ups]))
    eval_text = re.sub(r"\s+", "", "\n".join(evaluation_points))
    rule = _OFFICIAL_SAMPLE_FORMAT_RULES.get(method)
    if not rule:
        return False
    base_format_ok = (
        any(term in task_text for term in rule["task_any"])
        and all(term in task_text for term in rule["task_all"])
        and any(term in eval_text for term in rule["eval_any"])
    )
    if method != "경험면접":
        return base_format_ok
    # STAR의 T는 단순히 "본인"이라는 말만 있는 것으로 충분하지 않다.
    # 당시 맡은 과제·역할·목표·책임 중 하나를 실제 답변에서 끌어내야 한다.
    task_role_visible = any(
        term in task_text for term in ("역할", "과제", "목표", "책임", "담당")
    )
    return bool(base_format_ok and task_role_visible)


def _quality_check_statuses(
    method: str,
    question: dict[str, Any],
    checks: dict[str, bool],
) -> dict[str, str]:
    """Distinguish a passed gate from a gate that does not apply to the method."""

    debate_only = {
        "focus_scenario_coherence",
        "decision_dilemma_quality",
        "debate_option_defensibility",
        "debate_outcome_flexibility",
        "debate_case_neutrality",
    }
    material_only = {"case_materials_sufficient"}
    task_methods = {"상황면접", "발표면접", "토론면접", "인바스켓면접", "창의적 문제해결력면접"}
    authority_methods = {"상황면접", "토론면접", "인바스켓면접", "창의적 문제해결력면접"}
    source = str(question.get("question_source") or "").strip()
    subscription_cli_semantic_replacements = {
        "method_shape",
        "main_question_method_shape",
        "main_question_job_context",
        "follow_up_quality",
        "evaluation_points_quality",
        "ksa_grounded",
        "official_sample_format",
        "operating_conditions_separated",
        "job_specific_context",
    }
    strict_realism_sources = {
        "template_fallback",
        "rule_fallback",
        "simulation_candidate",
        "quality_orchestrator_repair",
        "deterministic",
        "deterministic_template",
    }
    statuses: dict[str, str] = {}
    for name, passed in checks.items():
        applicable = True
        if (
            _is_subscription_cli_source(source)
            and name in subscription_cli_semantic_replacements
        ):
            # These legacy gates infer quality from exact template keywords.
            # Subscription CLI drafts use an exact evidence id plus the independent field-
            # realism gate, so requiring those words would perversely reward
            # the canned questions we are replacing.
            applicable = False
        elif name in debate_only:
            applicable = method == "토론면접"
        elif name in material_only:
            applicable = method in task_methods
        elif name == "decision_authority_context":
            applicable = method in authority_methods
        elif name == "inbasket_authority_context":
            applicable = method == "인바스켓면접"
        elif name == "operating_conditions_separated":
            applicable = method in task_methods
        elif name == "main_question_job_context":
            applicable = bool(question.get("model_question_preserved"))
        elif name == "evidence_linked":
            applicable = bool(question.get("question_evidence_required"))
        elif name == "field_realism":
            # Enforce the panel-readiness gate for subscription CLI paths and for
            # wholly deterministic fallbacks.  Legacy OpenAI/model sources keep
            # the realism diagnostics in the report during migration without
            # changing their long-standing readiness contract.
            applicable = _is_subscription_cli_source(source) or source in strict_realism_sources
        statuses[name] = "pass" if applicable and passed else "fail" if applicable else "not_applicable"
    return statuses


def _attach_question_quality_report(strategy: dict[str, Any]) -> dict[str, Any]:
    if not isinstance(strategy, dict):
        strategy = {}
    questions = strategy.get("interview_questions")
    if not isinstance(questions, list):
        strategy["question_quality_report"] = {
            "policy": QUALITY_POLICY_VERSION,
            "precision_grounding_policy": PRECISION_GROUNDING_POLICY,
            "passed": False,
            "summary": {
                "question_count": 0,
                "expected_question_count": 0,
                "count_matches_plan": False,
                "average_score": 0.0,
                "ready_count": 0,
                "needs_review_count": 0,
                "precision_grounding_failed_count": 0,
            },
            "items": [],
        }
        return strategy

    plan = strategy.get("question_plan_used") if isinstance(strategy.get("question_plan_used"), dict) else {}
    try:
        expected_count = int(plan.get("total_main_count") or 0)
    except Exception:
        expected_count = 0

    items: list[dict[str, Any]] = []
    ready_count = 0
    seen_quality_question_keys: set[str] = set()
    seen_quality_repeat_items_by_signature: dict[str, list[dict[str, Any]]] = {}
    for idx, raw in enumerate(questions, start=1):
        q = raw if isinstance(raw, dict) else {}
        method = str(q.get("type") or q.get("method") or "").strip()
        question = str(q.get("question") or "").strip()
        follow_ups = [str(x).strip() for x in (q.get("follow_ups") or []) if str(x).strip()] if isinstance(q.get("follow_ups"), list) else []
        evaluation_points = [
            str(x).strip()
            for x in (q.get("evaluation_points") or [])
            if str(x).strip()
        ] if isinstance(q.get("evaluation_points"), list) else []
        if not q.get("assessment_guide"):
            q["assessment_guide"] = _behavior_anchored_evaluation(
                method,
                str(q.get("question_focus") or "").strip(),
                evaluation_points,
                str(q.get("question_focus_type") or "").strip(),
                surface_focus=str(q.get("question_focus_surface") or "").strip(),
            )
        if not q.get("task_conditions"):
            q["task_conditions"] = _task_conditions_for_method(
                method=method,
                subject=str(q.get("competency") or "").strip(),
                focus=str(q.get("question_focus") or "").strip(),
                detail=str(q.get("ncs_detail") or q.get("ncsSclasCdnm") or "").strip(),
                comp_def=str(q.get("compeUnitDef") or "").strip(),
                focus_type=str(q.get("question_focus_type") or "").strip(),
                variation_index=_clamp_int(q.get("question_variation_index"), 0, 0, 1_000_000),
            )
        ksa_refs = [str(x).strip() for x in (q.get("ksa_refs") or []) if str(x).strip()] if isinstance(q.get("ksa_refs"), list) else []
        ksa_evidence = q.get("ksa_evidence") if isinstance(q.get("ksa_evidence"), list) else []
        ncs_code = str(q.get("ncsClCd") or "").strip()
        matching_ksa_evidence = [
            row
            for row in ksa_evidence
            if isinstance(row, dict) and str(row.get("ncsClCd") or "").strip() == ncs_code
        ] if ncs_code else []
        merged = "\n".join([question, *follow_ups, *evaluation_points, *ksa_refs])
        q_key = normalize_question_dedup_key(question)
        repeat_signature = _question_repeat_signature(q)
        surface_duplicate = bool(q_key and q_key in seen_quality_question_keys)
        previous_repeat_items = seen_quality_repeat_items_by_signature.get(repeat_signature, []) if repeat_signature else []
        repeat_duplicate = any(_question_near_repeat(q, previous) for previous in previous_repeat_items)
        if repeat_duplicate and _raw_model_scenarios_are_distinct(q, previous_repeat_items):
            repeat_duplicate = False
        has_specific_context = not any(marker in question for marker in ("해당 직무", "핵심 수행기준"))
        realism = evaluate_question_realism(q)
        precision_grounding = evaluate_question_precision_grounding(q)
        evaluation_alignment = evaluate_evaluation_elicitation_alignment(q)
        checks = {
            "supported_method": method in QUALITY_INTERVIEW_METHODS,
            "method_shape": _method_shape_ok(method, merged),
            "main_question_method_shape": _method_shape_ok(method, question),
            "main_question_job_context": _main_question_job_context_ok(q, question),
            "follow_up_depth": len(follow_ups) >= 3,
            "follow_up_quality": _follow_ups_quality_ok(method, q, follow_ups),
            "evaluation_points": len(evaluation_points) == 4,
            "evaluation_points_quality": _evaluation_points_quality_ok(method, evaluation_points),
            "ncs_grounded": bool(ncs_code and str(q.get("competency") or "").strip()),
            "detail_grounded": bool(str(q.get("ncs_detail") or q.get("ncsSclasCdnm") or "").strip()),
            "ksa_grounded": _ksa_evidence_relevance_ok(question, follow_ups, evaluation_points, q, matching_ksa_evidence),
            "evidence_linked": bool(str(q.get("question_evidence_id") or "").strip())
            if q.get("question_evidence_required")
            else True,
            "ksa_measurement_task": bool(evaluate_ksa_measurement(q).get("passed")),
            "candidate_surface_safe": _candidate_surface_ksa_safe(q, question, follow_ups, evaluation_points),
            "official_sample_format": _official_sample_format_ok(method, question, follow_ups, evaluation_points),
            "blind_hiring_safe": not _contains_blind_hiring_cue(merged),
            "unique_question": bool(
                q_key
                and not surface_duplicate
                and not repeat_duplicate
            ),
            "specific_context": has_specific_context,
            "job_specific_context": _job_specific_context_ok(q, question, follow_ups),
            "natural_wording": _natural_question_wording_ok(q, question, follow_ups),
            "focus_scenario_coherence": _focus_scenario_coherence_ok(method, q, question),
            "decision_dilemma_quality": _decision_dilemma_quality_ok(method, question),
            "debate_option_defensibility": _debate_option_defensibility_ok(method, question),
            "debate_outcome_flexibility": _debate_outcome_flexibility_ok(method, question),
            "debate_case_neutrality": _debate_case_neutrality_ok(method, question, q.get("task_conditions")),
            "operating_conditions_separated": _operating_conditions_separated(method, question),
            "standardized_task_conditions": _task_conditions_ok(method, q.get("task_conditions")),
            "case_materials_sufficient": _case_materials_sufficient_ok(method, q.get("task_conditions")),
            "decision_authority_context": _decision_authority_context_ok(method, q.get("task_conditions")),
            "inbasket_authority_context": _inbasket_authority_context_ok(method, q.get("task_conditions")),
            "behavior_anchored_evaluation": _behavior_anchors_ok(q.get("assessment_guide")),
            "field_realism": bool(realism.get("passed")),
            "precision_grounding": bool(precision_grounding.get("passed")),
        }
        check_statuses = _quality_check_statuses(method, q, checks)
        applicable_statuses = [status for status in check_statuses.values() if status != "not_applicable"]
        issues = [name for name, status in check_statuses.items() if status == "fail"]
        score = round(applicable_statuses.count("pass") / max(1, len(applicable_statuses)), 2)
        ready = not issues
        if ready:
            ready_count += 1
        items.append(
            {
                "index": idx,
                "type": method,
                "competency": str(q.get("competency") or "").strip(),
                "ncsClCd": str(q.get("ncsClCd") or "").strip(),
                "ncs_detail": str(q.get("ncs_detail") or q.get("ncsSclasCdnm") or "").strip(),
                "question_focus": str(q.get("question_focus") or "").strip(),
                "question_focus_surface": str(q.get("question_focus_surface") or "").strip(),
                "question_evidence_id": str(q.get("question_evidence_id") or "").strip(),
                "question_focus_type": str(q.get("question_focus_type") or "").strip(),
                "assessment_scale": str((q.get("assessment_guide") or {}).get("scale") or "").strip()
                if isinstance(q.get("assessment_guide"), dict)
                else "",
                "question_intent": _question_intent_key(question),
                "question_repeat_signature": repeat_signature,
                "question_repeat_duplicate": repeat_duplicate,
                "question_source": str(q.get("question_source") or "").strip(),
                "realism_policy": REALISM_POLICY_VERSION,
                "realism_score": int(realism.get("score") or 0),
                "realism_issue_codes": list(realism.get("issue_codes") or []),
                "realism_issues": list(realism.get("issues") or []),
                "realism_checks": dict(realism.get("checks") or {}),
                "precision_grounding_policy": PRECISION_GROUNDING_POLICY,
                "precision_grounding_issue_codes": list(
                    precision_grounding.get("issue_codes") or []
                ),
                "precision_grounding_issues": list(
                    precision_grounding.get("issues") or []
                ),
                "precision_grounding_demands": list(
                    precision_grounding.get("demands") or []
                ),
                "precision_grounding_metrics": dict(
                    precision_grounding.get("metrics") or {}
                ),
                # v1 remains shadow-only while its bounded semantic lattice is
                # calibrated on additional real institutions.  Expose its
                # fail/review evidence now; do not silently convert it into a
                # readiness veto until the evaluation corpus is broad enough.
                "evaluation_alignment_policy": EVALUATION_ELICITATION_POLICY,
                "evaluation_alignment_decision": str(
                    evaluation_alignment.get("decision") or ""
                ),
                "evaluation_alignment_checks": dict(
                    evaluation_alignment.get("checks") or {}
                ),
                "evaluation_alignment_issues": list(
                    evaluation_alignment.get("issues") or []
                ),
                "evaluation_alignment_metrics": dict(
                    evaluation_alignment.get("metrics") or {}
                ),
                "score": score,
                "ready": ready,
                "checks": checks,
                "check_statuses": check_statuses,
                "issues": issues,
            }
        )
        if q_key:
            seen_quality_question_keys.add(q_key)
        if repeat_signature:
            seen_quality_repeat_items_by_signature.setdefault(repeat_signature, []).append(dict(q))

    avg = round(sum(float(item.get("score") or 0.0) for item in items) / max(1, len(items)), 2)
    count_matches_plan = expected_count <= 0 or len(items) == expected_count
    passed = bool(count_matches_plan and items and ready_count == len(items))
    strategy["question_quality_report"] = {
        "policy": QUALITY_POLICY_VERSION,
        "precision_grounding_policy": PRECISION_GROUNDING_POLICY,
        "evaluation_alignment_policy": EVALUATION_ELICITATION_POLICY,
        "evaluation_alignment_enforcement": "shadow",
        "passed": passed,
        "summary": {
            "question_count": len(items),
            "expected_question_count": expected_count,
            "count_matches_plan": count_matches_plan,
            "average_score": avg,
            "ready_count": ready_count,
            "needs_review_count": len(items) - ready_count,
            "precision_grounding_failed_count": sum(
                not bool((item.get("checks") or {}).get("precision_grounding"))
                for item in items
            ),
            "evaluation_alignment_pass_count": sum(
                item.get("evaluation_alignment_decision") == "pass"
                for item in items
            ),
            "evaluation_alignment_review_count": sum(
                item.get("evaluation_alignment_decision") == "review"
                for item in items
            ),
            "evaluation_alignment_fail_count": sum(
                item.get("evaluation_alignment_decision") == "fail"
                for item in items
            ),
        },
        "items": items,
    }
    return strategy


def _attach_ksa_evidence_to_strategy(strategy: dict[str, Any], ncs_ksa: list[dict[str, Any]] | None) -> dict[str, Any]:
    if not isinstance(strategy, dict):
        strategy = {}
    questions = strategy.get("interview_questions")
    if not isinstance(questions, list):
        return strategy

    evidence_rows: list[dict[str, str]] = []
    seen_rows: set[str] = set()
    for raw in ncs_ksa or []:
        if not isinstance(raw, dict):
            continue
        row = _clean_ksa_evidence_row(raw)
        if not row["factorName"]:
            continue
        key = row["evidence_id"]
        if key in seen_rows:
            continue
        seen_rows.add(key)
        evidence_rows.append(row)
    if not evidence_rows:
        return _attach_question_quality_report(strategy)

    def _pick_for_question(question: dict[str, Any]) -> list[dict[str, str]]:
        code = str(question.get("ncsClCd", "")).strip()
        refs = [str(x).strip() for x in (question.get("ksa_refs") or []) if str(x).strip()] if isinstance(question.get("ksa_refs"), list) else []
        focus_ref = str(question.get("question_focus") or "").strip()
        if focus_ref:
            refs = [focus_ref, *[ref for ref in refs if _ksa_key(ref) != _ksa_key(focus_ref)]]
        ref_keys = [_ksa_key(x) for x in refs]
        picked: list[dict[str, str]] = []
        picked_keys: set[tuple[str, str]] = set()

        def add(row: dict[str, str]) -> None:
            if len(picked) >= 4:
                return
            # Two official KSA rows may legitimately share the same visible
            # factor label while belonging to different elements/numbers.
            # Evidence identity—not display text—is the audit boundary.
            key = (row["ncsClCd"], row["evidence_id"])
            if key in picked_keys:
                return
            picked_keys.add(key)
            picked.append(row)

        preferred = [row for row in evidence_rows if code and row["ncsClCd"] == code]
        if not code or not preferred:
            return []
        fallback = preferred

        preferred_id = str(question.get("question_evidence_id") or "").strip()
        if _is_subscription_cli_source(question.get("question_source")):
            # Model-backed public candidates must preserve the exact evidence
            # assignment emitted for their prompt index.  Inferring the first
            # convenient row here would turn a missing, invented, or swapped ID
            # into apparently valid audit evidence.
            if not preferred_id:
                return []
            return [
                row
                for row in fallback
                if row["evidence_id"] == preferred_id
            ][:1]
        if preferred_id:
            for row in fallback:
                if row["evidence_id"] == preferred_id:
                    add(row)
                    break

        for ref_key in ref_keys:
            if not ref_key:
                continue
            for row in fallback:
                factor_key = _ksa_key(row["factorName"])
                if ref_key in factor_key or factor_key in ref_key:
                    add(row)
            if len(picked) >= 2:
                break
        for row in fallback:
            add(row)
            if len(picked) >= 3:
                break
        return picked[:3]

    enriched: list[dict[str, Any]] = []
    for item in questions:
        if not isinstance(item, dict):
            enriched.append(item)
            continue
        q = dict(item)
        evidence = _pick_for_question(q)
        if evidence:
            existing_refs = [
                str(x).strip()
                for x in (q.get("ksa_refs") or [])
                if str(x).strip()
            ] if isinstance(q.get("ksa_refs"), list) else []
            for row in evidence:
                factor = row.get("factorName", "")
                if factor and factor not in existing_refs:
                    existing_refs.append(factor)
            q["ksa_refs"] = existing_refs[:4]
            q["ksa_evidence"] = evidence
            q["evidence_ids"] = [row["evidence_id"] for row in evidence if row.get("evidence_id")]
            preferred_id = str(q.get("question_evidence_id") or "").strip()
            if preferred_id not in q["evidence_ids"]:
                preferred_id = q["evidence_ids"][0] if q["evidence_ids"] else ""
            q["question_evidence_id"] = preferred_id
            q["question_evidence_required"] = True
            primary_evidence = next(
                (
                    row
                    for row in evidence
                    if str(row.get("evidence_id") or "").strip() == preferred_id
                ),
                evidence[0],
            )
            q["ncs_traceability"] = _ncs_traceability_payload(q, primary_evidence)
        enriched.append(q)
    strategy["interview_questions"] = enriched
    strategy["question_evidence_policy"] = "ncs_mcp_ksa_attached_by_evidence_id"
    return _attach_question_quality_report(strategy)


def _run_runtime_question_quality_orchestration(
    strategy: dict[str, Any],
    *,
    question_plan: dict[str, Any],
    ncs_ksa: list[dict[str, Any]] | None,
    avoid_questions: list[str] | None = None,
    generation_offset: int | None = None,
    job_context_text: str = "",
) -> dict[str, Any]:
    """Recheck generated questions and repair shallow or repeated tasks.

    The model pass and the deterministic quality pass intentionally remain
    separate.  A repair failure is captured per question by the service layer
    and does not fail the whole generation request.
    """

    if not isinstance(strategy, dict):
        strategy = {}
    avoid_questions = [
        cleaned
        for cleaned in (_strip_question_material_reference(value) for value in (avoid_questions or []))
        if cleaned
    ]
    questions = [
        dict(item)
        for item in (strategy.get("interview_questions") or [])
        if isinstance(item, dict)
    ]
    if not questions:
        strategy["question_quality_orchestration"] = {
            "policy": RUNTIME_QUESTION_ORCHESTRATION_POLICY,
            "status": "needs_review",
            "question_count": 0,
            "initial_failure_count": 1,
            "repaired_count": 0,
            "repair_error_count": 0,
            "unresolved_count": 1,
            "full_quality_unresolved_count": 1,
            "items": [],
            "stages": [
                {"name": "candidate_generation", "status": "failed", "question_count": 0},
            ],
        }
        return _attach_ksa_evidence_to_strategy(strategy, ncs_ksa)

    # Normalize evidence/public wording before the lightweight orchestration
    # gate.  Without this bridge, a valid public task can be rejected because
    # the gate still sees only the internal factor label, then needlessly
    # replaces a model-authored question with a full template.
    normalized_questions: list[dict[str, Any]] = []
    for raw_item in questions:
        item = dict(raw_item)
        provider_authored = _is_subscription_cli_source(
            item.get("question_source")
        )
        method = str(item.get("type") or item.get("method") or "경험면접").strip()
        ncs_code = str(item.get("ncsClCd") or "").strip()
        provided_focus_evidence = _evidence_row_for_id(
            ncs_ksa,
            ncs_code,
            item.get("question_evidence_id"),
        )
        focus = _clean_question_text(
            provided_focus_evidence.get("factorName")
            or item.get("question_focus")
            or ((item.get("ksa_refs") or [""])[0] if isinstance(item.get("ksa_refs"), list) else ""),
            max_chars=60,
        )
        focus_type = _ksa_type_for_focus(ncs_ksa, ncs_code, focus)
        subject = str(item.get("competency") or item.get("compeUnitName") or "").strip()
        detail = str(
            item.get("ncs_detail")
            or item.get("ncsSubdCdnm")
            or item.get("ncsSclasCdnm")
            or ""
        ).strip()
        focus_evidence = provided_focus_evidence or _evidence_row_for_focus(ncs_ksa, ncs_code, focus)
        task_frame = _question_task_frame(
            focus=focus or "핵심 수행기준",
            focus_type=focus_type,
            subject=subject,
            detail=detail,
            comp_def=str(item.get("compeUnitDef") or "").strip(),
            evidence_row=focus_evidence,
        )
        item["question_focus"] = focus or str(item.get("question_focus") or "").strip()
        item["question_focus_type"] = focus_type
        item["question_focus_surface"] = task_frame["task_object"]
        item["question_task_frame"] = task_frame
        item["question_evidence_id"] = task_frame.get("evidence_id", "")
        item["question_evidence_required"] = bool(focus_evidence)

        surface_repairs: set[str] = set()
        cleaned_question = _sanitize_candidate_document_leaks(
            item.get("question"),
            subject=subject,
        )
        if provider_authored:
            repaired_question, repaired = cleaned_question, False
        else:
            repaired_question, repaired = _repair_candidate_surface_text(
                cleaned_question,
                focus,
                task_frame["task_object"],
            )
        item["question"] = repaired_question
        if repaired:
            surface_repairs.add("question")
        repaired_followups: list[str] = []
        for value in item.get("follow_ups") or []:
            cleaned_value = _sanitize_candidate_document_leaks(value, subject=subject)
            if provider_authored:
                repaired_value, repaired = cleaned_value, False
            else:
                repaired_value, repaired = _repair_candidate_surface_text(
                    cleaned_value,
                    focus,
                    task_frame["task_object"],
                )
            repaired_followups.append(repaired_value)
            if repaired:
                surface_repairs.add("follow_ups")
        if repaired_followups:
            item["follow_ups"] = repaired_followups
            item["follow_up"] = repaired_followups[0]
        repaired_points: list[str] = []
        for value in item.get("evaluation_points") or []:
            cleaned_value = _sanitize_candidate_document_leaks(value, subject=subject)
            if provider_authored:
                repaired_value, repaired = cleaned_value, False
            else:
                repaired_value, repaired = _repair_candidate_surface_text(
                    cleaned_value,
                    focus,
                    task_frame["task_object"],
                )
            repaired_points.append(repaired_value)
            if repaired:
                surface_repairs.add("evaluation_points")
        if repaired_points:
            item["evaluation_points"] = repaired_points

        quality_repairs: set[str] = set()
        if not _task_conditions_ok(method, item.get("task_conditions")):
            item["task_conditions"] = _task_conditions_for_method(
                method=method,
                subject=subject,
                focus=focus,
                detail=detail,
                comp_def=str(item.get("compeUnitDef") or "").strip(),
                focus_type=focus_type,
                variation_index=_clamp_int(item.get("question_variation_index"), 0, 0, 1_000_000),
            )
            quality_repairs.add("task_conditions")
        if not _behavior_anchors_ok(item.get("assessment_guide")):
            evaluation_points = [
                str(value or "").strip()
                for value in (item.get("evaluation_points") or [])
                if str(value or "").strip()
            ]
            item["assessment_guide"] = _behavior_anchored_evaluation(
                method,
                focus,
                evaluation_points,
                focus_type,
                surface_focus=task_frame["task_object"],
            )
            quality_repairs.add("assessment_guide")

        if surface_repairs:
            item["candidate_surface_repairs"] = sorted(
                set(item.get("candidate_surface_repairs") or []) | surface_repairs
            )
        if quality_repairs:
            item["quality_repaired_fields"] = sorted(
                set(item.get("quality_repaired_fields") or []) | quality_repairs
            )
        source = str(item.get("question_source") or "").strip()
        if (surface_repairs or quality_repairs) and source in MODEL_PRESERVED_QUESTION_SOURCES:
            if source == "model":
                item["question_source"] = "model_main_quality_repaired_fields"
            item["model_question_preserved"] = True
        normalized_questions.append(item)
    questions = normalized_questions

    default_follow_count = max(
        3,
        min(5, int((question_plan or {}).get("follow_up_count", 3) or 3)),
    )
    history_texts = [str(text or "").strip() for text in (avoid_questions or []) if str(text or "").strip()]
    history_size = len(history_texts)
    # A bounded rolling history eventually stays at a constant size.  Seeding
    # only from its length would make every later repair retry the same small
    # variation range.  Mix recent content into the start offset so the window
    # keeps moving without retaining server-side user history.
    history_digest_input = "\n".join(
        normalize_question_dedup_key(text) for text in history_texts[-40:]
    )
    history_variation_seed = (
        int(hashlib.sha256(history_digest_input.encode("utf-8")).hexdigest()[:12], 16)
        if history_digest_input
        else 0
    )
    if generation_offset is None:
        normalized_generation_offset: int | None = None
    else:
        try:
            normalized_generation_offset = max(0, min(1_000_000, int(generation_offset)))
        except (TypeError, ValueError):
            normalized_generation_offset = 0

    def repair_question(
        original: dict[str, Any],
        index: int,
        reasons: list[str],
        attempt: int,
    ) -> dict[str, Any] | None:
        item = dict(original)
        method = str(item.get("type") or item.get("method") or "경험면접").strip()
        if method not in QUALITY_INTERVIEW_METHODS:
            method = "경험면접"
        source_is_preserved_provider = _is_subscription_cli_source(
            original.get("question_source")
        )
        if source_is_preserved_provider:
            # Provider-authored wording is validate-only at this boundary.
            # Duplicates and quality failures remain unresolved so the outer
            # provider retry can ask the AI for a genuinely different question;
            # a deterministic rewrite would reintroduce the canned STAR text.
            return None
        ncs_code = str(item.get("ncsClCd") or "").strip()
        current_focus = _clean_question_text(item.get("question_focus"), max_chars=60)
        official_terms = _ksa_terms_for_question(
            ncs_ksa=ncs_ksa,
            ncs_code=ncs_code,
            fallback_terms=[current_focus] if current_focus else [],
        )
        focus_candidates = [current_focus, *official_terms]
        focus_candidates = [
            value
            for pos, value in enumerate(focus_candidates)
            if value and _ksa_key(value) not in {_ksa_key(x) for x in focus_candidates[:pos]}
        ]
        focus = focus_candidates[(max(1, attempt) - 1) % len(focus_candidates)] if focus_candidates else current_focus
        focus = focus or "핵심 수행기준"
        focus_type = _ksa_type_for_focus(ncs_ksa, ncs_code, focus)
        subject = str(item.get("competency") or item.get("compeUnitName") or "").strip()
        detail = str(
            item.get("ncs_detail")
            or item.get("ncsSubdCdnm")
            or item.get("ncsSclasCdnm")
            or ""
        ).strip()
        follow_count = len(item.get("follow_ups") or []) if isinstance(item.get("follow_ups"), list) else 0
        follow_count = max(3, min(5, follow_count or default_follow_count))
        if normalized_generation_offset is None:
            variation_index = history_size + history_variation_seed + (index * 37) + attempt
        else:
            # The browser sends a monotonic question offset for the current
            # generation context. Multiplying each question slot by the repair
            # block size prevents count changes from reusing another run's
            # deterministic variation range.
            variation_index = ((normalized_generation_offset + index) * 37) + attempt
        focus_evidence = _evidence_row_for_focus(ncs_ksa, ncs_code, focus)
        task_frame = _question_task_frame(
            focus=focus,
            focus_type=focus_type,
            subject=subject,
            detail=detail,
            comp_def=str(item.get("compeUnitDef") or "").strip(),
            evidence_row=focus_evidence,
        )
        evaluation_points = _method_evaluation_points(
            method,
            [focus, *[term for term in official_terms if _ksa_key(term) != _ksa_key(focus)]],
            focus_type,
            surface_focus=task_frame["task_object"],
        )
        repaired_question = _question_for_method(
            method=method,
            subject=subject,
            focus=focus,
            detail=detail,
            comp_def=str(item.get("compeUnitDef") or "").strip(),
            focus_type=focus_type,
            variation_index=variation_index,
            task_frame=task_frame,
            job_context_text=job_context_text,
        )
        repaired_source = "quality_orchestrator_repair"
        repaired_model_preserved = False
        repaired_followups = _followups_for_method(
            method=method,
            subject=subject,
            focus=focus,
            count=follow_count,
            variant_index=variation_index,
            focus_type=focus_type,
            task_frame=task_frame,
        )

        item.update(
            {
                "type": method,
                "method": method,
                "question_focus": focus,
                "question_focus_type": focus_type,
                "question_focus_source": "official_ksa",
                "question_focus_surface": task_frame["task_object"],
                "question_task_frame": task_frame,
                "question_evidence_id": task_frame.get("evidence_id", ""),
                "question_evidence_required": bool(focus_evidence),
                "ksa_refs": [focus, *[term for term in official_terms if _ksa_key(term) != _ksa_key(focus)]][:4],
                "question": repaired_question,
                "follow_ups": repaired_followups,
                "evaluation_points": evaluation_points,
                "question_source": repaired_source,
                "model_question_preserved": repaired_model_preserved,
            }
        )
        item["follow_up"] = item["follow_ups"][0] if item["follow_ups"] else ""
        item["task_conditions"] = _task_conditions_for_method(
            method=method,
            subject=subject,
            focus=focus,
            detail=detail,
            comp_def=str(item.get("compeUnitDef") or "").strip(),
            focus_type=focus_type,
            variation_index=variation_index,
        )
        item["question_variation_index"] = variation_index
        item["assessment_guide"] = _behavior_anchored_evaluation(
            method,
            focus,
            evaluation_points,
            focus_type,
            surface_focus=task_frame["task_object"],
        )
        existing_reasons = [
            str(reason).strip()
            for reason in (item.get("model_replacement_reasons") or [])
            if str(reason).strip()
        ] if isinstance(item.get("model_replacement_reasons"), list) else []
        item["model_replacement_reasons"] = list(
            dict.fromkeys([*existing_reasons, *[f"orchestrator_{reason}" for reason in reasons]])
        )
        return item

    deterministic_sources = {
        "template_fallback",
        "rule_fallback",
        "simulation_candidate",
        "quality_orchestrator_repair",
    }
    required_repair_reasons = {
        index: ["generation_offset_variation"]
        for index, item in enumerate(questions)
        if normalized_generation_offset is not None
        and normalized_generation_offset > 0
        and str(item.get("question_source") or "").strip() in deterministic_sources
    }

    def audit_question(item: dict[str, Any]) -> dict[str, Any]:
        source = str(item.get("question_source") or "").strip()
        measurement = evaluate_ksa_measurement(item)
        precision_grounding = evaluate_question_precision_grounding(item)
        precision_issues = [
            f"precision_grounding_{code}"
            for code in (precision_grounding.get("issue_codes") or [])
            if str(code).strip()
        ]
        measurement_issues = [
            str(code).strip()
            for code in (measurement.get("issues") or [])
            if str(code).strip()
        ]
        checks = dict(measurement.get("checks") or {})
        checks["precision_grounding"] = bool(precision_grounding.get("passed"))
        if not _is_subscription_cli_source(source):
            return {
                "passed": bool(measurement.get("passed"))
                and bool(precision_grounding.get("passed")),
                "issues": list(
                    dict.fromkeys([*measurement_issues, *precision_issues])
                ),
                "checks": checks,
            }
        realism = evaluate_question_realism(item)
        realism_issues = [
            f"field_realism_{code}"
            for code in (realism.get("issue_codes") or [])
            if str(code).strip()
        ]
        checks.update(
            {
                f"field_realism_{name}": passed
                for name, passed in (realism.get("checks") or {}).items()
            }
        )
        return {
            "passed": bool(realism.get("passed"))
            and bool(measurement.get("passed"))
            and bool(precision_grounding.get("passed")),
            "issues": list(
                dict.fromkeys(
                    [*measurement_issues, *realism_issues, *precision_issues]
                )
            ),
            "checks": checks,
        }

    orchestrated_questions, metadata = orchestrate_question_set(
        questions,
        avoid_questions=avoid_questions,
        repair_question=repair_question,
        audit_question=audit_question,
        max_repair_attempts=36,
        required_repair_reasons=required_repair_reasons,
    )
    _refresh_question_repeat_metadata(orchestrated_questions)
    strategy["interview_questions"] = orchestrated_questions
    strategy["interview_by_competency"] = _group_interview_questions_for_response(orchestrated_questions)
    strategy = _attach_ksa_evidence_to_strategy(strategy, ncs_ksa)
    report = strategy.get("question_quality_report") if isinstance(strategy.get("question_quality_report"), dict) else {}
    report_passed = report.get("passed") is True
    report_summary = report.get("summary") if isinstance(report.get("summary"), dict) else {}
    report_items = [item for item in (report.get("items") or []) if isinstance(item, dict)]
    try:
        expected_question_count = int(report_summary.get("expected_question_count") or 0)
    except (TypeError, ValueError):
        expected_question_count = 0
    question_count_gap = (
        abs(expected_question_count - len(orchestrated_questions))
        if expected_question_count > 0
        else 0
    )
    fallback_adjustment_degraded = (
        str(strategy.get("fallback_adjustment_status") or "").strip()
        == "degraded_to_runtime_recheck"
    )
    operational_warnings = ["fallback_adjustment_degraded"] if fallback_adjustment_degraded else []
    full_quality_unresolved_count = (
        sum(not bool(item.get("ready")) for item in report_items) + question_count_gap
    )
    metadata_items = [item for item in (metadata.get("items") or []) if isinstance(item, dict)]
    for report_item in report_items:
        try:
            metadata_index = int(report_item.get("index") or 0) - 1
        except (TypeError, ValueError):
            continue
        if not 0 <= metadata_index < len(metadata_items) or report_item.get("ready") is True:
            continue
        final_issues = list(metadata_items[metadata_index].get("final_issues") or [])
        final_issues.extend(
            f"full_quality_{issue}"
            for issue in (report_item.get("issues") or [])
            if str(issue or "").strip()
        )
        metadata_items[metadata_index]["final_issues"] = list(dict.fromkeys(final_issues))
    metadata["items"] = metadata_items
    metadata["full_quality_unresolved_count"] = full_quality_unresolved_count
    metadata["question_count_gap"] = question_count_gap
    metadata["unresolved_count"] = (
        sum(bool(item.get("final_issues")) for item in metadata_items) + question_count_gap
    )
    metadata["stages"] = [
        {
            "name": "candidate_generation",
            "status": (
                "degraded"
                if fallback_adjustment_degraded
                else "passed" if not question_count_gap else "partial"
            ),
            "question_count": len(questions),
            "expected_question_count": expected_question_count,
        },
        *list(metadata.get("stages") or []),
        {
            "name": "full_quality_gate",
            "status": "passed" if report_passed else "needs_review",
            "ready_count": int((report.get("summary") or {}).get("ready_count") or 0),
            "question_count": int((report.get("summary") or {}).get("question_count") or len(orchestrated_questions)),
        },
    ]
    if not report_passed or operational_warnings:
        metadata["status"] = "needs_review"
    metadata["quality_report_passed"] = report_passed
    metadata["operational_warnings"] = operational_warnings
    metadata["generation_offset"] = normalized_generation_offset
    strategy["question_quality_orchestration"] = metadata
    strategy["question_customization_policy"] = (
        "model_candidate_then_ksa_measurement_history_dedup_repair_and_full_recheck"
    )
    # Keep the quality report based on the substantive question text, then
    # expose the concrete case rows in the final candidate-facing question.
    # This avoids treating a long data appendix as a wording defect while
    # ensuring exports and print views remain self-contained.
    for item in strategy.get("interview_questions") or []:
        if isinstance(item, dict):
            method = str(item.get("type") or item.get("method") or "").strip()
            item["question"] = _ensure_question_material_reference(
                item.get("question"),
                method,
                item.get("task_conditions"),
            )
    strategy["interview_by_competency"] = _group_interview_questions_for_response(
        strategy.get("interview_questions") or []
    )
    return strategy


def _select_verified_sclass_candidates(
    candidates: list[dict[str, Any]] | None,
    max_keep: int = 4,
    min_keep: int = 1,
    score_margin: float = 0.18,
    min_confidence: float = 0.62,
) -> list[dict[str, Any]]:
    rows = [dict(x) for x in (candidates or []) if isinstance(x, dict)]
    if not rows:
        return []

    def _conf(v: dict[str, Any]) -> float:
        try:
            return float(v.get("confidence", 0.0) or 0.0)
        except Exception:
            return 0.0

    rows.sort(key=_conf, reverse=True)

    deduped: list[dict[str, Any]] = []
    seen_keys: set[tuple[str, str, str, str]] = set()
    for row in rows:
        key = (
            str(row.get("ncs_code_no", "")).strip(),
            str(row.get("ncs_lclass_code", "")).strip(),
            str(row.get("ncs_mclass_code", "")).strip(),
            str(row.get("ncs_sclass_code", "")).strip(),
        )
        if key in seen_keys:
            continue
        seen_keys.add(key)
        deduped.append(row)
    if not deduped:
        return []

    keep_limit = _clamp_sclass_limit(max_keep, default=4)
    min_keep = max(1, min(keep_limit, int(min_keep)))

    top_conf = _conf(deduped[0])
    threshold = max(float(min_confidence), top_conf - float(score_margin))

    selected: list[dict[str, Any]] = []
    for row in deduped:
        if _conf(row) >= threshold:
            selected.append(row)
        if len(selected) >= keep_limit:
            break

    if len(selected) < min_keep:
        for row in deduped:
            if row in selected:
                continue
            selected.append(row)
            if len(selected) >= min_keep:
                break

    return selected[:keep_limit]


def _merge_sclass_candidates(
    primary: list[dict[str, Any]] | None,
    secondary: list[dict[str, Any]] | None,
    max_items: int = 8,
) -> list[dict[str, Any]]:
    out: list[dict[str, Any]] = []
    index_by_key: dict[tuple[str, str, str, str] | tuple[str, str], int] = {}

    def _key(row: dict[str, Any]) -> tuple[str, str, str, str] | tuple[str, str]:
        code_key = (
            str(row.get("ncs_code_no", "")).strip(),
            str(row.get("ncs_lclass_code", "")).strip(),
            str(row.get("ncs_mclass_code", "")).strip(),
            str(row.get("ncs_sclass_code", "")).strip(),
        )
        if any(code_key):
            return code_key
        return ("name", str(row.get("sclass_name", "")).strip())

    def _conf(row: dict[str, Any]) -> float:
        try:
            return float(row.get("confidence", 0.0) or 0.0)
        except Exception:
            return 0.0

    for bucket in ((primary or []), (secondary or [])):
        for row in bucket:
            if not isinstance(row, dict):
                continue
            r = dict(row)
            key = _key(r)
            if key in index_by_key:
                i = index_by_key[key]
                if _conf(r) > _conf(out[i]):
                    out[i] = r
                continue
            index_by_key[key] = len(out)
            out.append(r)
            if len(out) >= max_items:
                return out
    return out


def _expand_direct_with_same_family_mentions(
    direct_candidates: list[dict[str, Any]] | None,
    reverse_candidates: list[dict[str, Any]] | None,
    jd_text: str,
    max_add: int = 2,
) -> list[dict[str, Any]]:
    direct = [dict(x) for x in (direct_candidates or []) if isinstance(x, dict)]
    if len(direct) != 1:
        return direct

    base = direct[0]
    base_l = str(base.get("ncs_lclass_code", "")).strip()
    base_m = str(base.get("ncs_mclass_code", "")).strip()
    base_s = str(base.get("ncs_sclass_code", "")).strip()
    if not (base_l and base_m):
        return direct

    norm_text = re.sub(r"\s+", "", str(jd_text or ""))
    if not norm_text:
        return direct

    seen: set[tuple[str, str, str]] = {(base_l, base_m, base_s)}
    add_count = 0
    for row in (reverse_candidates or []):
        if not isinstance(row, dict):
            continue
        l_cd = str(row.get("ncs_lclass_code", "")).strip()
        m_cd = str(row.get("ncs_mclass_code", "")).strip()
        s_cd = str(row.get("ncs_sclass_code", "")).strip()
        if (l_cd, m_cd) != (base_l, base_m):
            continue
        if (l_cd, m_cd, s_cd) in seen:
            continue

        name = str(row.get("sclass_name", "")).strip()
        if not name or re.sub(r"\s+", "", name) not in norm_text:
            continue
        try:
            conf = float(row.get("confidence", 0.0) or 0.0)
        except Exception:
            conf = 0.0
        if conf < 0.60:
            continue

        seen.add((l_cd, m_cd, s_cd))
        direct.append(dict(row))
        add_count += 1
        if add_count >= max(1, int(max_add)):
            break
    return direct


def _build_avoid_questions_context(questions: list[str], max_items: int = 12) -> str:
    seen: set[str] = set()
    items: list[str] = []
    for raw in reversed([str(q or "").strip() for q in questions]):
        key = normalize_question_dedup_key(raw)
        if not key or key in seen:
            continue
        seen.add(key)
        items.append(_clean_question_text(raw, max_chars=160))
        if len(items) >= max(1, int(max_items)):
            break
    if not items:
        return ""
    return (
        "[반복 금지 - 최근 생성 또는 채택된 질문, 최신순]\n"
        + "\n".join(f"- {item}" for item in items)
    )


_MAX_AVOID_QUESTION_ITEMS = 500
_MAX_AVOID_QUESTION_CHARS = 1600


def _extract_question_texts(value: Any) -> list[str]:
    if isinstance(value, str):
        raw = value.strip()
        if not raw:
            return []
        try:
            parsed = json.loads(raw)
        except Exception:
            parsed = [part.strip() for part in re.split(r"[\r\n]+", raw) if part.strip()]
    else:
        parsed = value

    if not isinstance(parsed, list):
        return []
    # History is client-provided operational context, not an unbounded archive.
    # Keep the newest unique entries so a long-running browser session cannot
    # make every subsequent generation progressively more expensive.
    recent_items = parsed[-_MAX_AVOID_QUESTION_ITEMS:]
    reversed_out: list[str] = []
    seen: set[str] = set()
    for item in reversed(recent_items):
        if isinstance(item, dict):
            text = _clean_question_text(
                item.get("question") or item.get("text") or "",
                max_chars=_MAX_AVOID_QUESTION_CHARS,
            )
        else:
            text = _clean_question_text(item, max_chars=_MAX_AVOID_QUESTION_CHARS)
        key = normalize_question_dedup_key(text)
        if not key or key in seen:
            continue
        seen.add(key)
        reversed_out.append(text)
    return list(reversed(reversed_out))


def _filter_questions_against_avoid_list(
    questions: list[dict[str, Any]],
    avoid_questions: list[str],
) -> list[dict[str, Any]]:
    if not questions or not avoid_questions:
        return questions
    avoid_texts = [str(q or "").strip() for q in avoid_questions if str(q or "").strip()]
    avoid_keys = {normalize_question_dedup_key(q) for q in avoid_texts if normalize_question_dedup_key(q)}
    if not avoid_texts and not avoid_keys:
        return questions

    out: list[dict[str, Any]] = []
    for item in questions:
        if not isinstance(item, dict):
            continue
        q_text = str(item.get("question") or "").strip()
        q_key = normalize_question_dedup_key(q_text)
        if q_key and q_key in avoid_keys:
            continue
        if any(is_similar_question_text(q_text, avoid) for avoid in avoid_texts):
            continue
        out.append(item)
    return out


def _filter_ncs_code_result_against_avoid_list(
    result: dict[str, Any],
    avoid_questions: list[str],
) -> None:
    """Filter NCS-code questions while keeping follow-up links internally consistent."""
    raw_main_questions = result.get("main_questions")
    if not isinstance(raw_main_questions, list):
        return

    main_questions: list[dict[str, Any]] = []
    old_to_new_index: dict[int, int] = {}
    for old_index, question in enumerate(raw_main_questions):
        if not isinstance(question, dict):
            continue
        if not _filter_questions_against_avoid_list([question], avoid_questions):
            continue
        old_to_new_index[old_index] = len(main_questions)
        main_questions.append(question)

    follow_up_questions: list[dict[str, Any]] = []
    raw_follow_up_questions = result.get("follow_up_questions")
    if isinstance(raw_follow_up_questions, list):
        for follow_up in raw_follow_up_questions:
            if not isinstance(follow_up, dict):
                continue
            try:
                old_index = int(follow_up.get("for_question_index"))
            except (TypeError, ValueError):
                continue
            if old_index not in old_to_new_index:
                continue
            remapped_follow_up = dict(follow_up)
            remapped_follow_up["for_question_index"] = old_to_new_index[old_index]
            follow_up_questions.append(remapped_follow_up)

    result["main_questions"] = main_questions
    result["follow_up_questions"] = follow_up_questions
    result["question_count"] = len(main_questions)
    result["follow_up_count"] = len(follow_up_questions)
    result["total_count"] = len(main_questions) + len(follow_up_questions)


def _generation_provider_descriptor(provider: str = "") -> dict[str, Any]:
    provider = normalize_generation_provider(
        provider,
        default=_configured_generation_provider(),
    )
    if provider == "codex_cli":
        return {
            "provider": "codex_cli",
            "provider_label": "Codex · ChatGPT 로그인",
            "auth_mode": "chatgpt_subscription",
            "requires_request_api_key": False,
            "local_only": True,
        }
    if provider == "claude_code":
        return {
            "provider": "claude_code",
            "provider_label": "Claude Code · Claude 로그인",
            "auth_mode": "claude_subscription",
            "requires_request_api_key": False,
            "local_only": True,
        }
    if provider in request_supported_generation_providers():
        config = generation_provider_config(provider)
        descriptor = {
            "provider": str(config.get("provider") or provider),
            "provider_label": str(config.get("provider_label") or provider),
            "auth_mode": str(config.get("auth_mode") or "request_scoped_api_key"),
            "requires_request_api_key": bool(config.get("requires_request_api_key")),
            "credential_managed_by": str(config.get("credential_managed_by") or "request"),
            "key_label": str(config.get("key_label") or "API 키"),
            "default_model": str(config.get("default_model") or ""),
            "supports_custom_model": bool(config.get("supports_custom_model")),
            "local_only": bool(config.get("local_only")),
        }
        if provider == "openai_api":
            descriptor["model_orchestration"] = {
                "policy": "role-based-openai-models-diversity-v2",
                "ncs_candidate_rerank": openai_role_model("ncs_rerank"),
                "question_authoring": openai_role_model("question_authoring"),
                "quality_review": openai_role_model("quality_review"),
                "quality_regeneration": openai_role_model("quality_regeneration"),
            }
        if provider == OPENROUTER_PROVIDER:
            server_key_state = settings.openrouter_server_key_state()
            recovery_model = openrouter_recovery_model()
            descriptor["server_key_enabled"] = settings.openrouter_server_key_enabled()
            descriptor["server_key_state"] = server_key_state
            descriptor["recovery_model"] = recovery_model
            descriptor["recovery_enabled"] = bool(recovery_model)
            standard_effort, _ = openrouter_reasoning_effort(
                interview_methods=[],
                target_count=1,
                follow_up_count=3,
                stage="primary",
            )
            high_risk_effort, _ = openrouter_reasoning_effort(
                interview_methods=["발표면접"],
                target_count=1,
                follow_up_count=3,
                stage="primary",
            )
            quality_retry_effort, _ = openrouter_reasoning_effort(
                interview_methods=[],
                target_count=1,
                follow_up_count=3,
                stage="quality_retry",
            )
            descriptor["reasoning_profiles"] = {
                "standard": standard_effort,
                "high_risk": high_risk_effort,
                "quality_retry": quality_retry_effort,
            }
            descriptor["timeout_profiles_sec"] = {
                "standard": provider_timeout_sec(provider, 0),
                "high_risk": provider_timeout_sec(
                    provider,
                    0,
                    openrouter_env_name="OPENROUTER_HIGH_RISK_TIMEOUT_SEC",
                ),
            }
            if server_key_state == "configured":
                descriptor.update(
                    {
                        "auth_mode": "server_env_api_key",
                        "requires_request_api_key": False,
                        "credential_managed_by": "server_env",
                    }
                )
        return descriptor
    return {
        "provider": provider,
        "provider_label": provider or "설정 오류",
        "auth_mode": "misconfigured",
        "requires_request_api_key": False,
        "local_only": False,
    }


def _subscription_cli_provider_http_error(
    exc: BaseException,
    *,
    provider: str,
) -> HTTPException | None:
    """Surface local subscription-CLI failures instead of degrading silently."""

    provider = _GENERATION_PROVIDER_ALIASES.get(
        str(provider or "").strip().lower(),
        str(provider or "").strip().lower(),
    )
    code = str(getattr(exc, "code", "") or "").strip()
    if provider not in _SUBSCRIPTION_CLI_PROVIDERS or not code:
        return None

    retryable = bool(getattr(exc, "retryable", False))
    provider_label = "Codex" if provider == "codex_cli" else "Claude Code"
    message = f"{provider_label} generation failed"
    status_code = 502

    if code.endswith("_usage_limit_reached"):
        status_code = 429
        message = (
            f"{provider_label} 구독 사용량 한도에 도달했습니다. "
            "잠시 후 다시 시도하거나 로그인 상태를 확인하세요."
        )
    elif code.endswith("_authentication_required"):
        status_code = 503
        login_command = "codex login" if provider == "codex_cli" else "claude auth login"
        message = (
            f"{provider_label} 로그인 상태를 확인할 수 없습니다. "
            f"PowerShell에서 {login_command}을 실행한 뒤 다시 시도하세요."
        )
    elif code.endswith("_unavailable"):
        status_code = 503
        message = f"{provider_label} CLI를 찾을 수 없습니다. 먼저 설치 상태를 확인하세요."
    elif code.endswith("_timeout"):
        status_code = 504
        message = f"{provider_label} 응답 시간이 초과되었습니다. 잠시 후 다시 시도하세요."
    elif code.endswith("_input_too_large"):
        status_code = 413
        message = f"{provider_label}에 전달할 문서 입력이 너무 큽니다. 입력 범위를 줄여 주세요."
    elif code.endswith("_invalid_output"):
        status_code = 502
        message = f"{provider_label}가 유효한 구조화 결과를 반환하지 않았습니다. 다시 시도해 주세요."
    elif code.endswith("_execution_failed"):
        status_code = 502
        message = f"{provider_label} 실행 중 오류가 발생했습니다. 잠시 후 다시 시도하세요."
    else:
        return None

    return HTTPException(
        status_code=status_code,
        detail={
            "code": code,
            "provider": provider,
            "message": message,
            "retryable": retryable,
        },
    )


_INSTITUTION_API_SAFE_MODEL_FAILURE_REASONS = frozenset(
    {
        "openai_request_timeout",
        "openai_network_unreachable",
        "openrouter_request_timeout",
        "openrouter_network_unreachable",
        "openrouter_request_failed",
        "model_response_not_object",
        "model_response_invalid_shape",
        "model_response_invalid_json",
        "model_response_truncated",
        "model_response_content_filtered",
        "model_response_refused",
        "model_question_count_mismatch",
        "model_question_content_missing",
        "model_question_diversity_mismatch",
        "question_set_count_or_diversity_failed",
        "institution_api_question_generation_failed",
        "institution_api_question_quality_rejected",
        "openai_http_400",
        "openai_http_401",
        "openai_http_403",
        "openai_http_404",
        "openai_http_408",
        "openai_http_409",
        "openai_http_422",
        "openai_http_425",
        "openai_http_429",
        "openai_http_500",
        "openai_http_502",
        "openai_http_503",
        "openai_http_504",
        "openrouter_http_400",
        "openrouter_http_401",
        "openrouter_http_403",
        "openrouter_http_404",
        "openrouter_http_408",
        "openrouter_http_409",
        "openrouter_http_422",
        "openrouter_http_425",
        "openrouter_http_429",
        "openrouter_http_500",
        "openrouter_http_502",
        "openrouter_http_503",
        "openrouter_http_504",
    }
)


def _institution_api_provider_http_error(
    exc: BaseException,
    *,
    provider: str = "openai_api",
) -> HTTPException:
    """Return a stable, non-sensitive failure for request-scoped API use."""

    resolved_provider = _request_generation_provider(provider)
    config = generation_provider_config(resolved_provider)
    provider_label = str(config.get("provider_label") or "API")
    error_prefix = "openrouter" if resolved_provider == "openrouter_api" else "openai"
    normalized = str(exc or "").casefold()
    if "ai_quality_review_timeout" in normalized or "generation request deadline exhausted" in normalized:
        return HTTPException(
            status_code=504,
            detail={
                "code": f"{resolved_provider}_timeout",
                "provider": resolved_provider,
                "message": "질문 생성 또는 독립 AI 품질검수가 제한 시간 안에 완료되지 않았습니다. 다시 시도해 주세요.",
                "retryable": True,
            },
        )
    if "ai_quality_review_network_failed" in normalized:
        return HTTPException(
            status_code=503,
            detail={
                "code": f"{resolved_provider}_unreachable",
                "provider": resolved_provider,
                "message": "독립 AI 품질검수 서비스에 연결하지 못했습니다. 네트워크 상태를 확인한 뒤 다시 시도해 주세요.",
                "retryable": True,
            },
        )
    if any(
        code in normalized
        for code in (
            "ai_quality_review_invalid_json",
            "ai_quality_review_invalid_shape",
            "ai_quality_review_invalid_score",
            "ai_quality_review_invalid_failure_codes",
            "ai_quality_review_provider_failed",
            "ai_quality_review_empty_input",
        )
    ):
        return HTTPException(
            status_code=502,
            detail={
                "code": f"{resolved_provider}_quality_rejected",
                "provider": resolved_provider,
                "message": "독립 AI 품질검수를 신뢰할 수 없어 질문을 반환하지 않았습니다. 다시 생성해 주세요.",
                "retryable": True,
                "quality_diagnostics": {
                    "failure_scope": "ai_quality_review",
                },
            },
        )
    if any(
        f"{error_prefix}_http_{status}" in normalized
        for status in (401, 403)
    ):
        return HTTPException(
            status_code=401,
            detail={
                "code": f"{resolved_provider}_authentication_failed",
                "provider": resolved_provider,
                "message": f"입력한 {provider_label} 키를 인증하지 못했습니다.",
                "retryable": False,
            },
        )
    if f"{error_prefix}_http_429" in normalized:
        return HTTPException(
            status_code=429,
            detail={
                "code": f"{resolved_provider}_usage_limit_reached",
                "provider": resolved_provider,
                "message": f"{provider_label} 사용량 또는 요청 한도를 확인해 주세요.",
                "retryable": True,
            },
        )
    if (
        f"{error_prefix}_request_timeout" in normalized
        or f"{error_prefix}_http_408" in normalized
    ):
        return HTTPException(
            status_code=504,
            detail={
                "code": f"{resolved_provider}_timeout",
                "provider": resolved_provider,
                "message": f"{provider_label} 응답 시간이 초과되었습니다. 문항 수를 줄이거나 잠시 후 다시 시도해 주세요.",
                "retryable": True,
            },
        )
    if f"{error_prefix}_network_unreachable" in normalized:
        return HTTPException(
            status_code=503,
            detail={
                "code": f"{resolved_provider}_unreachable",
                "provider": resolved_provider,
                "message": f"서버에서 {provider_label}에 연결하지 못했습니다. 네트워크 연결을 확인해 주세요.",
                "retryable": True,
            },
        )
    if any(
        reason in normalized
        for reason in (
            "model_response_not_object",
            "model_response_invalid_shape",
            "model_response_invalid_json",
            "model_response_truncated",
            "model_question_count_mismatch",
            "model_question_content_missing",
        )
    ):
        return HTTPException(
            status_code=502,
            detail={
                "code": f"{resolved_provider}_invalid_output",
                "provider": resolved_provider,
                "message": f"{provider_label} 응답이 잘렸거나 요청한 문항 수·형식을 충족하지 못했습니다.",
                "retryable": True,
            },
        )
    if any(
        reason in normalized
        for reason in (
            "model_question_diversity_mismatch",
            "question_set_count_or_diversity_failed",
            "institution_api_question_generation_failed",
            "institution_api_question_quality_rejected",
        )
    ):
        quality_diagnostics = getattr(exc, "quality_diagnostics", {})
        quality_diagnostics = (
            dict(quality_diagnostics) if isinstance(quality_diagnostics, dict) else {}
        )
        requested_count = int(
            quality_diagnostics.get("requested_question_count") or 0
        )
        failed_count = int(quality_diagnostics.get("failed_question_count") or 0)
        if requested_count > 0 and failed_count > 0:
            message = (
                f"문서 크기나 GPT 과부하가 아니라, 생성된 {requested_count}개 문항 중 "
                f"{failed_count}개가 KSA·구조·안전 필수 검사를 통과하지 못했습니다. "
                "통과 문항을 함께 반환할 수 없는 무결성 실패가 남았습니다."
            )
        else:
            message = (
                "문서 크기나 GPT 과부하가 아니라, 생성 결과의 KSA·구조·안전 "
                "필수 검사가 해결되지 않았습니다."
            )
        return HTTPException(
            status_code=502,
            detail={
                "code": f"{resolved_provider}_quality_rejected",
                "provider": resolved_provider,
                "message": message,
                "retryable": True,
                "quality_diagnostics": quality_diagnostics,
            },
        )
    if any(
        reason in normalized
        for reason in ("model_response_content_filtered", "model_response_refused")
    ):
        return HTTPException(
            status_code=422,
            detail={
                "code": f"{resolved_provider}_content_restricted",
                "provider": resolved_provider,
                "message": f"{provider_label}가 입력 내용에 대한 질문 생성을 완료하지 않았습니다. 입력 문서를 확인해 주세요.",
                "retryable": False,
            },
        )
    if any(
        f"{error_prefix}_http_{status}" in normalized
        for status in (400, 404, 409, 422, 425)
    ):
        return HTTPException(
            status_code=502,
            detail={
                "code": f"{resolved_provider}_request_rejected",
                "provider": resolved_provider,
                "message": f"{provider_label}가 모델 또는 요청 형식을 거절했습니다. 모델 설정을 확인해 주세요.",
                "retryable": False,
            },
        )
    if any(
        f"{error_prefix}_http_{status}" in normalized
        for status in (500, 502, 503, 504)
    ):
        return HTTPException(
            status_code=502,
            detail={
                "code": f"{resolved_provider}_upstream_unavailable",
                "provider": resolved_provider,
                "message": f"{provider_label} 서비스가 일시적으로 응답하지 않습니다. 잠시 후 다시 시도해 주세요.",
                "retryable": True,
            },
        )
    return HTTPException(
        status_code=502,
        detail={
            "code": f"{resolved_provider}_generation_failed",
            "provider": resolved_provider,
            "message": f"{provider_label}에서 질문을 생성하지 못했습니다. 잠시 후 다시 시도해 주세요.",
            "retryable": True,
        },
    )


def _require_institution_api_model_output(strategy: Any) -> None:
    """Reject an explicit model failure before any deterministic repair runs."""

    if not isinstance(strategy, dict):
        raise RuntimeError("institution_api_invalid_generation_result")
    questions = strategy.get("interview_questions")
    has_question = isinstance(questions, list) and any(
        isinstance(item, dict) and str(item.get("question") or "").strip()
        for item in questions
    )
    explicit_failure = bool(str(strategy.get("error") or "").strip()) or (
        str(strategy.get("question_generation_policy") or "").strip()
        == "model_only_no_template_fallback"
    )
    if explicit_failure or not has_question:
        raw_error = str(strategy.get("error") or "").strip()
        safe_reason = ""
        if raw_error.startswith("model_generation_failed:"):
            candidate = raw_error.split(":", 1)[1].strip()
            if candidate in _INSTITUTION_API_SAFE_MODEL_FAILURE_REASONS:
                safe_reason = candidate
        elif raw_error.startswith("model_generation_skipped:"):
            candidate = raw_error.split(":", 1)[1].strip()
            if candidate in _INSTITUTION_API_SAFE_MODEL_FAILURE_REASONS:
                safe_reason = candidate
        raise RuntimeError(safe_reason or "institution_api_empty_generation")


_INSTITUTION_QUALITY_RETRY_POLICY = "institution-openai-quality-retry-v1"
_INSTITUTION_RETRYABLE_QUALITY_CODES = frozenset(
    {
        "deterministic_fallback",
        "question_evidence_assignment_failed",
        "question_quality_report_failed",
        "question_quality_orchestration_failed",
        "precision_grounding_failed",
        "ai_quality_review_failed",
        "unsafe_question_surface",
        "follow_up_count_mismatch",
    }
)
_INSTITUTION_HARD_QUALITY_CHECKS = frozenset(
    {
        "supported_method",
        "method_shape",
        "main_question_method_shape",
        "main_question_job_context",
        "evaluation_points",
        "evaluation_points_quality",
        "ncs_grounded",
        "ksa_grounded",
        "ksa_measurement_task",
        "detail_grounded",
        "evidence_linked",
        "candidate_surface_safe",
        "official_sample_format",
        "blind_hiring_safe",
        "unique_question",
        "specific_context",
        "job_specific_context",
        "focus_scenario_coherence",
        "decision_dilemma_quality",
        "debate_option_defensibility",
        "debate_outcome_flexibility",
        "debate_case_neutrality",
        "operating_conditions_separated",
        "standardized_task_conditions",
        "case_materials_sufficient",
        "decision_authority_context",
        "inbasket_authority_context",
        "behavior_anchored_evaluation",
        "precision_grounding",
    }
)
_INSTITUTION_SOFT_QUALITY_CHECKS = frozenset(
    {
        "follow_up_depth",
        "follow_up_quality",
        "natural_wording",
        "field_realism",
    }
)
# These checks used to be returned as ``human_review_required``.  Public
# question output is now fail-closed, so editorial/realism failures are hard
# gates just like evidence and safety failures.
_INSTITUTION_PUBLIC_HARD_QUALITY_CHECKS = (
    _INSTITUTION_HARD_QUALITY_CHECKS | _INSTITUTION_SOFT_QUALITY_CHECKS
)
_INSTITUTION_SERVER_SAFETY_CHECKS = frozenset(
    {
        "supported_method",
        "evidence_linked",
        "candidate_surface_safe",
        "blind_hiring_safe",
        "unique_question",
    }
)
_BROKEN_PUBLIC_QUESTION_SURFACE_RE = re.compile(
    r"(?:지원가|지원와|KSA\s*가\s*드러난\s*장면|자료가\s*서로\s*달랐던\s*때)",
    re.IGNORECASE,
)
_INSTITUTION_HARD_REALISM_ISSUES = frozenset(
    {
        "instruction_injection_artifact",
        "candidate_visible_instruction_injection",
        "candidate_visible_ncs_label",
        "label_like_metadata_exposure",
    }
)
_SAFE_QUALITY_CODE_RE = re.compile(r"^[a-z][a-z0-9_]{0,79}$")


def _has_broken_public_question_surface(question: dict[str, Any]) -> bool:
    values = [
        question.get("question"),
        *(question.get("follow_ups") or [] if isinstance(question.get("follow_ups"), list) else []),
        *(
            question.get("evaluation_points") or []
            if isinstance(question.get("evaluation_points"), list)
            else []
        ),
    ]
    return bool(_BROKEN_PUBLIC_QUESTION_SURFACE_RE.search("\n".join(str(value or "") for value in values)))


def _planned_follow_up_counts(
    question_plan: Any,
    question_count: int,
) -> list[int | None]:
    """Resolve the server-locked 1..5 follow-up count for every question slot."""

    count = max(0, int(question_count or 0))
    if count <= 0 or not isinstance(question_plan, dict):
        return [None] * count

    def locked_count(value: Any, *, present: bool) -> int | None:
        if not present:
            return None
        try:
            parsed = int(value)
        except (TypeError, ValueError):
            return 0
        return parsed if 1 <= parsed <= 5 else 0

    default_present = "follow_up_count" in question_plan
    default_count = locked_count(
        question_plan.get("follow_up_count"),
        present=default_present,
    )
    sequence = [
        item
        for item in (question_plan.get("question_sequence") or [])
        if isinstance(item, dict)
    ]
    expected: list[int | None] = []
    for index in range(count):
        row = sequence[index] if index < len(sequence) else None
        if row is not None and "follow_up_count" in row:
            expected.append(
                locked_count(row.get("follow_up_count"), present=True)
            )
        elif default_count is not None:
            expected.append(default_count)
        elif sequence:
            # A locked sequence that omits a slot/count is malformed.  Keep a
            # fail-closed sentinel instead of silently accepting any 1..5.
            expected.append(0)
        else:
            expected.append(None)
    return expected


def _follow_up_count_mismatch_indexes(
    rows: list[dict[str, Any]],
    question_plan: Any,
) -> list[int]:
    """Return one-based slots whose authored array violates the request count."""

    expected_counts = _planned_follow_up_counts(question_plan, len(rows))
    mismatches: list[int] = []
    for index, (row, expected) in enumerate(
        zip(rows, expected_counts, strict=True),
        start=1,
    ):
        if expected is None:
            continue
        raw_follow_ups = row.get("follow_ups")
        actual = (
            len(
                [
                    value
                    for value in raw_follow_ups
                    if isinstance(value, str) and value.strip()
                ]
            )
            if isinstance(raw_follow_ups, list)
            else 0
        )
        if actual != expected:
            mismatches.append(index)
    return mismatches


def _result_follow_up_count_mismatch_indexes(result: Any) -> list[int]:
    if not isinstance(result, dict):
        return []
    rows = [
        row
        for row in (result.get("interview_questions") or [])
        if isinstance(row, dict)
    ]
    question_plan = result.get("question_plan_used")
    if not isinstance(question_plan, dict):
        question_plan = result.get("question_plan")
    return _follow_up_count_mismatch_indexes(rows, question_plan)


def _institution_question_rejection_codes(
    result: Any,
    *,
    require_quality_metadata: bool = False,
) -> list[str]:
    """Return stable server-owned reasons why public question output is unsafe.

    The codes intentionally carry no model text, KSA labels, exception strings,
    or precision evidence.  They can therefore be used in a bounded repair
    prompt and audit metadata without reflecting sensitive request data.
    """

    if not isinstance(result, dict):
        return ["invalid_question_result"]

    questions: list[dict[str, Any]] = []
    for key in ("main_questions", "questions", "interview_questions"):
        rows = result.get(key)
        if isinstance(rows, list):
            questions.extend(row for row in rows if isinstance(row, dict))

    codes: list[str] = []
    if result.get("error"):
        codes.append("result_error")
    if not any(str(row.get("question") or "").strip() for row in questions):
        codes.append("empty_question_set")

    if bool(result.get("template_fallback_used")):
        codes.append("deterministic_fallback")
    elif any(
        str(row.get("question_source") or "").strip()
        in {"template_fallback", "rule_fallback", "quality_orchestrator_repair"}
        for row in questions
    ):
        codes.append("deterministic_fallback")
    if any(
        str(row.get("question_source") or "").strip()
        != "openai_api"
        for row in questions
    ):
        codes.append("disallowed_question_source")
    if any(
        row.get("degraded") is True
        or row.get("human_review_required") is True
        or row.get("review_required") is True
        for row in questions
    ):
        codes.append("degraded_question")
    if _result_follow_up_count_mismatch_indexes(result):
        codes.append("follow_up_count_mismatch")

    evidence_assignment = result.get("question_evidence_assignment")
    if (
        isinstance(evidence_assignment, dict)
        and evidence_assignment.get("applicable") is True
        and evidence_assignment.get("passed") is not True
    ):
        codes.append("question_evidence_assignment_failed")

    quality_report = result.get("question_quality_report")
    if require_quality_metadata and not isinstance(quality_report, dict):
        codes.append("question_quality_report_missing")

    orchestration = result.get("question_quality_orchestration")
    if require_quality_metadata and not isinstance(orchestration, dict):
        codes.append("question_quality_orchestration_missing")
    elif (
        isinstance(orchestration, dict)
        and str(orchestration.get("status") or "").strip() != "passed"
        and any(
            isinstance(item, dict) and bool(item.get("final_issues"))
            for item in (orchestration.get("items") or [])
        )
    ):
        codes.append("question_quality_orchestration_failed")

    unsafe_surface = any(_has_broken_public_question_surface(row) for row in questions)
    if not unsafe_surface:
        unsafe_surface = any(
            {
                str(value or "").strip()
                for value in (evaluate_question_realism(row).get("issue_codes") or [])
                if str(value or "").strip()
            }
            & _INSTITUTION_HARD_REALISM_ISSUES
            for row in questions
        )
    if unsafe_surface:
        codes.append("unsafe_question_surface")

    ai_review = result.get("ai_quality_review")
    if not isinstance(ai_review, dict):
        if require_quality_metadata:
            codes.append("ai_quality_review_missing")
    elif str(ai_review.get("status") or "").strip() != "passed":
        codes.append("ai_quality_review_failed")

    # Precision diagnostics remain visible, but semantic grounding is decided
    # by the independent AI review against the official KSA and job context.
    return list(dict.fromkeys(codes))


def _institution_question_quality_issue_codes(
    result: Any,
    *,
    blocking_only: bool = False,
) -> list[str]:
    """Return bounded server-owned issue codes for a targeted model retry."""

    if not isinstance(result, dict):
        return []
    codes: list[str] = []
    report = result.get("question_quality_report")
    if isinstance(report, dict) and not blocking_only:
        for item in report.get("items") or []:
            if not isinstance(item, dict):
                continue
            for value in item.get("issues") or []:
                code = str(value or "").strip()
                if _SAFE_QUALITY_CODE_RE.fullmatch(code):
                    codes.append(code)
            for value in item.get("realism_issue_codes") or []:
                code = str(value or "").strip()
                if _SAFE_QUALITY_CODE_RE.fullmatch(code):
                    codes.append(f"field_realism_{code}")
            for value in item.get("precision_grounding_issue_codes") or []:
                code = str(value or "").strip()
                if _SAFE_QUALITY_CODE_RE.fullmatch(code):
                    codes.append(f"precision_grounding_{code}")
    orchestration = result.get("question_quality_orchestration")
    if isinstance(orchestration, dict):
        for item in orchestration.get("items") or []:
            if not isinstance(item, dict):
                continue
            for value in item.get("final_issues") or []:
                code = str(value or "").strip()
                if _SAFE_QUALITY_CODE_RE.fullmatch(code):
                    codes.append(code)
    ai_review = result.get("ai_quality_review")
    if isinstance(ai_review, dict):
        for item in ai_review.get("items") or []:
            if (
                not isinstance(item, dict)
                or (blocking_only and item.get("passed") is True)
            ):
                continue
            for value in (
                list(item.get("reason_codes") or [])
                + list(item.get("regeneration_guidance_codes") or [])
            ):
                code = str(value or "").strip()
                if _SAFE_QUALITY_CODE_RE.fullmatch(code):
                    codes.append(code)
    return list(dict.fromkeys(codes))[:30]


class InstitutionQuestionQualityRejected(RuntimeError):
    """Quality rejection with bounded, server-owned diagnostics only."""

    def __init__(self, diagnostics: dict[str, Any] | None = None) -> None:
        super().__init__("institution_api_question_quality_rejected")
        self.quality_diagnostics = dict(diagnostics or {})


def _institution_hard_question_indexes(result: Any) -> list[int]:
    """Return one-based slots that cannot be exposed as interview drafts."""

    if not isinstance(result, dict):
        return []
    questions = [
        item
        for item in (result.get("interview_questions") or [])
        if isinstance(item, dict)
    ]
    if not questions:
        return []
    hard_indexes: set[int] = set(
        _result_follow_up_count_mismatch_indexes(result)
    )
    orchestration = result.get("question_quality_orchestration")
    if isinstance(orchestration, dict):
        for fallback_index, item in enumerate(orchestration.get("items") or [], start=1):
            if not isinstance(item, dict) or not item.get("final_issues"):
                continue
            try:
                index = int(item.get("index") or fallback_index)
            except (TypeError, ValueError):
                index = fallback_index
            hard_indexes.add(index)
    elif questions:
        return list(range(1, len(questions) + 1))

    ai_review = result.get("ai_quality_review")
    if isinstance(ai_review, dict):
        for fallback_index, item in enumerate(ai_review.get("items") or [], start=1):
            if not isinstance(item, dict) or item.get("passed") is True:
                continue
            try:
                index = int(item.get("index") or fallback_index)
            except (TypeError, ValueError):
                index = fallback_index
            hard_indexes.add(index)

    evidence_assignment = result.get("question_evidence_assignment")
    if isinstance(evidence_assignment, dict):
        for value in evidence_assignment.get("mismatched_indexes") or []:
            try:
                index = int(value)
            except (TypeError, ValueError):
                continue
            if 1 <= index <= len(questions):
                hard_indexes.add(index)

    for index, question in enumerate(questions, start=1):
        source = str(question.get("question_source") or "").strip()
        if bool(result.get("template_fallback_used")) or source in {
            "template_fallback",
            "rule_fallback",
            "quality_orchestrator_repair",
        }:
            hard_indexes.add(index)
        if evaluate_question_precision_grounding(question).get("passed") is not True:
            # Independent AI grounding is authoritative for semantic quality;
            # the deterministic precision report remains diagnostic only.
            pass
        if _has_broken_public_question_surface(question):
            hard_indexes.add(index)
    return sorted(index for index in hard_indexes if 1 <= index <= len(questions))


def _institution_quality_failure_diagnostics(
    result: Any,
    *,
    attempt_count: int,
) -> dict[str, Any]:
    """Build a non-sensitive explanation for the UI and operational logs."""

    questions = [
        item
        for item in ((result or {}).get("interview_questions") or [])
        if isinstance(item, dict)
    ] if isinstance(result, dict) else []
    failed_indexes = _institution_hard_question_indexes(result)
    issue_counts: dict[str, int] = {}
    report = result.get("question_quality_report") if isinstance(result, dict) else None
    if isinstance(report, dict):
        for item in report.get("items") or []:
            if not isinstance(item, dict):
                continue
            for raw_code in item.get("issues") or []:
                code = str(raw_code or "").strip()
                if not _SAFE_QUALITY_CODE_RE.fullmatch(code):
                    continue
                public_code = (
                    code
                    if code in _INSTITUTION_HARD_QUALITY_CHECKS | _INSTITUTION_SOFT_QUALITY_CHECKS
                    else "unknown_quality_issue"
                )
                issue_counts[public_code] = issue_counts.get(public_code, 0) + 1
            for raw_code in item.get("realism_issue_codes") or []:
                code = str(raw_code or "").strip()
                if code in _INSTITUTION_HARD_REALISM_ISSUES:
                    public_code = f"field_realism_{code}"
                    issue_counts[public_code] = issue_counts.get(public_code, 0) + 1
    orchestration_issue_counts: dict[str, int] = {}
    orchestration = result.get("question_quality_orchestration") if isinstance(result, dict) else None
    if isinstance(orchestration, dict):
        for item in orchestration.get("items") or []:
            if not isinstance(item, dict):
                continue
            for raw_code in item.get("final_issues") or []:
                code = str(raw_code or "").strip()
                if not _SAFE_QUALITY_CODE_RE.fullmatch(code):
                    continue
                orchestration_issue_counts[code] = (
                    orchestration_issue_counts.get(code, 0) + 1
                )
    return {
        "requested_question_count": len(questions),
        "failed_question_count": len(failed_indexes),
        "failed_indexes": failed_indexes[:50],
        "issue_counts": dict(sorted(issue_counts.items())[:30]),
        "orchestration_issue_counts": dict(
            sorted(orchestration_issue_counts.items())[:30]
        ),
        "attempt_count": max(1, min(3, int(attempt_count or 1))),
        "failure_scope": "question_quality",
    }


def _targeted_quality_retry_plan(
    runtime_question_plan: dict[str, Any],
    failed_indexes: list[int],
) -> tuple[dict[str, Any], list[int], list[str]]:
    """Create a server-locked retry plan containing only failed slots."""

    sequence = [
        dict(item)
        for item in (runtime_question_plan.get("question_sequence") or [])
        if isinstance(item, dict)
    ]
    retry_sequence: list[dict[str, Any]] = []
    original_indexes: list[int] = []
    retry_methods: list[str] = []
    for original_index in sorted(set(failed_indexes)):
        if not 1 <= original_index <= len(sequence):
            continue
        item = dict(sequence[original_index - 1])
        method = str(item.get("type") or "").strip()
        item["retry_original_index"] = original_index
        item["index"] = len(retry_sequence) + 1
        retry_sequence.append(item)
        original_indexes.append(original_index)
        if method and method not in retry_methods:
            retry_methods.append(method)
    retry_plan = dict(runtime_question_plan)
    retry_plan["question_sequence"] = retry_sequence
    retry_plan["total_main_count"] = len(retry_sequence)
    retry_plan["targeted_retry"] = True
    retry_plan["targeted_retry_original_indexes"] = original_indexes
    return retry_plan, original_indexes, retry_methods


_INSTITUTION_GENERATION_BATCH_SIZE = 5
_INSTITUTION_GENERATION_BATCH_CONCURRENCY = 4


def _generation_batch_plans(
    runtime_question_plan: dict[str, Any],
    *,
    max_batch_size: int = _INSTITUTION_GENERATION_BATCH_SIZE,
) -> list[tuple[dict[str, Any], list[int], list[str]]]:
    """Split a locked server plan without changing slot, method, or KSA evidence."""

    sequence = [
        dict(item)
        for item in (runtime_question_plan.get("question_sequence") or [])
        if isinstance(item, dict)
    ]
    if not sequence:
        return [(dict(runtime_question_plan), [], [])]
    batch_size = max(1, min(20, int(max_batch_size or 1)))
    batches: list[tuple[dict[str, Any], list[int], list[str]]] = []
    for start in range(0, len(sequence), batch_size):
        source_rows = sequence[start : start + batch_size]
        local_rows: list[dict[str, Any]] = []
        original_indexes: list[int] = []
        methods: list[str] = []
        detail_counts: dict[str, dict[str, Any]] = {}
        detail_order: list[str] = []
        for local_index, source in enumerate(source_rows, start=1):
            row = dict(source)
            try:
                original_index = int(row.get("index") or (start + local_index))
            except (TypeError, ValueError):
                original_index = start + local_index
            row["generation_original_index"] = original_index
            row["index"] = local_index
            local_rows.append(row)
            original_indexes.append(original_index)
            method = str(row.get("type") or "").strip()
            if method and method not in methods:
                methods.append(method)
            detail = str(row.get("detail") or "").strip()
            if detail not in detail_counts:
                detail_order.append(detail)
                detail_counts[detail] = {
                    "detail": detail,
                    "enabled": True,
                    "main_count": 0,
                    "follow_up_count": max(
                        0,
                        min(5, int(row.get("follow_up_count") or 0)),
                    ),
                }
            detail_counts[detail]["main_count"] += 1

        selected_items = [detail_counts[detail] for detail in detail_order]
        batch_plan = dict(runtime_question_plan)
        batch_plan["question_sequence"] = local_rows
        batch_plan["total_main_count"] = len(local_rows)
        batch_plan["selected_items"] = selected_items
        batch_plan["items"] = [dict(item) for item in selected_items]
        batch_plan["selected_terms"] = [
            detail for detail in detail_order if detail
        ]
        batch_plan["generation_batch"] = True
        batch_plan["generation_batch_original_indexes"] = list(original_indexes)
        batches.append((batch_plan, original_indexes, methods))
    return batches




def _institution_hard_question_rejection_codes(
    result: Any,
    *,
    require_quality_metadata: bool = False,
) -> list[str]:
    """Keep provenance, safety, wording, and field-realism failures fail-closed."""

    if not isinstance(result, dict):
        return ["invalid_question_result"]
    questions = [
        row
        for row in (result.get("interview_questions") or [])
        if isinstance(row, dict)
    ]
    codes: list[str] = []
    if result.get("error"):
        codes.append("result_error")
    if not questions or not all(str(row.get("question") or "").strip() for row in questions):
        codes.append("empty_question_set")
    if bool(result.get("template_fallback_used")) or any(
        str(row.get("question_source") or "").strip()
        in {"template_fallback", "rule_fallback", "quality_orchestrator_repair"}
        for row in questions
    ):
        codes.append("deterministic_fallback")
    if any(
        str(row.get("question_source") or "").strip()
        != "openai_api"
        for row in questions
    ):
        codes.append("disallowed_question_source")
    if any(
        row.get("degraded") is True
        or row.get("human_review_required") is True
        or row.get("review_required") is True
        for row in questions
    ):
        codes.append("degraded_question")
    if _result_follow_up_count_mismatch_indexes(result):
        codes.append("follow_up_count_mismatch")
    evidence_assignment = result.get("question_evidence_assignment")
    if (
        isinstance(evidence_assignment, dict)
        and evidence_assignment.get("applicable") is True
        and evidence_assignment.get("passed") is not True
    ):
        codes.append("question_evidence_assignment_failed")

    report = result.get("question_quality_report")
    if not isinstance(report, dict):
        if require_quality_metadata:
            codes.append("question_quality_report_missing")

    ai_review = result.get("ai_quality_review")
    if not isinstance(ai_review, dict):
        if require_quality_metadata:
            codes.append("ai_quality_review_missing")
    elif str(ai_review.get("status") or "").strip() != "passed":
        codes.append("ai_quality_review_failed")

    orchestration = result.get("question_quality_orchestration")
    if require_quality_metadata and not isinstance(orchestration, dict):
        codes.append("question_quality_orchestration_missing")
    elif (
        isinstance(orchestration, dict)
        and str(orchestration.get("status") or "").strip() != "passed"
        and any(
            isinstance(item, dict) and bool(item.get("final_issues"))
            for item in (orchestration.get("items") or [])
        )
    ):
        codes.append("question_quality_orchestration_failed")
    if any(_has_broken_public_question_surface(row) for row in questions):
        codes.append("unsafe_question_surface")
    return list(dict.fromkeys(codes))


def _require_institution_api_question_output(
    result: Any,
    *,
    require_quality_metadata: bool = False,
) -> None:
    """Reject empty, deterministic, or failed-quality public API output."""

    if _institution_question_rejection_codes(
        result,
        require_quality_metadata=require_quality_metadata,
    ):
        raise RuntimeError("institution_api_question_generation_failed")


def _auxiliary_review_regeneration_context(
    review: dict[str, Any],
    *,
    server_codes: list[str] | None = None,
) -> str:
    """Build a bounded retry instruction without echoing draft wording."""

    reason_codes: list[str] = []
    guidance_codes: list[str] = []
    for item in review.get("items") or []:
        if not isinstance(item, dict) or item.get("passed") is True:
            continue
        for raw in item.get("reason_codes") or []:
            code = str(raw or "").strip()
            if _SAFE_QUALITY_CODE_RE.fullmatch(code):
                reason_codes.append(code)
        for raw in item.get("regeneration_guidance_codes") or []:
            code = str(raw or "").strip()
            if _SAFE_QUALITY_CODE_RE.fullmatch(code):
                guidance_codes.append(code)
    payload = {
        "reason_codes": list(dict.fromkeys(reason_codes))[:16],
        "regeneration_guidance_codes": list(dict.fromkeys(guidance_codes))[:16],
        "server_safety_codes": [
            code
            for code in dict.fromkeys(server_codes or [])
            if _SAFE_QUALITY_CODE_RE.fullmatch(str(code or ""))
        ][:8],
    }
    return (
        "[독립 AI 품질검수 재생성 지침]\n"
        "직전 문장을 고치거나 일부 표현을 재사용하지 말고, 동일한 공식 KSA와 "
        "직무 맥락에서 주질문·꼬리질문·평가포인트를 완전히 새로 작성하세요.\n"
        f"{json.dumps(payload, ensure_ascii=False, separators=(',', ':'))}"
    )


def _auxiliary_question_structure_codes(
    result: Any,
    *,
    expected_count: int,
) -> list[str]:
    """Validate the fixed auxiliary API shape without rewriting model copy."""

    questions = _public_question_rows(result)
    codes: list[str] = []
    if len(questions) != max(1, int(expected_count or 1)):
        codes.append("question_count_invalid")
    if any(
        not isinstance(question.get("follow_ups"), list)
        or not 1 <= len(question.get("follow_ups") or []) <= 5
        or any(
            not isinstance(value, str) or not value.strip()
            for value in question.get("follow_ups") or []
        )
        for question in questions
    ):
        codes.append("follow_up_count_invalid")
    if any(
        not isinstance(question.get("evaluation_points"), list)
        or not 1 <= len(question.get("evaluation_points") or []) <= 5
        or any(
            not isinstance(value, str) or not value.strip()
            for value in question.get("evaluation_points") or []
        )
        for question in questions
    ):
        codes.append("evaluation_point_count_invalid")
    return codes


def _quality_gate_auxiliary_question_result(
    *,
    generate_once: Callable[[str], dict[str, Any]],
    provider: str,
    api_key_override: str,
    generation_model: str,
    job_context: dict[str, Any],
    expected_count: int,
) -> dict[str, Any]:
    """Run the same independent two-pass AI gate for legacy public endpoints."""

    retry_context = ""
    last_result: dict[str, Any] = {}
    last_release_codes: list[str] = []
    for attempt_count in (1, 2):
        result = generate_once(retry_context)
        if not isinstance(result, dict):
            raise RuntimeError("institution_api_question_generation_failed")
        pre_review_codes = _institution_question_rejection_codes(result)
        fatal_pre_review_codes = [
            code
            for code in pre_review_codes
            if code not in {"unsafe_question_surface", "degraded_question"}
        ]
        if fatal_pre_review_codes:
            raise RuntimeError("institution_api_question_generation_failed")
        pre_review_structure_codes = _auxiliary_question_structure_codes(
            result,
            expected_count=expected_count,
        )
        try:
            _require_official_ksa_result(result)
        except HTTPException as exc:
            # A provider-authored shape error is itself regeneration input.
            # Let one bounded retry run from the server-owned official KSA
            # registry, but never bypass missing/unavailable official evidence.
            if (
                pre_review_structure_codes
                and result.get("ncs_ksa_available") is True
                and any(
                    isinstance(row, dict)
                    for row in (result.get("official_ksa_evidence") or [])
                )
            ):
                pass
            else:
                raise InstitutionQuestionQualityRejected(
                    {
                        "requested_question_count": len(_public_question_rows(result)),
                        "failed_question_count": len(_public_question_rows(result)),
                        "failure_scope": "official_ksa_evidence",
                    }
                ) from exc

        questions = _public_question_rows(result)
        official_ksa = [
            dict(row)
            for row in (result.get("official_ksa_evidence") or [])
            if isinstance(row, dict)
        ]
        unit_rows: list[dict[str, Any]] = []
        seen_units: set[tuple[str, str]] = set()
        for row in official_ksa:
            unit_key = (
                str(row.get("ncsClCd") or "").strip(),
                str(row.get("compeUnitName") or "").strip(),
            )
            if unit_key in seen_units:
                continue
            seen_units.add(unit_key)
            unit_rows.append(
                {
                    "ncsClCd": unit_key[0],
                    "compeUnitName": unit_key[1],
                    "compeUnitDef": str(row.get("compeUnitDef") or "").strip(),
                }
            )
        methods = list(
            dict.fromkeys(
                str(row.get("type") or row.get("question_type") or "면접질문").strip()
                for row in questions
                if str(row.get("type") or row.get("question_type") or "면접질문").strip()
            )
        )
        review = review_interview_questions_with_ai(
            questions=questions,
            ncs_matches=unit_rows,
            ncs_ksa=official_ksa,
            interview_methods=methods,
            job_context=job_context,
            provider=provider,
            api_key_override=api_key_override,
            generation_model=generation_model,
        )
        review["attempt_count"] = attempt_count
        result["ai_quality_review"] = review
        last_result = result
        release_codes = list(
            dict.fromkeys(
                [
                    *_institution_question_rejection_codes(result),
                    *pre_review_structure_codes,
                ]
            )
        )
        last_release_codes = release_codes
        if (
            str(review.get("status") or "").strip() == "passed"
            and not release_codes
        ):
            result["question_release_status"] = "ai_quality_review_passed"
            return result
        retry_context = _auxiliary_review_regeneration_context(
            review,
            server_codes=release_codes,
        )

    failed_indexes = [
        int(item.get("index") or index)
        for index, item in enumerate(
            ((last_result.get("ai_quality_review") or {}).get("items") or []),
            start=1,
        )
        if isinstance(item, dict) and item.get("passed") is not True
    ]
    if not failed_indexes and last_release_codes:
        failed_indexes = list(range(1, len(_public_question_rows(last_result)) + 1))
    raise InstitutionQuestionQualityRejected(
        {
            "requested_question_count": len(_public_question_rows(last_result)),
            "failed_question_count": len(failed_indexes),
            "failed_indexes": failed_indexes,
            "attempt_count": 2,
            "failure_scope": "ai_quality_review",
        }
    )


def _raw_model_question_texts(strategy: Any, *, max_items: int = 20) -> list[str]:
    if not isinstance(strategy, dict):
        return []
    questions = strategy.get("interview_questions")
    if not isinstance(questions, list):
        return []
    texts: list[str] = []
    for item in questions:
        if not isinstance(item, dict):
            continue
        text = _clean_question_text(
            item.get("model_question_raw") or item.get("question"),
            max_chars=160,
        )
        if text:
            texts.append(text)
        if len(texts) >= max(1, int(max_items)):
            break
    return texts


def _extract_generated_question_items(
    strategy: Any,
    *,
    max_items: int = 100,
) -> list[dict[str, Any]]:
    if not isinstance(strategy, dict):
        return []
    questions = strategy.get("interview_questions")
    if not isinstance(questions, list):
        return []

    output: list[dict[str, Any]] = []
    for idx, item in enumerate(questions, start=1):
        if not isinstance(item, dict):
            continue
        question = _clean_question_text(item.get("question"), max_chars=400)
        if not question:
            continue
        follow_ups = [
            _clean_question_text(fu, max_chars=200)
            for fu in (item.get("follow_ups") or [])
            if isinstance(fu, str) and _clean_question_text(fu, max_chars=200)
        ]
        row = {
            "index": idx,
            "question": question,
            "type": str(item.get("type") or item.get("method") or "").strip(),
            "ncsClCd": str(item.get("ncsClCd") or item.get("ncs_code") or "").strip(),
            "competency": str(item.get("competency") or "").strip(),
            "question_source": str(item.get("question_source") or "").strip(),
        }
        review_only = bool(
            item.get("human_review_required")
            or item.get("degraded")
            or row["question_source"] in {
                "server_ksa_fallback",
                "template_fallback",
                "rule_fallback",
            }
        )
        if review_only:
            row["review_required"] = True
        row["ready"] = bool(question) and not review_only
        if follow_ups:
            row["follow_ups"] = follow_ups
        output.append(row)
        if len(output) >= max(1, int(max_items)):
            break
    return output


def _extract_all_generated_question_items(strategy: Any) -> list[dict[str, Any]]:
    """Expose all generated question slots for full visibility, including empty slots."""

    if not isinstance(strategy, dict):
        return []
    questions = strategy.get("interview_questions")
    if not isinstance(questions, list):
        return []
    expected_count = 0
    for plan_key in ("question_plan_used", "question_plan"):
        plan_value = strategy.get(plan_key)
        if isinstance(plan_value, dict):
            try:
                expected_count = int(plan_value.get("total_main_count") or 0)
            except Exception:
                expected_count = 0
            if expected_count > 0:
                break
    if expected_count <= 0:
        expected_count = len(questions)

    output: list[dict[str, Any]] = []
    for idx in range(1, expected_count + 1):
        item = questions[idx - 1] if idx <= len(questions) else None
        row = {
            "index": idx,
            "question": "",
            "type": "",
            "ncsClCd": "",
            "competency": "",
            "question_source": "",
            "ready": False,
        }
        if isinstance(item, dict):
            question = _clean_question_text(item.get("question"), max_chars=400)
            row["question"] = question
            row["type"] = str(item.get("type") or item.get("method") or "").strip()
            row["ncsClCd"] = str(
                item.get("ncsClCd") or item.get("ncs_code") or ""
            ).strip()
            row["competency"] = str(item.get("competency") or "").strip()
            row["question_source"] = str(item.get("question_source") or "").strip()
            follow_ups = [
                _clean_question_text(fu, max_chars=200)
                for fu in (item.get("follow_ups") or [])
                if isinstance(fu, str) and _clean_question_text(fu, max_chars=200)
            ]
            if follow_ups:
                row["follow_ups"] = follow_ups
            evaluation_points = [
                _clean_question_text(point, max_chars=180)
                for point in (item.get("evaluation_points") or item.get("eval_points") or [])
                if isinstance(point, str) and _clean_question_text(point, max_chars=180)
            ]
            if evaluation_points:
                row["evaluation_points"] = evaluation_points
            if item.get("question_hash"):
                row["question_hash"] = str(item.get("question_hash"))
            if item.get("question_evidence_id"):
                row["question_evidence_id"] = str(item.get("question_evidence_id"))
            review_only = bool(
                item.get("human_review_required")
                or item.get("degraded")
                or row["question_source"] in {
                    "server_ksa_fallback",
                    "template_fallback",
                    "rule_fallback",
                }
            )
            if review_only:
                row["review_required"] = True
            row["ready"] = bool(question) and not review_only
            if "index" in item:
                try:
                    row["source_index"] = int(item.get("index"))
                except Exception:
                    pass
            if "generation_original_index" in item:
                try:
                    row["generation_original_index"] = int(
                        item.get("generation_original_index")
                    )
                except Exception:
                    pass
        output.append(row)
    return output


def _build_generated_question_text_rows(strategy: Any) -> list[dict[str, Any]]:
    """Build a flat, consumer-friendly question-list payload."""

    if isinstance(strategy, list):
        rows = strategy
    else:
        rows = _extract_all_generated_question_items(strategy)
    output: list[dict[str, Any]] = []
    for row in rows:
        if not isinstance(row, dict):
            continue
        question = str(row.get("question") or "").strip()
        raw_index = row.get("index")
        try:
            index = int(raw_index)
        except Exception:
            index = len(output) + 1
        output_row = {
            "index": index,
            "question": question,
            "type": str(row.get("type") or "").strip(),
            "ncsClCd": str(row.get("ncsClCd") or "").strip(),
            "competency": str(row.get("competency") or "").strip(),
            "question_source": str(row.get("question_source") or "").strip(),
            "ready": bool(row.get("ready")),
        }
        if row.get("review_required") is True:
            output_row["review_required"] = True
        output.append(output_row)
    return output


def _build_generated_question_text_payload(
    source: Any,
) -> tuple[list[dict[str, Any]], list[str]]:
    rows = _build_generated_question_text_rows(source)
    texts = [item.get("question") for item in rows if item.get("question")]
    return rows, texts


def _normalize_generated_questions(
    raw_questions: Any,
    *,
    expected_count: int | None = None,
    max_item_chars: int = 500,
) -> list[dict[str, Any]]:
    """Normalize heterogeneous question rows into a unified export shape."""

    if not isinstance(raw_questions, list):
        return []
    output: list[dict[str, Any]] = []
    for idx, item in enumerate(raw_questions, start=1):
        if not isinstance(item, dict):
            continue
        row = {
            "index": idx,
            "question": _clean_question_text(item.get("question"), max_chars=max_item_chars),
            "type": str(item.get("type") or item.get("question_type") or item.get("method") or "").strip(),
            "ncsClCd": str(
                item.get("ncsClCd") or item.get("ncs_code") or item.get("unit_code") or ""
            ).strip(),
            "competency": str(item.get("competency") or item.get("compeUnitName") or "").strip(),
            "question_source": str(item.get("question_source") or "").strip(),
        }
        review_only = bool(
            item.get("human_review_required")
            or item.get("degraded")
            or row["question_source"] in {
                "server_ksa_fallback",
                "template_fallback",
                "rule_fallback",
            }
        )
        if review_only:
            row["review_required"] = True
        follow_ups = [
            _clean_question_text(fu, max_chars=260)
            for fu in (item.get("follow_ups") or [])
            if isinstance(fu, str) and _clean_question_text(fu, max_chars=260)
        ]
        if follow_ups:
            row["follow_ups"] = follow_ups
        evaluation_points = [
            _clean_question_text(point, max_chars=260)
            for point in (item.get("evaluation_points") or item.get("eval_points") or [])
            if isinstance(point, str) and _clean_question_text(point, max_chars=260)
        ]
        if evaluation_points:
            row["evaluation_points"] = evaluation_points
        if item.get("question_hash"):
            row["question_hash"] = str(item.get("question_hash"))
        if item.get("question_evidence_id"):
            row["question_evidence_id"] = str(item.get("question_evidence_id"))
        # A degraded/provider-free draft must never look production-ready just
        # because it contains non-empty text. The full strategy carries the
        # detailed quality report; this compact export mirrors that release
        # boundary so the UI and downloads cannot silently approve a fallback.
        row["ready"] = bool(row["question"]) and not review_only
        output.append(row)

    target_count = int(expected_count) if expected_count and int(expected_count) > 0 else len(output)
    if target_count > len(output):
        for idx in range(len(output) + 1, target_count + 1):
            output.append(
                {
                    "index": idx,
                    "question": "",
                    "type": "",
                    "ncsClCd": "",
                    "competency": "",
                    "question_source": "",
                    "ready": False,
                }
            )
    return output


def _generated_question_preview_limit(
    question_plan: dict[str, Any] | None,
    strategy: dict[str, Any] | None,
    *,
    fallback: int = 4,
) -> int:
    """Calculate how many generated questions to expose in response previews."""

    candidate = 0
    try:
        candidate = int((question_plan or {}).get("total_main_count") or 0)
    except Exception:
        candidate = 0

    generated = []
    if isinstance(strategy, dict):
        maybe_generated = strategy.get("interview_questions")
        if isinstance(maybe_generated, list):
            generated = maybe_generated
    if candidate <= 0 and generated:
        candidate = len(generated)

    if generated and candidate > len(generated):
        candidate = len(generated)

    if candidate <= 0:
        candidate = max(1, int(fallback or 1))
    return min(max(1, candidate), 100)


def _build_generated_questions_payload(
    strategy: Any,
    question_plan: dict[str, Any] | None,
    *,
    include_all: bool,
    explicit_max_items: int | None = None,
    fallback: int = 4,
) -> list[dict[str, Any]]:
    """Build question list for response with optional full/limited mode."""

    if include_all:
        return _extract_all_generated_question_items(strategy)

    if isinstance(explicit_max_items, int) and explicit_max_items > 0:
        return _extract_generated_question_items(
            strategy,
            max_items=explicit_max_items,
        )

    return _extract_generated_question_items(
        strategy,
        max_items=_generated_question_preview_limit(
            question_plan,
            strategy,
            fallback=fallback,
        ),
    )


_QUESTION_EVIDENCE_ASSIGNMENT_POLICY = "planned-question-evidence-assignment-v1"


def _planned_question_evidence_assignments(
    *,
    question_plan: dict[str, Any],
    interview_methods: list[str],
    ncs_matches: list[dict[str, Any]],
    ncs_ksa: list[dict[str, Any]],
) -> tuple[dict[str, Any], list[tuple[int, str]]]:
    """Rebuild the exact server-side sequence that was embedded in the prompt."""

    target_count = int(question_plan.get("total_main_count") or 0)
    planned_sequence = _planned_question_sequence_for_prompt(
        question_plan,
        interview_methods,
        target_count,
        ncs_matches=ncs_matches,
        ncs_ksa=ncs_ksa,
    )
    if not planned_sequence:
        return dict(question_plan), []

    runtime_plan = dict(question_plan)
    runtime_plan["question_sequence"] = [dict(item) for item in planned_sequence]
    locks: list[tuple[int, str]] = []
    for fallback_index, item in enumerate(planned_sequence, start=1):
        evidence_id = str(item.get("evidence_id") or "").strip()
        try:
            index = int(item.get("index") or fallback_index)
        except (TypeError, ValueError):
            index = fallback_index
        if not re.fullmatch(r"ksa_[0-9a-f]{24}", evidence_id):
            continue
        if not _evidence_row_for_id(ncs_ksa, str(item.get("ncsClCd") or ""), evidence_id):
            continue
        locks.append((index, evidence_id))
    return runtime_plan, locks


def _raw_model_evidence_assignment_report(
    strategy: Any,
    evidence_locks: list[tuple[int, str]],
) -> dict[str, Any]:
    """Compare raw model IDs with the prompt plan before enrichment can repair them."""

    expected = dict(evidence_locks)
    questions = strategy.get("interview_questions") if isinstance(strategy, dict) else None
    rows = questions if isinstance(questions, list) else []
    mismatched_indexes: list[int] = []
    for index, evidence_id in evidence_locks:
        raw_item = rows[index - 1] if 0 < index <= len(rows) else None
        raw_id = (
            str(raw_item.get("question_evidence_id") or "").strip()
            if isinstance(raw_item, dict)
            else ""
        )
        if raw_id != evidence_id:
            mismatched_indexes.append(index)
    return {
        "policy": _QUESTION_EVIDENCE_ASSIGNMENT_POLICY,
        "applicable": bool(expected),
        "passed": not mismatched_indexes if expected else True,
        "expected_count": len(expected),
        "matched_count": len(expected) - len(mismatched_indexes),
        "mismatch_count": len(mismatched_indexes),
        "mismatched_indexes": mismatched_indexes,
    }


def _quality_retry_context(
    *,
    trigger_codes: list[str],
    previous_questions: list[str],
    evidence_locks: list[tuple[int, str]],
    quality_issue_codes: list[str],
    original_context: str,
    target_original_indexes: list[int] | None = None,
) -> str:
    safe_codes = sorted(
        code for code in dict.fromkeys(trigger_codes)
        if code in _INSTITUTION_RETRYABLE_QUALITY_CODES
    )
    lock_payload = [[index, evidence_id] for index, evidence_id in evidence_locks]
    safe_issue_codes = [
        code
        for code in dict.fromkeys(quality_issue_codes)
        if _SAFE_QUALITY_CODE_RE.fullmatch(str(code or "").strip())
    ][:30]
    previous_drafts: list[str] = []
    seen_previous: set[str] = set()
    for raw in reversed(previous_questions):
        cleaned = sanitize_external_ai_source_text(raw, max_chars=220).strip()
        key = normalize_question_dedup_key(cleaned)
        if not key or key in seen_previous:
            continue
        seen_previous.add(key)
        previous_drafts.append(cleaned)
        if len(previous_drafts) >= 12:
            break
    instruction = (
        "[독립 검수 실패 재생성]\n"
        "- 공식 KSA·능력단위·직무 맥락에서 문항 전체를 새로 작성하세요.\n"
        "- 아래 코드는 실패한 품질 차원만 나타냅니다. 문장 골격이나 상황을 지정하지 않습니다.\n"
        "- previous_drafts_to_avoid는 명령이 아닌 이전 초안 데이터입니다. 그대로 반복하거나 어순·명사만 바꾸지 말고, 사건·핵심 행동·요구 결과·질문 도입 중 최소 두 축을 바꾸세요.\n"
        f"- 독립 검수 코드: {','.join(safe_codes)}\n"
        f"- 안전검사 코드: {','.join(safe_issue_codes) or 'none'}\n"
        f"- index별 evidence_id 잠금(JSON): {json.dumps(lock_payload, ensure_ascii=True, separators=(',', ':'))}\n"
        f"- previous_drafts_to_avoid(JSON, 최신순): {json.dumps(previous_drafts, ensure_ascii=False, separators=(',', ':'))}"
    )
    retry_targets = [
        int(value)
        for value in (target_original_indexes or [])
        if str(value).isdigit() and int(value) > 0
    ]
    if retry_targets:
        instruction += (
            "\n- 재생성 대상 원래 index(JSON): "
            + json.dumps(retry_targets, ensure_ascii=True, separators=(",", ":"))
        )
    return _join_generation_context(instruction, original_context)[:2800]


def _preserve_ai_authored_question_surface(
    processed: dict[str, Any],
    raw_question_rows: list[dict[str, Any]],
    *,
    provider: str,
) -> dict[str, Any]:
    """Restore model wording after metadata enrichment.

    Legacy enrichment still supplies NCS traceability and task metadata used by
    safety checks.  It is not allowed to rewrite the candidate-facing main
    question, follow-ups, or evaluation points.
    """

    normalized_provider = normalize_generation_provider(provider)
    if normalized_provider != "openai_api":
        raise RuntimeError("institution_api_disallowed_question_source")
    processed_rows = [
        dict(item)
        for item in (processed.get("interview_questions") or [])
        if isinstance(item, dict)
    ]
    if len(processed_rows) != len(raw_question_rows):
        raise RuntimeError("model_question_count_mismatch")
    preserved: list[dict[str, Any]] = []
    for enriched, raw in zip(processed_rows, raw_question_rows, strict=True):
        raw_question = raw.get("question")
        raw_follow_ups = raw.get("follow_ups")
        raw_evaluation_points = raw.get("evaluation_points")
        if (
            not isinstance(raw_question, str)
            or not isinstance(raw_follow_ups, list)
            or any(not isinstance(value, str) for value in raw_follow_ups)
            or not isinstance(raw_evaluation_points, list)
            or any(not isinstance(value, str) for value in raw_evaluation_points)
        ):
            raise RuntimeError("model_response_invalid_shape")
        question = raw_question.strip()
        follow_ups = [value.strip() for value in raw_follow_ups if value.strip()]
        evaluation_points = [
            value.strip() for value in raw_evaluation_points if value.strip()
        ]
        if not question or not follow_ups or not evaluation_points:
            raise RuntimeError("model_question_content_missing")
        row = dict(enriched)
        row["question"] = question
        row["follow_ups"] = follow_ups
        row["follow_up"] = follow_ups[0]
        row["evaluation_points"] = evaluation_points
        row["question_source"] = normalized_provider
        row["model_question_preserved"] = True
        for key in (
            "model_replacement_reasons",
            "quality_repair_reasons",
            "model_question_replaced",
            "fallback_question",
            "fallback_used",
            "human_review_required",
        ):
            row.pop(key, None)
        preserved.append(row)
    processed["interview_questions"] = preserved
    processed["template_fallback_used"] = False
    processed["generation_mode"] = "ai_only"
    return processed


def _enrich_ai_authored_question_metadata(
    strategy: dict[str, Any],
    *,
    question_plan: dict[str, Any],
    interview_methods: list[str],
    ncs_matches: list[dict[str, Any]],
    ncs_ksa: list[dict[str, Any]],
) -> dict[str, Any]:
    """Attach locked NCS metadata without running legacy question rewriting.

    The provider owns every candidate-facing sentence. This function only
    restores the server-owned slot, ability-unit and KSA traceability fields
    that the UI and audit trail need. In particular it never reads or writes
    ``question``, ``follow_ups`` or ``evaluation_points``.
    """

    if not isinstance(strategy, dict):
        raise RuntimeError("model_response_not_object")
    provider_question_keys = {
        "type",
        "competency",
        "ncsClCd",
        "question",
        "follow_ups",
        "evaluation_points",
        "question_evidence_id",
        "question_focus_surface",
        "question_focus",
        "ksa_refs",
        "question_source",
    }
    source_rows = [
        {
            key: value
            for key, value in item.items()
            if key in provider_question_keys
        }
        for item in (strategy.get("interview_questions") or [])
        if isinstance(item, dict)
    ]
    sequence = [
        dict(item)
        for item in (question_plan.get("question_sequence") or [])
        if isinstance(item, dict)
    ]
    if sequence and len(source_rows) != len(sequence):
        raise RuntimeError("model_question_count_mismatch")

    enriched_rows: list[dict[str, Any]] = []
    for index, source in enumerate(source_rows):
        row = dict(source)
        planned = sequence[index] if index < len(sequence) else {}
        planned_code = str(planned.get("ncsClCd") or "").strip()
        planned_evidence_id = str(planned.get("evidence_id") or "").strip()
        raw_code = str(row.get("ncsClCd") or "").strip()
        raw_evidence_id = str(source.get("question_evidence_id") or "").strip()
        raw_evidence = _evidence_row_for_id(
            ncs_ksa,
            raw_code or planned_code,
            raw_evidence_id,
        )
        planned_evidence = _evidence_row_for_id(
            ncs_ksa,
            planned_code,
            planned_evidence_id,
        )
        evidence = raw_evidence or planned_evidence
        unit = next(
            (
                dict(item)
                for item in ncs_matches
                if isinstance(item, dict)
                and planned_code
                and str(item.get("ncsClCd") or "").strip() == planned_code
            ),
            {},
        )
        if not unit:
            unit = next(
                (
                    dict(item)
                    for item in ncs_matches
                    if isinstance(item, dict)
                    and raw_code
                    and str(item.get("ncsClCd") or "").strip() == raw_code
                ),
                {},
            )

        raw_method = str(row.get("type") or row.get("method") or "").strip()
        allowed_methods = {
            str(value or "").strip()
            for value in interview_methods
            if str(value or "").strip()
        }
        method = (
            raw_method
            if raw_method and (not allowed_methods or raw_method in allowed_methods)
            else str(planned.get("type") or "").strip()
        )
        if not method and interview_methods:
            method = str(interview_methods[index % len(interview_methods)] or "").strip()
        if method:
            row["type"] = method
            row["method"] = method

        ncs_code = (
            str(evidence.get("ncsClCd") or "").strip()
            or planned_code
            or str(unit.get("ncsClCd") or "").strip()
            or str(row.get("ncsClCd") or "").strip()
        )
        competency = (
            str(evidence.get("compeUnitName") or "").strip()
            or str(planned.get("compeUnitName") or "").strip()
            or str(unit.get("compeUnitName") or "").strip()
            or str(row.get("competency") or "").strip()
        )
        row["ncsClCd"] = ncs_code
        row["competency"] = competency
        row["compeUnitDef"] = (
            str(planned.get("compeUnitDef") or "").strip()
            or str(unit.get("compeUnitDef") or "").strip()
        )
        row["ncsSubdCdnm"] = str(unit.get("ncsSubdCdnm") or "").strip()
        row["ncsSclasCdnm"] = str(unit.get("ncsSclasCdnm") or "").strip()
        row["ncs_detail"] = (
            str(planned.get("detail") or "").strip()
            or str(unit.get("matchedDetailName") or "").strip()
            or row["ncsSubdCdnm"]
            or row["ncsSclasCdnm"]
        )

        if planned_evidence_id:
            row["question_evidence_assignment_valid"] = bool(
                planned_evidence and raw_evidence_id == planned_evidence_id
            )
        if raw_evidence_id or planned_evidence_id:
            row["question_evidence_id"] = raw_evidence_id or planned_evidence_id
        if evidence:
            official_factor = str(evidence.get("factorName") or "").strip()
            ksa_type = str(
                evidence.get("ksaTypeName")
                or evidence.get("factorType")
                or evidence.get("ksa_type")
                or ""
            ).strip()
            task_frame = _question_task_frame(
                focus=official_factor,
                focus_type=ksa_type,
                subject=competency or row["ncs_detail"] or "해당 직무",
                detail=row["ncs_detail"],
                comp_def=row["compeUnitDef"],
                evidence_row=evidence,
            )
            row["question_focus"] = official_factor
            row["question_focus_type"] = ksa_type
            row["question_focus_source"] = "official_ksa"
            row["question_focus_surface"] = str(
                task_frame.get("task_object") or ""
            ).strip()
            row["question_task_frame"] = task_frame
            row["question_evidence_required"] = True
        enriched_rows.append(row)

    # The model response is untrusted even after the generation service has
    # normalized it.  Keep only server-owned public diagnostics here so an
    # alternate builder, a future refactor, or a test double cannot reflect a
    # provider exception, API key, prompt, or arbitrary debug field into the
    # public strategy payload.
    public_strategy_keys = {
        "ncs_candidates_raw",
        "ncs_ksa_used",
        "ncs_context_used",
        "provider_generation_request_count",
        "provider_generation_request_limit",
        "transport_attempt_limit_per_generation_request",
        "provider_generation_model",
        "generation_provider",
        "provider_reasoning_effort",
        "provider_reasoning_stage",
        "provider_reasoning_reason",
        "provider_timeout_recovery_used",
        "provider_timeout_recovery_model",
        "provider_timeout_recovery_reasoning_effort",
        "provider_candidate_variant_count",
        "provider_candidate_variant_received_count",
        "question_candidate_selection",
        "provider_generation_notes",
        "model_question_generation_counts",
        "ncs_link",
        "question_generation_policy",
        "error",
        "warning",
    }
    enriched = {
        key: value
        for key, value in strategy.items()
        if key in public_strategy_keys
    }
    enriched["interview_questions"] = enriched_rows
    enriched["interview_by_competency"] = _group_interview_questions_for_response(
        enriched_rows
    )
    enriched["question_plan_used"] = dict(question_plan)
    enriched["interview_methods_used"] = list(interview_methods)
    enriched["question_customization_policy"] = "ai_authored_metadata_only_v1"
    return enriched


def _server_ai_question_safety_issues(
    rows: list[dict[str, Any]],
    avoid_questions: list[str] | None = None,
    question_plan: dict[str, Any] | None = None,
) -> list[dict[str, Any]]:
    """Return bounded structural/safety findings without semantic scoring."""

    findings: list[dict[str, Any]] = []
    follow_up_count_mismatches = set(
        _follow_up_count_mismatch_indexes(rows, question_plan)
    )
    seen_questions: set[str] = {
        key
        for value in (avoid_questions or [])
        if (key := normalize_question_dedup_key(str(value or "").strip()))
    }
    seen_question_texts: list[str] = [
        str(value or "").strip()
        for value in (avoid_questions or [])
        if str(value or "").strip()
    ]
    for index, row in enumerate(rows, start=1):
        issues: list[str] = []
        question = str(row.get("question") or "").strip()
        follow_ups = [
            str(value).strip()
            for value in (row.get("follow_ups") or [])
            if str(value).strip()
        ] if isinstance(row.get("follow_ups"), list) else []
        evaluation_points = [
            str(value).strip()
            for value in (row.get("evaluation_points") or [])
            if str(value).strip()
        ] if isinstance(row.get("evaluation_points"), list) else []
        if not question:
            issues.append("question_content_missing")
        if index in follow_up_count_mismatches:
            issues.append("follow_up_count_mismatch")
        elif not 1 <= len(follow_ups) <= 5:
            issues.append("follow_up_count")
        if not 1 <= len(evaluation_points) <= 5:
            issues.append("evaluation_point_count")
        if str(row.get("question_source") or "").strip() != "openai_api":
            issues.append("disallowed_question_source")
        if str(row.get("type") or row.get("method") or "").strip() not in SUPPORTED_INTERVIEW_METHODS:
            issues.append("unsupported_method")
        evidence_id = str(row.get("question_evidence_id") or "").strip()
        evidence_ids = {
            str(value or "").strip()
            for value in (row.get("evidence_ids") or [])
            if str(value or "").strip()
        } if isinstance(row.get("evidence_ids"), list) else set()
        if not evidence_id or evidence_id not in evidence_ids:
            issues.append("evidence_not_linked")
        visible_values = [question, *follow_ups, *evaluation_points]
        if _contains_blind_hiring_cue("\n".join(visible_values)):
            issues.append("blind_hiring_unsafe")
        if _has_broken_public_question_surface(row):
            issues.append("unsafe_question_surface")
        realism_issues = {
            str(value or "").strip()
            for value in (evaluate_question_realism(row).get("issue_codes") or [])
            if str(value or "").strip()
        }
        if realism_issues & _INSTITUTION_HARD_REALISM_ISSUES:
            issues.append("unsafe_metadata_exposure")
        question_key = normalize_question_dedup_key(question)
        if question_key and (
            question_key in seen_questions
            or any(
                is_similar_question_text(question, previous)
                for previous in seen_question_texts
            )
        ):
            issues.append("duplicate_question")
        if question_key:
            seen_questions.add(question_key)
            seen_question_texts.append(question)
        findings.append(
            {
                "index": index,
                "issues": list(dict.fromkeys(issues)),
                "passed": not issues,
            }
        )
    return findings


def _audit_ai_authored_strategy_without_repair(
    strategy: dict[str, Any],
    ncs_ksa: list[dict[str, Any]],
    *,
    avoid_questions: list[str] | None = None,
) -> dict[str, Any]:
    """Attach evidence and deterministic safety diagnostics without rewriting."""

    audited = _attach_ksa_evidence_to_strategy(strategy, ncs_ksa)
    rows = [
        item
        for item in (audited.get("interview_questions") or [])
        if isinstance(item, dict)
    ]
    safety_items = _server_ai_question_safety_issues(
        rows,
        avoid_questions=avoid_questions,
        question_plan=(
            audited.get("question_plan_used")
            if isinstance(audited.get("question_plan_used"), dict)
            else None
        ),
    )
    safety_passed = bool(rows) and all(item["passed"] for item in safety_items)
    audited["question_quality_orchestration"] = {
        "policy": "ai-authored-metadata-and-safety-audit-v2",
        "status": "passed" if safety_passed else "failed",
        "question_count": len(rows),
        "initial_failure_count": sum(not item["passed"] for item in safety_items),
        "repaired_count": 0,
        "repair_error_count": 0,
        "unresolved_count": sum(not item["passed"] for item in safety_items),
        "full_quality_unresolved_count": sum(not item["passed"] for item in safety_items),
        "items": [
            {
                "index": int(item["index"]),
                "initial_issues": list(item["issues"]),
                "final_issues": list(item["issues"]),
                "repaired": False,
            }
            for item in safety_items
        ],
        "stages": [
            {"name": "ai_generation", "status": "passed", "question_count": len(rows)},
            {
                "name": "read_only_safety_audit",
                "status": "passed" if safety_passed else "failed",
                "question_count": len(rows),
            },
        ],
    }
    return audited


def _apply_external_question_duplicate_safety(
    strategy: dict[str, Any],
    avoid_questions: list[str] | None,
) -> dict[str, Any]:
    """Merge history/rejected-draft duplicates into existing safety metadata."""

    avoided = [
        str(value or "").strip()
        for value in (avoid_questions or [])
        if str(value or "").strip()
    ]
    if not avoided:
        return strategy
    rows = [
        item
        for item in (strategy.get("interview_questions") or [])
        if isinstance(item, dict)
    ]
    duplicate_indexes = {
        index
        for index, row in enumerate(rows, start=1)
        if (question := str(row.get("question") or "").strip())
        and any(
            normalize_question_dedup_key(question)
            == normalize_question_dedup_key(previous)
            or is_similar_question_text(question, previous)
            for previous in avoided
        )
    }
    if not duplicate_indexes:
        return strategy
    orchestration = strategy.get("question_quality_orchestration")
    if not isinstance(orchestration, dict):
        orchestration = {
            "policy": "external-question-duplicate-safety-v1",
            "items": [],
        }
        strategy["question_quality_orchestration"] = orchestration
    items = orchestration.get("items")
    if not isinstance(items, list):
        items = []
        orchestration["items"] = items
    indexed_items = {
        int(item.get("index") or fallback_index): item
        for fallback_index, item in enumerate(items, start=1)
        if isinstance(item, dict)
        and str(item.get("index") or fallback_index).isdigit()
    }
    for index in range(1, len(rows) + 1):
        if index in indexed_items:
            continue
        item = {
            "index": index,
            "initial_issues": [],
            "final_issues": [],
            "repaired": False,
        }
        items.append(item)
        indexed_items[index] = item
    for fallback_index, item in enumerate(items, start=1):
        if not isinstance(item, dict):
            continue
        try:
            index = int(item.get("index") or fallback_index)
        except (TypeError, ValueError):
            index = fallback_index
        if index not in duplicate_indexes:
            continue
        for key in ("initial_issues", "final_issues"):
            values = [
                str(value or "").strip()
                for value in (item.get(key) or [])
                if str(value or "").strip()
            ]
            item[key] = list(dict.fromkeys([*values, "duplicate_question"]))
    unresolved_count = sum(
        bool(item.get("final_issues"))
        for item in items
        if isinstance(item, dict)
    )
    orchestration["status"] = "failed"
    orchestration["unresolved_count"] = unresolved_count
    orchestration["full_quality_unresolved_count"] = unresolved_count
    orchestration["history_duplicate_count"] = len(duplicate_indexes)
    return strategy


async def _generate_quality_gated_institution_strategy(
    *,
    build_kwargs: dict[str, Any],
    question_plan: dict[str, Any],
    interview_methods: list[str],
    ncs_matches: list[dict[str, Any]],
    ncs_ksa: list[dict[str, Any]],
    avoid_questions: list[str],
    generation_offset: int | None,
) -> dict[str, Any]:
    """Generate once and perform one bounded retry only for quality rejection."""

    active_generation_provider = normalize_generation_provider(
        build_kwargs.get("generation_provider", ""),
        default=_configured_generation_provider(),
    )
    model_requests_per_batch = _clamp_int(
        os.getenv("INSTITUTION_MODEL_REQUESTS_PER_BATCH"),
        default=1,
        lo=1,
        hi=1,
    )
    quality_retry_enabled = _coerce_bool_flag(
        os.getenv("INSTITUTION_QUALITY_RETRY_ENABLED"),
        default=True,
    )
    generation_batch_size = _clamp_int(
        os.getenv("INSTITUTION_GENERATION_BATCH_SIZE"),
        default=_INSTITUTION_GENERATION_BATCH_SIZE,
        lo=1,
        hi=20,
    )
    generation_batch_concurrency = _clamp_int(
        os.getenv("INSTITUTION_GENERATION_BATCH_CONCURRENCY"),
        default=_INSTITUTION_GENERATION_BATCH_CONCURRENCY,
        lo=1,
        hi=4,
    )
    generation_started_at = time.perf_counter()
    generation_attempt_elapsed_ms: list[int] = []

    def attach_model_orchestration(
        result: dict[str, Any],
        *,
        authoring_model: str,
        regeneration_model: str = "",
        regeneration_used: bool = False,
    ) -> dict[str, Any]:
        """Expose the role split without leaking credentials or provider errors."""

        review = result.get("ai_quality_review")
        review_model = (
            str(review.get("model") or "").strip()
            if isinstance(review, dict)
            else ""
        )
        rerank_used = any(
            str(row.get("rerank_method") or "").strip() == "ai"
            for row in ncs_matches
            if isinstance(row, dict)
        )
        configured_regeneration_model = (
            openai_role_model("quality_regeneration")
            if active_generation_provider == "openai_api"
            else regeneration_model or authoring_model
        )
        result["ai_model_orchestration"] = {
            "policy": "role-based-openai-models-diversity-v2",
            "provider": active_generation_provider,
            "roles": {
                "ncs_candidate_rerank": {
                    "model": (
                        openai_role_model("ncs_rerank")
                        if active_generation_provider == "openai_api"
                        else authoring_model
                    ),
                    "used": rerank_used,
                },
                "question_authoring": {
                    "model": authoring_model,
                    "used": True,
                    "guide": {
                        "id": "ncs-interviewer-2020-kordoc-v1",
                        "mode": "authoring_advice_only",
                        "star": "experience_probe_guidance_not_gate",
                    },
                    "candidate_variant_count": int(
                        result.get("provider_candidate_variant_count") or 1
                    ),
                    "diversity_cycle": int(
                        first_build_kwargs.get("diversity_cycle") or 0
                    ),
                },
                "quality_review": {
                    "model": review_model,
                    "used": bool(review_model),
                    "independent_call": True,
                    "acceptance_policy": (
                        dict(review.get("acceptance_policy") or {})
                        if isinstance(review, dict)
                        else {}
                    ),
                },
                "quality_regeneration": {
                    "model": regeneration_model or configured_regeneration_model,
                    "used": bool(regeneration_used),
                    "avoids_previous_drafts": bool(regeneration_used),
                },
            },
        }
        return result

    def with_generation_timing(result: dict[str, Any]) -> dict[str, Any]:
        result["generation_timing"] = {
            "total_elapsed_ms": max(
                0,
                int(round((time.perf_counter() - generation_started_at) * 1000)),
            ),
            "generation_attempt_count": len(generation_attempt_elapsed_ms),
            "generation_attempt_elapsed_ms": list(generation_attempt_elapsed_ms),
        }
        return result

    def rejection_diagnostics(
        result: Any,
        *,
        attempt_count: int,
    ) -> dict[str, Any]:
        diagnostics = _institution_quality_failure_diagnostics(
            result,
            attempt_count=attempt_count,
        )
        diagnostics["total_elapsed_ms"] = max(
            0,
            int(round((time.perf_counter() - generation_started_at) * 1000)),
        )
        diagnostics["generation_attempt_elapsed_ms"] = list(
            generation_attempt_elapsed_ms
        )
        return diagnostics

    async def attach_independent_ai_review(
        result: dict[str, Any],
        *,
        attempt_count: int,
        prior_questions_to_avoid: list[str] | None = None,
    ) -> dict[str, Any]:
        questions = [
            dict(item)
            for item in (result.get("interview_questions") or [])
            if isinstance(item, dict)
        ]
        review = await asyncio.to_thread(
            review_interview_questions_with_ai,
            questions=questions,
            ncs_matches=ncs_matches,
            ncs_ksa=ncs_ksa,
            interview_methods=interview_methods,
            job_context={
                "notice": str(build_kwargs.get("notice_text") or ""),
                "job_description": str(build_kwargs.get("jd_text") or ""),
                "duties": str(build_kwargs.get("duty_text") or ""),
                "evaluation": str(build_kwargs.get("evaluation_text") or ""),
            },
            provider=active_generation_provider,
            api_key_override=str(build_kwargs.get("api_key_override") or ""),
            generation_model=str(build_kwargs.get("generation_model") or ""),
            avoid_questions=(
                list(prior_questions_to_avoid)
                if prior_questions_to_avoid is not None
                else list(avoid_questions)
            ),
        )
        review["attempt_count"] = max(1, min(2, int(attempt_count or 1)))
        result["ai_quality_review"] = review
        return result

    runtime_question_plan, planned_evidence_locks = _planned_question_evidence_assignments(
        question_plan=question_plan,
        interview_methods=interview_methods,
        ncs_matches=ncs_matches,
        ncs_ksa=ncs_ksa,
    )
    initial_generation_batch_count = len(
        _generation_batch_plans(
            runtime_question_plan,
            max_batch_size=generation_batch_size,
        )
    )
    provider_generation_request_limit = max(
        1,
        initial_generation_batch_count
        * (model_requests_per_batch + int(quality_retry_enabled)),
    )

    async def run_once(
        local_build_kwargs: dict[str, Any],
        *,
        local_question_plan: dict[str, Any],
        local_interview_methods: list[str],
        local_evidence_locks: list[tuple[int, str]],
    ) -> tuple[dict[str, Any], list[str]]:
        attempt_started_at = time.perf_counter()
        # A retry batch already contains only failed original slots.  Generate
        # each one independently so one weak/duplicate draft cannot steer the
        # other repairs toward the same wording.  The calls remain bounded by
        # the same semaphore and are merged/revalidated globally afterward.
        local_count = int(local_question_plan.get("total_main_count") or 0)
        retry_batch_size = (
            1
            if local_question_plan.get("targeted_retry") and local_count <= 4
            else generation_batch_size
        )
        batch_specs = _generation_batch_plans(
            local_question_plan,
            max_batch_size=retry_batch_size,
        )
        semaphore = asyncio.Semaphore(generation_batch_concurrency)

        async def build_batch(
            batch_plan: dict[str, Any],
            original_indexes: list[int],
            batch_methods: list[str],
            *,
            recovery_attempt: int = 0,
        ) -> dict[str, Any]:
            effective_methods = batch_methods or list(local_interview_methods)
            batch_runtime_plan, batch_evidence_locks = (
                _planned_question_evidence_assignments(
                    question_plan=batch_plan,
                    interview_methods=effective_methods,
                    ncs_matches=ncs_matches,
                    ncs_ksa=ncs_ksa,
                )
            )
            batch_kwargs = dict(local_build_kwargs)
            batch_kwargs["question_plan"] = batch_runtime_plan
            batch_kwargs["interview_methods"] = effective_methods
            batch_kwargs["target_count_override"] = int(
                batch_runtime_plan.get("total_main_count") or 0
            )
            if recovery_attempt:
                # Preserve every completed batch and retry only the provider
                # batch whose response was malformed/truncated.  The retry is
                # still bounded by the request-wide deadline propagated by
                # ``to_thread`` and owns one semantic model request at most.
                batch_kwargs["max_model_requests"] = 1
                if active_generation_provider == "openai_api":
                    batch_kwargs["generation_model"] = openai_role_model(
                        "quality_regeneration"
                    )
                recovery_note = (
                    "[서버 배치 복구] 이전 응답이 JSON 형식, 필수 필드 또는 정확한 "
                    "문항 수 검사에 실패했습니다. 같은 잠금 슬롯을 새로 작성하고 "
                    "interview_questions 배열 개수를 정확히 지키세요."
                )
                batch_kwargs["extra_context"] = _join_generation_context(
                    recovery_note,
                    str(batch_kwargs.get("extra_context") or ""),
                )[:2000]
            async with semaphore:
                # ``run_in_executor`` does not propagate ContextVars.  The
                # request-wide generation deadline therefore used to vanish
                # exactly where the OpenRouter call started, allowing nested
                # timeout/recovery requests to outlive Vercel's proxy window.
                # ``to_thread`` copies the current context into the worker so
                # every provider attempt shares the same hard deadline.
                generated = await asyncio.to_thread(
                    build_jd_strategy_with_openai,
                    **batch_kwargs,
                )
            _require_institution_api_model_output(generated)
            evidence_assignment = _raw_model_evidence_assignment_report(
                generated,
                batch_evidence_locks,
            )
            raw_questions = _raw_model_question_texts(generated)
            raw_question_rows = [
                dict(item)
                for item in (generated.get("interview_questions") or [])
                if isinstance(item, dict)
            ]
            processed = _enrich_ai_authored_question_metadata(
                generated,
                question_plan=batch_runtime_plan,
                interview_methods=effective_methods,
                ncs_matches=ncs_matches,
                ncs_ksa=ncs_ksa,
            )
            processed = _preserve_ai_authored_question_surface(
                processed,
                raw_question_rows,
                provider=active_generation_provider,
            )
            processed = _audit_ai_authored_strategy_without_repair(
                processed,
                ncs_ksa,
            )
            processed = _apply_external_question_duplicate_safety(
                processed,
                avoid_questions,
            )
            return {
                "strategy": processed,
                "raw_questions": raw_questions,
                "raw_question_rows": raw_question_rows,
                "evidence_assignment": evidence_assignment,
                "original_indexes": list(original_indexes),
                "question_count": int(
                    batch_runtime_plan.get("total_main_count") or 0
                ),
                "batch_recovery_attempt": recovery_attempt,
            }

        try:
            batch_results_raw = await asyncio.gather(
                *(
                    build_batch(batch_plan, original_indexes, batch_methods)
                    for batch_plan, original_indexes, batch_methods in batch_specs
                ),
                return_exceptions=True,
            )
            failed_positions = [
                index
                for index, result in enumerate(batch_results_raw)
                if isinstance(result, BaseException)
            ]
            recoverable_batch_codes = {
                "model_response_not_object",
                "model_response_invalid_shape",
                "model_response_invalid_json",
                "model_response_truncated",
                "model_question_count_mismatch",
                "model_question_content_missing",
            }
            retry_positions = [
                index
                for index in failed_positions
                if str(batch_results_raw[index]) in recoverable_batch_codes
            ]
            if retry_positions:
                recovered_results = await asyncio.gather(
                    *(
                        build_batch(
                            *batch_specs[index],
                            recovery_attempt=1,
                        )
                        for index in retry_positions
                    ),
                    return_exceptions=True,
                )
                for index, recovered in zip(
                    retry_positions,
                    recovered_results,
                    strict=True,
                ):
                    batch_results_raw[index] = recovered

            remaining_errors = [
                result
                for result in batch_results_raw
                if isinstance(result, BaseException)
            ]
            if remaining_errors:
                raise remaining_errors[0]
            batch_results = [
                result for result in batch_results_raw if isinstance(result, dict)
            ]
        finally:
            generation_attempt_elapsed_ms.append(
                max(0, int(round((time.perf_counter() - attempt_started_at) * 1000)))
            )

        processed = dict(batch_results[0]["strategy"])
        merged_questions: list[dict[str, Any]] = []
        merged_raw_question_rows: list[dict[str, Any]] = []
        raw_questions: list[str] = []
        provider_generation_request_count = 0
        batch_question_counts: list[int] = []
        recovered_batch_count = 0
        for batch_result in batch_results:
            batch_strategy = batch_result["strategy"]
            merged_questions.extend(
                dict(item)
                for item in (batch_strategy.get("interview_questions") or [])
                if isinstance(item, dict)
            )
            merged_raw_question_rows.extend(batch_result["raw_question_rows"])
            raw_questions.extend(batch_result["raw_questions"])
            provider_generation_request_count += int(
                batch_strategy.get("provider_generation_request_count") or 0
            )
            batch_question_counts.append(int(batch_result["question_count"] or 0))
            recovered_batch_count += int(
                int(batch_result.get("batch_recovery_attempt") or 0) > 0
            )

        processed["interview_questions"] = merged_questions
        processed["interview_by_competency"] = _group_interview_questions_for_response(
            merged_questions
        )
        processed["question_plan_used"] = local_question_plan
        processed["interview_methods_used"] = list(local_interview_methods)
        processed["provider_generation_request_count"] = (
            provider_generation_request_count
        )
        processed["generation_batching"] = {
            "applied": len(batch_results) > 1,
            "policy": "locked-plan-parallel-batches-v1",
            "batch_count": len(batch_results),
            "batch_size_limit": generation_batch_size,
            "max_concurrency": generation_batch_concurrency,
            "batch_question_counts": batch_question_counts,
            "recovered_batch_count": recovered_batch_count,
        }
        processed = _audit_ai_authored_strategy_without_repair(
            processed,
            ncs_ksa,
        )
        processed = _apply_external_question_duplicate_safety(
            processed,
            avoid_questions,
        )
        processed["question_evidence_assignment"] = (
            _raw_model_evidence_assignment_report(
                {"interview_questions": merged_raw_question_rows},
                local_evidence_locks,
            )
        )
        return processed, raw_questions

    first_build_kwargs = dict(build_kwargs)
    history_digest_input = "\n".join(
        normalize_question_dedup_key(value)
        for value in avoid_questions[-50:]
        if normalize_question_dedup_key(value)
    )
    history_cycle = (
        int(hashlib.sha256(history_digest_input.encode("utf-8")).hexdigest()[:8], 16)
        if history_digest_input
        else 0
    )
    try:
        requested_cycle = max(0, int(generation_offset or 0))
    except (TypeError, ValueError):
        requested_cycle = 0
    first_build_kwargs.setdefault("avoid_questions", list(avoid_questions))
    first_build_kwargs.setdefault(
        "diversity_cycle",
        (history_cycle + requested_cycle) % 1_000_000,
    )
    first_build_kwargs.setdefault("max_model_requests", model_requests_per_batch)
    first_build_kwargs.setdefault("transport_max_attempts", 1)
    strategy, previous_questions = await run_once(
        first_build_kwargs,
        local_question_plan=runtime_question_plan,
        local_interview_methods=interview_methods,
        local_evidence_locks=planned_evidence_locks,
    )
    strategy = await attach_independent_ai_review(strategy, attempt_count=1)
    initial_authoring_model = str(
        strategy.get("provider_generation_model")
        or first_build_kwargs.get("generation_model")
        or openai_role_model("question_authoring")
    ).strip()
    first_generation_request_count = int(
        strategy.get("provider_generation_request_count") or 0
    )
    trigger_codes = _institution_question_rejection_codes(
        strategy,
        require_quality_metadata=True,
    )
    first_quality_issue_codes = _institution_question_quality_issue_codes(
        strategy,
        blocking_only=True,
    )
    if not trigger_codes:
        strategy["model_quality_retry"] = {
            "policy": _INSTITUTION_QUALITY_RETRY_POLICY,
            "provider": active_generation_provider,
            "attempted": False,
            "retry_count": 0,
            "attempt_count": 1,
            "outcome": "not_needed",
            "trigger_codes": [],
            "previous_candidate_count": 0,
            "evidence_lock_count": len(planned_evidence_locks),
            "provider_generation_request_count": first_generation_request_count,
            "provider_generation_request_limit": provider_generation_request_limit,
            "transport_attempt_limit_per_generation_request": 1,
        }
        strategy["question_release_status"] = "ai_quality_review_passed"
        strategy["human_review_required"] = False
        attach_model_orchestration(
            strategy,
            authoring_model=initial_authoring_model,
        )
        return with_generation_timing(strategy)

    first_hard_codes = _institution_hard_question_rejection_codes(
        strategy,
        require_quality_metadata=True,
    )
    if not quality_retry_enabled:
        logger.warning(
            "institution_question_quality_retry_disabled provider=%s codes=%s",
            active_generation_provider,
            ",".join(sorted(set(trigger_codes))),
        )
        raise InstitutionQuestionQualityRejected(
            rejection_diagnostics(strategy, attempt_count=1)
        )

    if any(code not in _INSTITUTION_RETRYABLE_QUALITY_CODES for code in trigger_codes):
        raise InstitutionQuestionQualityRejected(
            rejection_diagnostics(strategy, attempt_count=1)
        )

    first_question_count = len(
        [
            item
            for item in (strategy.get("interview_questions") or [])
            if isinstance(item, dict)
        ]
    )
    failed_indexes = _institution_hard_question_indexes(strategy)
    targeted_retry_forbidden = {
        "invalid_question_result",
        "result_error",
        "empty_question_set",
        "question_evidence_assignment_failed",
        "question_quality_report_missing",
        "question_quality_report_unclassified",
        "question_quality_unclassified_issue",
        "question_quality_orchestration_missing",
        "question_quality_orchestration_unclassified",
        "question_count_mismatch",
    }
    use_targeted_retry = bool(
        first_question_count > 1
        and failed_indexes
        and len(failed_indexes) < first_question_count
        and not (set(first_hard_codes) & targeted_retry_forbidden)
    )
    retry_runtime_plan = runtime_question_plan
    retry_interview_methods = interview_methods
    retry_original_indexes: list[int] = []
    retry_evidence_locks = list(planned_evidence_locks)
    if use_targeted_retry:
        retry_plan, retry_original_indexes, retry_interview_methods = (
            _targeted_quality_retry_plan(runtime_question_plan, failed_indexes)
        )
        retry_interview_methods = retry_interview_methods or list(interview_methods)
        retry_runtime_plan, retry_evidence_locks = _planned_question_evidence_assignments(
            question_plan=retry_plan,
            interview_methods=retry_interview_methods,
            ncs_matches=ncs_matches,
            ncs_ksa=ncs_ksa,
        )

    original_candidate_rows = [
        dict(item)
        for item in (strategy.get("interview_questions") or [])
        if isinstance(item, dict)
    ]
    content_retry_requested = bool(
        set(trigger_codes)
        & {
            "deterministic_fallback",
            "question_quality_report_failed",
            "question_quality_orchestration_failed",
            "precision_grounding_failed",
            "ai_quality_review_failed",
            "unsafe_question_surface",
            "follow_up_count_mismatch",
        }
    )
    rejected_previous_questions = (
        (
            [
                str(original_candidate_rows[index - 1].get("question") or "").strip()
                for index in retry_original_indexes
                if 1 <= index <= len(original_candidate_rows)
            ]
            if use_targeted_retry
            else list(previous_questions)
        )
        if content_retry_requested
        else []
    )
    final_retry_avoid_questions = [
        value
        for value in [*avoid_questions, *rejected_previous_questions]
        if str(value or "").strip()
    ]
    author_retry_avoid_questions = [
        value
        for value in [*avoid_questions, *rejected_previous_questions]
        if str(value or "").strip()
    ]

    logger.warning(
        "institution_question_quality_retry_started provider=%s codes=%s candidates=%s evidence_locks=%s targeted=%s retry_count=%s",
        active_generation_provider,
        ",".join(sorted(trigger_codes)),
        len(previous_questions),
        len(retry_evidence_locks),
        use_targeted_retry,
        len(retry_original_indexes) if use_targeted_retry else first_question_count,
    )
    retry_kwargs = dict(first_build_kwargs)
    retry_context = _quality_retry_context(
        trigger_codes=trigger_codes,
        previous_questions=rejected_previous_questions,
        evidence_locks=retry_evidence_locks,
        quality_issue_codes=first_quality_issue_codes,
        original_context=str(first_build_kwargs.get("extra_context") or ""),
        target_original_indexes=retry_original_indexes,
    )
    request_secret = str(first_build_kwargs.get("api_key_override") or "").strip()
    if request_secret:
        retry_context = retry_context.replace(request_secret, "[redacted]")
    retry_kwargs["extra_context"] = retry_context
    retry_kwargs["avoid_questions"] = author_retry_avoid_questions
    retry_kwargs["diversity_cycle"] = (
        int(first_build_kwargs.get("diversity_cycle") or 0) + 1
    ) % 1_000_000
    retry_kwargs["question_plan"] = retry_runtime_plan
    retry_kwargs["interview_methods"] = retry_interview_methods
    retry_kwargs["target_count_override"] = int(
        retry_runtime_plan.get("total_main_count") or 0
    )
    # The first builder invocation already owns its one provider-recovery
    # (slim prompt) budget.  The outer retry is specifically the single
    # quality regeneration, so it must not silently open another slim retry.
    retry_kwargs["max_model_requests"] = 1
    retry_kwargs["transport_max_attempts"] = 1
    if active_generation_provider == "openai_api":
        retry_kwargs["generation_model"] = openai_role_model(
            "quality_regeneration"
        )
    retried, _unused_questions = await run_once(
        retry_kwargs,
        local_question_plan=retry_runtime_plan,
        local_interview_methods=retry_interview_methods,
        local_evidence_locks=retry_evidence_locks,
    )
    retry_generation_request_count = int(
        retried.get("provider_generation_request_count") or 0
    )
    total_generation_request_count = first_generation_request_count + int(
        retry_generation_request_count
    )
    if use_targeted_retry:
        original_questions = [
            dict(item)
            for item in (strategy.get("interview_questions") or [])
            if isinstance(item, dict)
        ]
        retried_questions = [
            dict(item)
            for item in (retried.get("interview_questions") or [])
            if isinstance(item, dict)
        ]
        if len(retried_questions) != len(retry_original_indexes):
            raise InstitutionQuestionQualityRejected(
                rejection_diagnostics(retried, attempt_count=2)
            )
        for original_index, question in zip(
            retry_original_indexes,
            retried_questions,
            strict=True,
        ):
            original_questions[original_index - 1] = question
        merged = dict(strategy)
        merged["interview_questions"] = original_questions
        merged["question_plan_used"] = runtime_question_plan
        merged["interview_methods_used"] = list(interview_methods)
        merged["provider_generation_request_count"] = total_generation_request_count
        retry_assignment = retried.get("question_evidence_assignment")
        retry_mismatches = (
            list(retry_assignment.get("mismatched_indexes") or [])
            if isinstance(retry_assignment, dict)
            else []
        )
        mapped_mismatches = [
            retry_original_indexes[int(index) - 1]
            for index in retry_mismatches
            if str(index).isdigit()
            and 1 <= int(index) <= len(retry_original_indexes)
        ]
        merged["question_evidence_assignment"] = {
            "policy": _QUESTION_EVIDENCE_ASSIGNMENT_POLICY,
            "applicable": bool(planned_evidence_locks),
            "passed": not mapped_mismatches if planned_evidence_locks else True,
            "expected_count": len(planned_evidence_locks),
            "matched_count": len(planned_evidence_locks) - len(mapped_mismatches),
            "mismatch_count": len(mapped_mismatches),
            "mismatched_indexes": mapped_mismatches,
        }
        retried = _audit_ai_authored_strategy_without_repair(
            merged,
            ncs_ksa,
        )
        retried = _apply_external_question_duplicate_safety(
            retried,
            final_retry_avoid_questions,
        )
        retried["provider_generation_request_count"] = total_generation_request_count
        # The orchestration pass preserves evidence IDs, while the raw assignment
        # attestation above proves that the targeted model response did not drift.
        retried["question_evidence_assignment"] = merged[
            "question_evidence_assignment"
        ]
    else:
        retried = _apply_external_question_duplicate_safety(
            retried,
            final_retry_avoid_questions,
        )
    retried = await attach_independent_ai_review(
        retried,
        attempt_count=2,
        prior_questions_to_avoid=final_retry_avoid_questions,
    )
    quality_regeneration_model = str(
        retried.get("provider_generation_model")
        or retry_kwargs.get("generation_model")
        or ""
    ).strip()
    targeted_retry_metadata = (
        {
            "retry_scope": "failed_questions",
            "retried_question_count": len(retry_original_indexes),
            "retried_indexes": list(retry_original_indexes),
            "retry_evidence_lock_count": len(retry_evidence_locks),
        }
        if use_targeted_retry
        else {}
    )
    retry_codes = _institution_question_rejection_codes(
        retried,
        require_quality_metadata=True,
    )
    if retry_codes:
        hard_retry_codes = _institution_hard_question_rejection_codes(
            retried,
            require_quality_metadata=True,
        )
        logger.warning(
            "institution_question_quality_retry_failed provider=%s codes=%s hard_codes=%s",
            active_generation_provider,
            ",".join(sorted(set(retry_codes))),
            ",".join(sorted(set(hard_retry_codes))),
        )
        raise InstitutionQuestionQualityRejected(
            rejection_diagnostics(retried, attempt_count=2)
        )

    logger.info(
        "institution_question_quality_retry_passed provider=%s evidence_locks=%s",
        active_generation_provider,
        len(planned_evidence_locks),
    )
    retried["model_quality_retry"] = {
        "policy": _INSTITUTION_QUALITY_RETRY_POLICY,
        "provider": active_generation_provider,
        "attempted": True,
        "retry_count": 1,
        "attempt_count": 2,
        "outcome": "passed_after_retry",
        "trigger_codes": sorted(trigger_codes),
        "previous_candidate_count": len(previous_questions),
        "evidence_lock_count": len(planned_evidence_locks),
        "provider_generation_request_count": total_generation_request_count,
        "provider_generation_request_limit": provider_generation_request_limit,
        "transport_attempt_limit_per_generation_request": 1,
        **targeted_retry_metadata,
    }
    retried["question_release_status"] = "ai_quality_review_passed"
    retried["human_review_required"] = False
    attach_model_orchestration(
        retried,
        authoring_model=initial_authoring_model,
        regeneration_model=quality_regeneration_model,
        regeneration_used=True,
    )
    return with_generation_timing(retried)


def _verify_institution_openai_api(api_key: str) -> tuple[bool, str]:
    """One-shot credential check retained for internal diagnostics and tests.

    Public status and health endpoints never call this function because a
    request-scoped secret must not be submitted to those GET endpoints.
    """

    from app.services.openai_http import check_openai_connectivity_with_retries

    return check_openai_connectivity_with_retries(api_key, max_attempts=1)


@app.get("/api/generation-provider/status")
def generation_provider_status(
    request: Request,
    provider: str = Query(default=""),
) -> dict[str, Any]:
    """Report credential readiness without receiving or verifying a secret."""

    _reject_sensitive_query_params(request, destination="generation request body")
    active_provider = _request_generation_provider(provider) if str(provider or "").strip() else _configured_generation_provider()
    descriptor = _generation_provider_descriptor(active_provider)
    descriptor["configured_default"] = _configured_generation_provider()
    descriptor["supported_providers"] = list(request_supported_generation_providers())
    try:
        configured_maximum = int(
            str(
                os.getenv(
                    "GENERATION_MAX_MAIN_QUESTIONS",
                    str(GENERATION_MAX_MAIN_QUESTIONS),
                )
            ).strip()
        )
    except (TypeError, ValueError):
        configured_maximum = GENERATION_MAX_MAIN_QUESTIONS
    descriptor["generation_limits"] = {
        "max_main_questions_per_request": max(
            1,
            min(GENERATION_MAX_MAIN_QUESTIONS, configured_maximum),
        ),
        "max_follow_up_questions_per_main": 5,
        "max_ncs_details_per_request": 1,
        "max_interview_methods_per_request": 1,
        "request_budget_sec": int(_generation_request_budget_sec()),
    }
    return {
        **descriptor,
        "status": "key_required",
        "available": True,
        "authenticated": False,
        "credential_configured": False,
        "message": "본인의 OpenAI API 키(sk-...)를 입력해 주세요. 키는 해당 생성 요청에만 사용됩니다.",
        "login_command": "",
    }


@app.get("/health")
def health(request: Request) -> dict:
    del request
    mcp = ncs_mcp_status()
    mcp_ready = bool(mcp.get("configured") and mcp.get("reachable") and mcp.get("ksaAvailable"))
    return {
        "version": app.version,
        "status": "ok" if mcp_ready else "degraded",
        "keys": {
            "public_inst": bool(settings.public_inst_key()),
            "ncs": bool(settings.ncs_key()),
            "openai": False,
            "openai_institution_managed": False,
            "openai_request_scoped": True,
            "openai_authenticated": False,
        },
        "question_generation": _generation_provider_descriptor(),
        "ncs_source": "remote-mcp",
        "ncs_mcp": mcp,
    }


@app.get("/")
def ui() -> FileResponse:
    return FileResponse(UI_INDEX)


def _presentation_docx_text(value: Any, limit: int = 2000) -> str:
    return re.sub(r"[\x00-\x08\x0b\x0c\x0e-\x1f]", " ", str(value or "").strip())[:limit]


@app.post("/api/presentation-material/docx")
def presentation_material_docx(payload: dict[str, Any] = Body(...)) -> StreamingResponse:
    """Return the generated presentation packet as a real Word document."""
    packet = payload.get("presentation_material") if isinstance(payload, dict) else None
    if not isinstance(packet, dict) or packet.get("generated") is not True:
        raise HTTPException(status_code=422, detail="presentation_material is not a generated packet")
    document = Document()
    normal = document.styles["Normal"]
    normal.font.name = "Malgun Gothic"
    normal.font.size = Pt(10)
    document.add_heading(_presentation_docx_text(packet.get("title"), 180) or "발표 과제 자료", level=0)
    intro = document.add_paragraph()
    intro.add_run("생성 방식: ").bold = True
    intro.add_run("서버 자동 구성 · 공고문·직무기술서·NCS KSA 기반 · 사람 검토 필요")
    document.add_heading("발표 메인 과제", level=1)
    document.add_paragraph(_presentation_docx_text(packet.get("task_prompt"), 2200))
    document.add_paragraph(_presentation_docx_text(packet.get("scenario_label"), 240))
    document.add_paragraph(
        f"NCS 세분류: {_presentation_docx_text(packet.get('ncs_detail'), 120)} | "
        f"능력단위: {_presentation_docx_text(packet.get('competency'), 180)} | "
        f"평가 초점: {_presentation_docx_text(packet.get('focus'), 180)}"
    )
    document.add_heading("발표 과제 자료", level=1)
    rows = [row for row in (packet.get("case_materials") or []) if isinstance(row, dict)][:10]
    table = document.add_table(rows=1, cols=3)
    table.style = "Table Grid"
    for cell, label in zip(table.rows[0].cells, ("출처", "항목", "내용")):
        cell.text = label
    for row in rows:
        cells = table.add_row().cells
        cells[0].text = _presentation_docx_text(row.get("source"), 120)
        cells[1].text = _presentation_docx_text(row.get("field"), 160)
        cells[2].text = _presentation_docx_text(row.get("value"), 1200)
    facts = [_presentation_docx_text(value, 500) for value in (packet.get("case_facts") or []) if str(value or "").strip()][:6]
    if facts:
        document.add_heading("사례 사실", level=1)
        for fact in facts:
            document.add_paragraph(fact, style="List Bullet")
    document.add_heading("슬라이드 구성", level=1)
    for slide in (packet.get("slide_outline") or [])[:4]:
        if not isinstance(slide, dict):
            continue
        paragraph = document.add_paragraph(style="List Number")
        paragraph.add_run(_presentation_docx_text(slide.get("title"), 160)).bold = True
        paragraph.add_run(f" — {_presentation_docx_text(slide.get('instruction'), 700)}")
    document.add_heading("발표 제약조건", level=1)
    for constraint in (packet.get("constraints") or [])[:6]:
        document.add_paragraph(_presentation_docx_text(constraint, 600), style="List Bullet")
    document.add_heading("필수 산출물", level=1)
    for deliverable in (packet.get("required_deliverables") or [])[:6]:
        document.add_paragraph(_presentation_docx_text(deliverable, 600), style="List Bullet")
    document.add_heading("사용·검토 원칙", level=1)
    for rule in (packet.get("use_rules") or [])[:6]:
        document.add_paragraph(_presentation_docx_text(rule, 600), style="List Bullet")
    output = io.BytesIO()
    document.save(output)
    output.seek(0)
    return StreamingResponse(
        iter([output.getvalue()]),
        media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        headers={"Content-Disposition": 'attachment; filename="presentation-material.docx"'},
    )


@app.get("/api/ncs/sclass/options")
def ncs_sclass_options() -> dict:
    global _SCLASS_OPTIONS_CACHE
    if _SCLASS_OPTIONS_CACHE is not None:
        return {"count": len(_SCLASS_OPTIONS_CACHE), "items": _SCLASS_OPTIONS_CACHE}

    if not NCS_SCLASS_CSV.exists():
        raise HTTPException(status_code=404, detail=f"csv not found: {NCS_SCLASS_CSV}")

    by_name: dict[str, dict] = {}
    try:
        with NCS_SCLASS_CSV.open("r", encoding="utf-8-sig", newline="") as f:
            reader = csv.DictReader(f)
            for row in reader:
                name = str(row.get("NCS_SCLAS_CDNM", "")).strip()
                if not name or name in by_name:
                    continue
                by_name[name] = {
                    "name": name,
                    "ncs_code_no": str(row.get("NCS_CODE_NO", "")).strip(),
                    "lclass_code": str(row.get("NCS_LCLAS_CD", "")).strip(),
                    "mclass_code": str(row.get("NCS_MCLAS_CD", "")).strip(),
                    "sclass_code": str(row.get("NCS_SCLAS_CD", "")).strip(),
                }
    except Exception as e:
        raise _internal_http_error("ncs_catalog_read_failed", e) from e

    items = [by_name[k] for k in sorted(by_name.keys())]
    _SCLASS_OPTIONS_CACHE = items
    return {"count": len(items), "items": items}


def _find_sclass_code_tuple(sclass_name: str) -> dict[str, str] | None:
    def _norm_key(v: str) -> str:
        n = re.sub(r"\s+", "", str(v or "").strip()).lower()
        return re.sub(r"[·･ㆍ•∙⋅\-\_/|(),.\[\]{}]", "", n)

    name = str(sclass_name or "").strip()
    name_key = _norm_key(name)
    if not name or not NCS_SCLASS_CSV.exists():
        return None
    try:
        with NCS_SCLASS_CSV.open("r", encoding="utf-8-sig", newline="") as f:
            reader = csv.DictReader(f)
            for row in reader:
                row_name = str(row.get("NCS_SCLAS_CDNM", "")).strip()
                if row_name != name and _norm_key(row_name) != name_key:
                    continue
                l_cd = str(row.get("NCS_LCLAS_CD", "")).strip()
                m_cd = str(row.get("NCS_MCLAS_CD", "")).strip()
                s_cd = str(row.get("NCS_SCLAS_CD", "")).strip()
                if l_cd and m_cd and s_cd:
                    return {
                        "ncs_lclass_code": l_cd,
                        "ncs_mclass_code": m_cd,
                        "ncs_sclass_code": s_cd,
                    }
    except Exception:
        return None
    return None


@app.get("/api/ncs/units/options")
def ncs_unit_options(
    q: str = Query(default="", description="NCS detail classification search text"),
    limit: int = Query(default=300, ge=1, le=1000),
) -> dict:
    _require_ncs_mcp_url()
    term = str(q or "").strip()
    if not term:
        return {"count": 0, "items": [], "source": "ncs-mcp", "message": "Enter a confirmed NCS detail classification."}
    terms = _canonicalize_detail_lookup_terms(_parse_sclass_terms(term))
    try:
        items = search_units_by_detail(terms, max_units=limit)
        source = "ncs-mcp"
        message = ""
        if not items:
            items = suggest_units_by_text(terms, max_units=min(limit, 50))
            source = "ncs-mcp-suggest"
            message = "Exact detail-class match was not found. Review suggested NCS units manually."
    except NcsMcpError as exc:
        raise _internal_http_error(
            "ncs_mcp_search_failed",
            exc,
            status_code=502,
        ) from exc
    return {"count": len(items), "items": items, "source": source, "message": message}


@app.get("/api/ncs/sclass/ksa")
def ncs_sclass_ksa(
    sclass_name: str = Query(default="", alias="sclassName", description="소분류명(예: 총무)"),
    ncs_lclass_code: str = Query(default="", alias="ncsLclasCd"),
    ncs_mclass_code: str = Query(default="", alias="ncsMclasCd"),
    ncs_sclass_code: str = Query(default="", alias="ncsSclasCd"),
    max_units: int = Query(default=80, ge=1, le=200),
) -> dict:
    _require_legacy_ncs_api_enabled()
    l_cd = str(ncs_lclass_code or "").strip()
    m_cd = str(ncs_mclass_code or "").strip()
    s_cd = str(ncs_sclass_code or "").strip()
    s_nm = str(sclass_name or "").strip()

    # If caller passed only 소분류명, resolve code tuple from local catalog CSV.
    if not (l_cd and m_cd and s_cd) and s_nm:
        row = _find_sclass_code_tuple(s_nm)
        if row:
            l_cd = row["ncs_lclass_code"]
            m_cd = row["ncs_mclass_code"]
            s_cd = row["ncs_sclass_code"]

    if not (l_cd and m_cd and s_cd):
        raise HTTPException(
            status_code=400,
            detail="ncsLclasCd/ncsMclasCd/ncsSclasCd or sclassName is required",
        )

    try:
        result = fetch_ncs_ksa_by_sclass_code(
            ncs_lclass_code=l_cd,
            ncs_mclass_code=m_cd,
            ncs_sclass_code=s_cd,
            sclass_name=s_nm,
            max_units=max_units,
        )
        units = result.get("units", [])
        ksa = result.get("ksa", [])
        return {
            "query": {
                "sclassName": s_nm,
                "ncsLclasCd": l_cd,
                "ncsMclasCd": m_cd,
                "ncsSclasCd": s_cd,
            },
            "counts": {"units": len(units), "ksa": len(ksa)},
            "units": units,
            "ksa": ksa,
        }
    except Exception as e:
        raise _internal_http_error(
            "ncs_sclass_ksa_failed",
            e,
            status_code=502,
        ) from e


@app.get("/api/ops/metrics")
def ops_metrics(x_admin_token: str | None = Header(default=None)) -> dict:
    _require_admin(x_admin_token)
    return {"queue": queue.stats()}


def _require_admin(x_admin_token: str | None) -> None:
    if not settings.enable_admin_endpoints():
        raise HTTPException(status_code=403, detail="admin endpoints are disabled")
    expected = settings.admin_token()
    if not expected:
        raise HTTPException(status_code=403, detail="ADMIN_TOKEN is required")
    if not secrets.compare_digest(str(x_admin_token or ""), expected):
        raise HTTPException(status_code=401, detail="invalid admin token")


@app.get("/api/ops/quality-metrics")
def ops_question_quality_metrics(x_admin_token: str | None = Header(default=None)) -> dict:
    _require_admin(x_admin_token)
    try:
        return {
            "quality_policy_version": QUALITY_POLICY_VERSION,
            "metrics": question_quality_metrics(),
            "evidence_note": "운영 리뷰 DB에 기록된 생성·검토·에스컬레이션 이벤트 기준",
        }
    except Exception as exc:
        raise _internal_http_error("question_quality_metrics_failed", exc, status_code=503) from exc


@app.get("/api/ops/ax-readiness")
def ops_ax_readiness(x_admin_token: str | None = Header(default=None)) -> dict:
    _require_admin(x_admin_token)
    metrics = question_quality_metrics()
    mcp = ncs_mcp_status()
    state_path = BASE_DIR.parent / "reports" / "question_quality_loop" / "state.json"
    loop_state: dict[str, Any] = {}
    try:
        if state_path.is_file():
            candidate = json.loads(state_path.read_text(encoding="utf-8"))
            if isinstance(candidate, dict):
                loop_state = candidate
    except Exception:
        loop_state = {}
    decisions = metrics.get("decisions") if isinstance(metrics.get("decisions"), dict) else {}
    negative_reviews = int(decisions.get("reject") or 0) + int(decisions.get("needs_edit") or 0)
    signals = {
        "ready_asset_evidence": bool(mcp.get("configured") and mcp.get("reachable") and mcp.get("ksaAvailable")),
        "ready_asset_pilot": bool(mcp.get("configured")),
        "ready_asset_ref": "GET /health · NCS_MCP configured/reachable/ksaAvailable",
        "enabled_output_evidence": int(metrics.get("runs") or 0) > 0,
        "enabled_output_pilot": True,
        "enabled_output_ref": f"question_quality_runs={int(metrics.get('runs') or 0)}",
        "enabled_review_evidence": int(metrics.get("reviews") or 0) > 0,
        "enabled_review_pilot": True,
        "enabled_review_ref": f"question_quality_reviews={int(metrics.get('reviews') or 0)}",
        "first_redesign_evidence": False,
        "first_redesign_pilot": True,
        "first_redesign_ref": "docs/QUESTION_QUALITY_ORCHESTRATION.md",
        "first_auto_evidence": False,
        "first_auto_pilot": True,
        "first_auto_ref": "scripts/run_question_quality_loop.py (운영 자동처리 증거는 별도 필요)",
        "first_escalation_evidence": int(metrics.get("escalation_runs") or 0) > 0,
        "first_escalation_pilot": True,
        "first_escalation_ref": f"escalation_runs={int(metrics.get('escalation_runs') or 0)}",
        "first_exception_evidence": False,
        "first_exception_pilot": True,
        "first_exception_ref": "승인·반려는 구현됨; 이의·롤백 훈련 로그는 아직 없음",
        "first_metrics_evidence": int(metrics.get("runs") or 0) > 0,
        "first_metrics_pilot": True,
        "first_metrics_ref": "GET /api/ops/quality-metrics",
        "first_feedback_evidence": negative_reviews > 0,
        "first_feedback_pilot": True,
        "first_feedback_ref": f"reject_or_needs_edit_reviews={negative_reviews}",
        "native_core_evidence": False,
        "native_core_pilot": True,
        "native_core_ref": "AI 중단 영향분석·핵심 KPI 연계 운영 증거 필요",
        "native_loop_evidence": bool(int(metrics.get("active_eval_cases") or 0) > 0 and loop_state.get("success") is True),
        "native_loop_pilot": True,
        "native_loop_ref": (
            f"active_eval_cases={int(metrics.get('active_eval_cases') or 0)} · "
            f"latest_loop_success={loop_state.get('success') is True}"
        ),
        "native_resilience_evidence": False,
        "native_resilience_pilot": True,
        "native_resilience_ref": "템플릿 강등 경로는 구현됨; 실제 장애훈련·RTO 기록 필요",
        "native_scope_evidence": False,
        "native_scope_pilot": True,
        "native_scope_ref": "적용·제외 범위, SLA, 감사 주기의 운영 승인 문서 필요",
    }
    return {
        "assessment": assess_ax_readiness(signals),
        "quality_metrics": metrics,
        "ncs_mcp": mcp,
        "latest_quality_loop": loop_state,
        "disclaimer": "AX체크 기준의 자체 점검이며 공식 인증·감사 결과가 아닙니다.",
    }


@app.get("/api/quality/runs/{run_id}")
def get_quality_run_endpoint(
    run_id: str,
    x_review_token: str | None = Header(default=None, alias="X-Review-Token"),
) -> dict:
    if not verify_question_quality_run_token(str(run_id or "").strip(), str(x_review_token or "")):
        raise HTTPException(status_code=401, detail="invalid review token")
    run = get_question_quality_run(str(run_id or "").strip())
    if not run:
        raise HTTPException(status_code=404, detail="quality run not found")
    return {"status": "success", "data": run}


@app.post("/api/quality/runs/{run_id}/review")
def review_quality_run_endpoint(
    run_id: str,
    request: Request,
    payload: dict[str, Any] | None = Body(default=None),
) -> dict:
    candidate = dict(payload or {})
    candidate["run_id"] = str(run_id or "").strip()
    try:
        feedback = sanitize_feedback_payload(candidate)
        if feedback["verdict"] in {"reject", "needs_edit"}:
            if not feedback.get("issue_codes"):
                raise ValueError("reject/needs_edit feedback requires at least one issue code")
            if not str(feedback.get("question_text") or "").strip():
                raise ValueError("reject/needs_edit feedback requires question_text for the improvement loop")
        result = record_question_quality_review(feedback)
    except LookupError as exc:
        raise HTTPException(status_code=404, detail=str(exc)) from exc
    except PermissionError as exc:
        raise HTTPException(status_code=401, detail=str(exc)) from exc
    except QuestionQualityReviewConflictError as exc:
        raise HTTPException(status_code=409, detail=str(exc)) from exc
    except ValueError as exc:
        raise HTTPException(status_code=422, detail=str(exc)) from exc
    _record_audit_event(
        request,
        action=f"question_quality_{feedback['verdict']}",
        resource_type="question_quality_run",
        resource_id=str(run_id or ""),
    )
    return {"status": "success", "data": result}


@app.post("/api/quality/reviews/{review_id}/promote-to-eval")
def promote_quality_review_endpoint(
    review_id: int,
    payload: dict[str, Any] | None = Body(default=None),
    x_admin_token: str | None = Header(default=None),
) -> dict:
    _require_admin(x_admin_token)
    case_type = str((payload or {}).get("case_type") or "").strip()
    if case_type not in {"golden", "negative", "regression"}:
        raise HTTPException(status_code=422, detail="case_type must be golden, negative, or regression")
    try:
        result = promote_question_quality_eval_case(review_id, case_type, QUALITY_POLICY_VERSION)
    except LookupError as exc:
        raise HTTPException(status_code=404, detail=str(exc)) from exc
    except ValueError as exc:
        raise HTTPException(status_code=422, detail=str(exc)) from exc
    return {"status": "success", "data": result}


@app.post("/api/quality/runs/{run_id}/questions/{question_hash}/rollback-review")
def rollback_quality_review_endpoint(
    run_id: str,
    question_hash: str,
    request: Request,
    payload: dict[str, Any] | None = Body(default=None),
) -> dict:
    candidate = dict(payload or {})
    candidate.update(
        {
            "run_id": str(run_id or "").strip(),
            "question_hash": str(question_hash or "").strip(),
            "verdict": "approve",
            "issue_codes": [],
        }
    )
    try:
        sanitized = sanitize_feedback_payload(candidate)
        if sanitized.get("expected_review_id") == 0:
            raise ValueError("rollback expected_review_id must be greater than zero")
        result = rollback_question_quality_review(
            run_id=sanitized["run_id"],
            review_token=str(sanitized.get("review_token") or ""),
            question_hash=sanitized["question_hash"],
            reviewer_ref=str(sanitized.get("reviewer_ref") or "local-reviewer"),
            note=str(sanitized.get("note") or ""),
            expected_review_id=sanitized.get("expected_review_id"),
        )
    except LookupError as exc:
        raise HTTPException(status_code=404, detail=str(exc)) from exc
    except PermissionError as exc:
        raise HTTPException(status_code=401, detail=str(exc)) from exc
    except QuestionQualityReviewConflictError as exc:
        raise HTTPException(status_code=409, detail=str(exc)) from exc
    except ValueError as exc:
        raise HTTPException(status_code=422, detail=str(exc)) from exc
    _record_audit_event(
        request,
        action="question_quality_review_rollback",
        resource_type="question_quality_run",
        resource_id=str(run_id or ""),
    )
    return {"status": "success", "data": result}


@app.get("/api/quality/eval-cases")
def list_quality_eval_cases_endpoint(
    active_only: bool = Query(default=True),
    x_admin_token: str | None = Header(default=None),
) -> dict:
    _require_admin(x_admin_token)
    return {"status": "success", "data": list_question_quality_eval_cases(active_only=active_only)}


@app.post("/api/admin/sync/public-inst")
def admin_sync_public_inst(
    max_pages: int = Query(default=5, ge=1, le=100),
    num_of_rows: int = Query(default=100, ge=1, le=1000),
    x_admin_token: str | None = Header(default=None),
) -> dict:
    _require_admin(x_admin_token)
    try:
        return sync_public_institutions(max_pages=max_pages, num_of_rows=num_of_rows)
    except Exception as e:
        raise _internal_http_error(
            "public_institution_sync_failed",
            e,
            status_code=502,
        ) from e


@app.post("/api/admin/sync/ncs")
def admin_sync_ncs(
    path: str = Query(..., description="NCS API relative path for units"),
    pages: int = Query(default=20, ge=1, le=500),
    num_of_rows: int = Query(default=100, ge=1, le=1000),
    x_admin_token: str | None = Header(default=None),
) -> dict:
    _require_admin(x_admin_token)
    _require_legacy_ncs_api_enabled()
    try:
        return sync_ncs_units(path=path, pages=pages, num_of_rows=num_of_rows)
    except Exception as e:
        raise _internal_http_error("ncs_sync_failed", e, status_code=502) from e


@app.get("/api/integrations/public-inst/{resource}")
def public_inst_proxy(
    resource: str,
    page_no: int = Query(default=1, ge=1),
    num_of_rows: int = Query(default=20, ge=1, le=100),
    data_type: str = Query(default="json", pattern="^(json|xml|JSON|XML)$"),
) -> dict:
    try:
        return fetch_public_inst(resource=resource, page_no=page_no, num_of_rows=num_of_rows, data_type=data_type)
    except ValueError as e:
        raise HTTPException(status_code=400, detail="invalid public institution request") from e
    except RuntimeError as e:
        raise _internal_http_error("public_institution_proxy_failed", e) from e
    except Exception as e:
        raise _internal_http_error(
            "public_institution_proxy_failed",
            e,
            status_code=502,
        ) from e


@app.get("/api/integrations/ncs")
def ncs_proxy(
    path: str = Query(..., description="NCS API relative path"),
    page_no: int | None = Query(default=None, ge=1),
    num_of_rows: int | None = Query(default=None, ge=1, le=1000),
    data_type: str | None = Query(default=None, alias="type"),
    ncs_job_cd: str | None = Query(default=None),
    ncs_cl_cd: str | None = Query(default=None),
    x_admin_token: str | None = Header(default=None),
) -> dict:
    _require_legacy_ncs_api_enabled()
    _require_admin(x_admin_token)
    query: dict = {}
    if page_no is not None:
        query["pageNo"] = page_no
    if num_of_rows is not None:
        query["numOfRows"] = num_of_rows
    if data_type:
        query["type"] = data_type
    if ncs_job_cd:
        query["ncsJobCd"] = ncs_job_cd
    if ncs_cl_cd:
        query["ncsClCd"] = ncs_cl_cd
    try:
        return fetch_ncs(path=path, query=query)
    except ValueError as e:
        raise HTTPException(status_code=400, detail="invalid legacy NCS request") from e
    except RuntimeError as e:
        raise _internal_http_error("legacy_ncs_proxy_failed", e) from e
    except Exception as e:
        raise _internal_http_error(
            "legacy_ncs_proxy_failed",
            e,
            status_code=502,
        ) from e


@app.get("/api/integrations/ncs/highschool")
def ncs_highschool_proxy(
    mcd_nm: str = Query(..., alias="mcdNm", description="고교 교과목 명"),
    targ_yy: str = Query(..., alias="targYy", description="개정년도 (2015/2018)"),
    cd_name: str | None = Query(default=None, alias="cdName", description="고교 능력단위명(옵션)"),
    return_type: str = Query(default="xml", alias="returnType", pattern="^(xml|json|XML|JSON)$"),
) -> dict:
    _require_legacy_ncs_api_enabled()
    try:
        return fetch_ncs_highschool_course(
            mcd_nm=mcd_nm,
            targ_yy=targ_yy,
            cd_name=cd_name,
            return_type=return_type,
        )
    except ValueError as e:
        raise HTTPException(status_code=400, detail="invalid NCS high-school request") from e
    except RuntimeError as e:
        raise _internal_http_error("ncs_highschool_proxy_failed", e) from e
    except Exception as e:
        raise _internal_http_error(
            "ncs_highschool_proxy_failed",
            e,
            status_code=502,
        ) from e


@app.get("/api/ncs/diagnose")
def ncs_diagnose(sample_job_cd: str = Query(default="02020101")) -> dict:
    try:
        _ = sample_job_cd
        status = ncs_mcp_status()
        ok = bool(status.get("reachable") and status.get("ksaAvailable"))
        return {
            "provider": "ncs-mcp",
            "ok": ok,
            "ncs_mcp": status,
            "message": "NCS_MCP local NCS DB server is ready" if ok else "NCS_MCP local NCS DB server is not ready",
        }
    except Exception as e:
        raise _internal_http_error("ncs_diagnose_failed", e) from e


def _extract_result_items(data: dict) -> list[dict]:
    if isinstance(data.get("result"), list):
        return data["result"]
    body = (data.get("response") or {}).get("body") or {}
    items = (body.get("items") or {}).get("item")
    if items is None:
        return []
    if isinstance(items, list):
        return items
    return [items]


def _valid_option_text(v: str) -> bool:
    if not v:
        return False
    bad_markers = ["???", "占", "챙"]
    if any(b in v for b in bad_markers):
        return False
    return True


@app.get("/api/alio/recommend")
def alio_recommend(
    desired_job: str = Query(..., min_length=2),
    desired_region: str = Query(default=""),
    strengths: str = Query(..., min_length=5),
    pages: int = Query(default=2, ge=1, le=5),
    per_page: int = Query(default=100, ge=10, le=300),
) -> dict:
    strengths = _validate_generation_text_input(
        strengths,
        field_name="strengths",
        max_chars=_MAX_STRENGTHS_CHARS,
    )
    candidates: list[dict] = []
    # 1) Try recruitment API
    try:
        for page in range(1, pages + 1):
            resp = fetch_recruitment("list", page_no=page, num_of_rows=per_page, data_type="json")
            for row in _extract_result_items(resp.get("data", {})):
                pid = str(row.get("recrtPbancTtlPc", "")) + "_" + str(row.get("instCd", ""))
                candidates.append(
                    {
                        "posting_id": pid,
                        "title": row.get("recrtPbancTtl", "") or row.get("title", ""),
                        "institution_name": row.get("instNm", "") or row.get("institutionName", ""),
                        "region": row.get("workRgnNm", "") or row.get("ctpvNm", ""),
                        "r6000": row.get("ncsLclasCd", "") or "R6000_MANAGEMENT",
                        "description": row.get("recrtPbancCn", "") or row.get("dutyCn", ""),
                        "jd_text": row.get("dutyCn", "") or row.get("jobDc", ""),
                    }
                )
            if not _extract_result_items(resp.get("data", {})):
                break
    except Exception:
        candidates = []

    # 2) Fallback to DB postings if recruitment API not available
    if not candidates:
        local = repo_recommend_postings(desired_job=desired_job, desired_region=desired_region, limit=30)
        for p in local:
            details = repo_get_posting(p["posting_id"]) or {}
            req_text = " ".join([r.get("item", "") for r in details.get("requirements_top", [])])
            candidates.append(
                {
                    "posting_id": p["posting_id"],
                    "title": p["title"],
                    "institution_name": p.get("institution_name", ""),
                    "region": p.get("region_code", ""),
                    "r6000": p.get("r6000", "R6000_MANAGEMENT"),
                    "description": req_text,
                    "jd_text": req_text,
                }
            )

    ranked = rank_postings_with_openai(desired_job, desired_region, strengths, candidates)
    result_items = []
    for item in ranked:
        pid = str(item.get("posting_id"))
        src = next((c for c in candidates if str(c.get("posting_id")) == pid), None)
        if src:
            _ALIO_CACHE[pid] = src
        result_items.append(item)
    return {"count": len(result_items), "items": result_items}


@app.post("/api/alio/attachments")
def alio_attachments(payload: dict[str, Any]) -> dict[str, Any]:
    """Inspect a public ALIO list/detail URL and return selectable metadata.

    This endpoint only discovers metadata.  A selected notice/JD may then be
    fetched through ``/api/alio/attachment`` and is still sent through the
    existing upload/review endpoints where the signed human-review gate,
    upload-size limit, and document parser apply.
    """
    raw_url = str(payload.get("url", "")).strip() if isinstance(payload, dict) else ""
    if not raw_url:
        raise HTTPException(status_code=422, detail="url is required")
    try:
        return inspect_alio_url(raw_url)
    except AlioIngestionError as exc:
        status = 504 if exc.code == "upstream_timeout" else 502 if exc.retryable else 422
        raise HTTPException(
            status_code=status,
            detail={"code": exc.code, "message": exc.message, "retryable": exc.retryable},
        ) from exc


@app.post("/api/alio/attachment")
def alio_attachment(payload: dict[str, Any], request: Request) -> StreamingResponse:
    """Proxy one public ALIO document into the normal browser upload flow."""

    raw_url = str(payload.get("url", "")).strip() if isinstance(payload, dict) else ""
    expected_name = str(payload.get("name", "")).strip() if isinstance(payload, dict) else ""
    if not raw_url:
        raise HTTPException(status_code=422, detail="url is required")
    try:
        downloaded = download_alio_attachment(
            raw_url,
            expected_name=expected_name,
            max_bytes=settings.max_upload_bytes(),
        )
    except AlioIngestionError as exc:
        status = 504 if exc.code == "upstream_timeout" else 413 if exc.code == "attachment_too_large" else 502 if exc.retryable else 422
        raise HTTPException(
            status_code=status,
            detail={"code": exc.code, "message": exc.message, "retryable": exc.retryable},
        ) from exc

    _record_audit_event(
        request,
        action="alio_attachment_import",
        resource_type="public_alio_attachment",
        resource_id=_sha256_text(downloaded.url)[:24],
    )
    ascii_name = re.sub(r"[^A-Za-z0-9._-]+", "_", downloaded.filename).strip("._") or "alio_attachment"
    disposition = (
        f'attachment; filename="{ascii_name[:100]}"; '
        f"filename*=UTF-8''{quote(downloaded.filename, safe='')}"
    )
    return StreamingResponse(
        io.BytesIO(downloaded.data),
        media_type=downloaded.content_type or "application/octet-stream",
        headers={
            "Content-Disposition": disposition,
            "Content-Length": str(len(downloaded.data)),
            "Cache-Control": "no-store",
            "X-Content-Type-Options": "nosniff",
        },
    )


@app.get("/api/alio/options")
def alio_options(
    pages: int = Query(default=1, ge=1, le=3),
    per_page: int = Query(default=50, ge=10, le=150),
) -> dict:
    regions: set[str] = set()
    jobs: set[str] = set()

    try:
        for page in range(1, pages + 1):
            resp = fetch_public_inst("list", page_no=page, num_of_rows=per_page, data_type="json", timeout_sec=4.0)
            rows = _extract_result_items(resp.get("data", {}))
            if not rows:
                break
            for r in rows:
                ctpv = str(r.get("ctpvNm", "")).strip()
                if _valid_option_text(ctpv):
                    regions.add(ctpv)
    except Exception:
        pass

    try:
        for page in range(1, pages + 1):
            resp = fetch_recruitment("list", page_no=page, num_of_rows=per_page, data_type="json", timeout_sec=4.0)
            rows = _extract_result_items(resp.get("data", {}))
            if not rows:
                break
            for r in rows:
                title = str(r.get("recrtPbancTtl", "") or r.get("title", "")).strip()
                ncs_name = str(r.get("ncsLclasNm", "")).strip()
                if _valid_option_text(title):
                    jobs.add(title)
                if _valid_option_text(ncs_name):
                    jobs.add(ncs_name)
                region = str(r.get("workRgnNm", "") or r.get("ctpvNm", "")).strip()
                if _valid_option_text(region):
                    regions.add(region)
    except Exception:
        pass

    if not jobs:
        for p in repo_list_postings()[:100]:
            t = str(p.get("title", "")).strip()
            if _valid_option_text(t):
                jobs.add(t)
    if not regions:
        regions.update(["서울", "부산", "대구", "인천", "광주", "대전", "울산", "경기", "강원"])

    return {"jobs": sorted(jobs)[:200], "regions": sorted(regions)[:100]}


@app.post("/api/alio/strategy")
def alio_strategy(payload: dict) -> dict:
    desired_job = str(payload.get("desired_job", "")).strip()
    strengths = _validate_generation_text_input(
        payload.get("strengths", ""),
        field_name="strengths",
        max_chars=_MAX_STRENGTHS_CHARS,
    )
    posting_id = str(payload.get("posting_id", "")).strip()
    if not (desired_job and strengths and posting_id):
        raise HTTPException(status_code=400, detail="desired_job, strengths, posting_id are required")

    posting = _ALIO_CACHE.get(posting_id)
    if not posting:
        # Try local posting id fallback
        try:
            details = repo_get_posting(int(posting_id))
        except Exception:
            details = None
        if not details:
            raise HTTPException(status_code=404, detail="selected posting not found in cache")
        posting = {
            "posting_id": posting_id,
            "title": details.get("title", ""),
            "institution_name": details.get("institution_name", ""),
            "region": details.get("codes", {}).get("R3000", ""),
            "r6000": details.get("codes", {}).get("R6000", "R6000_MANAGEMENT"),
            "jd_text": " ".join([r.get("item", "") for r in details.get("requirements_top", [])]),
        }

    strategy = build_strategy_with_openai(
        desired_job=desired_job,
        strengths=strengths,
        posting_data=posting,
    )
    return {"posting": posting, "strategy": strategy}


@app.get("/api/ncs/sclass-list")
def get_ncs_sclass_list() -> dict:
    """NCS 소분류 283개 목록 반환 (자동완성용)."""
    cats: list[str] = []
    if NCS_SCLASS_CSV.exists():
        with open(NCS_SCLASS_CSV, encoding="utf-8-sig") as f:
            for row in csv.DictReader(f):
                n = row.get("NCS_SCLAS_CDNM", "").strip()
                if n and n not in cats:
                    cats.append(n)
    return {"ncs_cats": cats}


@app.post("/api/jd/extract-sclass")
async def extract_sclass_endpoint(jd_file: UploadFile = File(...)) -> dict:
    """
    PDF 직무기술서를 받아 소분류 목록을 추출해서 반환.
    matched  : NCS 282개 사전에 있는 공식 소분류
    unmatched: 사전에 없는 자체 명칭 (NCS 미개발 등)
    """
    name = (jd_file.filename or "").lower()
    _reject_hwp_upload("직무기술서 파일", jd_file.filename)
    data = await _read_upload_limited(jd_file, "jd_file")
    if not data:
        raise HTTPException(status_code=400, detail="uploaded file is empty")

    if name.endswith(".pdf"):
        try:
            return extract_sclass_from_pdf_bytes(data, filename=(jd_file.filename or ""))
        except RuntimeError as e:
            raise _internal_http_error("jd_sclass_extract_failed", e) from e
    elif name.endswith(".txt"):
        text = data.decode("utf-8", errors="ignore")
        return extract_sclass_from_text(text, filename=(jd_file.filename or ""))
    else:
        raise HTTPException(status_code=400, detail="only .pdf or .txt supported")


@app.post("/api/jd/parse-review")
async def parse_jd_review_endpoint(request: Request, jd_file: UploadFile = File(...)) -> dict:
    """Parse a JD with Kordoc and return editable human-review fields."""

    data = await _read_upload_limited(jd_file, "jd_file")
    if not data:
        raise HTTPException(status_code=400, detail="uploaded file is empty")
    parsed = _parse_upload_document(data, jd_file.filename or "", "jd_file")
    structured = structure_job_description(parsed, filename=jd_file.filename or "")
    structured_fields = structured.get("fields") if isinstance(structured.get("fields"), dict) else None
    if structured_fields is not None:
        _sanitize_parse_review_ability_artifacts(structured)
        _recover_hangul_fallback_detail_candidates(parsed, structured_fields)
        # Kordoc can return an entire numbered table row as one candidate.
        # Normalize it before the browser renders the review dropdown and
        # before the signed review session is created.
        normalized_detail_candidates = _canonicalize_detail_lookup_terms(
            _parse_sclass_terms(
                "\n".join(
                    str(value or "").strip()
                    for value in (structured_fields.get("ncs_detail_candidates") or [])
                    if str(value or "").strip()
                )
            )
        )
        # Kordoc's table projection can lose the classification columns on
        # some ALIO/HWP-to-PDF layouts.  Keep the normal Kordoc result as the
        # primary source, but recover official NCS labels from the structural
        # PDF parser when that result is empty.  This is deliberately
        # fallback-only so a human-review candidate list is never silently
        # replaced by a heuristic list when Kordoc already found candidates.
        structured_fields["ncs_detail_candidates"] = normalized_detail_candidates
        # A positive "declared no mapping" signal means the document itself
        # explicitly says that no NCS detail classification applies. Re-running
        # the slower structural PDF parser cannot recover a valid detail in that
        # case and risks promoting a neighbouring classification column. Keep
        # structural recovery for genuinely missing/ambiguous Kordoc output.
        kordoc_declared_no_mapping = bool(
            structured_fields.get("ncs_detail_absence_declared_no_mapping")
        )
        if (
            not normalized_detail_candidates
            and not kordoc_declared_no_mapping
            and str(jd_file.filename or "").lower().endswith(".pdf")
        ):
            try:
                structural_result = extract_sclass_from_pdf_bytes(
                    data,
                    filename=jd_file.filename or "",
                )
                # The legacy structural endpoint returns 소분류 matches for
                # compatibility, but the review dropdown must contain the
                # actual 세분류 column. Prefer the table-level detail trace;
                # only fall back to legacy matches when no detail column was
                # recoverable at all.
                detail_candidates = _canonicalize_detail_lookup_terms(
                    _parse_sclass_terms(
                        "\n".join(
                            str(value or "").strip()
                            for value in (structural_result.get("detail_candidates") or [])
                            if str(value or "").strip()
                        )
                    )
                )
                detail_table_found = bool(structural_result.get("detail_table_found"))
                recovered_candidates = _canonicalize_detail_lookup_terms(
                    _parse_sclass_terms(
                        "\n".join(
                            str(value or "").strip()
                            for value in (structural_result.get("matched") or [])
                            if str(value or "").strip()
                        )
                    )
                )
                # An explicit 세분류 column is authoritative, including when
                # its value is "해당사항 없음". Do not fall back to the
                # neighbouring 소분류 (for example 법무/인사·조직) in that
                # case; an empty list correctly signals human review.
                if detail_table_found or detail_candidates:
                    recovered_candidates = detail_candidates
                if recovered_candidates:
                    structured_fields["ncs_detail_candidates"] = recovered_candidates
                    detail_evidence = structural_result.get("detail_candidate_evidence") or []
                    if detail_candidates and isinstance(detail_evidence, list):
                        evidence_by_key = {
                            _norm_detail_coverage_key(str(item.get("label", ""))): item
                            for item in detail_evidence
                            if isinstance(item, dict) and str(item.get("label", "")).strip()
                        }
                        structured_fields["ncs_detail_source"] = "pdf_table_detail"
                        structured_fields["ncs_detail_candidate_evidence"] = [
                            {
                                "detail": candidate,
                                "text": candidate,
                                "page": int(
                                    (evidence_by_key.get(_norm_detail_coverage_key(candidate)) or {}).get("page", 0)
                                    or 0
                                ),
                                "source": "pdf_table_detail",
                                "raw": candidate,
                                "snippet": candidate,
                            }
                            for candidate in recovered_candidates
                        ]
                    else:
                        structured_fields["ncs_detail_source"] = "pdf_structural_fallback"
                        structured_fields["ncs_detail_candidate_evidence"] = [
                            {
                                "detail": candidate,
                                "text": candidate,
                                "page": 0,
                                "source": "pdf_structural_fallback",
                                "raw": candidate,
                                "snippet": candidate,
                            }
                            for candidate in recovered_candidates
                        ]
                elif detail_table_found:
                    structured_fields["ncs_detail_candidates"] = []
                    structured_fields["ncs_detail_source"] = "pdf_table_detail_empty"
                    structured_fields["ncs_detail_candidate_evidence"] = []
                    structured_fields["ncs_detail_absence_reason"] = ""
                    structured_fields["ncs_detail_absence_state"] = ""
                    structured_fields["ncs_detail_absence_evidence"] = ""
                    structured_fields["ncs_detail_absence_filtered_candidate_reason"] = ""
                    structured_fields["ncs_detail_absence_saw_ncs_table"] = True
                    structured_fields["ncs_detail_absence_saw_detail_header"] = True
                    structured_fields["ncs_detail_absence_blank_or_dash_detail_cell"] = False
                    structured_fields["ncs_detail_absence_declared_no_mapping"] = False
            except Exception as exc:
                # Structural recovery is an availability enhancement.  A
                # Kordoc response (including an empty review state) must still
                # be returned when the optional parser is unavailable.
                logger.warning("jd_parse_review_structural_sclass_fallback_failed: %s", exc)
        # A locally collected ALIO profile can suggest official labels when a
        # document has no explicit 세분류 cell.  These rows are trained only
        # from single-label, explicitly classified ALIO JDs and remain inside
        # the same signed human-review gate; an explicit "NCS 미개발/매핑 없음"
        # declaration always suppresses the suggestion path.
        if (
            not structured_fields.get("ncs_detail_candidates")
            and not structured_fields.get("ncs_detail_absence_declared_no_mapping")
        ):
            profile_context_parts = [
                str(parsed.get("markdown") or "")[:60000],
                *[
                    str(value or "")
                    for key in ("duties", "knowledge", "skills", "attitudes", "qualifications", "preferences")
                    for value in (
                        structured_fields.get(key)
                        if isinstance(structured_fields.get(key), list)
                        else [structured_fields.get(key)]
                    )
                    if str(value or "").strip()
                ],
            ]
            profile_suggestions = suggest_sclass_from_profile(
                "\n".join(profile_context_parts),
                max_items=5,
            )
            structured_fields["ncs_detail_suggestions"] = profile_suggestions
            if profile_suggestions:
                suggested_names = _canonicalize_detail_lookup_terms(
                    [
                        str(item.get("sclass_name") or "").strip()
                        for item in profile_suggestions
                        if str(item.get("sclass_name") or "").strip()
                    ]
                )
                structured_fields["ncs_detail_candidates"] = suggested_names
                structured_fields["ncs_detail_source"] = "alio_corpus_review_suggestion"
                suggestion_by_name = {
                    _norm_sclass_key(str(item.get("sclass_name") or "")): item
                    for item in profile_suggestions
                }
                structured_fields["ncs_detail_candidate_evidence"] = [
                    {
                        "detail": name,
                        "text": name,
                        "page": 0,
                        "source": "alio_corpus_review_suggestion",
                        "raw": name,
                        "snippet": str(
                            (suggestion_by_name.get(_norm_sclass_key(name)) or {}).get("evidence")
                            or "ALIO 명시 세분류 코퍼스 어휘 매칭"
                        )[:500],
                        "confidence": float(
                            (suggestion_by_name.get(_norm_sclass_key(name)) or {}).get("confidence")
                            or 0.0
                        ),
                        "review_required": True,
                    }
                    for name in suggested_names
                ]
        detail_states = classify_official_detail_names(
            list(structured_fields.get("ncs_detail_candidates") or []),
            self_developed_names=list(
                structured_fields.get("ncs_self_developed_detail_candidates")
                or []
            ),
        )
        ability_states = classify_official_ability_unit_names(
            list(structured_fields.get("ability_units") or []),
            selected_detail_names=list(
                structured_fields.get("ncs_detail_candidates") or []
            ),
        )
        structured_fields["ncs_detail_mapping_states"] = detail_states
        structured_fields["ability_unit_mapping_states"] = ability_states
        structured_fields["ncs_unmapped_detail_candidates"] = [
            row["sourceName"]
            for row in detail_states
            if row.get("mappingState") == "not_in_current_official_catalog"
        ]
        structured_fields["ncs_unmapped_ability_units"] = [
            row["sourceName"]
            for row in ability_states
            if row.get("mappingState") == "not_in_current_official_catalog"
        ]
        source_ability_scopes = (
            structured_fields.get("ability_units_by_detail")
            if isinstance(structured_fields.get("ability_units_by_detail"), dict)
            else {}
        )
        if not source_ability_scopes and structured_fields.get("ability_units"):
            source_ability_scopes = {
                "": list(structured_fields.get("ability_units") or [])
            }
        structured_fields["ncs_detail_convergence_suggestions"] = (
            derive_detail_candidates_from_exact_ability_scopes(
                source_ability_scopes
            )
        )
    review_session = _create_review_session(data, structured, jd_file.filename or "")
    _record_audit_event(
        request,
        action="jd_parse_review",
        resource_type="jd_review_session",
        resource_id=review_session["id"],
    )
    structured["review_session_id"] = review_session["id"]
    structured["review_session"] = review_session
    return structured


@app.post("/api/notice/parse-review")
async def parse_notice_review_endpoint(
    request: Request,
    notice_file: UploadFile = File(...),
) -> dict:
    """Parse a job notice and return editable duty/evaluation text candidates."""

    data = await _read_upload_limited(notice_file, "notice_file")
    if not data:
        raise HTTPException(status_code=400, detail="notice_file is empty")
    filename = notice_file.filename or ""
    parsed = _parse_upload_document(data, filename, "notice_file")
    structured = structure_job_notice(parsed, filename=filename)
    review_session = _create_review_session(data, structured, filename)
    _record_audit_event(
        request,
        action="notice_parse_review",
        resource_type="notice_review_session",
        resource_id=review_session["id"],
    )
    structured["review_session_id"] = review_session["id"]
    structured["review_session"] = review_session
    return structured


@app.post("/api/jd/strategy/upload")
async def jd_strategy_upload(
    request: Request,
    jd_file: UploadFile = File(...),
    notice_file: UploadFile | None = File(default=None),
    presentation_material_file: UploadFile | None = File(default=None),
    strengths: str = Form(default=""),
    generation_api_key: str = Form(default=""),
    openai_api_key: str = Form(default=""),
    generation_provider: str = Form(default=""),
    manual_sclass: str = Form(default=""),
    manual_sclass_add: str = Form(default=""),
    manual_sclass_remove: str = Form(default=""),
    duty_text: str = Form(default=""),
    qualification_text: str = Form(default=""),
    preference_text: str = Form(default=""),
    evaluation_text: str = Form(default=""),
    presentation_material_text: str = Form(default=""),
    question_plan_json: str = Form(default=""),
    interview_methods_json: str = Form(default=""),
    avoid_questions_json: str = Form(default=""),
    include_all_questions: str | None = Form(default=None),
    generated_questions_max_items: int | None = Form(default=None),
    generation_offset: int | None = Form(default=None),
    jd_review_json: str = Form(default=""),
    notice_review_json: str = Form(default=""),
) -> dict:
    # 최적값 고정 (사용자 노출 제거)
    _reject_sensitive_query_params(request, destination="form data")
    # Read the removed legacy field from the parsed form only to reject stale
    # clients safely.  It is intentionally absent from the public OpenAPI
    # contract and is never accepted as a generation credential.
    compatibility_form = await request.form()
    openrouter_api_key = str(compatibility_form.get("openrouter_api_key", "") or "")
    generation_model = str(compatibility_form.get("generation_model", "") or "")
    strengths = _validate_generation_text_input(
        strengths,
        field_name="strengths",
        max_chars=_MAX_STRENGTHS_CHARS,
    )
    duty_text = _validate_generation_text_input(
        duty_text,
        field_name="duty_text",
        max_chars=_MAX_DUTY_TEXT_CHARS,
    )
    qualification_text = _validate_generation_text_input(
        qualification_text,
        field_name="qualification_text",
        max_chars=_MAX_QUALIFICATION_TEXT_CHARS,
    )
    preference_text = _validate_generation_text_input(
        preference_text,
        field_name="preference_text",
        max_chars=_MAX_PREFERENCE_TEXT_CHARS,
    )
    evaluation_text = _validate_generation_text_input(
        evaluation_text,
        field_name="evaluation_text",
        max_chars=_MAX_EVALUATION_TEXT_CHARS,
    )
    presentation_material_text = _validate_generation_text_input(
        presentation_material_text,
        field_name="presentation_material_text",
        max_chars=_MAX_PRESENTATION_MATERIAL_TEXT_CHARS,
    )
    request_avoid_questions = _validate_and_extract_generation_avoid_questions(
        avoid_questions_json,
        field_name="avoid_questions_json",
    )
    interview_methods = _parse_interview_methods(interview_methods_json)
    if (
        (presentation_material_file is not None or presentation_material_text.strip())
        and "발표면접" not in interview_methods
    ):
        raise HTTPException(
            status_code=422,
            detail={
                "code": "presentation_material_requires_presentation_method",
                "message": "발표 자료는 발표면접을 선택한 경우에만 입력·첨부할 수 있습니다.",
                "retryable": False,
            },
        )
    run_top_k, run_ksa_units, run_ksa_factors = FAST_NCS_TOP_K, FAST_KSA_UNITS, FAST_KSA_FACTORS_PER_UNIT
    (
        request_generation_provider,
        request_generation_model,
        request_generation_api_key,
    ) = _resolve_request_generation(
        generation_api_key=generation_api_key,
        openai_api_key=openai_api_key,
        openrouter_api_key=openrouter_api_key,
        provider=generation_provider,
        generation_model=generation_model,
    )
    _require_allowed_openai_key(
        request_generation_api_key,
        request,
        provider=request_generation_provider,
    )

    async def _read_text(
        upload: UploadFile | None,
        label: str,
        *,
        parse_document: bool = True,
    ) -> tuple[str, bytes, str]:
        if not upload:
            return "", b"", ""
        name = (upload.filename or "").lower()
        data = await _read_upload_limited(upload, label)
        if not data:
            raise HTTPException(status_code=400, detail=f"{label} is empty")
        if not parse_document:
            return "", data, name
        parsed = _parse_upload_document(data, upload.filename or "", label)
        text = str(parsed.get("markdown") or "")
        if text.strip():
            return text, data, name
        raise HTTPException(status_code=400, detail=f"{label} could not be parsed")

    # Enforce upload size/type/emptiness before validating review metadata, while
    # avoiding a second Kordoc parse for the already reviewed JD.
    jd_text, jd_bytes, jd_name = await _read_text(
        jd_file,
        "jd_file",
        parse_document=False,
    )

    review_payload: dict[str, Any] = {}
    if jd_review_json.strip():
        try:
            candidate = json.loads(jd_review_json)
            if isinstance(candidate, dict):
                review_payload = candidate
        except json.JSONDecodeError as exc:
            raise HTTPException(status_code=400, detail=f"jd_review_json is invalid: {exc}") from exc

    if review_payload.get("review_confirmed") is not True:
        raise HTTPException(
            status_code=400,
            detail="jd_review_json.review_confirmed must be the boolean true",
        )
    reviewed_fields = review_payload.get("fields") if isinstance(review_payload.get("fields"), dict) else {}
    _validate_generation_text_collection_input(
        reviewed_fields.get("duties") or [],
        field_name="jd_review_json.fields.duties",
        max_chars=_MAX_DUTY_TEXT_CHARS,
    )
    _validate_generation_text_collection_input(
        reviewed_fields.get("qualifications") or [],
        field_name="jd_review_json.fields.qualifications",
        max_chars=_MAX_QUALIFICATION_TEXT_CHARS,
    )
    _validate_generation_text_collection_input(
        reviewed_fields.get("preferences") or [],
        field_name="jd_review_json.fields.preferences",
        max_chars=_MAX_PREFERENCE_TEXT_CHARS,
    )
    _validate_generation_text_collection_input(
        reviewed_fields.get("ability_units") or [],
        field_name="jd_review_json.fields.ability_units",
        max_chars=_MAX_ABILITY_UNIT_TEXT_CHARS,
    )
    review_session: dict[str, Any] | None = _validate_review_session(
        review_payload,
        jd_bytes,
        jd_file.filename or "",
    )
    reviewed_markdown = str((review_session or {}).get("markdown") or "").strip()
    if reviewed_markdown:
        jd_text = reviewed_markdown
    reviewed_detail_terms = _parse_sclass_terms(
        "\n".join(
            str(value).strip()
            for value in (reviewed_fields.get("ncs_detail_candidates") or [])
            if str(value).strip()
        )
    )
    reviewed_ability_units = _reviewed_ability_unit_names(
        reviewed_fields.get("ability_units") or []
    )
    reviewed_ability_unit_ordinals = _reviewed_ability_unit_ordinals(reviewed_fields)
    question_plan = _parse_question_plan_json(question_plan_json, reviewed_detail_terms)
    _enforce_question_plan_capacity(question_plan)
    if question_plan["selected_terms"]:
        reviewed_detail_terms = list(question_plan["selected_terms"])
    if not reviewed_detail_terms:
        raise HTTPException(
            status_code=422,
            detail="Human-reviewed NCS detail candidates are required",
        )

    # Only inspect the second document after deterministic capacity validation.
    # The UI already parsed it for human review, so a signed/hash-bound review
    # session lets generation reuse that markdown without a second Kordoc call.
    notice_text = ""
    notice_payload: dict[str, Any] = {}
    if notice_review_json.strip():
        try:
            candidate = json.loads(notice_review_json)
            if isinstance(candidate, dict):
                notice_payload = candidate
        except json.JSONDecodeError as exc:
            raise HTTPException(
                status_code=400,
                detail=f"notice_review_json is invalid: {exc}",
            ) from exc
        if notice_payload.get("review_confirmed") is not True:
            raise HTTPException(
                status_code=400,
                detail="notice_review_json.review_confirmed must be the boolean true",
            )
    if notice_file:
        _, notice_bytes, _ = await _read_text(
            notice_file,
            "notice_file",
            parse_document=False,
        )
        if notice_payload:
            notice_session = _validate_review_session(
                notice_payload,
                notice_bytes,
                notice_file.filename or "",
                field_name="notice_review_json",
                upload_label="notice_file",
                parse_endpoint="/api/notice/parse-review",
            )
            notice_text = str((notice_session or {}).get("markdown") or "").strip()
        else:
            parsed_notice = _parse_upload_document(
                notice_bytes,
                notice_file.filename or "",
                "notice_file",
            )
            notice_text = str(parsed_notice.get("markdown") or "").strip()
            if not notice_text:
                raise HTTPException(status_code=400, detail="notice_file could not be parsed")
    elif notice_payload:
        raise HTTPException(
            status_code=400,
            detail="notice_file is required when notice_review_json is provided",
        )
    presentation_material_filename = ""
    if presentation_material_file is not None:
        presentation_material_filename = str(presentation_material_file.filename or "").strip()
        _reject_hwp_upload("발표 자료 파일", presentation_material_filename)
        presentation_material_bytes = await _read_upload_limited(
            presentation_material_file,
            "presentation_material_file",
        )
        if len(presentation_material_bytes) > _MAX_PRESENTATION_MATERIAL_FILE_BYTES:
            raise HTTPException(
                status_code=413,
                detail=(
                    "presentation_material_file exceeds the 2 MiB presentation-material limit"
                ),
            )
        if not presentation_material_bytes:
            raise HTTPException(status_code=400, detail="presentation_material_file is empty")
        parsed_presentation_material = _parse_upload_document(
            presentation_material_bytes,
            presentation_material_filename,
            "presentation_material_file",
        )
        parsed_material_text = str(parsed_presentation_material.get("markdown") or "").strip()
        if not parsed_material_text:
            raise HTTPException(status_code=422, detail="presentation_material_file could not be parsed")
        presentation_material_text = _validate_generation_text_input(
            "\n\n".join(
                value
                for value in (presentation_material_text, parsed_material_text)
                if str(value or "").strip()
            ),
            field_name="presentation_material_text",
            max_chars=_MAX_PRESENTATION_MATERIAL_TEXT_CHARS,
        )
    duty_text_clean = _merge_review_text(
        duty_text,
        reviewed_fields.get("duties") or [],
        max_chars=3000,
    )
    qualification_text_clean = _merge_review_text(
        qualification_text,
        reviewed_fields.get("qualifications") or [],
        max_chars=2400,
    )
    preference_text_clean = _merge_review_text(
        preference_text,
        reviewed_fields.get("preferences") or [],
        max_chars=2400,
    )
    evaluation_text_clean = str(evaluation_text or "").strip()
    include_all_generated_questions = _coerce_bool_flag(
        include_all_questions,
        default=True,
    )
    explicit_generated_questions_max = _coerce_positive_int(
        generated_questions_max_items
    )
    manual_sclass_final_terms = _parse_sclass_terms(manual_sclass)
    manual_sclass_add_terms = _parse_sclass_terms(manual_sclass_add)
    manual_sclass_remove_terms = _parse_sclass_terms(manual_sclass_remove)

    if not jd_text.strip():
        fallback_terms = manual_sclass_final_terms or manual_sclass_add_terms
        if fallback_terms:
            jd_text = "소분류: " + ", ".join(fallback_terms)
        else:
            raise HTTPException(status_code=400, detail="no readable text in jd_file")

    vision_terms: list[str] = []
    use_vision_ocr = (str(__import__("os").getenv("ENABLE_VISION_OCR", "false")).strip().lower() in {"1", "true", "yes", "y"})
    # Vision OCR은 텍스트 추출이 실패했을 때만 실행 (텍스트가 있으면 오히려 NCS 매칭 오염 가능)
    if use_vision_ocr and jd_name.endswith(".pdf") and len(jd_text.strip()) < 50:
        with use_generation_request(
            provider=request_generation_provider,
            generation_model=request_generation_model,
        ):
            vision_terms = extract_focus_terms_from_pdf_vision(
                jd_bytes,
                max_pages=2,
                api_key_override=request_generation_api_key,
                generation_provider=request_generation_provider,
            )
    prompt_notice_text = _build_priority_notice_text(
        notice_text=notice_text,
        duty_text=duty_text_clean,
        qualification_text=qualification_text_clean,
        preference_text=preference_text_clean,
        evaluation_text=evaluation_text_clean,
    )
    notice_context = build_notice_context_from_jd(jd_text=jd_text, notice_text=prompt_notice_text, max_chars=5000)

    _require_ncs_mcp_url()
    mcp_only = False
    ncs_source = "ncs-mcp"
    ncs_error = ""

    jd_for_match = jd_text
    if vision_terms:
        jd_for_match = " ".join(vision_terms)

    subcategory_text = extract_subcategory_text(jd_text) if jd_text.strip() else " ".join(vision_terms)
    extracted_small_categories = extract_small_categories_from_jd(jd_text) if jd_text.strip() else []

    # 소분류 확정 규칙:
    # 1) 문서 추출 결과를 기본으로 사용
    # 2) 수기 추가가 있으면 append
    # 3) 수기 삭제가 있으면 제거
    # 4) 레거시 호환: add/remove가 없고 manual_sclass가 있으면 최종 확정 목록으로 간주
    if manual_sclass_add_terms or manual_sclass_remove_terms:
        small_categories = _merge_sclass_terms(
            base_terms=extracted_small_categories,
            add_terms=manual_sclass_add_terms,
            remove_terms=manual_sclass_remove_terms,
        )
    elif manual_sclass_final_terms:
        small_categories = _merge_sclass_terms(
            base_terms=[],
            add_terms=manual_sclass_final_terms,
            remove_terms=[],
        )
    else:
        small_categories = list(extracted_small_categories)

    manual_terms = list(manual_sclass_final_terms or manual_sclass_add_terms)
    if small_categories:
        subcategory_text = f"소분류 후보: {', '.join(small_categories)}\n{subcategory_text}".strip()
    core_small_categories = small_categories[:6]
    inferred_keywords: list[str] = []
    reviewed_keywords: list[str] = []
    ai_sclass_candidates: list[dict] = []
    ai_ncs_code_candidates: list[dict] = []

    # 1) seed 구성: extract_small_categories_from_jd() 결과 + 소분류 텍스트 토큰 + vision 키워드
    import re as _re
    seeds: list[str] = []
    sub_tokens = _re.findall(r"[\uAC00-\uD7A3]{2,12}", subcategory_text or "")
    raw_tokens = _re.findall(r"[\uAC00-\uD7A3]{2,12}", jd_text or "")
    for term in (small_categories + sub_tokens + raw_tokens[:40] + vision_terms):
        t = str(term).strip()
        if t and t not in seeds:
            seeds.append(t)

    show_all_from_small_categories = bool(
        os.getenv("NCS_SHOW_ALL_FROM_SMALL_CATEGORIES", "true").strip().lower() in {"1", "true", "yes", "y"}
    )
    sclass_bundle = resolve_sclass_candidates_bundle(
        jd_text=jd_text,
        small_categories=small_categories,
        manual_terms=manual_terms,
        subcategory_text=subcategory_text,
        doc_name=jd_name,
        show_all_from_small_categories=show_all_from_small_categories,
        enable_ai_fallback=True,
        verified_sclass_limit=_clamp_sclass_limit(os.getenv("NCS_VERIFIED_SCLASS_LIMIT", "4"), default=4),
        verified_min_keep=_clamp_sclass_limit(os.getenv("NCS_VERIFIED_SCLASS_MIN_KEEP", "1"), default=1),
        score_margin=_to_float_or(os.getenv("NCS_VERIFIED_SCORE_MARGIN", "0.18"), 0.18),
        min_confidence=_to_float_or(os.getenv("NCS_VERIFIED_MIN_CONFIDENCE", "0.62"), 0.62),
    )
    csv_sclass_candidates: list[dict] = sclass_bundle["csv_sclass_candidates"]
    verified_sclass: list[dict] = sclass_bundle["verified_sclass"]

    # CSV 실패 시 keywords 기반 fallback (NCS API 호출)
    if not verified_sclass and seeds:
        with use_generation_request(
            provider=request_generation_provider,
            generation_model=request_generation_model,
        ):
            inferred_keywords = infer_keywords_from_subcategory_ai(subcategory_text=subcategory_text, jd_text=jd_for_match)
            reviewed_keywords = review_ocr_terms_with_openai(terms=(inferred_keywords or seeds[:12]), jd_text=jd_for_match)

    ncs_query_terms = [str(v.get("sclass_name", "")).strip() for v in verified_sclass if str(v.get("sclass_name", "")).strip()]
    if not ncs_query_terms:
        ncs_query_terms = [t for t in (reviewed_keywords or inferred_keywords or seeds[:8]) if t]

    mcp_only = True

    ncs_items: list[dict[str, Any]] = []
    # Preferred path: use authoritative NCS-MCP only when review-confirmed detail labels
    # are available. If this path is not usable, continue with public/API fallback.
    mcp_lookup_terms = _canonicalize_detail_lookup_terms(reviewed_detail_terms)
    if mcp_only:
        detail_search_unit_limit = max(20, run_top_k * 12)
        if reviewed_ability_units:
            detail_search_unit_limit = min(
                200,
                max(80, len(reviewed_ability_units) * 4, detail_search_unit_limit),
            )
        try:
            ncs_items = search_units_by_detail(
                mcp_lookup_terms,
                max_units=detail_search_unit_limit,
            )
        except NcsMcpError as exc:
            logger.error(
                "ncs_mcp_search_failed",
                exc_info=(type(exc), exc, exc.__traceback__),
            )
            # A reviewed exact detail label must never be downgraded to a
            # keyword, public API, or local-map guess when the authoritative
            # NCS MCP is unavailable. Preserve the user's review state and let
            # the same request be retried without exposing the provider error.
            raise HTTPException(
                status_code=503,
                detail={
                    "code": "ncs_mcp_unavailable",
                    "message": (
                        "NCS 공식 세분류·능력단위 조회가 일시적으로 불가능합니다. "
                        "선택 상태를 유지한 채 다시 시도해 주세요."
                    ),
                    "retryable": True,
                    "lookup_terms": list(mcp_lookup_terms),
                },
            ) from exc
        else:
            matched_lookup_terms, unmatched_lookup_terms = _detail_lookup_coverage(
                mcp_lookup_terms,
                ncs_items,
            )
            if not ncs_items or unmatched_lookup_terms:
                suggestion_terms = unmatched_lookup_terms or mcp_lookup_terms
                try:
                    suggestions = suggest_units_by_text(
                        suggestion_terms,
                        max_units=12,
                    )
                except NcsMcpError:
                    suggestions = []
                message = (
                    "Local NCS DB returned partial exact coverage for reviewed "
                    "detail-class terms. Human selection is required."
                    if ncs_items
                    else "Local NCS DB returned no exact competency units for "
                    "the reviewed detail candidates. Human selection is required."
                )
                raise HTTPException(
                    status_code=422,
                    detail={
                        "message": message,
                        "lookup_terms": list(mcp_lookup_terms),
                        "matched_detail_terms": list(matched_lookup_terms),
                        "unmatched_detail_terms": list(
                            unmatched_lookup_terms or mcp_lookup_terms
                        ),
                        "suggested_ncs_units": list(suggestions or []),
                    },
                )
            if reviewed_ability_units:
                locked_units, missing_ability_units = _lock_units_to_reviewed_ability_units(
                    ncs_items,
                    reviewed_ability_units,
                )
                if missing_ability_units:
                    ordinal_recovered, missing_ability_units = (
                        _recover_ordinal_scoped_reviewed_ability_units(
                            ncs_items,
                            missing_ability_units,
                            reviewed_ability_unit_ordinals,
                            already_locked=locked_units,
                        )
                    )
                    locked_units.extend(ordinal_recovered)
                if missing_ability_units:
                    catalog_recovered, missing_ability_units = (
                        _recover_code_scoped_reviewed_ability_units(
                            ncs_items,
                            missing_ability_units,
                            exact_official_units_by_name(missing_ability_units),
                            already_locked=locked_units,
                        )
                    )
                    locked_units.extend(catalog_recovered)
                if missing_ability_units:
                    recovery_candidates: list[dict[str, Any]] = []
                    try:
                        for missing_name in missing_ability_units:
                            recovery_candidates.extend(
                                suggest_units_by_text([missing_name], max_units=20)
                            )
                    except NcsMcpError:
                        recovery_candidates = []
                    recovered_units, missing_ability_units = (
                        _recover_code_scoped_reviewed_ability_units(
                            ncs_items,
                            missing_ability_units,
                            recovery_candidates,
                            already_locked=locked_units,
                        )
                    )
                    locked_units.extend(recovered_units)
                if missing_ability_units:
                    raise HTTPException(
                        status_code=422,
                        detail={
                            "code": "ncs_required_ability_unit_mismatch",
                            "message": (
                                "검토·확정한 요구능력단위 일부가 선택한 세분류의 공식 NCS 능력단위와 "
                                "정확히 일치하지 않습니다. 직무기술서 표의 능력단위명을 다시 확인해 주세요."
                            ),
                            "reviewed_detail_terms": list(mcp_lookup_terms),
                            "reviewed_ability_units": list(reviewed_ability_units),
                            "matched_ability_units": [
                                str(unit.get("compeUnitName") or "").strip()
                                for unit in locked_units
                                if str(unit.get("compeUnitName") or "").strip()
                            ],
                            "unmatched_ability_units": list(missing_ability_units),
                            "retryable": False,
                        },
                    )
                ncs_items = locked_units
                ncs_source = "ncs-mcp+required-unit-lock"
            else:
                ncs_source = "ncs-mcp"
            ncs_query_terms = list(mcp_lookup_terms)

    # 4) 코드 기반 조회 후, 키워드 기반으로 순차 fallback
    max_sclass_verified = _clamp_sclass_limit(os.getenv("NCS_API_MAX_SCLASS_VERIFIED", "4"), default=4)
    max_sclass_name = _clamp_sclass_limit(os.getenv("NCS_API_MAX_SCLASS_NAME", "4"), default=4)
    if not mcp_only and verified_sclass:
        fetch_limit = min(max_sclass_verified, max(1, len(verified_sclass)))
        ncs_items = fetch_ncs_units_hrdk_by_verified_sclass(verified_sclass, max_sclass=fetch_limit)
        if ncs_items:
            ncs_source = "api-hrdk-sclass-verified"

    if not mcp_only and not ncs_items and ncs_query_terms:
        fetch_limit = min(max_sclass_name, max(1, len(ncs_query_terms)))
        ncs_items = fetch_ncs_units_hrdk_by_sclass_names(ncs_query_terms, max_sclass=fetch_limit)
        if ncs_items:
            ncs_source = "api-hrdk-sclass-name"

    if not mcp_only and not ncs_items and ncs_query_terms:
        ncs_items = fetch_ncs_units_hrdk_by_keywords(ncs_query_terms, max_items=60)
        if ncs_items:
            ncs_source = "api-hrdk-keyword"

    if not mcp_only and not ncs_items and seeds:
        ai_ncs_code_candidates = ai_extract_ncs_cl_codes(seed_terms=seeds[:18], jd_text=jd_for_match, max_items=8)
        if ai_ncs_code_candidates:
            ncs_items = fetch_ncs_units_hrdk_by_cl_codes(ai_ncs_code_candidates, max_items=40)
            if ncs_items:
                ncs_source = "api-hrdk-clcode"

    if not mcp_only and not ncs_items:
        ncs_source = "fallback-local-map"
        ncs_error = "외부 NCS 조회가 불안정하여 로컬 매핑으로 대체했습니다."

    ncs_matches = []
    unit_rank_query_text = _build_priority_query_text(
        base_text=jd_for_match,
        duty_text=duty_text_clean,
        qualification_text=qualification_text_clean,
        preference_text=preference_text_clean,
        evaluation_text=evaluation_text_clean,
    )
    if ncs_items and (
        ncs_source.startswith("ncs-mcp")
        or ncs_source
        in {
            "api-hrdk-code-first",
            "api-hrdk-clcode",
            "api-hrdk-keyword",
            "api-hrdk-sclass-verified",
            "api-hrdk-sclass-name",
        }
    ):
        # The upload flow has already required a human-confirmed exact detail
        # classification before this point. Calling a second model merely to
        # reorder those authoritative local-DB candidates adds latency and can
        # dilute detail coverage, so keep this ranking deterministic. The
        # request-scoped key remains reserved for the actual question draft.
        unit_rerank_api_key = (
            ""
            if mcp_only
            and review_payload.get("review_confirmed") is True
            and reviewed_detail_terms
            else request_generation_api_key
        )
        with use_generation_request(
            provider=request_generation_provider,
            generation_model=request_generation_model,
        ):
            ncs_matches, rerank_mode = rerank_ncs_matches(
                jd_text=unit_rank_query_text or jd_for_match,
                ncs_items=ncs_items,
                top_k=run_top_k,
                preferred_sclass=ncs_query_terms,
                openai_api_key=unit_rerank_api_key,
                generation_provider=request_generation_provider,
            )
        if ncs_matches:
            ncs_source = f"{ncs_source}+ai-rerank" if rerank_mode == "ai" else f"{ncs_source}+rerank"
        else:
            # rank 결과가 비어도 상위 원본을 안전 fallback으로 사용
            for it in ncs_items[:8]:
                ncs_matches.append(
                    {
                        "ncsClCd": str(it.get("ncsClCd", "")).strip(),
                        "compeUnitName": str(it.get("compeUnitName", "")).strip(),
                        "compeUnitLevel": str(it.get("compeUnitLevel", "")).strip(),
                        "ncsSubdCdnm": str(it.get("ncsSubdCdnm", "")).strip(),
                        "compeUnitDef": str(it.get("compeUnitDef", "")).strip(),
                        "score": float(it.get("score", 0.5) or 0.5),
                        "matched_keywords": list(it.get("matched_keywords", []) or []),
                    }
                )
    elif ncs_items and ncs_query_terms:
        def _norm(v: str) -> str:
            return (v or "").replace(" ", "").strip().lower()

        def _canon(v: str) -> str:
            n = _norm(v)
            if n.endswith("사") and len(n) >= 3:
                return n[:-1]
            return n

        soc_norm = {_canon(x) for x in ncs_query_terms}
        matched = []
        for it in ncs_items:
            sclas = _canon(str(it.get("ncsSclasCdnm", "")))
            exact_soc = bool(sclas and sclas in soc_norm)
            if exact_soc:
                matched.append(
                    {
                        "ncsClCd": it.get("ncsClCd", ""),
                        "compeUnitName": it.get("compeUnitName", ""),
                        "compeUnitLevel": it.get("compeUnitLevel", ""),
                        "ncsSubdCdnm": it.get("ncsSubdCdnm", ""),
                        "compeUnitDef": it.get("compeUnitDef", ""),
                        "score": 9.999,
                        "matched_keywords": [it.get("ncsSclasCdnm", "")],
                    }
                )
        seen = set()
        dedup = []
        for m in matched:
            code = str(m.get("ncsClCd", "")).strip()
            if not code or code in seen:
                continue
            dedup.append(m)
            seen.add(code)
            if len(dedup) >= 5:
                break
        ncs_matches = dedup
        if ncs_matches:
            ncs_source = "api-soclass"

    if not mcp_only and not ncs_matches:
        # 마지막 fallback: 내부 샘플 매퍼로 최소 매핑 확보
        local_items = map_ncs(category="R6000_MANAGEMENT", text=jd_for_match, top_k=8)
        for it in (local_items or [])[:8]:
            code = str(it.get("ncsClCd", "")).strip()
            if not code:
                continue
            ncs_matches.append(
                {
                    "ncsClCd": code,
                    "compeUnitName": str(it.get("compeUnitName", "")).strip(),
                    "compeUnitLevel": str(it.get("compeUnitLevel", "")).strip(),
                    "ncsSubdCdnm": str(it.get("ncsSubdCdnm", "")).strip(),
                    "compeUnitDef": str(it.get("compeUnitDef", "")).strip(),
                    "score": float(it.get("score", 0.3) or 0.3),
                    "matched_keywords": list(it.get("matched_keywords", []) or []),
                }
            )
        if ncs_matches:
            ncs_source = "fallback-local-map+rerank"
            ncs_error = "외부 NCS 매핑 실패로 로컬 매퍼를 사용했습니다."
        elif not ncs_error:
            ncs_error = f"NCS 매핑 결과가 없어 JD 기반 질문으로 대체합니다. query={ncs_query_terms[:8]}"
    ncs_matches = _ensure_question_plan_unit_coverage(
        question_plan,
        ncs_matches,
        ncs_items,
    )
    if reviewed_ability_units and ncs_matches:
        # Ranking may rebuild row dictionaries. Re-attach the exact-lock audit
        # fields by official unit code so the response proves which reviewed
        # cell selected every KSA lookup target.
        locked_by_code = {
            str(unit.get("ncsClCd") or unit.get("unit_code") or "").strip(): unit
            for unit in ncs_items
            if str(unit.get("ncsClCd") or unit.get("unit_code") or "").strip()
        }
        ncs_matches = [
            {
                **match,
                "requiredAbilityUnitName": str(
                    locked_by_code.get(
                        str(match.get("ncsClCd") or match.get("unit_code") or "").strip(),
                        {},
                    ).get("requiredAbilityUnitName")
                    or ""
                ).strip(),
                "requiredAbilityUnitMatch": str(
                    locked_by_code.get(
                        str(match.get("ncsClCd") or match.get("unit_code") or "").strip(),
                        {},
                    ).get("requiredAbilityUnitMatch")
                    or ""
                ).strip(),
            }
            for match in ncs_matches
        ]

    # NCS 평가요소를 수집해 OpenAI 입력에 함께 전달한다.
    # 질문계획에서 실제 배정될 능력단위만 조회한 뒤, JD 핵심 +
    # 담당업무 텍스트 기준 TF-IDF로 KSA를 선별한다. 한 문항을 위해
    # 사용되지 않을 상위 능력단위 전체를 원격 조회하지 않는다.
    ncs_ksa: list[dict[str, Any]] = []
    ncs_ksa_candidates: list[dict[str, Any]] = []
    ksa_units: list[dict[str, Any]] = []
    ksa_query_text = _build_priority_query_text(
        base_text=jd_text,
        duty_text=duty_text_clean,
        qualification_text=qualification_text_clean,
        preference_text=preference_text_clean,
        evaluation_text=evaluation_text_clean,
    )
    if ncs_matches:
        ksa_rank_top_n = _clamp_int(os.getenv("KSA_RANK_TOP_N", "12"), default=12, lo=6, hi=20)
        ksa_rank_per_unit = _clamp_int(os.getenv("KSA_RANK_PER_UNIT_LIMIT", "3"), default=3, lo=1, hi=4)
        ksa_rank_units = _clamp_int(os.getenv("KSA_RANK_MAX_UNITS", "5"), default=5, lo=2, hi=5)
        ksa_candidate_per_unit = _clamp_int(os.getenv("KSA_CANDIDATE_PER_UNIT", "12"), default=12, lo=3, hi=24)
        ksa_sim_weight = _to_float_or(os.getenv("KSA_SIMILARITY_WEIGHT", "0.75"), 0.75)
        ksa_unit_weight = _to_float_or(os.getenv("KSA_UNIT_WEIGHT", "0.25"), 0.25)

        planned_ksa_units = _select_units_for_question_plan(
            question_plan,
            ncs_matches,
        )
        ksa_units = _collect_ksa_candidate_units(
            primary_units=planned_ksa_units or ncs_matches,
            secondary_units=None if planned_ksa_units else ncs_items,
            max_units=ksa_rank_units,
        )
        if not ksa_units:
            ksa_units = _collect_ksa_candidate_units(
                primary_units=ncs_matches[:run_top_k],
                secondary_units=None,
                max_units=max(1, run_top_k),
            )

        ncs_ksa_candidates = _fetch_ncs_ksa_or_502(
            ncs_matches=ksa_units,
            max_units=len(ksa_units),
            max_factors_per_unit=ksa_candidate_per_unit,
        )
        unit_scores: dict[str, float] = {}
        for x in (ksa_units or []):
            code = str(x.get("ncsClCd", "")).strip()
            if not code:
                continue
            try:
                unit_scores[code] = float(x.get("score", 1.0) or 1.0)
            except Exception:
                unit_scores[code] = 1.0
        ncs_ksa = rank_ksa_factors_by_query(
            ksa_rows=ncs_ksa_candidates,
            query_text=ksa_query_text,
            unit_scores=unit_scores,
            target_count=ksa_rank_top_n,
            per_unit_limit=ksa_rank_per_unit,
            similarity_weight=ksa_sim_weight,
            unit_weight=ksa_unit_weight,
            ngram_min=2,
            ngram_max=4,
        )
        if not ncs_ksa:
            ncs_ksa = _fetch_ncs_ksa_or_502(
                ncs_matches=ncs_matches[:run_top_k],
                max_units=min(run_ksa_units, len(ncs_matches)),
                max_factors_per_unit=run_ksa_factors,
            )
        ncs_ksa = _supplement_ksa_for_question_plan(
            question_plan=question_plan,
            ncs_matches=ncs_matches,
            ncs_ksa=ncs_ksa,
            max_factors_per_unit=run_ksa_factors,
        )
    ncs_factor_sources = sorted(
        {
            str(x.get("factorSource", "")).strip()
            for x in (ncs_ksa or [])
            if str(x.get("factorSource", "")).strip()
        }
    )
    ncs_context = build_ncs_context_pack(
        jd_text=jd_for_match,
        notice_text=notice_context,
        ncs_items=ncs_items,
        ncs_matches=ncs_matches,
    )
    presentation_material_packet = _build_presentation_material_packet(
        interview_methods=interview_methods,
        jd_text=jd_text,
        notice_text=notice_context,
        duty_text=duty_text_clean,
        question_plan=question_plan,
        ncs_matches=ncs_matches,
        ncs_ksa=ncs_ksa,
        supplemental_text=presentation_material_text,
    )
    presentation_material_prompt = _presentation_material_prompt_text(presentation_material_packet)
    avoid_context = _build_avoid_questions_context(request_avoid_questions, max_items=20)
    avoid_context = _join_generation_context(avoid_context, _quality_feedback_context(ncs_matches))
    avoid_context = _join_generation_context(avoid_context, presentation_material_prompt)
    enable_ai_refine = bool(inferred_keywords or reviewed_keywords or ai_ncs_code_candidates)

    build_kwargs = {
        "jd_text": jd_text,
        "notice_text": notice_context,
        "strengths": strengths,
        "region": "",
        "ncs_matches": ncs_matches,
        "ncs_ksa": ncs_ksa,
        "ncs_context": ncs_context,
        "duty_text": duty_text_clean,
        "evaluation_text": evaluation_text_clean,
        "desired_job": "",
        "api_key_override": request_generation_api_key,
        "target_count_override": question_plan["total_main_count"],
        "follow_up_count": question_plan["follow_up_count"],
        "question_plan": question_plan,
        "interview_methods": interview_methods,
        "extra_context": avoid_context,
        "generation_provider": request_generation_provider,
        "generation_model": request_generation_model,
    }
    try:
        with use_generation_request(
            provider=request_generation_provider,
            generation_model=request_generation_model,
        ):
            strategy = await _generate_quality_gated_institution_strategy(
                build_kwargs=build_kwargs,
                question_plan=question_plan,
                interview_methods=interview_methods,
                ncs_matches=ncs_matches,
                ncs_ksa=ncs_ksa,
                avoid_questions=request_avoid_questions,
                generation_offset=generation_offset,
            )
    except HTTPException:
        raise
    except Exception as e:
        logger.warning(
            "jd_strategy_ai_generation_failed_closed provider=%s",
            request_generation_provider,
        )
        raise _institution_api_provider_http_error(
            e,
            provider=request_generation_provider,
        ) from e
    strategy = _attach_presentation_material_packet(strategy, presentation_material_packet)
    _register_question_quality_evidence(
        strategy,
        source_endpoint="/api/jd/strategy/upload",
        ncs_matches=ncs_matches,
    )
    generated_questions_all = _extract_all_generated_question_items(strategy)
    generated_questions = _build_generated_questions_payload(
        strategy,
        question_plan,
        include_all=include_all_generated_questions,
        explicit_max_items=explicit_generated_questions_max,
        fallback=run_top_k,
    )
    generated_question_text_rows, generated_question_texts = _build_generated_question_text_payload(
        strategy
    )
    generated_questions_mode = (
        "all" if include_all_generated_questions else (
            "custom_limit" if explicit_generated_questions_max is not None else "preview"
        )
    )
    generated_questions_limit = (
        explicit_generated_questions_max
        if explicit_generated_questions_max is not None
        else _generated_question_preview_limit(
            question_plan,
            strategy,
            fallback=run_top_k,
        )
    )

    if review_session:
        _record_audit_event(
            request,
            action="jd_strategy_generate",
            resource_type="jd_review_session",
            resource_id=str(review_session.get("id") or ""),
        )

    return {
        "filename": jd_file.filename,
        "notice_filename": notice_file.filename if notice_file else "",
        "presentation_material_filename": presentation_material_filename,
        "presentation_material_text_preview": presentation_material_text[:1200],
        "presentation_material": presentation_material_packet,
        "jd_text_preview": jd_text[:1200],
        "notice_text_preview": notice_text[:1200],
        "notice_context_preview": notice_context[:1200],
        "duty_text_preview": duty_text_clean[:1200],
        "qualification_text_preview": qualification_text_clean[:1200],
        "preference_text_preview": preference_text_clean[:1200],
        "evaluation_text_preview": evaluation_text_clean[:1200],
        "jd_review_confirmed": review_payload.get("review_confirmed") is True,
        "jd_review_session_id": (review_session or {}).get("id", ""),
        "jd_review_document_sha256": (review_session or {}).get("document_sha256", ""),
        "jd_review": (
            {
                "review_confirmed": review_payload.get("review_confirmed") is True,
                "review_session_id": (review_session or {}).get("id", review_payload.get("review_session_id", "")),
                "fields": reviewed_fields,
            }
            if review_payload
            else None
        ),
        "question_plan": question_plan,
        "interview_methods": interview_methods,
        "operational_notice": OPERATIONAL_REVIEW_NOTICE,
        "profile_used": bool((strengths or "").strip()),
        "ncs_source": ncs_source,
        "ncs_error": ncs_error,
        "question_generation": {
            **_generation_provider_descriptor(request_generation_provider),
            "resolved_model": str(
                strategy.get("provider_generation_model")
                or request_generation_model
                or generation_provider_config(request_generation_provider).get("default_model")
                or ""
            ),
        },
        "provider_key_source": _generation_key_source(
            request_generation_provider,
            generation_api_key=generation_api_key,
            openai_api_key=openai_api_key,
            openrouter_api_key=openrouter_api_key,
        ),
        "openai_key_source": (
            _openai_key_source(request_generation_api_key, request)
            if request_generation_provider == "openai_api"
            else "not_selected"
        ),
        "extracted_focus_terms": vision_terms,
        "subcategory_text_preview": subcategory_text[:800],
        "small_categories_extracted": extracted_small_categories,
        "small_categories": small_categories,
        "core_small_categories": core_small_categories,
        "inferred_keywords": inferred_keywords,
        "reviewed_keywords": reviewed_keywords,
        "pipeline_mode": ("direct-ncs" if not enable_ai_refine else "ai-refine+ncs"),
        "manual_sclass": manual_terms,
        "manual_sclass_add": manual_sclass_add_terms,
        "manual_sclass_remove": manual_sclass_remove_terms,
        "manual_sclass_final": manual_sclass_final_terms,
        "required_ability_units_reviewed": reviewed_ability_units,
        "required_ability_unit_lock_applied": bool(reviewed_ability_units),
        "ai_sclass_candidates": ai_sclass_candidates,
        "csv_sclass_candidates": csv_sclass_candidates,
        "ai_ncs_code_candidates": ai_ncs_code_candidates,
        "verified_sclass": verified_sclass,
        "ncs_code_nos": [str(x.get("ncs_code_no", "")) for x in verified_sclass if str(x.get("ncs_code_no", ""))],
        "ncs_matches": ncs_matches,
        "ncs_ksa": ncs_ksa,
        "ncs_ksa_candidate_count": len(ncs_ksa_candidates),
        "ncs_ksa_queried_unit_count": len(ksa_units),
        "ncs_factor_sources": ncs_factor_sources,
        "runtime_knobs": {
            "ncs_top_k": run_top_k,
            "ksa_units": run_ksa_units,
            "ksa_factors_per_unit": run_ksa_factors,
        },
        "ncs_context": ncs_context,
        "strategy": strategy,
        "generated_question_text_rows": generated_question_text_rows,
        "generated_question_texts": generated_question_texts,
        "generated_questions_all": generated_questions_all,
        "generated_questions": generated_questions,
        "generated_questions_mode": generated_questions_mode,
        "generated_questions_limit": generated_questions_limit,
        "generated_questions_count": len(generated_questions),
        "generated_questions_total_count": len(generated_questions_all),
    }


@app.post("/api/questions/generate-from-text")
async def generate_questions_from_text(request: Request, payload: dict) -> dict:
    _reject_sensitive_query_params(request, destination="JSON body")
    notice_text = _validate_generation_text_input(
        payload.get("notice_text", ""),
        field_name="notice_text",
        max_chars=_MAX_NOTICE_TEXT_CHARS,
    )
    duty_text = _validate_generation_text_input(
        payload.get("duty_text", ""),
        field_name="duty_text",
        max_chars=_MAX_DUTY_TEXT_CHARS,
    )
    _validate_generation_text_input(
        payload.get("qualification_text", ""),
        field_name="qualification_text",
        max_chars=_MAX_QUALIFICATION_TEXT_CHARS,
    )
    _validate_generation_text_input(
        payload.get("preference_text", ""),
        field_name="preference_text",
        max_chars=_MAX_PREFERENCE_TEXT_CHARS,
    )
    evaluation_text = _validate_generation_text_input(
        payload.get("evaluation_text", ""),
        field_name="evaluation_text",
        max_chars=_MAX_EVALUATION_TEXT_CHARS,
    )
    presentation_material_text = _validate_generation_text_input(
        payload.get("presentation_material_text", ""),
        field_name="presentation_material_text",
        max_chars=_MAX_PRESENTATION_MATERIAL_TEXT_CHARS,
    )
    _validate_generation_text_input(
        payload.get("strengths", ""),
        field_name="strengths",
        max_chars=_MAX_STRENGTHS_CHARS,
    )
    selected_ncs = payload.get("selected_ncs", [])
    _validate_selected_ncs_generation_input(selected_ncs)
    raw_avoid_questions = payload.get("avoid_questions")
    avoid_field_name = "avoid_questions"
    if raw_avoid_questions is None:
        raw_avoid_questions = payload.get("current_questions")
        avoid_field_name = "current_questions"
    if raw_avoid_questions is None:
        raw_avoid_questions = payload.get("currentQuestions")
        avoid_field_name = "currentQuestions"
    if raw_avoid_questions is None:
        raw_avoid_questions = payload.get("avoid_questions_json")
        avoid_field_name = "avoid_questions_json"
    request_avoid_questions = _validate_and_extract_generation_avoid_questions(
        raw_avoid_questions,
        field_name=avoid_field_name,
    )
    raw_interview_methods = payload.get("interview_methods_json", payload.get("interview_methods", ""))
    if not isinstance(raw_interview_methods, str):
        raw_interview_methods = json.dumps(raw_interview_methods, ensure_ascii=False)
    interview_methods = _parse_interview_methods(raw_interview_methods)
    if presentation_material_text.strip() and "발표면접" not in interview_methods:
        raise HTTPException(
            status_code=422,
            detail={
                "code": "presentation_material_requires_presentation_method",
                "message": "발표 자료는 발표면접을 선택한 경우에만 입력할 수 있습니다.",
                "retryable": False,
            },
        )
    (
        request_generation_provider,
        request_generation_model,
        request_generation_api_key,
    ) = _resolve_request_generation(
        generation_api_key=payload.get("generation_api_key", ""),
        openai_api_key=payload.get("openai_api_key", ""),
        openrouter_api_key=payload.get("openrouter_api_key", ""),
        provider=payload.get("generation_provider", ""),
        generation_model=payload.get("generation_model", ""),
    )
    _require_allowed_openai_key(
        request_generation_api_key,
        request,
        provider=request_generation_provider,
    )
    raw_generation_offset = payload.get("generation_offset")
    generation_offset = (
        _clamp_int(raw_generation_offset, default=0, lo=0, hi=1_000_000)
        if raw_generation_offset is not None
        else None
    )
    knobs = payload.get("runtime_knobs", {}) if isinstance(payload.get("runtime_knobs", {}), dict) else {}
    run_top_k, run_ksa_units, run_ksa_factors = _clamp_runtime_knobs(
        ncs_top_k=knobs.get("ncs_top_k"),
        ksa_units=knobs.get("ksa_units"),
        ksa_factors_per_unit=knobs.get("ksa_factors_per_unit"),
    )
    include_all_generated_questions = _coerce_bool_flag(
        payload.get("include_all_questions"),
        default=True,
    )
    explicit_generated_questions_max = _coerce_positive_int(
        payload.get("generated_questions_max_items")
    )
    if not notice_text:
        raise HTTPException(status_code=400, detail="notice_text is required")
    if not isinstance(selected_ncs, list) or not selected_ncs:
        raise HTTPException(status_code=400, detail="selected_ncs is required")

    ncs_matches: list[dict[str, Any]] = []
    seen_codes: set[str] = set()
    for row in selected_ncs:
        if not isinstance(row, dict):
            continue
        code = str(row.get("ncsClCd", "")).strip()
        if not code or code in seen_codes:
            continue
        seen_codes.add(code)
        ncs_matches.append(
            {
                "ncsClCd": code,
                "compeUnitName": str(row.get("compeUnitName", "")).strip() or f"NCS-{code}",
                "compeUnitLevel": str(row.get("compeUnitLevel", "")).strip(),
                "ncsSubdCdnm": str(row.get("ncsSubdCdnm", "")).strip(),
                "compeUnitDef": str(row.get("compeUnitDef", "")).strip(),
                "score": 1.0,
                "matched_keywords": [code],
            }
        )
    if not ncs_matches:
        raise HTTPException(status_code=400, detail="selected_ncs has no valid ncsClCd")

    plan_terms = []
    seen_plan_terms: set[str] = set()
    for row in ncs_matches:
        term = str(row.get("ncsSubdCdnm") or row.get("compeUnitName") or "").strip()
        key = _norm_sclass_key(term)
        if term and key and key not in seen_plan_terms:
            seen_plan_terms.add(key)
            plan_terms.append(term)
    raw_question_plan = payload.get("question_plan_json", payload.get("question_plan", ""))
    if not isinstance(raw_question_plan, str):
        raw_question_plan = json.dumps(raw_question_plan, ensure_ascii=False)
    question_plan = _parse_question_plan_json(raw_question_plan, plan_terms)
    question_plan = _restrict_question_plan_to_terms(question_plan, plan_terms)
    _enforce_question_plan_capacity(question_plan)
    _require_ncs_mcp_url()

    prompt_notice_text = _build_priority_notice_text(
        notice_text=notice_text,
        duty_text=duty_text,
        evaluation_text=evaluation_text,
    )

    # NCS 평가요소를 수집해 OpenAI 입력에 함께 전달한다.
    ksa_rank_top_n = _clamp_int(os.getenv("KSA_RANK_TOP_N", "12"), default=12, lo=6, hi=20)
    ksa_rank_per_unit = _clamp_int(os.getenv("KSA_RANK_PER_UNIT_LIMIT", "3"), default=3, lo=1, hi=4)
    ksa_rank_units = _clamp_int(os.getenv("KSA_RANK_MAX_UNITS", "5"), default=5, lo=2, hi=5)
    ksa_candidate_per_unit = _clamp_int(os.getenv("KSA_CANDIDATE_PER_UNIT", "12"), default=12, lo=3, hi=24)
    ksa_sim_weight = _to_float_or(os.getenv("KSA_SIMILARITY_WEIGHT", "0.75"), 0.75)
    ksa_unit_weight = _to_float_or(os.getenv("KSA_UNIT_WEIGHT", "0.25"), 0.25)

    planned_ksa_units = _select_units_for_question_plan(
        question_plan,
        ncs_matches,
    )
    ksa_units = _collect_ksa_candidate_units(
        primary_units=planned_ksa_units or ncs_matches,
        secondary_units=None,
        max_units=min(ksa_rank_units, len(ncs_matches)),
    )
    ncs_ksa_candidates = _fetch_ncs_ksa_or_502(
        ncs_matches=ksa_units,
        max_units=len(ksa_units),
        max_factors_per_unit=ksa_candidate_per_unit,
    )
    unit_scores = {str(x.get("ncsClCd", "")).strip(): 1.0 for x in (ksa_units or []) if str(x.get("ncsClCd", "")).strip()}
    ksa_query_text = _build_priority_query_text(
        base_text=notice_text,
        duty_text=duty_text,
        evaluation_text=evaluation_text,
    )
    ncs_ksa = rank_ksa_factors_by_query(
        ksa_rows=ncs_ksa_candidates,
        query_text=ksa_query_text,
        unit_scores=unit_scores,
        target_count=ksa_rank_top_n,
        per_unit_limit=ksa_rank_per_unit,
        similarity_weight=ksa_sim_weight,
        unit_weight=ksa_unit_weight,
        ngram_min=2,
        ngram_max=4,
    )
    if not ncs_ksa:
        ncs_ksa = _fetch_ncs_ksa_or_502(
            ncs_matches=ncs_matches[:run_top_k],
            max_units=min(run_ksa_units, len(ncs_matches)),
            max_factors_per_unit=run_ksa_factors,
        )
    ncs_ksa = _supplement_ksa_for_question_plan(
        question_plan=question_plan,
        ncs_matches=ncs_matches,
        ncs_ksa=ncs_ksa,
        max_factors_per_unit=run_ksa_factors,
    )
    ncs_factor_sources = sorted(
        {
            str(x.get("factorSource", "")).strip()
            for x in (ncs_ksa or [])
            if str(x.get("factorSource", "")).strip()
        }
    )
    ncs_context = build_ncs_context_pack(
        jd_text=notice_text,
        notice_text=prompt_notice_text,
        ncs_items=ncs_matches,
        ncs_matches=ncs_matches,
    )
    presentation_material_packet = _build_presentation_material_packet(
        interview_methods=interview_methods,
        jd_text=notice_text,
        notice_text=prompt_notice_text,
        duty_text=duty_text,
        question_plan=question_plan,
        ncs_matches=ncs_matches,
        ncs_ksa=ncs_ksa,
        supplemental_text=presentation_material_text,
    )
    presentation_material_prompt = _presentation_material_prompt_text(presentation_material_packet)
    avoid_context = _build_avoid_questions_context(request_avoid_questions, max_items=20)
    avoid_context = _join_generation_context(avoid_context, _quality_feedback_context(ncs_matches))
    avoid_context = _join_generation_context(avoid_context, presentation_material_prompt)

    build_kwargs = {
        "jd_text": notice_text,
        "notice_text": prompt_notice_text,
        "strengths": "",
        "region": "",
        "ncs_matches": ncs_matches,
        "ncs_ksa": ncs_ksa,
        "ncs_context": ncs_context,
        "duty_text": duty_text,
        "evaluation_text": evaluation_text,
        "desired_job": "",
        "api_key_override": request_generation_api_key,
        "target_count_override": question_plan["total_main_count"] or None,
        "follow_up_count": question_plan["follow_up_count"],
        "question_plan": question_plan,
        "interview_methods": interview_methods,
        "extra_context": avoid_context,
        "generation_provider": request_generation_provider,
        "generation_model": request_generation_model,
    }
    try:
        with use_generation_request(
            provider=request_generation_provider,
            generation_model=request_generation_model,
        ):
            strategy = await _generate_quality_gated_institution_strategy(
                build_kwargs=build_kwargs,
                question_plan=question_plan,
                interview_methods=interview_methods,
                ncs_matches=ncs_matches,
                ncs_ksa=ncs_ksa,
                avoid_questions=request_avoid_questions,
                generation_offset=generation_offset,
            )
    except HTTPException:
        raise
    except Exception as e:
        logger.warning(
            "manual_strategy_ai_generation_failed_closed provider=%s",
            request_generation_provider,
        )
        raise _institution_api_provider_http_error(
            e,
            provider=request_generation_provider,
        ) from e
    strategy = _attach_presentation_material_packet(strategy, presentation_material_packet)
    _register_question_quality_evidence(
        strategy,
        source_endpoint="/api/questions/generate-from-text",
        ncs_matches=ncs_matches,
    )
    generated_questions_all = _extract_all_generated_question_items(strategy)
    generated_questions = _build_generated_questions_payload(
        strategy,
        question_plan,
        include_all=include_all_generated_questions,
        explicit_max_items=explicit_generated_questions_max,
        fallback=run_top_k,
    )
    generated_question_text_rows, generated_question_texts = _build_generated_question_text_payload(
        strategy
    )
    generated_questions_mode = (
        "all" if include_all_generated_questions else (
            "custom_limit" if explicit_generated_questions_max is not None else "preview"
        )
    )
    generated_questions_limit = (
        explicit_generated_questions_max
        if explicit_generated_questions_max is not None
        else _generated_question_preview_limit(
            question_plan,
            strategy,
            fallback=run_top_k,
        )
    )

    _record_audit_event(
        request,
        action="manual_ncs_generate",
        resource_type="selected_ncs",
        resource_id=_sha256_text(",".join(str(x.get("ncsClCd", "")) for x in ncs_matches)),
    )

    return {
        "input_mode": "manual_text+ncs_click",
        "filename": "",
        "notice_filename": "직무내용 직접입력",
        "jd_text_preview": notice_text[:1200],
        "notice_text_preview": notice_text[:1200],
        "notice_context_preview": notice_text[:1200],
        "duty_text_preview": duty_text[:1200],
        "evaluation_text_preview": evaluation_text[:1200],
        "presentation_material_text_preview": presentation_material_text[:1200],
        "presentation_material": presentation_material_packet,
        "profile_used": False,
        "ncs_source": "manual-selected",
        "ncs_error": "",
        "question_generation": {
            **_generation_provider_descriptor(request_generation_provider),
            "resolved_model": str(
                strategy.get("provider_generation_model")
                or request_generation_model
                or generation_provider_config(request_generation_provider).get("default_model")
                or ""
            ),
        },
        "provider_key_source": _generation_key_source(
            request_generation_provider,
            generation_api_key=payload.get("generation_api_key", ""),
            openai_api_key=payload.get("openai_api_key", ""),
            openrouter_api_key=payload.get("openrouter_api_key", ""),
        ),
        "openai_key_source": (
            _openai_key_source(request_generation_api_key, request)
            if request_generation_provider == "openai_api"
            else "not_selected"
        ),
        "extracted_focus_terms": [],
        "subcategory_text_preview": "",
        "small_categories": [],
        "core_small_categories": [],
        "inferred_keywords": [],
        "reviewed_keywords": [],
        "pipeline_mode": "manual-ncs-select",
        "manual_sclass": [],
        "ai_sclass_candidates": [],
        "csv_sclass_candidates": [],
        "ai_ncs_code_candidates": [],
        "verified_sclass": [],
        "ncs_code_nos": [],
        "question_plan": question_plan,
        "interview_methods": interview_methods,
        "operational_notice": OPERATIONAL_REVIEW_NOTICE,
        "ncs_matches": ncs_matches,
        "ncs_ksa": ncs_ksa,
        "ncs_ksa_candidate_count": len(ncs_ksa_candidates),
        "ncs_ksa_queried_unit_count": len(ksa_units),
        "ncs_factor_sources": ncs_factor_sources,
        "runtime_knobs": {
            "ncs_top_k": run_top_k,
            "ksa_units": run_ksa_units,
            "ksa_factors_per_unit": run_ksa_factors,
        },
        "ncs_context": ncs_context,
        "generated_question_text_rows": generated_question_text_rows,
        "generated_question_texts": generated_question_texts,
        "strategy": strategy,
        "generated_questions_all": generated_questions_all,
        "generated_questions": generated_questions,
        "generated_questions_mode": generated_questions_mode,
        "generated_questions_limit": generated_questions_limit,
        "generated_questions_count": len(generated_questions),
        "generated_questions_total_count": len(generated_questions_all),
    }


@app.post("/api/questions/generate-personalized")
def generate_questions_personalized(
    request: Request,
    payload: dict[str, Any] | None = Body(default=None),
    ncs_code: str = Query("", description="NCS competency code (e.g., 02020302)"),
    competency_name: str = Query("", description="NCS competency unit name (optional)"),
    job_posting_query: str = Query("", alias="job_posting", include_in_schema=False),
    user_profile_query: str = Query("", alias="user_profile", include_in_schema=False),
    target_count: int = Query(1, description="Number of question templates (operational maximum 1)"),
) -> dict:
    """Generate personalized interview question templates based on job posting and user profile.

    IMPROVEMENT: Questions are TEMPLATES that incorporate job posting and user profile content.
    All questions are examples meant to guide actual interview preparation.

    JSON Body:
        generation_provider: Must be ``openai_api`` for the public endpoint.
        generation_api_key: Request-scoped key; provider is verified from its prefix.
        ncs_code: NCS competency code (required, e.g., '02020302')
        competency_name: Optional competency name for better context
        job_posting: Job posting/recruitment info text (company, position, requirements)
        user_profile: User profile/resume text (experience, skills, achievements)
        target_count: Number of question templates (currently exactly 1)

    Returns:
        - ncs_code: Input NCS code
        - competency_name: Competency unit name
        - company_from_posting: Extracted company/organization name
        - questions: Array of personalized question templates
        - note: Reminder that these are templates for adaptation
    """
    try:
        _reject_sensitive_query_params(request, destination="JSON body")
        body = payload if isinstance(payload, dict) else {}
        requested_provider = body.get("generation_provider", "")
        if str(job_posting_query or "").strip() or str(user_profile_query or "").strip():
            raise HTTPException(
                status_code=400,
                detail="job_posting and user_profile must be sent in the JSON body, not the query string",
            )
        target_count = _validated_auxiliary_generation_count(
            body.get("target_count", target_count),
            field_name="target_count",
        )
        ncs_code = _validate_generation_ncs_code(
            body.get("ncs_code") or ncs_code or "",
        )
        competency_name = _validate_generation_text_input(
            body.get("competency_name") or competency_name or "",
            field_name="competency_name",
            max_chars=_MAX_NCS_NAME_CHARS,
        )
        job_posting = _validate_generation_text_input(
            body.get("job_posting", ""),
            field_name="job_posting",
            max_chars=_MAX_NOTICE_TEXT_CHARS,
        )
        user_profile = _validate_generation_text_input(
            body.get("user_profile", ""),
            field_name="user_profile",
            max_chars=_MAX_STRENGTHS_CHARS,
        )
        (
            request_generation_provider,
            request_generation_model,
            request_generation_api_key,
        ) = _resolve_request_generation(
            generation_api_key=body.get("generation_api_key", ""),
            openai_api_key=body.get("openai_api_key", ""),
            openrouter_api_key=body.get("openrouter_api_key", ""),
            provider=requested_provider,
            generation_model=body.get("generation_model", ""),
        )
        _require_allowed_openai_key(
            request_generation_api_key,
            request,
            provider=request_generation_provider,
        )
        if not ncs_code or not ncs_code.strip():
            raise HTTPException(status_code=400, detail="ncs_code is required")

        if len(ncs_code.strip()) < 4:
            raise HTTPException(status_code=400, detail="ncs_code format invalid (e.g., 02020302)")

        _require_ncs_mcp_url()

        def _generate_once(regeneration_context: str) -> dict[str, Any]:
            effective_model = (
                openai_role_model("quality_regeneration")
                if regeneration_context and request_generation_provider == "openai_api"
                else request_generation_model
            )
            with use_generation_request(
                provider=request_generation_provider,
                generation_model=effective_model,
            ):
                return generate_personalized_interview_questions(
                    ncs_code=ncs_code.strip(),
                    competency_name=competency_name.strip() or "",
                    job_posting=job_posting.strip() or "",
                    user_profile=user_profile.strip() or "",
                    target_count=target_count,
                    extra_context=regeneration_context,
                    api_key_override=request_generation_api_key,
                    generation_model=effective_model,
                    generation_provider=request_generation_provider,
                )

        result = _quality_gate_auxiliary_question_result(
            generate_once=_generate_once,
            provider=request_generation_provider,
            api_key_override=request_generation_api_key,
            generation_model=request_generation_model,
            job_context={
                "job_description": job_posting,
                "duties": job_posting,
            },
            expected_count=target_count,
        )

        result_questions = result.get("questions") if isinstance(result, dict) else []
        generated_questions = _normalize_generated_questions(
            result_questions,
            expected_count=target_count,
        )
        generated_questions_count = len(generated_questions)
        generated_questions_total_count = len(result_questions)
        generated_question_text_rows, generated_question_texts = _build_generated_question_text_payload(
            result_questions
        )

        return {
            "status": "success",
            "data": result,
            "generated_question_text_rows": generated_question_text_rows,
            "generated_question_texts": generated_question_texts,
            "generated_questions": generated_questions,
            "generated_questions_mode": "all",
            "generated_questions_limit": target_count,
            "generated_questions_count": generated_questions_count,
            "generated_questions_total_count": generated_questions_total_count,
            "timestamp": datetime.now(timezone.utc).isoformat(),
        }

    except HTTPException:
        raise
    except Exception as e:
        logger.error("personalized_question_generation_failed provider=%s", request_generation_provider)
        raise _institution_api_provider_http_error(
            e,
            provider=request_generation_provider,
        ) from e


@app.post("/api/questions/generate-by-ncs-code")
def generate_questions_by_ncs_code(
    request: Request,
    payload: dict[str, Any] | None = Body(default=None),
    ncs_code: str = Query("", description="NCS competency code (e.g., 02020302)"),
    competency_name: str = Query("", description="NCS competency unit name (optional)"),
    target_count: int = Query(1, description="Number of questions to generate (operational maximum 1)"),
    include_followups: bool = Query(True, description="Include follow-up questions (default True)"),
) -> dict:
    """Generate interview questions using only NCS code (no job description file required).

    IMPROVEMENT: New endpoint for generating diverse interview questions directly from NCS codes.
    Supports 4 question types: behavioral, situational, technical, development-oriented.

    JSON Body:
        generation_provider: Must be ``openai_api`` for the public endpoint.
        generation_api_key: Request-scoped key; provider is verified from its prefix.
        ncs_code: NCS competency code (required, e.g. '02020302')
        competency_name: Optional competency name for context
        target_count: Number of main questions (currently exactly 1)
        include_followups: Include follow-up questions (default True)

    Returns:
        - ncs_code: Input NCS code
        - competency_name: Competency unit name
        - main_questions: Array of behavioral/situational/technical questions
        - follow_up_questions: Array of follow-up questions for depth
        - total_count: Total question count
    """
    try:
        _reject_sensitive_query_params(request, destination="JSON body")
        body = payload if isinstance(payload, dict) else {}
        requested_provider = body.get("generation_provider", "")
        target_count = _validated_auxiliary_generation_count(
            body.get("target_count", target_count),
            field_name="target_count",
        )
        ncs_code = _validate_generation_ncs_code(
            body.get("ncs_code") or ncs_code or "",
        )
        competency_name = _validate_generation_text_input(
            body.get("competency_name") or competency_name or "",
            field_name="competency_name",
            max_chars=_MAX_NCS_NAME_CHARS,
        )
        raw_avoid_questions = body.get("avoid_questions")
        avoid_field_name = "avoid_questions"
        if raw_avoid_questions is None:
            raw_avoid_questions = body.get("current_questions")
            avoid_field_name = "current_questions"
        if raw_avoid_questions is None:
            raw_avoid_questions = body.get("currentQuestions")
            avoid_field_name = "currentQuestions"
        if raw_avoid_questions is None:
            raw_avoid_questions = body.get("avoid_questions_json")
            avoid_field_name = "avoid_questions_json"
        request_avoid_questions = _validate_and_extract_generation_avoid_questions(
            raw_avoid_questions,
            field_name=avoid_field_name,
        )
        (
            request_generation_provider,
            request_generation_model,
            request_generation_api_key,
        ) = _resolve_request_generation(
            generation_api_key=body.get("generation_api_key", ""),
            openai_api_key=body.get("openai_api_key", ""),
            openrouter_api_key=body.get("openrouter_api_key", ""),
            provider=requested_provider,
            generation_model=body.get("generation_model", ""),
        )
        _require_allowed_openai_key(
            request_generation_api_key,
            request,
            provider=request_generation_provider,
        )
        if "include_followups" in body:
            raw_include_followups = body.get("include_followups")
            if not isinstance(raw_include_followups, bool):
                raise HTTPException(status_code=422, detail="include_followups must be a boolean")
            include_followups = raw_include_followups
        # Validate inputs
        if not ncs_code or not ncs_code.strip():
            raise HTTPException(status_code=400, detail="ncs_code is required")

        if len(ncs_code.strip()) < 4:
            raise HTTPException(status_code=400, detail="ncs_code format invalid (e.g., 02020302)")

        _require_ncs_mcp_url()

        avoid_questions = request_avoid_questions
        avoid_context = _build_avoid_questions_context(avoid_questions, max_items=16)

        def _generate_once(regeneration_context: str) -> dict[str, Any]:
            combined_context = "\n\n".join(
                value
                for value in (avoid_context, regeneration_context)
                if str(value or "").strip()
            )
            effective_model = (
                openai_role_model("quality_regeneration")
                if regeneration_context and request_generation_provider == "openai_api"
                else request_generation_model
            )
            with use_generation_request(
                provider=request_generation_provider,
                generation_model=effective_model,
            ):
                return generate_interview_questions_by_ncs_code(
                    ncs_code=ncs_code.strip(),
                    competency_name=competency_name.strip() or "",
                    target_count=target_count,
                    include_followups=include_followups,
                    extra_context=combined_context,
                    api_key_override=request_generation_api_key,
                    generation_model=effective_model,
                    generation_provider=request_generation_provider,
                )

        result = _quality_gate_auxiliary_question_result(
            generate_once=_generate_once,
            provider=request_generation_provider,
            api_key_override=request_generation_api_key,
            generation_model=request_generation_model,
            job_context={
                "duties": competency_name.strip() or f"NCS-{ncs_code.strip()}",
            },
            expected_count=target_count,
        )

        if str(result.get("generation_mode", "")).strip() == "ai_generation_empty_no_fallback":
            raise HTTPException(
                status_code=503,
                detail=(
                    "AI question generation failed without template fallback. "
                    "Check OpenAI network/socket permissions and retry."
                ),
            )

        _filter_ncs_code_result_against_avoid_list(result, avoid_questions)
        result_main_questions = result.get("main_questions") if isinstance(result, dict) else []
        generated_questions = _normalize_generated_questions(
            result_main_questions,
            expected_count=target_count,
        )
        if not any(
            isinstance(row, dict) and str(row.get("question") or "").strip()
            for row in result_main_questions
        ):
            # The provider result was valid before history filtering, but every
            # item may have matched the caller's avoid list. Never report a
            # successful generation with zero usable questions.
            raise RuntimeError("institution_api_question_generation_failed")
        generated_questions_count = len(generated_questions)
        generated_questions_total_count = len(result_main_questions)
        generated_question_text_rows, generated_question_texts = _build_generated_question_text_payload(
            result_main_questions
        )

        return {
            "status": "success",
            "data": result,
            "generated_question_text_rows": generated_question_text_rows,
            "generated_question_texts": generated_question_texts,
            "generated_questions": generated_questions,
            "generated_questions_mode": "all",
            "generated_questions_limit": target_count,
            "generated_questions_count": generated_questions_count,
            "generated_questions_total_count": generated_questions_total_count,
            "timestamp": datetime.now(timezone.utc).isoformat(),
        }

    except HTTPException:
        raise
    except Exception as e:
        logger.error("ncs_code_question_generation_failed provider=%s", request_generation_provider)
        raise _institution_api_provider_http_error(
            e,
            provider=request_generation_provider,
        ) from e


@app.get("/api/questions/templates")
def get_question_templates() -> dict:
    """Get available question types and templates.

    Returns information about supported question types, evaluation criteria, etc.
    """
    return {
        "question_types": [
            {
                "type": "행동기반",
                "description": "과거 경험에 기반한 행동 사례 질문",
                "count": 5,
                "key_focus": ["근거기반 판단", "문제해결", "실행력"]
            },
            {
                "type": "상황면접",
                "description": "특정 상황에서의 대응 방식을 묻는 질문",
                "count": 3,
                "key_focus": ["의사결정", "우선순위", "위기대응"]
            },
            {
                "type": "직무지식",
                "description": "직무 관련 지식과 학습 현황을 묻는 질문",
                "count": 2,
                "key_focus": ["전문성", "학습력", "이해도"]
            },
            {
                "type": "개발지향",
                "description": "개인 개발과 미래 계획에 관한 질문",
                "count": 2,
                "key_focus": ["성장의욕", "학습계획", "비전"]
            },
            {
                "type": "협업성",
                "description": "팀 협업과 상호작용에 관한 질문",
                "count": 2,
                "key_focus": ["의사소통", "협업", "갈등해결"]
            },
            {
                "type": "미래대비",
                "description": "향후 역할 수행 능력에 관한 질문",
                "count": 2,
                "key_focus": ["준비도", "전략성", "실행가능성"]
            },
            {
                "type": "성찰",
                "description": "경험 성찰과 성장에 관한 질문",
                "count": 2,
                "key_focus": ["자기인식", "학습", "개선"]
            }
        ],
        "total_templates": 18,
        "follow_up_templates": 8,
        "typical_question_count": 20,
        "estimated_interview_time_minutes": 60,
    }


@app.post("/api/questions/generate-batch")
def generate_batch_diverse_questions(
    request: Request,
    payload: dict[str, Any] | None = Body(default=None),
    ncs_code: str = Query("", description="NCS code (e.g., 0202010203_19v2)"),
    competency_name: str = Query("", description="Competency name"),
    batch_count: int = Query(1, description="Number of questions to generate (operational maximum 1)"),
) -> dict:
    """Generate one diverse interview question in the current production mode.

    IMPORTANT: AI generates completely different questions EVERY TIME - no caching!
    Each request = Fresh questions. No repetition ever.

    Format: #1, #2, #3... with question type, competency, NCS code, question, follow-up, eval points

    JSON Body:
        generation_provider: Must be ``openai_api`` for the public endpoint.
        generation_api_key: Request-scoped key; provider is verified from its prefix.
        ncs_code: NCS code (required)
        competency_name: Competency name (optional)
        batch_count: Total questions (currently exactly 1)

    Returns:
        Batch of diverse questions in numbered format
    """
    import uuid

    try:
        _reject_sensitive_query_params(request, destination="JSON body")
        body = payload if isinstance(payload, dict) else {}
        requested_provider = body.get("generation_provider", "")
        batch_count = _validated_auxiliary_generation_count(
            body.get("batch_count", batch_count),
            field_name="batch_count",
        )
        ncs_code = _validate_generation_ncs_code(
            body.get("ncs_code") or ncs_code or "",
        )
        competency_name = _validate_generation_text_input(
            body.get("competency_name") or competency_name or "",
            field_name="competency_name",
            max_chars=_MAX_NCS_NAME_CHARS,
        )
        raw_avoid_questions = body.get("avoid_questions")
        avoid_field_name = "avoid_questions"
        if raw_avoid_questions is None:
            raw_avoid_questions = body.get("current_questions")
            avoid_field_name = "current_questions"
        if raw_avoid_questions is None:
            raw_avoid_questions = body.get("currentQuestions")
            avoid_field_name = "currentQuestions"
        if raw_avoid_questions is None:
            raw_avoid_questions = body.get("avoid_questions_json")
            avoid_field_name = "avoid_questions_json"
        request_avoid_questions = _validate_and_extract_generation_avoid_questions(
            raw_avoid_questions,
            field_name=avoid_field_name,
        )
        (
            request_generation_provider,
            request_generation_model,
            request_generation_api_key,
        ) = _resolve_request_generation(
            generation_api_key=body.get("generation_api_key", ""),
            openai_api_key=body.get("openai_api_key", ""),
            openrouter_api_key=body.get("openrouter_api_key", ""),
            provider=requested_provider,
            generation_model=body.get("generation_model", ""),
        )
        _require_allowed_openai_key(
            request_generation_api_key,
            request,
            provider=request_generation_provider,
        )
        if not ncs_code or not ncs_code.strip():
            raise HTTPException(status_code=400, detail="ncs_code required")

        _require_ncs_mcp_url()

        # Generate multiple rounds of diverse questions with strict deduplication
        final_questions = []
        seen_questions = set()
        seen_question_texts = []
        final_official_ksa_evidence: dict[str, dict[str, Any]] = {}
        avoid_keys = {
            normalize_question_dedup_key(q)
            for q in request_avoid_questions
            if normalize_question_dedup_key(q)
        }
        max_attempts = 2
        attempt = 0
        final_ai_quality_review: dict[str, Any] = {}

        while len(final_questions) < batch_count and attempt < max_attempts:
            attempt += 1
            avoid_context = _build_avoid_questions_context(
                [*request_avoid_questions, *seen_question_texts],
                max_items=16,
            )
            def _generate_once(regeneration_context: str) -> dict[str, Any]:
                combined_context = "\n\n".join(
                    value
                    for value in (avoid_context, regeneration_context)
                    if str(value or "").strip()
                )
                effective_model = (
                    openai_role_model("quality_regeneration")
                    if regeneration_context and request_generation_provider == "openai_api"
                    else request_generation_model
                )
                with use_generation_request(
                    provider=request_generation_provider,
                    generation_model=effective_model,
                ):
                    return generate_diverse_interview_questions(
                        ncs_code=ncs_code.strip(),
                        competency_name=competency_name.strip() or "",
                        target_count=batch_count,
                        extra_context=combined_context,
                        api_key_override=request_generation_api_key,
                        generation_model=effective_model,
                        generation_provider=request_generation_provider,
                    )

            result = _quality_gate_auxiliary_question_result(
                generate_once=_generate_once,
                provider=request_generation_provider,
                api_key_override=request_generation_api_key,
                generation_model=request_generation_model,
                job_context={
                    "duties": competency_name.strip() or f"NCS-{ncs_code.strip()}",
                },
                expected_count=batch_count,
            )
            final_ai_quality_review = dict(result.get("ai_quality_review") or {})
            for evidence_row in result.get("official_ksa_evidence", []):
                if not isinstance(evidence_row, dict):
                    continue
                evidence_id = str(evidence_row.get("evidence_id") or "").strip()
                if evidence_id:
                    final_official_ksa_evidence[evidence_id] = dict(evidence_row)

            for q in result["questions"]:
                if len(final_questions) >= batch_count:
                    break

                # Normalize question key and block exact duplicates.
                q_text = str(q.get("question", "")).strip()
                q_key = normalize_question_dedup_key(q_text)
                if not q_key:
                    continue

                if q_key in seen_questions or q_key in avoid_keys:
                    continue

                # Block near-duplicate questions with minor wording changes.
                if any(is_similar_question_text(q_text, prev) for prev in seen_question_texts):
                    continue
                if any(is_similar_question_text(q_text, prev) for prev in request_avoid_questions):
                    continue

                final_questions.append(q)
                seen_questions.add(q_key)
                seen_question_texts.append(q_text)
        if len(final_questions) != batch_count:
            raise InstitutionQuestionQualityRejected(
                {
                    "requested_question_count": batch_count,
                    "failed_question_count": batch_count - len(final_questions),
                    "attempt_count": attempt,
                    "failure_scope": "question_deduplication",
                }
            )
        _refresh_question_repeat_metadata(final_questions)
        for i, q in enumerate(final_questions, 1):
            q["number"] = i
        _require_official_ksa_result(
            {
                "ncs_ksa_available": True,
                "questions": final_questions,
                "official_ksa_evidence": list(final_official_ksa_evidence.values()),
            }
        )

        response_data = {
            "status": "success",
            "data": {
                "ncs_code": ncs_code,
                "competency_name": competency_name or f"NCS-{ncs_code}",
                "ncs_ksa_available": True,
                "batch_count": len(final_questions),
                "questions": final_questions,
                "official_ksa_evidence": list(final_official_ksa_evidence.values()),
                "ai_quality_review": final_ai_quality_review,
                "question_release_status": "ai_quality_review_passed",
                "note": "각 질문은 AI가 생성한 고유한 질문입니다. 매 요청마다 다른 질문이 생성됩니다.",
            },
            "timestamp": datetime.now(timezone.utc).isoformat(),
            "request_id": str(uuid.uuid4()),  # Unique ID to prevent caching
        }
        batch_questions = response_data.get("data", {}).get("questions", [])
        normalized_batch_questions = _normalize_generated_questions(
            batch_questions,
            expected_count=batch_count,
        )
        response_data["generated_questions"] = normalized_batch_questions
        response_data["generated_questions_mode"] = "all"
        response_data["generated_questions_limit"] = batch_count
        response_data["generated_questions_count"] = len(normalized_batch_questions)
        response_data["generated_questions_total_count"] = len(batch_questions)
        response_data["generated_question_text_rows"], response_data["generated_question_texts"] = (
            _build_generated_question_text_payload(batch_questions)
        )

        # Return with NO-CACHE headers
        return JSONResponse(
            content=response_data,
            headers={
                "Cache-Control": "no-store, no-cache, must-revalidate, max-age=0",
                "Pragma": "no-cache",
                "Expires": "0",
                "X-Content-Type-Options": "nosniff",
            }
        )

    except HTTPException:
        raise
    except Exception as e:
        logger.error("batch_question_generation_failed provider=%s", request_generation_provider)
        raise _institution_api_provider_http_error(
            e,
            provider=request_generation_provider,
        ) from e


@app.post("/api/questions/generate-diverse")
def generate_diverse_questions(
    request: Request,
    payload: dict[str, Any] | None = Body(default=None),
    ncs_code: str = Query("", description="NCS competency code (e.g., 0202010102_19v2)"),
    competency_name: str = Query("", description="Competency unit name (optional)"),
    job_posting_query: str = Query("", alias="job_posting", include_in_schema=False),
    target_count: int = Query(1, description="Number of diverse question types (operational maximum 1)"),
) -> dict:
    """Generate one diverse interview question in the current production mode.

    IMPROVEMENT: Creates highly varied questions with 6 different formats:
    1. STAR (행동기반) - Specific past success with STAR structure
    2. BEI (행동사건) - Most difficult moment experienced
    3. 케이스 (Case analysis) - Problem-solving scenario
    4. SJT (Situational) - Decision-making under pressure
    5. 압박면접 (Pressure) - Handling failure/criticism
    6. 비판적사건 (Critical incident) - Learning moment/awareness change

    Each question includes:
    - Question number and type
    - Main question
    - Follow-up question
    - Evaluation points (역량 평가 포인트)

    JSON Body:
        generation_provider: Must be ``openai_api`` for the public endpoint.
        generation_api_key: Request-scoped key; provider is verified from its prefix.
        ncs_code: NCS competency code (required, e.g., '0202010102_19v2')
        competency_name: Competency unit name (optional)
        job_posting: Job posting context text (optional)
        target_count: How many diverse types (currently exactly 1)

    Returns:
        One diverse question with an evaluation angle
    """
    try:
        _reject_sensitive_query_params(request, destination="JSON body")
        body = payload if isinstance(payload, dict) else {}
        requested_provider = body.get("generation_provider", "")
        if str(job_posting_query or "").strip():
            raise HTTPException(
                status_code=400,
                detail="job_posting must be sent in the JSON body, not the query string",
            )
        target_count = _validated_auxiliary_generation_count(
            body.get("target_count", target_count),
            field_name="target_count",
        )
        ncs_code = _validate_generation_ncs_code(
            body.get("ncs_code") or ncs_code or "",
        )
        competency_name = _validate_generation_text_input(
            body.get("competency_name") or competency_name or "",
            field_name="competency_name",
            max_chars=_MAX_NCS_NAME_CHARS,
        )
        job_posting = _validate_generation_text_input(
            body.get("job_posting", ""),
            field_name="job_posting",
            max_chars=_MAX_NOTICE_TEXT_CHARS,
        )
        raw_avoid_questions = body.get("avoid_questions")
        avoid_field_name = "avoid_questions"
        if raw_avoid_questions is None:
            raw_avoid_questions = body.get("current_questions")
            avoid_field_name = "current_questions"
        if raw_avoid_questions is None:
            raw_avoid_questions = body.get("currentQuestions")
            avoid_field_name = "currentQuestions"
        if raw_avoid_questions is None:
            raw_avoid_questions = body.get("avoid_questions_json")
            avoid_field_name = "avoid_questions_json"
        request_avoid_questions = _validate_and_extract_generation_avoid_questions(
            raw_avoid_questions,
            field_name=avoid_field_name,
        )
        (
            request_generation_provider,
            request_generation_model,
            request_generation_api_key,
        ) = _resolve_request_generation(
            generation_api_key=body.get("generation_api_key", ""),
            openai_api_key=body.get("openai_api_key", ""),
            openrouter_api_key=body.get("openrouter_api_key", ""),
            provider=requested_provider,
            generation_model=body.get("generation_model", ""),
        )
        _require_allowed_openai_key(
            request_generation_api_key,
            request,
            provider=request_generation_provider,
        )
        if not ncs_code or not ncs_code.strip():
            raise HTTPException(status_code=400, detail="ncs_code is required")

        _require_ncs_mcp_url()

        ncs_code_clean = ncs_code.strip()
        competency_name_clean = competency_name.strip()
        avoid_keys = {
            normalize_question_dedup_key(q)
            for q in request_avoid_questions
            if normalize_question_dedup_key(q)
        }

        final_questions = []
        seen_keys = set()
        seen_texts = []
        final_official_ksa_evidence: dict[str, dict[str, Any]] = {}
        max_attempts = 2
        attempt = 0
        raw_result = {
            "ncs_code": ncs_code_clean,
            "competency_name": competency_name_clean or f"NCS-{ncs_code_clean}",
            "generation_mode": "ai_powered_diverse",
        }

        while len(final_questions) < target_count and attempt < max_attempts:
            attempt += 1
            needed = min(6, max(target_count - len(final_questions), 1))
            avoid_context = _build_avoid_questions_context(
                [*request_avoid_questions, *seen_texts],
                max_items=16,
            )
            def _generate_once(regeneration_context: str) -> dict[str, Any]:
                combined_context = "\n\n".join(
                    value
                    for value in (avoid_context, regeneration_context)
                    if str(value or "").strip()
                )
                effective_model = (
                    openai_role_model("quality_regeneration")
                    if regeneration_context and request_generation_provider == "openai_api"
                    else request_generation_model
                )
                with use_generation_request(
                    provider=request_generation_provider,
                    generation_model=effective_model,
                ):
                    return generate_diverse_interview_questions(
                        ncs_code=ncs_code_clean,
                        competency_name=competency_name_clean or "",
                        job_posting=job_posting.strip() or "",
                        target_count=needed,
                        extra_context=combined_context,
                        api_key_override=request_generation_api_key,
                        generation_model=effective_model,
                        generation_provider=request_generation_provider,
                    )

            raw_result = _quality_gate_auxiliary_question_result(
                generate_once=_generate_once,
                provider=request_generation_provider,
                api_key_override=request_generation_api_key,
                generation_model=request_generation_model,
                job_context={
                    "job_description": job_posting,
                    "duties": job_posting or competency_name_clean or f"NCS-{ncs_code_clean}",
                },
                expected_count=needed,
            )
            for evidence_row in raw_result.get("official_ksa_evidence", []):
                if not isinstance(evidence_row, dict):
                    continue
                evidence_id = str(evidence_row.get("evidence_id") or "").strip()
                if evidence_id:
                    final_official_ksa_evidence[evidence_id] = dict(evidence_row)

            for q in raw_result.get("questions", []):
                if len(final_questions) >= target_count:
                    break
                q_text = str(q.get("question", "")).strip()
                q_key = normalize_question_dedup_key(q_text)
                if not q_key:
                    continue
                if q_key in seen_keys or q_key in avoid_keys:
                    continue
                if any(is_similar_question_text(q_text, prev) for prev in seen_texts):
                    continue
                if any(is_similar_question_text(q_text, prev) for prev in request_avoid_questions):
                    continue
                final_questions.append(q)
                seen_keys.add(q_key)
                seen_texts.append(q_text)

        if len(final_questions) != target_count:
            raise InstitutionQuestionQualityRejected(
                {
                    "requested_question_count": target_count,
                    "failed_question_count": target_count - len(final_questions),
                    "attempt_count": attempt,
                    "failure_scope": "question_deduplication",
                }
            )
        _refresh_question_repeat_metadata(final_questions)
        for i, q in enumerate(final_questions, 1):
            q["number"] = i

        result = {
            "ncs_code": raw_result.get("ncs_code", ncs_code_clean),
            "competency_name": raw_result.get("competency_name", competency_name_clean or f"NCS-{ncs_code_clean}"),
            "generation_mode": raw_result.get("generation_mode", "ai_powered_diverse"),
            "ncs_ksa_available": True,
            "questions": final_questions,
            "question_count": len(final_questions),
            "official_ksa_evidence": list(final_official_ksa_evidence.values()),
            "ai_quality_review": dict(raw_result.get("ai_quality_review") or {}),
            "question_release_status": "ai_quality_review_passed",
            "note": "동일/유사 질문을 제거한 결과입니다. 매 요청마다 새 질문을 우선 생성합니다.",
        }
        _require_official_ksa_result(
            {
                "generation_mode": result.get("generation_mode", "ai_autonomous_ncs"),
                "ncs_ksa_available": True,
                "questions": final_questions,
                "official_ksa_evidence": list(final_official_ksa_evidence.values()),
            }
        )
        diverse_questions = result.get("questions") if isinstance(result, dict) else []
        generated_questions = _normalize_generated_questions(
            diverse_questions,
            expected_count=target_count,
        )
        generated_questions_count = len(generated_questions)
        generated_questions_total_count = len(diverse_questions)
        generated_question_text_rows, generated_question_texts = _build_generated_question_text_payload(
            diverse_questions
        )

        return {
            "status": "success",
            "data": result,
            "generated_question_text_rows": generated_question_text_rows,
            "generated_question_texts": generated_question_texts,
            "generated_questions": generated_questions,
            "generated_questions_mode": "all",
            "generated_questions_limit": target_count,
            "generated_questions_count": generated_questions_count,
            "generated_questions_total_count": generated_questions_total_count,
            "timestamp": datetime.now(timezone.utc).isoformat(),
        }

    except HTTPException:
        raise
    except Exception as e:
        logger.error("diverse_question_generation_failed provider=%s", request_generation_provider)
        raise _institution_api_provider_http_error(
            e,
            provider=request_generation_provider,
        ) from e
